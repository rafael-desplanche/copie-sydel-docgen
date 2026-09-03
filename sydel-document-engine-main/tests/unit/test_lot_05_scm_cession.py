from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    DocumentGenerationContext,
    DossierOptions,
    Person,
    ScmCessionAgrement,
    ScmCessionAssocie,
    ScmCessionCedant,
    ScmCessionConjoint,
    ScmCessionContext,
    ScmCessionCreditVendeur,
    ScmCessionEnregistrement,
    ScmCessionOrdre,
    ScmCessionPartsAttribution,
    ScmCessionPartsCedees,
    ScmCessionPrix,
    ScmCessionRepresentant,
    ScmCessionSignataire,
    ScmCessionSociete,
    Signature,
)
from sydel_doc_engine.generators.lot_05.acte_cession_parts_scm import (
    ActeCessionPartsScmGenerator,
)
from sydel_doc_engine.generators.lot_05.courrier_sde_cession_scm import (
    CourrierSdeCessionScmGenerator,
)
from sydel_doc_engine.generators.lot_05.pv_age_cession_scm import PvAgeCessionScmGenerator
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text
    ]
    return "\n".join(paragraphs + cells)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text
    assert "Ajouter en cas de CV" not in text


def _associe(
    name: str,
    parts: int,
    plage: str,
    *,
    morale: bool = False,
) -> ScmCessionAssocie:
    if morale:
        return ScmCessionAssocie(
            type_personne="personne_morale",
            denomination=name,
            forme_juridique="SELARL",
            parts=ScmCessionPartsAttribution(nb=parts, plage=plage),
        )
    prenom, nom = name.split(" ", 1)
    return ScmCessionAssocie(
        civilite_affichage="Monsieur" if prenom != "Anne" else "Madame",
        prenom=prenom,
        nom=nom,
        parts=ScmCessionPartsAttribution(nb=parts, plage=plage),
    )


def _base_context(structure: str = "SELARL") -> DocumentGenerationContext:
    variant = structure.lower()
    is_selas = structure == "SELAS"
    cessionnaire_forme = "SELAS" if is_selas else "SELARL"
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(scm_cession=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Dupont",
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 15),
            prestataire_signature_electronique="DocuSign",
        ),
        scm_cession=ScmCessionContext(
            variante_structure=variant,
            scm_cedee=ScmCessionSociete(
                denomination="SCM CABINET CENTRAL",
                forme_juridique="Société Civile de Moyens",
                capital_social="3 000",
                siege=Address(adresse_affichee="12 rue des Soins, 75008 Paris"),
                ville_rcs="Paris",
                numero_rcs="900 111 222",
                nb_parts_total=300,
                valeur_nominale_part="10",
                plage_parts_total="1 à 300",
                cogerants=["Monsieur Paul Bernard", "Monsieur Jean Dupont", "Madame Anne Martin"],
            ),
            cessionnaire=ScmCessionSociete(
                denomination=f"{cessionnaire_forme} CABINET DUPONT",
                forme_juridique=cessionnaire_forme,
                capital_social="10 000",
                siege=Address(adresse_affichee="20 avenue des Praticiens, 75008 Paris"),
                ville_rcs="Paris",
                representant=ScmCessionRepresentant(
                    civilite_affichage="Monsieur",
                    civilite_courte="M.",
                    prenom="Jean",
                    nom="Dupont",
                    fonction="président" if is_selas else "gérant",
                ),
            ),
            cedant=ScmCessionCedant(
                civilite_affichage="Monsieur",
                prenom="Jean",
                nom="Dupont",
                profession="chirurgien-dentiste",
                profession_reglementee_pluriel="chirurgiens-dentistes",
                date_naissance="1er janvier 1980",
                ville_naissance="Paris",
                departement_naissance="75",
                nationalite="française",
                adresse_affichee="1 rue du Cédant, 75008 Paris",
                situation_maritale="marié",
                ordre=ScmCessionOrdre(departemental="Paris", numero="12345"),
                numero_rpps="10000000001",
                conjoint=ScmCessionConjoint(
                    civilite_affichage="Madame",
                    prenom="Claire",
                    nom="Dupont",
                ),
            ),
            agrement=ScmCessionAgrement(
                date_pv="15 mai 2026",
                date_pv_lettres="deux mille vingt-six, le quinze mai",
                delai_mois="3" if is_selas else None,
                date_limite="15 août 2026" if is_selas else None,
            ),
            associes_presents=[
                _associe("Paul Bernard", 100, "1 à 100"),
                _associe("Jean Dupont", 100, "101 à 200"),
                _associe("Anne Martin", 100, "201 à 300"),
            ],
            associes_avant_cession=[
                _associe("Paul Bernard", 100, "1 à 100"),
                _associe("Jean Dupont", 100, "101 à 200"),
                _associe("Anne Martin", 100, "201 à 300"),
            ],
            associes_apres_cession=[
                _associe("Paul Bernard", 100, "1 à 100"),
                _associe("Jean Dupont", 50, "101 à 150"),
                _associe(f"{cessionnaire_forme} CABINET DUPONT", 50, "151 à 200", morale=True),
                _associe("Anne Martin", 100, "201 à 300"),
            ],
            signataires_pv=["M. Jean Dupont", "M. Paul Bernard", "Mme Anne Martin"],
            parts_cedees=ScmCessionPartsCedees(nb=50, plage="151 à 200"),
            prix=ScmCessionPrix(
                unitaire="100",
                unitaire_lettres="cent",
                global_="5 000",
                global_lettres="cinq mille",
            ),
            paiement_mode="pret_bancaire",
            credit_vendeur=ScmCessionCreditVendeur(actif=False),
            enregistrement=ScmCessionEnregistrement(
                service="SERVICE DEPARTEMENTAL DE L'ENREGISTREMENT",
                centre_finances_publiques="Centre des finances publiques de Paris",
                adresse_service="6 rue Paganini",
                cp_ville_service="75020 Paris",
                nombre_exemplaires="3",
                montant_droits="150",
            ),
            signataire_sde=ScmCessionSignataire(prenom="Sarah", nom="Durand"),
            nombre_exemplaires_lettres="quatre",
            prestataire_signature_electronique="DocuSign",
            date_acte_affichee="15 mai 2026",
            representant_cessionnaire_confirme=True,
        ),
    )


def test_scm_cession_selarl_generates_three_clean_docx(tmp_path: Path) -> None:
    ctx = _base_context("SELARL")

    outputs = [
        PvAgeCessionScmGenerator().generate(ctx, tmp_path),
        CourrierSdeCessionScmGenerator().generate(ctx, tmp_path),
        ActeCessionPartsScmGenerator().generate(ctx, tmp_path),
    ]
    texts = {path.name: _docx_text(path) for path in outputs}

    assert {path.name for path in outputs} == {
        "pv_age_cession_parts_scm.docx",
        "courrier_sde_cession_scm.docx",
        "acte_cession_parts_scm.docx",
    }
    assert "à compter de ce jour" in texts["pv_age_cession_parts_scm.docx"]
    assert "dans un délai de" not in texts["pv_age_cession_parts_scm.docx"]
    assert "4 exemplaires" in texts["courrier_sde_cession_scm.docx"]
    assert "SERVICE DEPARTEMENTAL" not in texts["courrier_sde_cession_scm.docx"]
    assert "chirurgiens-dentistes" in texts["acte_cession_parts_scm.docx"]
    assert "Yousign" in texts["acte_cession_parts_scm.docx"]
    for text in texts.values():
        _assert_clean(text)


def test_scm_cession_selas_generates_overlays(tmp_path: Path) -> None:
    ctx = _base_context("SELAS")

    pv_path = PvAgeCessionScmGenerator().generate(ctx, tmp_path)
    courrier_path = CourrierSdeCessionScmGenerator().generate(ctx, tmp_path)
    acte_path = ActeCessionPartsScmGenerator().generate(ctx, tmp_path)

    pv_text = _docx_text(pv_path)
    courrier_text = _docx_text(courrier_path)
    acte_text = _docx_text(acte_path)
    assert "dans un délai de 3 mois" in pv_text
    assert "15 août 2026" in pv_text
    assert "SERVICE DEPARTEMENTAL DE L'ENREGISTREMENT" in courrier_text
    assert "3 exemplaires" in courrier_text
    assert "SELAS au capital de 10 000" in acte_text
    assert "président" in acte_text
    assert "DocuSign" in acte_text
    for text in [pv_text, courrier_text, acte_text]:
        _assert_clean(text)


def test_orchestrator_selects_scm_cession_block_only_when_enabled() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())
    enabled = orchestrator.select_documents_for_context(_base_context("SELARL"))
    enabled_ids = [document.doc_id for document in enabled]

    disabled_ctx = _base_context("SELARL")
    disabled_ctx.dossier_options = DossierOptions(scm_cession=False)
    disabled = orchestrator.select_documents_for_context(disabled_ctx)
    disabled_ids = [document.doc_id for document in disabled]

    assert {"DOC-031", "DOC-032", "DOC-033"}.issubset(enabled_ids)
    assert {"DOC-031", "DOC-032", "DOC-033"}.isdisjoint(disabled_ids)


def test_pv_blocks_incoherent_parts_after_cession(tmp_path: Path) -> None:
    ctx = _base_context("SELARL")
    ctx.scm_cession.associes_apres_cession[0].parts.nb = 99

    with pytest.raises(ValueError, match="totaliser"):
        PvAgeCessionScmGenerator().generate(ctx, tmp_path)


def test_pv_blocks_missing_explicit_roles(tmp_path: Path) -> None:
    ctx = _base_context("SELARL")
    ctx.scm_cession.associes_apres_cession.pop()

    with pytest.raises(ValueError, match="exactement 4 associes"):
        PvAgeCessionScmGenerator().generate(ctx, tmp_path)


def test_acte_blocks_incomplete_credit_vendeur(tmp_path: Path) -> None:
    ctx = _base_context("SELAS")
    ctx.scm_cession.credit_vendeur = ScmCessionCreditVendeur(
        actif=True,
        montant="1 000",
        duree="2 ans",
        taux=None,
        majoration_interet_retard="4 points",
    )

    with pytest.raises(ValueError, match="credit_vendeur.taux"):
        ActeCessionPartsScmGenerator().generate(ctx, tmp_path)


def test_pv_age_renders_hyphen_bullets_on_two_lists(tmp_path: Path) -> None:
    # Retour UAT Rafael (DOC-031) : les 2 listes (documents deposes + ordre du
    # jour) doivent etre rendues en puces tiret « - ... », texte inchange.
    ctx = _base_context("SELARL")
    document = Document(PvAgeCessionScmGenerator().generate(ctx, tmp_path))
    para_texts = [p.text for p in document.paragraphs]

    # Liste A (documents deposes) : chaque item prefixe d'un tiret.
    assert "- Les copies des convocations des associés ;" in para_texts
    assert "- Un exemplaire du compromis de cession des parts sociales ;" in para_texts
    assert "- Le rapport de la gérance ;" in para_texts
    assert "- Le texte des résolutions proposées." in para_texts
    # Liste B (ordre du jour) : chaque item prefixe d'un tiret.
    assert "- Lecture du rapport de la gérance ;" in para_texts
    assert "- Agrément d'un nouvel associé, la SELARL CABINET DUPONT ;" in para_texts
    assert "- Modification corrélative des statuts." in para_texts
    # Garde-fou : les phrases hors liste ne sont PAS transformees en puces.
    assert (
        "Le Président dépose et met à la disposition des associés les documents suivants :"
        in para_texts
    )


def test_courrier_sde_objet_bold_underline_and_signataire_right(tmp_path: Path) -> None:
    # Retour UAT Rafael (DOC-032) : objet en gras + souligne, signataire a droite.
    ctx = _base_context("SELARL")
    document = Document(CourrierSdeCessionScmGenerator().generate(ctx, tmp_path))

    objet = next(p for p in document.paragraphs if p.text.startswith("Objet :"))
    assert objet.runs[0].bold is True
    assert objet.runs[0].underline is True

    signataire = next(p for p in document.paragraphs if p.text == "Sarah Durand")
    assert signataire.alignment == WD_ALIGN_PARAGRAPH.RIGHT
