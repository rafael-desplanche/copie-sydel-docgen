from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Associe,
    BienImmobilier,
    CapitalContext,
    Company,
    DecisionContext,
    DirigeantNomine,
    DocumentGenerationContext,
    Emprunt,
    Person,
    ReunionContext,
    ReunionPresident,
    Signature,
)
from sydel_doc_engine.generators.lot_02.pv_nomination_gerant import (
    VOTE_FORMULA,
    PvNominationGerantGenerator,
)


def _associes(count: int = 2) -> list[Associe]:
    if count == 1:
        return [
            Associe(
                genre=Gender.FEMININ,
                civilite_affichage="Madame",
                prenom="Alice",
                nom="Durand",
                nb_parts=1,
            )
        ]
    return [
        Associe(
            genre=Gender.FEMININ,
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Durand",
            nb_parts=60,
        ),
        Associe(
            genre=Gender.MASCULIN,
            civilite_affichage="Monsieur",
            prenom="Bruno",
            nom="Martin",
            nb_parts=40,
        ),
    ]


def _context(
    *,
    associes: list[Associe] | None = None,
    emprunt_actif: bool = False,
) -> DocumentGenerationContext:
    associes = associes or _associes()
    return DocumentGenerationContext(
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Signataire",
        ),
        societe=Company(
            forme_sociale_affichage="Société civile immobilière",
            forme_sociale_libelle_long="société civile immobilière",
            denomination="SCI TEST",
            capital_social="1 000",
            capital_variable=True,
            siege=Address(
                num_voie="10",
                voie="rue du Siège",
                cp="75001",
                ville="Paris",
            ),
            ville_rcs="Paris",
        ),
        decision=DecisionContext(date="13 mai 2026"),
        reunion=ReunionContext(
            date_lettres="treize mai deux mille vingt-six",
            president=ReunionPresident(
                civilite_president_seance="Madame",
                prenom_president_seance="Alice",
                nom_personne_seance="Durand",
            ),
        ),
        capital=CapitalContext(
            nb_parts_total=sum(associe.nb_parts for associe in associes),
            valeur_nominale_part="1",
        ),
        associes=associes,
        dirigeant_nomine=DirigeantNomine(
            genre=Gender.FEMININ,
            civilite_affichage="Madame",
            prenom="Claire",
            nom="Bernard",
            date_naissance=date(1985, 4, 3),
            ville_naissance="Lyon",
            departement_naissance="Rhône",
            nationalite="française",
            adresse_personnelle=Address(
                num_voie="22",
                voie="avenue des Fleurs",
                cp="69002",
                ville="Lyon",
            ),
            fonction_affichage="gérant",
        ),
        emprunt=Emprunt(
            actif=emprunt_actif,
            montant_max="250 000" if emprunt_actif else None,
        ),
        bien_immobilier=(
            BienImmobilier(
                adresse=Address(
                    num_voie="5",
                    voie="rue du Bien",
                    cp="33000",
                    ville="Bordeaux",
                )
            )
            if emprunt_actif
            else None
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 13),
            nombre_exemplaires="3",
        ),
    )


def _generate(tmp_path: Path, ctx: DocumentGenerationContext | None = None) -> Path:
    return PvNominationGerantGenerator().generate(ctx or _context(), tmp_path)


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text]


def _find_paragraph(document: Document, text: str):
    return next(paragraph for paragraph in document.paragraphs if paragraph.text == text)


def test_pv_nomination_gerant_creates_docx(tmp_path: Path) -> None:
    output_path = _generate(tmp_path)

    assert output_path == tmp_path / "pv_nomination_gerant.docx"
    assert output_path.is_file()


def test_pv_nomination_gerant_selarl_header_uses_written_form_and_simple_capital(
    tmp_path: Path,
) -> None:
    ctx = _context(associes=_associes(1))
    ctx.societe.forme_sociale = "SELARL"
    ctx.societe.forme_sociale_affichage = "SELARL"
    ctx.societe.forme_sociale_complete = "société d’exercice libéral à responsabilité limitée"
    ctx.societe.forme_sociale_abregee = "SELARL"
    ctx.societe.denomination = "SELARL MARTIN"
    ctx.societe.capital_social = "5 000"
    ctx.associes[0].profession_reglementee = "médecin"

    text = _docx_text(_generate(tmp_path, ctx))
    paragraphs = _paragraphs(_generate(tmp_path / "second", ctx))

    assert "SELARL MARTIN" in paragraphs
    assert "Société d’exercice libéral à responsabilité limitée de médecin" in paragraphs
    assert "Au capital de 5 000 euros" in paragraphs
    assert "SELARL à capital variable" not in text
    assert "Au capital minimum et effectif" not in text


def test_pv_nomination_gerant_repeats_two_associes(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))
    paragraphs = _paragraphs(_generate(tmp_path / "second"))

    assert "Les associés de la Société civile immobilière SCI TEST" in text
    assert "composé de 100 parts de 1 euro chacune, se sont réunis au siège social." in text
    assert "Sont présents ou représentés :" in text
    assert "Madame Alice Durand, détenant 60 parts," in text
    assert "Monsieur Bruno Martin, détenant 40 parts," in text
    assert "- Madame Alice Durand, détenant 60 parts," in paragraphs
    assert "- Monsieur Bruno Martin, détenant 40 parts," in paragraphs
    assert (
        "Les associés présents ou représentés disposent ensemble de la totalité des parts "
        "sociales. Cet ensemble est habilité à prendre des décisions."
    ) in text
    assert "Madame Alice Durand préside la séance." in text
    assert "Le président rappelle l’ordre du jour :" in text
    assert "· Nomination du gérant" in text
    assert "· Pouvoirs" in text
    assert "RCS de Paris" not in text
    assert "En cours d’immatriculation" in text
    assert "EXTRAORDINAIRE" not in text
    assert "extraordinaire" not in text
    assert "10 heures" not in text
    assert "De tout ce qui a été décidé" not in text
    assert "L’ordre du jour étant épuisé" not in text
    assert "Alice Durand" in paragraphs
    assert "Bruno Martin" in paragraphs


def test_pv_nomination_gerant_repeats_one_associe_with_singular_variants(
    tmp_path: Path,
) -> None:
    ctx = _context(associes=_associes(1))
    text = _docx_text(_generate(tmp_path, ctx))

    assert "Les associés de la Société civile immobilière SCI TEST" in text
    assert "se sont réunis au siège social." in text
    assert "Madame Alice Durand, détenant 1 part," in text
    assert (
        "Les associés présents ou représentés disposent ensemble de la totalité des parts "
        "sociales. Cet ensemble est habilité à prendre des décisions."
    ) in text
    assert "· Nomination du gérant" in text
    assert "- Madame Alice Durand, détenant 1 part," in text


def test_pv_nomination_gerant_without_emprunt_omits_borrowing_decision(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path, _context(emprunt_actif=False)))

    assert "Autorisation de  contracter un emprunt" not in text
    assert "contracter un emprunt d’un montant maximum" not in text
    assert "DEUXIEME DECISION" in text
    assert "TROISIEME DECISION" not in text


def test_pv_nomination_gerant_with_emprunt_writes_borrowing_decision(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path, _context(emprunt_actif=True)))

    assert (
        "· Autorisation de contracter un emprunt pour l’achat d’un bien immobilier sis "
        "5 rue du Bien, 33000 Bordeaux"
    ) in text
    assert (
        "L’assemblée générale décide de contracter un emprunt d’un montant "
        "maximum de 250 000 euros pour l’acquisition d’un bien immobilier sis "
        "5 rue du Bien, 33000 Bordeaux."
    ) in text
    assert "TROISIEME DECISION" in text


def test_pv_nomination_gerant_uses_plural_agenda_for_plural_function(
    tmp_path: Path,
) -> None:
    ctx = _context()
    ctx.dirigeant_nomine.fonction_affichage = "gérants"

    text = _docx_text(_generate(tmp_path, ctx))

    assert "· Nomination des premiers gérants" in text


def test_pv_nomination_gerant_uses_distinct_dirigeant_nomine(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "Madame Claire Bernard, née le 03/04/1985 à Lyon (Rhône)" in text
    assert "demeurant au 22 avenue des Fleurs, 69002 Lyon." in text
    assert "Madame Alice Durand, née le" not in text
    assert "Monsieur Bruno Martin, né le" not in text


def test_pv_nomination_gerant_uses_feminine_birth_variant(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "née le 03/04/1985" in text
    assert "né le 03/04/1985" not in text


def test_pv_nomination_gerant_restores_essential_docx_structure(tmp_path: Path) -> None:
    document = Document(_generate(tmp_path, _context(emprunt_actif=True)))

    company_name = _find_paragraph(document, "SCI TEST")
    assert company_name.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert company_name.runs[0].bold is True

    title_paragraph = document.tables[0].cell(0, 0).paragraphs[0]
    assert title_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert title_paragraph.text == (
        "PROCES-VERBAL DES DECISIONS\n"
        " DE L’ASSEMBLEE GENERALE\n"
        " DU 13 mai 2026"
    )
    assert all(run.bold for run in title_paragraph.runs if run.text.strip())

    first_decision = _find_paragraph(document, "PREMIERE DECISION")
    assert first_decision.runs[0].bold is True
    assert first_decision.runs[0].underline is True
    assert first_decision.paragraph_format.space_before is not None

    vote_formula = _find_paragraph(document, VOTE_FORMULA)
    assert vote_formula.runs[0].italic is True

    decision_item = _find_paragraph(document, "· Nomination du gérant")
    assert decision_item.text == "· Nomination du gérant"

    signature_name = _find_paragraph(document, "Alice Durand")
    assert signature_name.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert signature_name.runs[0].bold is True

    acceptance = _find_paragraph(
        document,
        (
            "Faire précéder la signature de la mention "
            "« Bon pour acceptation des fonctions de gérant »"
        ),
    )
    assert acceptance.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert acceptance.runs[0].italic is True
