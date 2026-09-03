from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from sydel_doc_engine.rendering.docx_builder import (
    BAIL_COMPACT_STYLE_PROFILE,
    DEROGATION_CUMUL_STYLE_PROFILE,
    LETTER_WIDE_STYLE_PROFILE,
    add_article_heading,
    add_bordered_data_table,
    add_centered_amount,
    add_checkbox_line,
    add_company_identity_block,
    add_form_section_heading,
    add_framed_section_title,
    add_framed_signature_block,
    add_framed_title,
    add_hyphen_list_item,
    add_italic_instruction,
    add_legal_reminder,
    add_letter_place_date,
    add_notice_box,
    add_paragraph,
    add_party_marker,
    add_right_aligned_lines,
    add_right_indented_block,
    add_signature_table,
    add_subject_heading,
    new_document,
)


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and all(
        borders.find(qn(f"w:{edge}")) is not None
        for edge in ("top", "left", "bottom", "right")
    )


def test_new_document_applies_global_style_profile() -> None:
    document = new_document()
    section = document.sections[0]

    assert abs(section.top_margin - Cm(2.5)) < 300
    assert abs(section.bottom_margin - Cm(2.5)) < 300
    assert abs(section.left_margin - Cm(2.5)) < 300
    assert abs(section.right_margin - Cm(2.5)) < 300
    assert document.styles["Normal"].font.name == "Roboto"
    assert document.styles["Normal"].font.size == Pt(10)


def test_framed_title_uses_centered_bordered_table() -> None:
    document = new_document()

    table = add_framed_title(document, ["TITRE", "SOUS-TITRE"])

    assert table.style.name == "Table Grid"
    assert _table_has_explicit_borders(table)
    paragraph = table.cell(0, 0).paragraphs[0]
    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert paragraph.text == "TITRE\nSOUS-TITRE"
    assert all(run.bold for run in paragraph.runs if run.text.strip())


def test_framed_signature_block_uses_explicit_borders() -> None:
    document = new_document()

    table = add_framed_signature_block(document, ["Fait à Paris", "Le 12/05/2026"])

    assert table.style.name == "Table Grid"
    assert _table_has_explicit_borders(table)
    assert table.cell(0, 0).paragraphs[0].text == "Fait à Paris"
    assert table.cell(0, 0).paragraphs[1].text == "Le 12/05/2026"


def test_hyphen_list_item_uses_visible_marker_and_hanging_indent() -> None:
    document = new_document()

    paragraph = add_hyphen_list_item(document, "Nomination du gérant ;")

    assert paragraph.text == "- Nomination du gérant ;"
    assert paragraph.paragraph_format.left_indent is not None
    assert paragraph.paragraph_format.first_line_indent is not None
    assert paragraph.paragraph_format.first_line_indent < 0


def test_paragraph_and_legal_reminder_helpers_apply_text_styles() -> None:
    document = new_document()

    paragraph = add_paragraph(document, "Texte", bold=True)
    add_legal_reminder(
        document,
        title="Rappel",
        title_suffix=" : Article",
        paragraphs=["Texte légal"],
    )

    assert paragraph.runs[0].bold is True
    reminder_title = document.paragraphs[2]
    assert reminder_title.runs[0].text == "Rappel"
    assert reminder_title.runs[0].italic is True
    assert reminder_title.runs[0].underline is True
    assert reminder_title.runs[1].italic is True
    reminder_text = document.paragraphs[3]
    assert reminder_text.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert reminder_text.runs[0].italic is True


def test_letter_helpers_apply_structural_alignment_and_emphasis() -> None:
    document = new_document(style_profile=LETTER_WIDE_STYLE_PROFILE)

    place_date = add_letter_place_date(document, "Paris, le 15/05/2026")
    right_lines = add_right_aligned_lines(document, ["Signature"])
    recipient = add_right_indented_block(
        document,
        ["Conseil departemental"],
        left_indent_cm=8.7,
        first_line_indent_cm=1.2,
    )
    subject = add_subject_heading(document, "Objet : test")
    company = add_company_identity_block(document, ["SOCIETE", "SELAS"])
    instruction = add_italic_instruction(document, "Mention manuscrite")

    section = document.sections[0]
    assert abs(section.left_margin - Cm(3.17)) < 300
    assert abs(section.right_margin - Cm(3.17)) < 300
    assert place_date.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert right_lines[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert abs(recipient[0].paragraph_format.left_indent - Cm(8.7)) < 300
    assert abs(recipient[0].paragraph_format.first_line_indent - Cm(1.2)) < 300
    assert subject.runs[0].bold is True
    assert subject.runs[0].underline is True
    assert company[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert company[0].runs[0].bold is True
    assert instruction.runs[0].italic is True


def test_lot03_style_profiles_and_helpers_apply_structure() -> None:
    document = new_document(style_profile=BAIL_COMPACT_STYLE_PROFILE)

    party = add_party_marker(document, "De premiere part")
    article = add_article_heading(document, "ARTICLE 1")
    section = add_framed_section_title(document, "SECTION")
    data_table = add_bordered_data_table(document, ["A", "B"], [["1", "2"]])
    amount = add_centered_amount(document, ["150 000", "EUR"])
    signatures = add_signature_table(document, [["Le vendeur", "L'acquereur"]])

    assert abs(document.sections[0].top_margin - Cm(1.75)) < 300
    assert abs(document.sections[0].bottom_margin - Cm(0.5)) < 300
    assert party.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert party.runs[0].bold is True
    assert party.runs[0].underline is True
    assert article.runs[0].underline is True
    assert _table_has_explicit_borders(section)
    assert _table_has_explicit_borders(data_table)
    assert amount[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert _table_has_explicit_borders(signatures)


def test_form_helpers_apply_checkbox_notice_and_cumul_profile() -> None:
    document = new_document(style_profile=DEROGATION_CUMUL_STYLE_PROFILE)

    heading = add_form_section_heading(document, "Identification")
    checkbox = add_checkbox_line(document, "OUI", checked=True)
    notice = add_notice_box(document, ["Rappel", "Piece a joindre"])

    assert abs(document.sections[0].top_margin - Cm(3.25)) < 500
    assert heading.runs[0].bold is True
    assert heading.runs[0].underline is True
    assert checkbox.text == "☒ OUI"
    assert _table_has_explicit_borders(notice)
