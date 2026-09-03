from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import Address, DocumentGenerationContext, Person, Signature
from sydel_doc_engine.generators.lot_01.declaration_non_condamnation import (
    DeclarationNonCondamnationGenerator,
)


def _context(genre: Gender = Gender.MASCULIN) -> DocumentGenerationContext:
    civilite = "Madame" if genre == Gender.FEMININ else "Monsieur"
    prenom = "Marie" if genre == Gender.FEMININ else "Jean"
    return DocumentGenerationContext(
        personne_signataire=Person(
            genre=genre,
            civilite=civilite,
            prenom=prenom,
            nom="Durand",
            adresse_perso=Address(
                num_voie="12",
                voie="rue des Lilas",
                cp="75008",
                ville="Paris",
            ),
            date_naissance=date(1990, 2, 3),
            ville_naissance="Paris",
            nationalite="française",
            nom_pere="Pierre Durand",
            nom_mere="Anne Martin",
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 12),
        ),
    )


def _generate(tmp_path: Path, genre: Gender = Gender.MASCULIN) -> Path:
    return DeclarationNonCondamnationGenerator().generate(_context(genre), tmp_path)


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _all_paragraphs(document: Document) -> list:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _find_paragraph(document: Document, expected_text: str):
    for paragraph in _all_paragraphs(document):
        if expected_text in paragraph.text:
            return paragraph
    raise AssertionError(f"Paragraphe introuvable : {expected_text}")


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and borders.find(qn("w:top")) is not None


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_declaration_non_condamnation_creates_docx(tmp_path: Path) -> None:
    output_path = _generate(tmp_path)

    assert output_path == tmp_path / "declaration_non_condamnation.docx"
    assert output_path.is_file()


def test_declaration_non_condamnation_contains_essential_texts(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "DECLARATION DE NON CONDAMNATION" in text
    assert "EN APPLICATION DE L’ARTICLE A.123-51 du Code de Commerce" in text
    assert "Je soussigné Monsieur Jean Durand" in text
    assert "Né le 03/02/1990 à Paris." in text
    assert "de nationalité française" in text
    assert "fils de Monsieur Pierre Durand" in text
    assert "et de Madame Anne Martin" in text
    assert "Déclare sur l’honneur, conformément aux dispositions de l’article A.123-51" in text
    assert "Fait à Paris" in text
    assert "Le 12/05/2026" in text
    assert "Rappel : Article L123-5 du code de commerce" in text
    assert "Les dispositions des deuxième et troisième alinéas de l’article L.123-4" in text


def test_declaration_non_condamnation_uses_feminine_agreements(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path, Gender.FEMININ))

    assert "Je soussignée Madame Marie Durand" in text
    assert "Née le 03/02/1990 à Paris." in text
    assert "fille de Monsieur Pierre Durand" in text


def test_declaration_non_condamnation_can_use_au_before_birth_city(tmp_path: Path) -> None:
    ctx = _context()
    ctx.personne_signataire.ville_naissance = "Bourget"
    ctx.personne_signataire.ville_naissance_article_au = True

    text = _docx_text(DeclarationNonCondamnationGenerator().generate(ctx, tmp_path))

    assert "Né le 03/02/1990 au Bourget." in text
    assert "Né le 03/02/1990 à Bourget." not in text


def test_declaration_non_condamnation_composes_personal_address(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "demeurant au 12 rue des Lilas, 75008 Paris" in text
    assert "demeurant au 12 rue des Lilas, Paris 75008" not in text


def test_declaration_non_condamnation_matches_source_visual_formatting(tmp_path: Path) -> None:
    document = Document(_generate(tmp_path))

    normal_style = document.styles["Normal"]
    assert normal_style.font.name == "Roboto"
    assert normal_style.font.size == Pt(10)

    title_table = document.tables[0]
    assert len(title_table.rows) == 1
    assert len(title_table.rows[0].cells) == 1
    assert title_table.style.name == "Table Grid"

    title_paragraph = title_table.cell(0, 0).paragraphs[0]
    assert title_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert all(run.bold for run in title_paragraph.runs if run.text.strip())
    assert _table_has_explicit_borders(title_table)

    assert len(document.tables) == 1
    signature_paragraph = _find_paragraph(document, "Fait à Paris")
    assert signature_paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    subject_paragraph = _find_paragraph(document, "Je soussigné Monsieur Jean Durand")
    assert all(run.bold for run in subject_paragraph.runs if run.text.strip())

    declaration_paragraph = _find_paragraph(document, "Déclare sur l’honneur")
    assert declaration_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert all(run.bold for run in declaration_paragraph.runs if run.text.strip())

    first_reminder_paragraph = _find_paragraph(document, "Le fait de donner")
    assert first_reminder_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert all(run.italic for run in first_reminder_paragraph.runs if run.text.strip())

    second_reminder_paragraph = _find_paragraph(document, "Les dispositions des deuxième")
    assert second_reminder_paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert all(run.italic for run in second_reminder_paragraph.runs if run.text.strip())
