from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    ApportTitres,
    CapitalSouscripteur,
    CapitalSouscription,
    DocumentGenerationContext,
    DossierOptions,
    ExerciceSocial,
    Person,
    RemunerationPresident,
    Signature,
    SocieteCible,
    SocieteSpfpl,
    SpfplPerson,
    StatutsPresident,
    StatutsSas,
)
from sydel_doc_engine.generators.lot_05.attestation_capital_liste_souscripteurs_sas import (
    AttestationCapitalListeSouscripteursSasGenerator,
)
from sydel_doc_engine.generators.lot_05.pv_remuneration_president import (
    PvRemunerationPresidentGenerator,
)


def _base_context() -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SAS",
        dossier_options=DossierOptions(associe_unique=True, apport=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 15)),
        statuts_sas=StatutsSas(type="spfpl_medecins", profession="medecin"),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            forme_sociale="Société par actions simplifiée",
            capital_social="60 000",
            nb_actions_total=600,
            valeur_nominale_action="100",
            profession="Médecins",
            ville_rcs="Paris",
            siege=Address(
                num_voie="10",
                voie="rue de la Paix",
                cp="75002",
                ville="Paris",
                adresse_affichee="10 rue de la Paix, 75002 Paris",
            ),
        ),
        actionnaire_unique=SpfplPerson(
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            genre=Gender.MASCULIN,
            profession="médecin",
            qualite_associe="actionnaire unique",
            adresse_personnelle=Address(
                num_voie="5",
                voie="rue Royale",
                cp="75008",
                ville="Paris",
            ),
            adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ),
        president=StatutsPresident(
            ref_associe_index=0,
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            fonction="Président",
            adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ),
        exercice_social=ExerciceSocial(
            date_cloture_premier_exercice="31 décembre 2026",
        ),
        remuneration_president=RemunerationPresident(
            type="absence_remuneration",
            date_fin_non_remuneree="31 décembre 2026",
        ),
        capital_souscription=CapitalSouscription(
            nb_actions_total=600,
            valeur_nominale_action="100",
            apports_nature_montant="50 000",
            apports_numeraire_montant="10 000",
            souscripteurs=[
                CapitalSouscripteur(
                    civilite_affichage="Docteur",
                    prenom="Camille",
                    nom="Martin",
                    profession="médecin",
                    adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
                    nb_actions=600,
                    qualite="actionnaire unique",
                )
            ],
        ),
        apport_titres=ApportTitres(nb_parts=50),
        societe_cible=SocieteCible(
            denomination="SELARL CABINET MARTIN",
            forme_sociale="SELARL",
            siege=Address(adresse_affichee="12 avenue des Ternes, 75017 Paris"),
            ville_rcs="Paris",
            numero_rcs="900 000 001",
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def test_pv_remuneration_president_generates_source_wording(tmp_path: Path) -> None:
    output_path = PvRemunerationPresidentGenerator().generate(_base_context(), tmp_path)

    text = _docx_text(output_path)

    assert output_path.name == "pv_remuneration_president.docx"
    assert "PROCES-VERBAL DES DECISIONS" in text
    assert "Docteur Camille Martin, actionnaire unique, décide qu'il ne percevra" in text
    assert "Fait à Paris en trois exemplaires" in text
    _assert_clean(text)


def test_pv_remuneration_president_blocks_feminine_wording(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.actionnaire_unique.genre = Gender.FEMININ

    with pytest.raises(ValueError, match="president masculin"):
        PvRemunerationPresidentGenerator().generate(ctx, tmp_path)


def test_attestation_capital_sas_generates_unique_subscriber_wording(
    tmp_path: Path,
) -> None:
    output_path = AttestationCapitalListeSouscripteursSasGenerator().generate(
        _base_context(),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "attestation_capital_liste_souscripteurs_sas.docx"
    assert "Liste des souscripteurs" in text
    assert "Répartition : 600 actions attribuées au Dr Camille Martin, actionnaire unique" in text
    assert "Apports en nature" in text
    _assert_clean(text)


def test_attestation_capital_sas_blocks_multiple_subscribers(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.capital_souscription.souscripteurs.append(
        CapitalSouscripteur(
            civilite_affichage="Docteur",
            prenom="Louise",
            nom="Bernard",
            profession="médecin",
            adresse_personnelle_affichee="9 rue Bleue, 75009 Paris",
            nb_actions=1,
        )
    )

    with pytest.raises(ValueError, match="exactement un souscripteur"):
        AttestationCapitalListeSouscripteursSasGenerator().generate(ctx, tmp_path)


def test_attestation_capital_sas_blocks_capital_mismatch(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.capital_souscription.apports_numeraire_montant = "9 000"

    with pytest.raises(ValueError, match="apports en nature et en numeraire"):
        AttestationCapitalListeSouscripteursSasGenerator().generate(ctx, tmp_path)
