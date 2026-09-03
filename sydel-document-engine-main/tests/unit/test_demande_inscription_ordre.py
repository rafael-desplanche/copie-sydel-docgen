from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Company,
    DocumentGenerationContext,
    DossierOptions,
    Mandataire,
    OrdreAddress,
    OrdreProfessionnel,
    Person,
    Signature,
)
from sydel_doc_engine.generators.lot_02.demande_inscription_ordre import (
    DemandeInscriptionOrdreGenerator,
)


def _context(
    structure: str,
    *,
    ordre: OrdreProfessionnel | None = None,
    mandataire: Mandataire | None = None,
    derogation: bool = False,
) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(derogation=derogation),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            titre_affichage="Dr",
            prenom="Jean",
            nom="Durand",
            adresse_personnelle_affichee="12 rue des Lilas\n75008 Paris",
        ),
        societe=Company(denomination="SEL EXEMPLE"),
        ordre=ordre or _selarl_selas_ordre(),
        mandataire=mandataire or _detailed_mandataire(),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
    )


def _selarl_selas_ordre(
    *,
    departement_inscription: str | None = "la Loire-Atlantique",
    derogation_mention_manuelle: str | None = None,
) -> OrdreProfessionnel:
    return OrdreProfessionnel(
        conseil_departemental_libelle="Conseil départemental de l’Ordre",
        departement_inscription=departement_inscription,
        destinataire_appel="Monsieur le Président",
        profession_signataire_affichee="chirurgien-dentiste",
        profession_ligne_destinataire="chirurgiens-dentistes",
        profession_reglementee_pluriel="chirurgiens-dentistes",
        adresse=OrdreAddress(
            ligne_1="6 rue du Conseil",
            cp="75001",
            ville="Paris",
        ),
        derogation_mention_manuelle=derogation_mention_manuelle,
    )


def _spfpl_ordre(
    *,
    derogation_mention_manuelle: str | None = None,
) -> OrdreProfessionnel:
    return OrdreProfessionnel(
        conseil_departemental_libelle="Conseil départemental de l’Ordre",
        destinataire_appel="Monsieur le Président",
        profession_signataire_affichee="chirurgien-dentiste",
        profession_ligne_destinataire="chirurgiens-dentistes",
        profession_reglementee_pluriel="chirurgiens-dentistes",
        adresse_affichee="6 rue du Conseil\n75001 Paris",
        derogation_mention_manuelle=derogation_mention_manuelle,
    )


def _scm_ordre() -> OrdreProfessionnel:
    return OrdreProfessionnel(
        conseil_departemental_libelle="Conseil départemental de l’Ordre",
        destinataire_appel="Monsieur le Président",
        profession_signataire_affichee="médecin",
        profession_ligne_destinataire="médecins",
        profession_reglementee_pluriel="médecins",
        adresse_bloc_affiche="4 avenue Ordinale\n69002 Lyon",
    )


def _detailed_mandataire() -> Mandataire:
    return Mandataire(
        civilite_affichage="Madame",
        prenom="Sophie",
        nom="Martin",
        fonction="juriste",
        cabinet="DAAT",
    )


def _configured_mandataire() -> Mandataire:
    return Mandataire(libelle_affiche="Monsieur Paul Bernard, mandataire du cabinet CONFIG")


def _generate(tmp_path: Path, ctx: DocumentGenerationContext) -> Path:
    return DemandeInscriptionOrdreGenerator().generate(ctx, tmp_path)


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text]


def _matching_paragraphs(path: Path, text: str):
    return [paragraph for paragraph in Document(path).paragraphs if paragraph.text == text]


def _assert_no_source_placeholders(text: str) -> None:
    assert "[" not in text
    assert "]" not in text
    assert "Dérogation ?" not in text


def test_demande_inscription_ordre_selarl_uses_structured_ordinal_address(
    tmp_path: Path,
) -> None:
    output_path = _generate(tmp_path, _context("SELARL"))
    text = _docx_text(output_path)
    paragraphs = _paragraphs(output_path)

    assert output_path == tmp_path / "demande_inscription_ordre.docx"
    assert "Dr Jean Durand" in text
    assert (
        "Conseil départemental de l'Ordre des chirurgiens-dentistes "
        "de la Loire-Atlantique"
    ) in paragraphs
    assert "Des chirurgiens-dentistes" not in paragraphs
    assert "6 rue du Conseil" in paragraphs
    assert "75001 Paris" in paragraphs
    assert (
        "Je donne pouvoir à Madame Sophie Martin, juriste du cabinet DAAT pour effectuer "
        "les formalités."
    ) in text
    assert _matching_paragraphs(output_path, "Paris, le 14/05/2026")[0].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    subject = _matching_paragraphs(
        output_path,
        "Objet : Demande d’inscription au tableau de l’Ordre",
    )[0]
    assert subject.runs[0].bold is True
    assert subject.runs[0].underline is True
    recipient = _matching_paragraphs(
        output_path,
        "Conseil départemental de l'Ordre des chirurgiens-dentistes de la Loire-Atlantique",
    )[0]
    assert recipient.paragraph_format.left_indent > Cm(8)
    assert _matching_paragraphs(output_path, "Dr Jean Durand")[-1].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_selas_uses_same_overlay_as_selarl(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path, _context("SELAS")))

    assert "Objet : Demande d’inscription au tableau de l’Ordre" in text
    assert (
        "Je sollicite l’inscription de ma société au tableau de l’Ordre des "
        "chirurgiens-dentistes."
    ) in text
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_spfpl_cession_accepts_configured_mandataire(
    tmp_path: Path,
) -> None:
    ctx = _context(
        "SPFPL cession",
        ordre=_spfpl_ordre(),
        mandataire=_configured_mandataire(),
    )
    text = _docx_text(_generate(tmp_path, ctx))

    assert "Monsieur Paul Bernard, mandataire du cabinet CONFIG" in text
    assert "Jordan ELBAZ" not in text
    assert "SYDEL" not in text
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_spfpl_apport_uses_spfpl_address_block(
    tmp_path: Path,
) -> None:
    text = _docx_text(
        _generate(
            tmp_path,
            _context("SPFPL apport", ordre=_spfpl_ordre(), mandataire=_configured_mandataire()),
        )
    )

    assert "6 rue du Conseil" in text
    assert "75001 Paris" in text
    assert "SPFPL apport" not in text
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_scm_requires_explicit_ordinal_data(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path, _context("SCM", ordre=_scm_ordre())))

    assert "médecin" in text
    assert "Des médecins" in text
    assert "4 avenue Ordinale" in text
    assert "69002 Lyon" in text
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_renders_manual_derogation_only_when_provided(
    tmp_path: Path,
) -> None:
    ctx = _context(
        "SPFPL_CESSION",
        ordre=_spfpl_ordre(derogation_mention_manuelle="Mention manuelle validée."),
        mandataire=_configured_mandataire(),
        derogation=True,
    )
    text = _docx_text(_generate(tmp_path, ctx))

    assert "une seule structure. Mention manuelle validée." in text
    _assert_no_source_placeholders(text)


def test_demande_inscription_ordre_blocks_derogation_without_manual_mention(
    tmp_path: Path,
) -> None:
    ctx = _context(
        "SPFPL_APPORT",
        ordre=_spfpl_ordre(),
        mandataire=_configured_mandataire(),
        derogation=True,
    )

    with pytest.raises(ValueError, match="ordre.derogation_mention_manuelle"):
        _generate(tmp_path, ctx)
