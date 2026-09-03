from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Company,
    DocumentGenerationContext,
    Person,
    Signature,
)
from sydel_doc_engine.generators.lot_01.procuration import ProcurationGenerator


def _context(
    genre: Gender = Gender.MASCULIN,
    *,
    image_optionnelle: Path | None = None,
) -> DocumentGenerationContext:
    civilite = "Madame" if genre == Gender.FEMININ else "Monsieur"
    prenom = "Marie" if genre == Gender.FEMININ else "Jean"
    return DocumentGenerationContext(
        personne_signataire=Person(
            genre=genre,
            civilite=civilite,
            prenom=prenom,
            nom="Durand",
            adresse_perso=Address(
                num_voie="12",
                voie="rue des Lilas",
                cp="75008",
                ville="Paris",
            ),
            fonction_dirigeant="Président",
        ),
        societe=Company(
            forme_sociale="SAS",
            denomination="DURAND CONSEIL",
            siege=Address(
                num_voie="80",
                voie="avenue Marceau",
                cp="75008",
                ville="Paris",
            ),
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 12),
            image_optionnelle=image_optionnelle,
        ),
    )


def _generate(
    tmp_path: Path,
    genre: Gender = Gender.MASCULIN,
    *,
    image_optionnelle: Path | None = None,
) -> Path:
    return ProcurationGenerator().generate(
        _context(genre, image_optionnelle=image_optionnelle),
        tmp_path,
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _document_paragraphs(path: Path) -> list[str]:
    document = Document(path)
    return [paragraph.text for paragraph in document.paragraphs if paragraph.text]


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and borders.find(qn("w:top")) is not None


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_procuration_creates_docx(tmp_path: Path) -> None:
    output_dir = tmp_path / "nested"

    output_path = _generate(output_dir)

    assert output_path == output_dir / "procuration.docx"
    assert output_path.is_file()


def test_procuration_contains_essential_texts(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "Procuration" in text
    assert "Je soussigné Monsieur Jean Durand" in text
    assert (
        "demeurant au 12 rue des Lilas, 75008 Paris, agissant en qualité de "
        "Président de la SAS DURAND CONSEIL, dont le siège est situé "
        "80 avenue Marceau, 75008 Paris"
    ) in text
    assert ". Agissant en qualité" not in text
    assert "Donne par les présentes pouvoir à :" in text
    assert (
        "De pour moi et en mon nom faire tous dépôts, immatriculations, modifications, "
        "radiations et de recevoir le registre des bénéficiaires effectifs, concernant mon "
        "entreprise auprès des registres."
    ) in text
    assert (
        "En conséquence, faire toutes déclarations et démarches, produire toutes pièces "
        "justificatives, effectuer tout dépôt de pièces, signer tous documents, requêtes et "
        "documents utiles, élire domicile, substituer en totalité ou en partie, et en général "
        "faire tout ce qui sera nécessaire."
    ) in text
    assert "L’exécution de ce mandat vaudra décharge au mandataire." in text
    assert "Fait pour servir et valoir ce que de droit." in text
    assert "RCS PARIS 788 531 432" not in text
    assert "0153814303" not in text
    assert "Fait à Paris" in text
    assert "Le 12/05/2026" in text
    assert "Jean Durand" in text


def test_procuration_does_not_duplicate_form_when_denomination_contains_it(
    tmp_path: Path,
) -> None:
    ctx = _context()
    assert ctx.societe is not None
    ctx.societe.forme_sociale = "SELARL"
    ctx.societe.forme_sociale_abregee = "SELARL"
    ctx.societe.denomination = "SELARL MARTIN"

    output_path = ProcurationGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)
    assert "agissant en qualité de Président de la SELARL MARTIN" in text
    assert "SELARL SELARL" not in text


def test_procuration_uses_feminine_agreement(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path, Gender.FEMININ))

    assert "Je soussignée Madame Marie Durand" in text


def test_procuration_composes_personal_address_with_postal_code_before_city(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "demeurant au 12 rue des Lilas, 75008 Paris, agissant" in text
    assert "demeurant au 12 rue des Lilas, Paris 75008" not in text


def test_procuration_composes_company_address_with_postal_code_before_city(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path))

    assert "dont le siège est situé 80 avenue Marceau, 75008 Paris" in text
    assert "dont le siège est situé au 80 avenue Marceau" not in text
    assert "dont le siège est situé 80 avenue Marceau, Paris 75008" not in text


def test_procuration_contains_exact_sydel_block(tmp_path: Path) -> None:
    paragraphs = _document_paragraphs(_generate(tmp_path))

    start = paragraphs.index("SYDEL")
    assert paragraphs[start : start + 2] == [
        "SYDEL",
        "80 avenue Marceau, 75008 PARIS",
    ]


def test_procuration_does_not_use_signature_image(tmp_path: Path) -> None:
    missing_image = tmp_path / "signature_absente.png"

    output_path = _generate(tmp_path, image_optionnelle=missing_image)

    assert len(Document(output_path).inline_shapes) == 0


def test_procuration_uses_signature_paragraphs_without_table(tmp_path: Path) -> None:
    document = Document(_generate(tmp_path))

    assert len(document.tables) == 1
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert "Fait à Paris" in paragraphs
    assert "Le 12/05/2026" in paragraphs
    assert "Jean Durand" in paragraphs
