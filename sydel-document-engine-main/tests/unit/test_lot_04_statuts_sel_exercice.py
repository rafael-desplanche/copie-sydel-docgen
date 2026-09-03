from __future__ import annotations

import re
from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Apport,
    Associe,
    CapitalContext,
    CessionBanque,
    Company,
    DepotFonds,
    DirigeantNomine,
    DocumentContext,
    DocumentGenerationContext,
    DocumentSignataire,
    DossierOptions,
    ExerciceLieu,
    ExerciceSocial,
    GeranceContext,
    Person,
    Signature,
    SpfplConjoint,
    SpfplOrdre,
    StatutsSel,
)
from sydel_doc_engine.generators.lot_04.statuts_selarl_dentiste import (
    StatutsSelarlDentisteGenerator,
)
from sydel_doc_engine.generators.lot_04.statuts_selarl_medecin import (
    StatutsSelarlMedecinGenerator,
)
from sydel_doc_engine.generators.lot_04.statuts_selas_medecin import (
    StatutsSelasMedecinGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _associate(*, gender: Gender = Gender.MASCULIN) -> Associe:
    return Associe(
        genre=gender,
        civilite_affichage="Docteur",
        prenom="Camille",
        nom="Martin",
        nb_parts=1000,
        profession="medecin",
        profession_reglementee="medecin",
        profession_reglementee_pluriel="medecins",
        qualification_principale="cardiologue",
        titre_professionnel="Docteur",
        qualite="associe unique",
        date_naissance=date(1980, 1, 2),
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="francaise",
        situation_maritale="marie",
        regime_matrimonial="communaute legale",
        conjoint=SpfplConjoint(
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Martin",
        ),
        adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ordre=SpfplOrdre(
            professionnel="Ordre des medecins",
            departement="Paris",
            ville="Paris",
            numero="12345",
            numero_rpps="10000000001",
        ),
        apport_numeraire="1 000",
        apport_numeraire_lettres="mille",
    )


def _context(*, overlay: str, gender: Gender = Gender.MASCULIN) -> DocumentGenerationContext:
    structure = "SELAS" if overlay == "selas_medecin" else "SELARL"
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(associe_unique=True),
        personne_signataire=Person(
            genre=gender,
            civilite="Monsieur" if gender == Gender.MASCULIN else "Madame",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 14),
            prestataire_signature_electronique="Yousign",
        ),
        societe=Company(
            denomination="SEL MARTIN",
            forme_sociale="societe d'exercice liberal par actions simplifiee",
            forme_sociale_complete=(
                "Societe d'exercice liberal par actions simplifiee"
                if overlay == "selas_medecin"
                else "Societe d'exercice liberal a responsabilite limitee"
            ),
            forme_sociale_abregee="SELAS",
            capital_social="1 000",
            capital_social_lettres="mille",
            duree="99 ans",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
        ),
        statuts_sel=StatutsSel(overlay=overlay, profession="medecin"),
        associes=[_associate(gender=gender)],
        dirigeant_nomine=DirigeantNomine(
            genre=gender,
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            fonction_affichage="President",
            ref_associe_index=0,
            duree_mandat="illimitee",
        ),
        capital=CapitalContext(
            montant="1 000",
            montant_lettres="mille",
            nombre_titres_total=1000,
            nombre_titres_total_lettres="mille",
            valeur_nominale_titre="1",
            valeur_nominale_titre_lettres="un euro",
            type_titre="actions" if overlay == "selas_medecin" else "parts_sociales",
        ),
        apport=Apport(montant="1 000", montant_lettres="mille"),
        depot_fonds=DepotFonds(
            banque=CessionBanque(
                nom="BANQUE EXEMPLE",
                adresse_affichee="1 boulevard Haussmann, 75009 Paris",
            )
        ),
        exercice_social=ExerciceSocial(
            debut="1er janvier",
            fin="31 decembre",
            date_cloture_premier_exercice="31 decembre 2026",
            lieux=[ExerciceLieu(adresse_affichee="12 avenue de la Republique, 75011 Paris")],
        ),
        gerance=GeranceContext(
            seuil_achat_materiel="10 000 euros",
            seuil_emprunt="50 000 euros",
        ),
        document=DocumentContext(
            nombre_exemplaires_lettres="trois",
            signataire=DocumentSignataire(prenom="Camille", nom="Martin"),
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return "\n".join(texts)


def _docx_paragraphs(path: Path) -> list[str]:
    return [paragraph.text for paragraph in Document(path).paragraphs if paragraph.text.strip()]


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _assert_annex_starts_next_page(path: Path) -> None:
    document = Document(path)
    annex_index = next(
        index for index, paragraph in enumerate(document.paragraphs) if paragraph.text == "ANNEXE"
    )
    preceding_xml = "\n".join(
        paragraph._p.xml for paragraph in document.paragraphs[max(0, annex_index - 2) : annex_index]
    )
    assert 'w:type="page"' in preceding_xml


def test_statuts_selarl_dentiste_generates_unique_associate_docx(tmp_path: Path) -> None:
    ctx = _context(overlay="selarl_dentiste")
    ctx.associes[0].profession = "chirurgien-dentiste"
    ctx.associes[0].profession_reglementee = "chirurgien-dentiste"
    ctx.associes[0].profession_reglementee_pluriel = "chirurgiens-dentistes"
    ctx.associes[0].ordre.professionnel = "Ordre des chirurgiens-dentistes"

    output_path = StatutsSelarlDentisteGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)
    document = Document(output_path)
    table_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    article_1 = next(p for p in document.paragraphs if p.text.startswith("ARTICLE 1"))

    assert output_path.name == "statuts_selarl_chirurgien_dentiste.docx"
    assert "SEL MARTIN" in text
    assert "Au capital de 1 000 euros" in text
    assert (
        "sous le numéro RPPS 10000000001, marié sous le régime de la communauté "
        "avec Madame Alice Martin."
    ) in text
    assert "marié sous le régime de la communauté légale" not in text
    assert "ARTICLE 5 - LIEU(X) D’EXERCICE" in text
    assert (
        "Le lieu d’exercice de la société est situé au "
        "12 avenue de la Republique, 75011 Paris. Il constitue le lieu d’exercice unique "
        "de la société"
    ) in text
    assert "Total des apports en numéraire : ci- 1 000." in text
    assert (
        "à Docteur Camille Martin, mille parts sociales en pleine propriété, ci"
    ) in text
    assert "1000 parts" in text
    assert "chirurgiens-dentistes" in text
    assert "Yousign" in text
    assert "STATUTS" in table_text
    assert "- Ouverture d’un compte bancaire" in text
    assert any(run.underline for run in article_1.runs)
    _assert_annex_starts_next_page(output_path)
    _assert_clean(text)


def test_statuts_selarl_medecin_skips_personne_2_source_alias(tmp_path: Path) -> None:
    output_path = StatutsSelarlMedecinGenerator().generate(
        _context(overlay="selarl_medecin"),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "statuts_selarl_medecin.docx"
    assert "Conseil" in text
    assert "personne_2" not in text
    assert "50 000 euros" in text
    assert (
        "sous le numéro national 12345 et sous le numéro RPPS 10000000001, "
        "marié sous le régime de la communauté avec Madame Alice Martin."
    ) in text
    assert "Docteur Camille Martin, associé unique." in text
    assert "- Ouverture d’un compte bancaire" in text
    _assert_annex_starts_next_page(output_path)
    _assert_clean(text)


def test_statuts_selarl_medecin_renders_separation_de_biens_clause(
    tmp_path: Path,
) -> None:
    ctx = _context(overlay="selarl_medecin")
    ctx.associes[0].regime_matrimonial = "separation de biens"

    output_path = StatutsSelarlMedecinGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)

    assert (
        "marié sous le régime de la séparation de biens avec Madame Alice Martin"
    ) in text
    assert "marié sous le régime de separation de biens" not in text
    _assert_clean(text)


def test_statuts_selarl_medecin_article_8_agrees_female_unique(
    tmp_path: Path,
) -> None:
    ctx = _context(overlay="selarl_medecin", gender=Gender.FEMININ)

    output_path = StatutsSelarlMedecinGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)

    assert "Docteur Camille Martin, associée unique." in text
    _assert_clean(text)


def test_statuts_selarl_dentiste_deposit_agrees_female_unique(
    tmp_path: Path,
) -> None:
    ctx = _context(overlay="selarl_dentiste", gender=Gender.FEMININ)

    output_path = StatutsSelarlDentisteGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)

    assert "déposée par l’associée unique conformément à la loi" in text
    assert "déposée par l’associé unique conformément à la loi" not in text
    _assert_clean(text)


def test_statuts_selarl_medecin_matches_source_docx_line_by_line(
    tmp_path: Path,
) -> None:
    ctx = _context(overlay="selarl_medecin")
    output_path = StatutsSelarlMedecinGenerator().generate(ctx, tmp_path)
    source_path = next(Path("project/source_documents/lot_04").glob("*SELARL*decins.docx"))

    source_article = _article_paragraphs(
        _render_source_medecin_paragraph(paragraph, ctx)
        for paragraph in _docx_paragraphs(source_path)
        if "[civilite_personne_2]" not in paragraph
    )
    generated_article = _article_paragraphs(_docx_paragraphs(output_path))

    assert generated_article == source_article


def test_statuts_selas_medecin_generates_without_second_lieu_by_default(
    tmp_path: Path,
) -> None:
    output_path = StatutsSelasMedecinGenerator().generate(
        _context(overlay="selas_medecin"),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "statuts_selas_medecin.docx"
    assert "Societe d'exercice liberal par actions simplifiee" in text
    assert "President" in text
    assert "nom_lieu_exercice_2" not in text
    _assert_clean(text)


def test_statuts_selas_medecin_renders_complete_second_lieu(tmp_path: Path) -> None:
    ctx = _context(overlay="selas_medecin")
    ctx.exercice_social.lieux.append(
        ExerciceLieu(
            nom="Cabinet secondaire",
            adresse_affichee="20 rue Bleue, 75009 Paris",
        )
    )

    output_path = StatutsSelasMedecinGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)

    assert "Cabinet secondaire, 20 rue Bleue, 75009 Paris" in text
    _assert_clean(text)


def test_statuts_sel_blocks_multi_associes(tmp_path: Path) -> None:
    ctx = _context(overlay="selarl_medecin")
    ctx.associes.append(_associate())

    with pytest.raises(ValueError, match="multi-associes"):
        StatutsSelarlMedecinGenerator().generate(ctx, tmp_path)


def test_statuts_selarl_dentiste_blocks_multi_associes(tmp_path: Path) -> None:
    # SELARL = unipersonnelle (decision Gad 2026-06-04) : deux associes -> ValueError,
    # meme si l'ancien flag PARTIAL est present dans metadata (sous-cas abandonne).
    ctx = _context(overlay="selarl_dentiste")
    ctx.metadata["selarl_dentiste_multi_associes_statuts_partial"] = "true"
    ctx.associes.append(_associate())

    with pytest.raises(ValueError, match="multi-associes"):
        StatutsSelarlDentisteGenerator().generate(ctx, tmp_path)


def test_statuts_selas_blocks_partial_second_lieu(tmp_path: Path) -> None:
    ctx = _context(overlay="selas_medecin")
    ctx.exercice_social.lieux.append(ExerciceLieu(nom="Cabinet secondaire"))

    with pytest.raises(ValueError, match="doivent etre fournis ensemble"):
        StatutsSelasMedecinGenerator().generate(ctx, tmp_path)


def test_statuts_selas_blocks_dirigeant_non_associe_signature(tmp_path: Path) -> None:
    ctx = _context(overlay="selas_medecin")
    ctx.dirigeant_nomine.nom = "Bernard"
    ctx.dirigeant_nomine.ref_associe_index = None

    with pytest.raises(ValueError, match="dirigeant non associe"):
        StatutsSelasMedecinGenerator().generate(ctx, tmp_path)


def test_statuts_sel_applies_female_birth_agreement(tmp_path: Path) -> None:
    ctx = _context(overlay="selas_medecin", gender=Gender.FEMININ)

    output_path = StatutsSelasMedecinGenerator().generate(ctx, tmp_path)

    text = _docx_text(output_path)

    assert "nÃ©e le" in text or "née le" in text
    _assert_clean(text)


def test_statuts_selas_header_agrees_masculine(tmp_path: Path) -> None:
    # Entete figee au masculin dans les blocs : conservee telle quelle pour un homme.
    output_path = StatutsSelasMedecinGenerator().generate(
        _context(overlay="selas_medecin", gender=Gender.MASCULIN),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert "LE SOUSSIGNE" in text
    assert "LA SOUSSIGNÉE" not in text
    # Ligne d'identification masculine.
    assert "né le 02/01/1980" in text
    assert "née le 02/01/1980" not in text
    _assert_clean(text)


def test_statuts_selas_header_agrees_feminine(tmp_path: Path) -> None:
    # Entete figee au masculin -> accordee au feminin pour une associee.
    output_path = StatutsSelasMedecinGenerator().generate(
        _context(overlay="selas_medecin", gender=Gender.FEMININ),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert "LA SOUSSIGNÉE" in text
    assert "LE SOUSSIGNE\xa0:" not in text
    assert "née le 02/01/1980" in text
    _assert_clean(text)


def test_statuts_selas_article_8_uses_dynamic_associate_label(tmp_path: Path) -> None:
    # Le bloc fige « à l'associé unique » est remplace par le token dynamique :
    # masculin -> « associé unique », feminin -> « associée unique ».
    masc = _docx_text(
        StatutsSelasMedecinGenerator().generate(
            _context(overlay="selas_medecin", gender=Gender.MASCULIN),
            tmp_path / "masc",
        )
    )
    fem = _docx_text(
        StatutsSelasMedecinGenerator().generate(
            _context(overlay="selas_medecin", gender=Gender.FEMININ),
            tmp_path / "fem",
        )
    )

    assert "attribuées en totalité à l’associé unique, Docteur Camille Martin." in masc
    assert "attribuées en totalité à l’associée unique, Docteur Camille Martin." in fem
    _assert_clean(masc)
    _assert_clean(fem)


def test_statuts_sel_orchestrator_selects_only_requested_overlay() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected = orchestrator.select_documents_for_context(_context(overlay="selarl_medecin"))

    selected_ids = {document.doc_id for document in selected}
    assert "DOC-017" in selected_ids
    assert "DOC-016" not in selected_ids
    assert "DOC-018" not in selected_ids


def test_statuts_sel_orchestrator_ignores_sel_statuts_without_overlay() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())
    ctx = _context(overlay="selarl_medecin")
    ctx.statuts_sel = None

    selected = orchestrator.select_documents_for_context(ctx)

    assert {"DOC-016", "DOC-017", "DOC-018"}.isdisjoint(
        {document.doc_id for document in selected}
    )


def _article_paragraphs(paragraphs) -> list[str]:
    normalized = [_normalize_source_line(paragraph) for paragraph in paragraphs]
    for index, paragraph in enumerate(normalized):
        if paragraph.startswith("ARTICLE 1"):
            return normalized[index:]
    raise AssertionError("ARTICLE 1 introuvable dans les statuts SEL.")


def _normalize_source_line(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("\u00a0", " ").strip())


def _render_source_medecin_paragraph(
    paragraph: str,
    ctx: DocumentGenerationContext,
) -> str:
    associate = ctx.associes[0]
    replacements = {
        "[denomination_societe]": ctx.societe.denomination,
        "[capital_social]": ctx.capital.montant,
        "[adresse_siege]": ctx.societe.siege.adresse_affichee,
        "[civilite]": associate.civilite_affichage,
        "[prenom]": associate.prenom,
        "[nom]": associate.nom,
        "[profession]": associate.profession,
        "[date_naissance]": associate.date_naissance.strftime("%d/%m/%Y"),
        "[ville_naissance]": associate.ville_naissance,
        "[departement_naissance]": associate.departement_naissance,
        "[nationalite]": associate.nationalite,
        "[adresse_personnelle]": associate.adresse_personnelle_affichee,
        "[ville_ordre]": associate.ordre.ville,
        "[numero_ordre]": associate.ordre.numero,
        "[numero_rpps]": associate.ordre.numero_rpps,
        "[situation_maritale]": (
            "marié sous le régime de la communauté avec Madame Alice Martin"
        ),
        "[forme_sociale_complete]": ctx.societe.forme_sociale_complete,
        "[capital_lettres]": ctx.capital.montant_lettres,
        "[nom_banque]": ctx.depot_fonds.banque.nom,
        "[adresse_banque]": ctx.depot_fonds.banque.adresse_affichee,
        "[nb_parts_total]": str(ctx.capital.nombre_titres_total),
        "[valeur_nominale_part]": ctx.capital.valeur_nominale_titre,
        "[seuil_achat_materiel]": ctx.gerance.seuil_achat_materiel,
        "[seuil_emprunt_gerance]": ctx.gerance.seuil_emprunt,
        "[date_cloture_exercice_1]": ctx.exercice_social.date_cloture_premier_exercice,
        "[lieu_signature]": ctx.signature.lieu,
        "[date_signature]": ctx.signature.date.strftime("%d/%m/%Y"),
        "[nombre_exemplaires_lettres]": ctx.document.nombre_exemplaires_lettres,
        "[prenom_signataire]": ctx.document.signataire.prenom,
        "[nom_signataire]": ctx.document.signataire.nom,
    }
    rendered = paragraph
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    rendered = rendered.replace("associée unique", "associé unique")
    if rendered == "Ouverture d’un compte bancaire":
        return "- Ouverture d’un compte bancaire"
    return rendered
