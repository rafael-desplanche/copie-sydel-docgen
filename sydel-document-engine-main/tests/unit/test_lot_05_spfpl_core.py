from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    ApportTitres,
    AssocieCible,
    CapitalSouscripteur,
    CapitalSouscription,
    CessionParts,
    DocumentContext,
    DocumentGenerationContext,
    DossierOptions,
    OperationSpfpl,
    Person,
    ProfessionalEntity,
    Signature,
    SocieteCible,
    SocieteSpfpl,
    SpfplConjoint,
    SpfplDirigeant,
    SpfplOrdre,
    SpfplPerson,
    SpfplRepresentant,
)
from sydel_doc_engine.generators.lot_05.acte_cession_parts_spfpl import (
    ActeCessionPartsSpfplGenerator,
)
from sydel_doc_engine.generators.lot_05.attestation_capital_liste_souscripteurs import (
    AttestationCapitalListeSouscripteursGenerator,
)
from sydel_doc_engine.generators.lot_05.attestation_commissaire_apports import (
    AttestationCommissaireApportsGenerator,
)
from sydel_doc_engine.generators.lot_05.contrat_apport_spfpl import (
    ContratApportSpfplGenerator,
)


def _base_context(*, operation: str = "cession") -> DocumentGenerationContext:
    is_apport = operation == "apport"
    return DocumentGenerationContext(
        structure="SPFPL apport" if is_apport else "SPFPL cession",
        dossier_options=DossierOptions(apport=is_apport, cession=not is_apport),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14), nombre_exemplaires="trois"),
        operation_spfpl=OperationSpfpl(type=operation),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            forme_sociale="SPFPLAS",
            capital_social="60 000",
            activite="participations financieres de profession liberale",
            profession="chirurgien-dentiste",
            ville_rcs="Paris",
            numero_rcs="en cours",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
            dirigeant=SpfplDirigeant(fonction="President"),
            representant=SpfplRepresentant(
                civilite_affichage="Monsieur",
                civilite_courte="M.",
                prenom="Camille",
                nom="Martin",
                fonction="President",
            ),
        ),
        cedant=_spfpl_person(),
        apporteur=_spfpl_person(),
        societe_cible=SocieteCible(
            denomination="SELARL CABINET MARTIN",
            forme_sociale="SELARL",
            forme_sociale_complete=(
                "societe d'exercice liberal a responsabilite limitee"
            ),
            profession_reglementee="chirurgien-dentiste",
            profession_reglementee_pluriel="chirurgiens-dentistes",
            capital_social="10 000",
            capital_social_lettres="dix mille euros",
            nb_parts_total=100,
            valeur_nominale_part="100",
            valeur_nominale_part_lettres="cent euros",
            siege=Address(adresse_affichee="12 avenue des Ternes, 75017 Paris"),
            ville_rcs="Paris",
            numero_rcs="900 000 001",
        ),
        associes_cible=[
            AssocieCible(
                civilite_affichage="Docteur",
                prenom="Camille",
                nom="Martin",
                nb_parts_avant=70,
                nb_parts_apres=10,
                plage_parts="1 a 10",
            ),
            AssocieCible(
                civilite_affichage="Docteur",
                prenom="Louise",
                nom="Bernard",
                nb_parts_avant=30,
                nb_parts_apres=30,
                plage_parts="11 a 40",
            ),
        ],
        cession_parts=CessionParts(
            nb_parts=60,
            nb_parts_lettres="soixante",
            plage_parts="41 a 100",
            prix_unitaire="1 000",
            prix_unitaire_lettres="mille euros",
            prix_total="60 000",
            prix_total_lettres="soixante mille euros",
            nombre_exemplaires_lettres="trois",
        ),
        apport_titres=ApportTitres(
            nb_parts=60,
            nb_parts_lettres="soixante",
            nature_titres="parts sociales",
            plage_parts="41 a 100",
            valeur_par_titre="1 000",
            valeur_par_titre_lettres="mille",
            valeur_globale="60 000",
            valeur_globale_lettres="soixante mille",
            nb_actions_attribuees=600,
            nb_actions_attribuees_lettres="six cents",
            valeur_nominale_action="100",
            valeur_nominale_action_lettres="cent",
        ),
        capital_souscription=CapitalSouscription(
            nb_actions_total=600,
            valeur_nominale_action="100",
            apports_nature_montant="60 000",
            apports_numeraire_montant="0 euro",
            souscripteurs=[
                CapitalSouscripteur(
                    civilite_affichage="Docteur",
                    prenom="Camille",
                    nom="Martin",
                    profession="chirurgien-dentiste",
                    adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
                    nb_actions=600,
                    qualite="actionnaire unique",
                )
            ],
        ),
        evaluateur_apport=_entity("EVAL CONSEIL", "Madame", "Eva", "Lemoine"),
        commissaire_aux_apports=_entity("CAA EXPERTISE", "Monsieur", "Nabil", "Saidi"),
        document=DocumentContext(nombre_exemplaires_lettres="trois"),
    )


def _spfpl_person() -> SpfplPerson:
    return SpfplPerson(
        civilite_affichage="Docteur",
        prenom="Camille",
        nom="Martin",
        genre=Gender.MASCULIN,
        profession="chirurgien-dentiste",
        profession_reglementee="chirurgien-dentiste",
        profession_reglementee_pluriel="chirurgiens-dentistes",
        date_naissance=date(1980, 1, 2),
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="francaise",
        situation_maritale="marie",
        conjoint=SpfplConjoint(
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Martin",
        ),
        adresse_personnelle=Address(adresse_affichee="5 rue Royale, 75008 Paris"),
        adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ordre=SpfplOrdre(
            professionnel="Ordre des chirurgiens-dentistes",
            departement="Paris",
            numero="12345",
            numero_rpps="10000000001",
        ),
    )


def _entity(
    denomination: str,
    civilite: str,
    prenom: str,
    nom: str,
) -> ProfessionalEntity:
    return ProfessionalEntity(
        denomination=denomination,
        forme_sociale="SAS",
        capital_social="1 000 euros",
        siege=Address(adresse_affichee="1 rue Scheffer, 75016 Paris"),
        ville_rcs="Paris",
        numero_rcs="948 483 730",
        representant=SpfplRepresentant(
            civilite_affichage=civilite,
            prenom=prenom,
            nom=nom,
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text
    assert "\nOU\n" not in text
    assert "d'acquerir/de recevoir" not in text


def test_acte_cession_parts_generates_dynamic_capital_and_preserves_source_frais(
    tmp_path: Path,
) -> None:
    output_path = ActeCessionPartsSpfplGenerator().generate(_base_context(), tmp_path)

    text = _docx_text(output_path)

    assert output_path.name == "acte_cession_parts_spfpl.docx"
    assert "Docteur Camille Martin detenant 70 parts" in text
    assert "Docteur Louise Bernard detenant 30 parts" in text
    assert "mille euros (1 000) euro par part cedee" in text
    assert "cession d'action consentie" in text
    _assert_clean(text)


def test_contrat_apport_uses_context_evaluateur_and_commissaire(tmp_path: Path) -> None:
    output_path = ContratApportSpfplGenerator().generate(
        _base_context(operation="apport"),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "contrat_apport_spfpl.docx"
    assert "EVAL CONSEIL" in text
    assert "CAA EXPERTISE" in text
    assert "SYDEL, Societe a responsabilite limitee" not in text
    _assert_clean(text)


def test_contrat_apport_blocks_missing_commissaire(tmp_path: Path) -> None:
    ctx = _base_context(operation="apport")
    ctx.commissaire_aux_apports = None

    with pytest.raises(ValueError, match="commissaire_aux_apports"):
        ContratApportSpfplGenerator().generate(ctx, tmp_path)


def test_attestation_capital_is_limited_to_unique_souscripteur(tmp_path: Path) -> None:
    ctx = _base_context(operation="apport")
    ctx.capital_souscription.souscripteurs.append(
        CapitalSouscripteur(
            civilite_affichage="Docteur",
            prenom="Louise",
            nom="Bernard",
            profession="chirurgien-dentiste",
            adresse_personnelle_affichee="9 rue Bleue, 75009 Paris",
            nb_actions=1,
        )
    )

    with pytest.raises(ValueError, match="exactement un souscripteur"):
        AttestationCapitalListeSouscripteursGenerator().generate(ctx, tmp_path)


def test_attestation_capital_generates_unique_shareholder_wording(tmp_path: Path) -> None:
    output_path = AttestationCapitalListeSouscripteursGenerator().generate(
        _base_context(operation="apport"),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "attestation_capital_liste_souscripteurs.docx"
    assert "actionnaire unique" in text
    assert "Apports en numeraire : 0 euro" in text
    _assert_clean(text)


def test_attestation_commissaire_apports_renders_single_selected_commissaire(
    tmp_path: Path,
) -> None:
    output_path = AttestationCommissaireApportsGenerator().generate(
        _base_context(operation="apport"),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "attestation_commissaire_apports.docx"
    assert "Aux fins de realisation de cet apport en nature" in text
    assert "CAA EXPERTISE" in text
    assert "ADVENSO" not in text
    assert "TS EXPERTISE" not in text
    _assert_clean(text)
