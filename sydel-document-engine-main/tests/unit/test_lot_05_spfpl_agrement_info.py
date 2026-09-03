from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    AssocieCible,
    CessionParts,
    DecisionContext,
    DocumentGenerationContext,
    DossierOptions,
    OperationSpfpl,
    OperationTitres,
    Person,
    ReunionContext,
    ReunionPresident,
    Signature,
    SocieteCible,
    SocieteSpfpl,
    SpfplDirigeant,
    SpfplPerson,
)
from sydel_doc_engine.generators.lot_05.note_information import NoteInformationGenerator
from sydel_doc_engine.generators.lot_05.pv_agrement_cession_spfpl_associe_unique import (
    PvAgrementCessionSpfplAssocieUniqueGenerator,
)
from sydel_doc_engine.generators.lot_05.pv_agrement_cession_spfpl_plusieurs_associes import (
    PvAgrementCessionSpfplPlusieursAssociesGenerator,
)


def _base_context(*, associe_unique: bool = True) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SPFPL cession",
        dossier_options=DossierOptions(cession=True, associe_unique=associe_unique),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        operation_spfpl=OperationSpfpl(type="cession"),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            forme_sociale="SPFPL",
            capital_social="1 000",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
            dirigeant=SpfplDirigeant(fonction="President"),
        ),
        cedant=SpfplPerson(
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            genre=Gender.MASCULIN,
        ),
        societe_cible=SocieteCible(
            denomination="SELARL CABINET MARTIN",
            forme_sociale="SELARL",
            profession_reglementee="chirurgiens-dentistes",
            capital_social="10 000",
            capital_social_lettres="dix mille euros",
            nb_parts_total=100,
            valeur_nominale_part="100",
            siege=Address(
                num_voie="12",
                voie="avenue des Ternes",
                cp="75017",
                ville="Paris",
                adresse_affichee="12 avenue des Ternes, 75017 Paris",
            ),
            ville_rcs="Paris",
            numero_rcs="900 000 001",
        ),
        cession_parts=CessionParts(
            nb_parts=60,
            nb_parts_lettres="soixante",
            plage_parts="41 a 100",
        ),
        operation_titres=OperationTitres(nb_titres=60),
        reunion=ReunionContext(
            annee_lettres="deux mille vingt-six",
            date_lettres="quatorze mai deux mille vingt-six",
            heure="10 heures",
            president=ReunionPresident(
                civilite_affichage="Docteur",
                prenom="Camille",
                nom="Martin",
                qualite="gerant associe",
            ),
        ),
    )


def _unique_context() -> DocumentGenerationContext:
    ctx = _base_context(associe_unique=True)
    ctx.associes_cible = [
        AssocieCible(
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            nb_parts_avant=100,
            nb_parts_apres=40,
            plage_parts="1 a 40",
        ),
        AssocieCible(
            type="personne_morale",
            denomination="SPFPL MARTIN",
            nb_parts_avant=0,
            nb_parts_apres=60,
            plage_parts="41 a 100",
            est_present_ou_represente=False,
        ),
    ]
    ctx.decision = DecisionContext(date="14/05/2026")
    return ctx


def _plural_context() -> DocumentGenerationContext:
    ctx = _base_context(associe_unique=False)
    ctx.associes_cible = [
        AssocieCible(
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            nb_parts_avant=70,
            nb_parts_apres=10,
            plage_parts="1 a 10",
        ),
        AssocieCible(
            civilite_affichage="Docteur",
            prenom="Louise",
            nom="Bernard",
            nb_parts_avant=30,
            nb_parts_apres=30,
            plage_parts="11 a 40",
        ),
        AssocieCible(
            type="personne_morale",
            denomination="SPFPL MARTIN",
            nb_parts_avant=0,
            nb_parts_apres=60,
            plage_parts="41 a 100",
            est_present_ou_represente=False,
        ),
    ]
    ctx.decision = DecisionContext(date="14/05/2026")
    return ctx


def _apport_note_context() -> DocumentGenerationContext:
    ctx = _unique_context()
    ctx.structure = "SPFPL apport"
    ctx.dossier_options = DossierOptions(apport=True)
    ctx.operation_spfpl = OperationSpfpl(type="apport")
    ctx.apporteur = SpfplPerson(
        civilite_affichage="Docteur",
        prenom="Camille",
        nom="Martin",
        genre=Gender.MASCULIN,
    )
    return ctx


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_no_placeholders_or_options(text: str) -> None:
    assert "[" not in text
    assert "]" not in text
    assert "\nOU\n" not in text
    assert "d'acquerir/de recevoir" not in text


def test_note_information_generates_cession_wording(tmp_path: Path) -> None:
    output_path = NoteInformationGenerator().generate(_unique_context(), tmp_path)

    text = _docx_text(output_path)

    assert output_path.name == "note_information.docx"
    assert "prevoit d'acquerir" in text
    assert "Apres ladite cession" in text
    assert "SPFPL MARTIN, titulaire de 60 parts sociales" in text
    _assert_no_placeholders_or_options(text)


def test_note_information_generates_apport_wording(tmp_path: Path) -> None:
    output_path = NoteInformationGenerator().generate(_apport_note_context(), tmp_path)

    text = _docx_text(output_path)

    assert "prevoit de recevoir en apport en nature" in text
    assert "Apres ledit apport" in text
    _assert_no_placeholders_or_options(text)


def test_pv_associe_unique_generates_cession_wording(tmp_path: Path) -> None:
    output_path = PvAgrementCessionSpfplAssocieUniqueGenerator().generate(
        _unique_context(),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "pv_agrement_cession_spfpl_associe_unique.docx"
    assert "L'associe unique autorise la cession" in text
    assert "contrat d'apport" not in text
    assert "autorise l'apport" not in text
    assert "parts apportees" not in text
    _assert_no_placeholders_or_options(text)


def test_pv_plusieurs_associes_generates_presence_and_signatures(tmp_path: Path) -> None:
    output_path = PvAgrementCessionSpfplPlusieursAssociesGenerator().generate(
        _plural_context(),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "pv_agrement_cession_spfpl_plusieurs_associes.docx"
    assert "Docteur Camille Martin detenant 70 parts" in text
    assert "Docteur Louise Bernard detenant 30 parts" in text
    assert "Projet du contrat de cession" in text
    assert "Camille Martin" in text
    assert "Louise Bernard" in text
    _assert_no_placeholders_or_options(text)


def test_pv_plusieurs_associes_blocks_missing_total_presence(tmp_path: Path) -> None:
    ctx = _plural_context()
    ctx.associes_cible[1].nb_parts_avant = 20

    with pytest.raises(ValueError, match="totalite des parts"):
        PvAgrementCessionSpfplPlusieursAssociesGenerator().generate(ctx, tmp_path)
