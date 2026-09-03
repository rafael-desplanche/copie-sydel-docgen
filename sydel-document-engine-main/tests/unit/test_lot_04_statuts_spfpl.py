from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Apport,
    ApportTitres,
    CapitalSouscripteur,
    CapitalSouscription,
    CessionBanque,
    DepotFonds,
    DocumentGenerationContext,
    DossierOptions,
    ExerciceSocial,
    OperationSpfpl,
    Person,
    ProfessionalEntity,
    Signature,
    SocieteCible,
    SocieteSpfpl,
    SpfplConjoint,
    SpfplOrdre,
    SpfplPerson,
    SpfplRepresentant,
)
from sydel_doc_engine.generators.lot_04.statuts_spfpl_apport import (
    StatutsSpfplApportGenerator,
)
from sydel_doc_engine.generators.lot_04.statuts_spfpl_cession import (
    StatutsSpfplCessionGenerator,
)


def _base_context(*, operation: str) -> DocumentGenerationContext:
    is_apport = operation == "apport"
    founder = SpfplPerson(
        civilite_affichage="Docteur",
        prenom="Camille",
        prenoms="Camille Andre",
        nom="Martin",
        genre=Gender.MASCULIN,
        profession="chirurgien-dentiste",
        profession_reglementee="chirurgiens-dentistes",
        date_naissance=date(1980, 1, 2),
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="francaise",
        situation_maritale="marie",
        regime_matrimonial="la communaute legale",
        conjoint=SpfplConjoint(
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Martin",
        ),
        adresse_personnelle=Address(adresse_affichee="5 rue Royale, 75008 Paris"),
        adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ordre=SpfplOrdre(
            professionnel="Ordre des chirurgiens-dentistes",
            departement="Paris",
            ville="Paris",
            numero="12345",
            numero_rpps="10000000001",
        ),
    )
    return DocumentGenerationContext(
        structure="SPFPL apport" if is_apport else "SPFPL cession",
        dossier_options=DossierOptions(
            apport=is_apport,
            cession=not is_apport,
            associe_unique=True,
        ),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        operation_spfpl=OperationSpfpl(type=operation),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            forme_sociale="par actions simplifiee",
            capital_social="60 000",
            capital_social_lettres="soixante mille",
            valeur_nominale_action="100",
            valeur_nominale_action_lettres="cent",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
        ),
        actionnaire_unique=founder,
        apport=Apport(montant="60 000", montant_lettres="soixante mille euros"),
        depot_fonds=DepotFonds(
            banque=CessionBanque(
                nom="BANQUE EXEMPLE",
                adresse_affichee="1 boulevard Haussmann, 75009 Paris",
            )
        ),
        apport_titres=ApportTitres(
            nb_parts=60,
            nb_parts_lettres="soixante",
            plage_parts="41 a 100",
            valeur_globale="60 000",
            valeur_nominale_action="100",
            valeur_nominale_action_lettres="cent",
        ),
        societe_cible=SocieteCible(
            denomination="SELARL CABINET MARTIN",
            siege=Address(adresse_affichee="12 avenue des Ternes, 75017 Paris"),
            ville_rcs="Paris",
            numero_rcs="900 000 001",
        ),
        capital_souscription=CapitalSouscription(
            nb_actions_total=600,
            valeur_nominale_action="100",
        ),
        exercice_social=None,
        commissaire_aux_apports=_entity(),
    )


def _entity() -> ProfessionalEntity:
    return ProfessionalEntity(
        denomination="CAA EXPERTISE",
        forme_sociale="SAS",
        capital_social="1 000 euros",
        siege=Address(adresse_affichee="1 rue Scheffer, 75016 Paris"),
        ville_rcs="Paris",
        numero_rcs="948 483 730",
        representant=SpfplRepresentant(
            civilite_affichage="Monsieur",
            prenom="Nabil",
            nom="Saidi",
        ),
    )


def _with_exercice(ctx: DocumentGenerationContext) -> DocumentGenerationContext:
    ctx.exercice_social = ExerciceSocial(
        debut="1er janvier",
        fin="31 decembre",
        date_cloture_premier_exercice="31 decembre 2026",
    )
    return ctx


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return "\n".join(texts)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def test_statuts_spfpl_cession_generates_source_overlay_without_signature_date(
    tmp_path: Path,
) -> None:
    output_path = StatutsSpfplCessionGenerator().generate(
        _with_exercice(_base_context(operation="cession")),
        tmp_path,
    )

    text = _docx_text(output_path)
    document = Document(output_path)
    table_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    acceptance = next(
        p for p in document.paragraphs if "Bon pour acceptation des fonctions" in p.text
    )

    assert output_path.name == "statuts_spfpl_cession.docx"
    assert "Société de Participations Financières de Profession Libérale" in text
    assert "BANQUE EXEMPLE sise 1 boulevard Haussmann, 75009 Paris" in text
    assert "Le\nDocteur Camille Martin" in text
    assert "Le 14/05/2026" not in text
    assert "Nomination d’un commissaire aux apports" not in text
    assert "DECISIONS DES ACTIONNAIRES" in table_text
    assert any(run.italic for run in acceptance.runs)
    _assert_clean(text)


def test_statuts_spfpl_apport_generates_nature_overlay_and_signature_date(
    tmp_path: Path,
) -> None:
    output_path = StatutsSpfplApportGenerator().generate(
        _with_exercice(_base_context(operation="apport")),
        tmp_path,
    )

    text = _docx_text(output_path)

    assert output_path.name == "statuts_spfpl_apport.docx"
    assert "Apports en nature" in text
    assert "SELARL CABINET MARTIN" in text
    assert "ayant son siège 12 avenue des Ternes, 75017 Paris" in text
    assert "Le 14/05/2026" in text
    assert "Ouverture d'un compte bancaire auprès de la Banque" not in text
    _assert_clean(text)


def test_statuts_spfpl_blocks_multi_associes(tmp_path: Path) -> None:
    ctx = _with_exercice(_base_context(operation="cession"))
    ctx.capital_souscription.souscripteurs.append(CapitalSouscripteur(prenom="Camille"))
    ctx.capital_souscription.souscripteurs.append(CapitalSouscripteur(prenom="Louise"))

    with pytest.raises(ValueError, match="multi-associes"):
        StatutsSpfplCessionGenerator().generate(ctx, tmp_path)
