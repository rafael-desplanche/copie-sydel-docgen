from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Company,
    DocumentGenerationContext,
    Person,
    Signature,
    StatutsCivilsApport,
    StatutsCivilsAssocie,
    StatutsCivilsCapitalDepot,
    StatutsCivilsContext,
    StatutsCivilsParts,
    StatutsCivilsRepresentant,
)
from sydel_doc_engine.generators.lot_04.statuts_scm import StatutsScmGenerator


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return "\n".join(texts)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _morale_associe(*, parts: int = 70, apport: str = "700") -> StatutsCivilsAssocie:
    return StatutsCivilsAssocie(
        type_personne="personne_morale",
        denomination="SELARL DURAND",
        forme_juridique="SELARL",
        profession="chirurgien-dentiste",
        capital_social="1 000 euros",
        siege=Address(adresse_affichee="5 rue Royale, 75008 Paris"),
        numero_rcs="900 000 001",
        ville_rcs="Paris",
        representant=StatutsCivilsRepresentant(
            civilite_affichage="Monsieur",
            prenom="Jean",
            nom="Durand",
            fonction="gerant",
        ),
        apport=StatutsCivilsApport(montant=apport, montant_lettres="sept cents euros"),
        parts=StatutsCivilsParts(nb=parts),
    )


def _person_associe(*, parts: int = 50, apport: str = "500") -> StatutsCivilsAssocie:
    return StatutsCivilsAssocie(
        type_personne="personne_physique",
        genre=Gender.FEMININ,
        civilite_affichage="Madame",
        prenom="Alice",
        prenoms="Alice",
        nom="Martin",
        profession="chirurgien-dentiste",
        date_naissance="2 fevrier 1982",
        ville_naissance="Lyon",
        nationalite="francaise",
        situation_maritale="celibataire",
        adresse_personnelle_affichee="2 rue Exemple, 69000 Lyon",
        apport=StatutsCivilsApport(montant=apport, montant_lettres="cinq cents euros"),
        parts=StatutsCivilsParts(nb=parts),
    )


def _context() -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SCM",
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Durand",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 15)),
        societe=Company(
            denomination="SCM CABINET DURAND MARTIN",
            denomination_courte="CABINET DURAND MARTIN",
            forme_sociale="Societe civile de moyens",
            siege=Address(
                num_voie="10",
                voie="rue de la Paix",
                cp="75002",
                ville="Paris",
            ),
        ),
        statuts_civils=StatutsCivilsContext(
            type="scm",
            forme_sociale="Societe civile de moyens",
            capital_social="1200",
            capital_social_lettres="mille deux cents euros",
            nb_parts_total=120,
            valeur_nominale_part="10 euros",
            capital_depot=StatutsCivilsCapitalDepot(
                banque_nom="BANQUE EXEMPLE",
                banque_adresse="1 rue Banque, 75009 Paris",
            ),
            associes=[
                _morale_associe(),
                _person_associe(),
            ],
        ),
    )


def test_statuts_scm_generates_dynamic_associates_apports_parts_and_signatures(
    tmp_path: Path,
) -> None:
    output_path = StatutsScmGenerator().generate(_context(), tmp_path)
    text = _docx_text(output_path)
    document = Document(output_path)
    table_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    lu_approuve = next(p for p in document.paragraphs if "Lu et approuvé" in p.text)

    assert output_path.name == "statuts_scm.docx"
    assert "Article 4 ‐ Objet social" in text
    assert "SELARL DURAND, représentée par Monsieur Jean Durand 70 parts" in text
    assert "Madame Alice Martin 50 parts" in text
    assert "ci- 500." in text
    assert "510" not in text
    assert "« Lu et approuvé »" in text
    assert "STATUTS" in table_text
    assert any(run.italic for run in lu_approuve.runs)
    _assert_clean(text)


def test_statuts_scm_blocks_when_parts_total_is_ambiguous(tmp_path: Path) -> None:
    ctx = _context()
    ctx.statuts_civils.associes[1].parts.nb = None

    with pytest.raises(ValueError, match="associes\\[\\]\\.parts\\.nb"):
        StatutsScmGenerator().generate(ctx, tmp_path)


def test_statuts_scm_blocks_when_capital_total_is_inconsistent(tmp_path: Path) -> None:
    ctx = _context()
    ctx.statuts_civils.associes[1].apport.montant = "510"

    with pytest.raises(ValueError, match="capital_social"):
        StatutsScmGenerator().generate(ctx, tmp_path)


def test_statuts_scm_requires_morale_representant_for_signature(tmp_path: Path) -> None:
    ctx = _context()
    ctx.statuts_civils.associes[0].representant = None

    with pytest.raises(ValueError, match="representant"):
        StatutsScmGenerator().generate(ctx, tmp_path)
