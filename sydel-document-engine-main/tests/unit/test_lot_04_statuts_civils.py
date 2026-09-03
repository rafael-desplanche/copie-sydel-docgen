from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Company,
    DocumentGenerationContext,
    Person,
    Signature,
    StatutsCivilsApport,
    StatutsCivilsAssocie,
    StatutsCivilsCapitalDepot,
    StatutsCivilsContext,
    StatutsCivilsGroupeParts,
    StatutsCivilsParts,
    StatutsCivilsRepresentant,
)
from sydel_doc_engine.generators.lot_04.statuts_sci import StatutsSciGenerator
from sydel_doc_engine.generators.lot_04.statuts_sci_iris import StatutsSciIrisGenerator
from sydel_doc_engine.generators.lot_04.statuts_scs import StatutsScsGenerator


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs if paragraph.text)
    return "\n".join(texts)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _person_associe(
    *,
    prenom: str,
    nom: str,
    nb_parts: int,
    debut: int,
    fin: int,
    role: str | None = None,
    montant: str | None = None,
) -> StatutsCivilsAssocie:
    amount = montant or str(nb_parts * 10)
    return StatutsCivilsAssocie(
        type_personne="personne_physique",
        role_statutaire=role,
        genre=Gender.MASCULIN,
        civilite_affichage="Monsieur",
        prenom=prenom,
        prenoms=prenom,
        nom=nom,
        date_naissance="1 janvier 1980",
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="francaise",
        situation_maritale="celibataire",
        adresse_personnelle=Address(
            num_voie="1",
            voie="rue Exemple",
            cp="75000",
            ville="Paris",
        ),
        apport=StatutsCivilsApport(montant=amount, montant_lettres=amount),
        parts=StatutsCivilsParts(
            nb=nb_parts,
            nb_lettres=str(nb_parts),
            plage_affichee=f"{debut} a {fin}",
            debut=debut,
            fin=fin,
            qualite_associe=role,
        ),
    )


def _morale_associe(*, nb_parts: int, debut: int, fin: int) -> StatutsCivilsAssocie:
    amount = str(nb_parts * 10)
    return StatutsCivilsAssocie(
        type_personne="personne_morale",
        denomination="SEL IRIS",
        forme_juridique="SELARL",
        capital_social="1 000 euros",
        siege=Address(adresse_affichee="2 rue Pro, 75000 Paris"),
        numero_rcs="900 000 001",
        ville_rcs="Paris",
        representant=StatutsCivilsRepresentant(
            civilite_affichage="Monsieur",
            prenom="Jean",
            nom="Durand",
            fonction="gerant",
        ),
        apport=StatutsCivilsApport(montant=amount, montant_lettres=amount),
        parts=StatutsCivilsParts(
            nb=nb_parts,
            nb_lettres=str(nb_parts),
            debut=debut,
            fin=fin,
        ),
    )


def _base_context(
    *,
    structure: str,
    statuts_type: str,
    associes: list[StatutsCivilsAssocie],
) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure=structure,
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Durand",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 15)),
        societe=Company(
            denomination=f"{structure} EXEMPLE",
            forme_sociale=structure,
            siege=Address(
                num_voie="10",
                voie="rue de la Paix",
                cp="75002",
                ville="Paris",
                adresse_affichee="10 rue de la Paix, 75002 Paris",
            ),
            ville_rcs="Paris",
        ),
        statuts_civils=StatutsCivilsContext(
            type=statuts_type,
            forme_sociale="societe civile",
            mention_capital_variable="a capital variable",
            capital_social="1000",
            capital_social_lettres="mille",
            capital_autorise="10000",
            capital_autorise_lettres="dix mille",
            capital_maximal="10000",
            capital_maximal_lettres="dix mille",
            nb_parts_total=100,
            nb_parts_total_lettres="cent",
            valeur_nominale_part="10",
            valeur_nominale_part_lettres="dix",
            plage_parts_totale="1 a 100",
            duree_societe="99",
            capital_depot=StatutsCivilsCapitalDepot(
                banque_nom="BANQUE EXEMPLE",
                banque_adresse="1 rue Banque, 75009 Paris",
            ),
            associes=associes,
            date_cloture_premier_exercice="31 decembre 2026",
            nombre_exemplaires_lettres="trois",
            denomination_cabinet_mandataire="DAAT",
        ),
    )


def test_statuts_sci_generates_dynamic_associates(tmp_path: Path) -> None:
    ctx = _base_context(
        structure="SCI",
        statuts_type="sci",
        associes=[
            _person_associe(prenom="Jean", nom="Durand", nb_parts=40, debut=1, fin=40),
            _person_associe(prenom="Alice", nom="Martin", nb_parts=60, debut=41, fin=100),
        ],
    )

    output_path = StatutsSciGenerator().generate(ctx, tmp_path)
    text = _docx_text(output_path)

    assert output_path.name == "statuts_sci.docx"
    assert "ARTICLE 1 - FORME" in text
    assert "Monsieur Jean Durand" in text
    assert "Monsieur Alice Martin" in text
    assert "A Paris, le 15/05/2026" in text
    _assert_clean(text)


def test_statuts_scs_requires_commandite_and_commanditaire(tmp_path: Path) -> None:
    ctx = _base_context(
        structure="SCS",
        statuts_type="scs",
        associes=[
            _person_associe(
                prenom="Jean",
                nom="Durand",
                nb_parts=100,
                debut=1,
                fin=100,
                role="commandite",
                montant="1000",
            )
        ],
    )
    ctx.statuts_civils.total_apports_commandites = "1000"

    with pytest.raises(ValueError, match="commanditaire"):
        StatutsScsGenerator().generate(ctx, tmp_path)


def test_statuts_scs_generates_roles_and_lu_approuve(tmp_path: Path) -> None:
    ctx = _base_context(
        structure="SCS",
        statuts_type="scs",
        associes=[
            _person_associe(
                prenom="Jean",
                nom="Durand",
                nb_parts=60,
                debut=1,
                fin=60,
                role="commandite",
                montant="600",
            ),
            _person_associe(
                prenom="Alice",
                nom="Martin",
                nb_parts=40,
                debut=61,
                fin=100,
                role="commanditaire",
                montant="400",
            ),
        ],
    )
    ctx.statuts_civils.total_apports_commandites = "600"

    output_path = StatutsScsGenerator().generate(ctx, tmp_path)
    text = _docx_text(output_path)
    document = Document(output_path)
    signature_table_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    assert output_path.name == "statuts_scs.docx"
    assert "Associes commandites" in text
    assert "Associes commanditaires" in text
    assert "Lu et approuve" in text
    assert "Monsieur Jean Durand" in signature_table_text
    assert "Monsieur Alice Martin" in signature_table_text
    _assert_clean(text)


def test_statuts_sci_iris_requires_result_groups(tmp_path: Path) -> None:
    ctx = _base_context(
        structure="SCI IRIS",
        statuts_type="sci_iris",
        associes=[
            _morale_associe(nb_parts=40, debut=1, fin=40),
            _person_associe(prenom="Alice", nom="Martin", nb_parts=60, debut=41, fin=100),
        ],
    )

    with pytest.raises(ValueError, match="resultat_groupes_parts"):
        StatutsSciIrisGenerator().generate(ctx, tmp_path)


def test_statuts_sci_iris_generates_morale_and_result_groups(tmp_path: Path) -> None:
    ctx = _base_context(
        structure="SCI IRIS",
        statuts_type="sci_iris",
        associes=[
            _morale_associe(nb_parts=40, debut=1, fin=40),
            _person_associe(prenom="Alice", nom="Martin", nb_parts=60, debut=41, fin=100),
        ],
    )
    ctx.statuts_civils.resultat_groupes_parts = [
        StatutsCivilsGroupeParts(
            parts_debut=1,
            parts_fin=40,
            quote_part_resultat_exceptionnel="40 %",
        ),
        StatutsCivilsGroupeParts(
            parts_debut=41,
            parts_fin=100,
            quote_part_resultat_exceptionnel="60 %",
        ),
    ]
    ctx.statuts_civils.resultat_quote_part_exceptionnel_total = "100 %"

    output_path = StatutsSciIrisGenerator().generate(ctx, tmp_path)
    text = _docx_text(output_path)
    document = Document(output_path)
    matrix_table_text = "\n".join(
        paragraph.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    assert output_path.name == "statuts_sci_iris.docx"
    assert "SCI IRIS" in text
    assert "SEL IRIS, representee par Monsieur Jean Durand" in text
    assert "Parts 1 a 40" in matrix_table_text
    assert "40 %" in matrix_table_text
    assert "Total" in matrix_table_text
    assert "100 %" in matrix_table_text
    _assert_clean(text)
