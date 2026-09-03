from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Company,
    DocumentGenerationContext,
    Domiciliation,
    Person,
    Signature,
)
from sydel_doc_engine.generators.lot_01.autorisation_domiciliation import (
    AutorisationDomiciliationGenerator,
)


def _context(
    genre: Gender = Gender.MASCULIN,
    *,
    adresse_domiciliation_affichee: str = "15 rue du Libre, Lyon 69002",
    image_optionnelle: Path | None = None,
) -> DocumentGenerationContext:
    civilite = "Madame" if genre == Gender.FEMININ else "Monsieur"
    prenom = "Marie" if genre == Gender.FEMININ else "Jean"
    return DocumentGenerationContext(
        personne_signataire=Person(
            genre=genre,
            civilite=civilite,
            prenom=prenom,
            nom="Durand",
        ),
        societe=Company(
            denomination="DURAND CONSEIL",
            capital="1 000",
            siege=Address(
                num_voie="80",
                voie="avenue Marceau",
                cp="75008",
                ville="Paris",
            ),
        ),
        domiciliation=Domiciliation(
            adresse_domiciliation_affichee=adresse_domiciliation_affichee,
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 12),
            image_optionnelle=image_optionnelle,
        ),
    )


def _generate(
    tmp_path: Path,
    genre: Gender = Gender.MASCULIN,
    *,
    adresse_domiciliation_affichee: str = "15 rue du Libre, Lyon 69002",
    image_optionnelle: Path | None = None,
) -> Path:
    return AutorisationDomiciliationGenerator().generate(
        _context(
            genre,
            adresse_domiciliation_affichee=adresse_domiciliation_affichee,
            image_optionnelle=image_optionnelle,
        ),
        tmp_path,
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts: list[str] = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and borders.find(qn("w:top")) is not None


def _table_text(table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def test_autorisation_domiciliation_creates_docx(tmp_path: Path) -> None:
    output_dir = tmp_path / "nested"

    output_path = _generate(output_dir)

    assert output_path == output_dir / "autorisation_domiciliation.docx"
    assert output_path.is_file()


def test_autorisation_domiciliation_contains_essential_texts(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    # Le titre vit dans le tableau d'en-tete du modele source.
    assert "AUTORISATION DE DOMICILIATION" in text
    # Phrase fidele au modele tokenise corrige (retour humain LOCK V1) : adresse
    # complete du cabinet/siege « du cabinet au [num_voie] [voie], [cp] [ville] » ;
    # terme juridique « pour une durée indéterminée » conserve (et non « pour 99 ans »).
    assert (
        "autorise la domiciliation de DURAND CONSEIL au capital de 1 000 € "
        "en cours de formation, dans les locaux du cabinet au 80 avenue Marceau, "
        "75008 Paris, pour une durée indéterminée."
    ) in text
    assert "Fait à Paris" in text
    assert "Le 12 mai 2026" in text
    assert "Monsieur Jean Durand" in text


def test_autorisation_domiciliation_preserves_legal_term(tmp_path: Path) -> None:
    text = _docx_text(_generate(tmp_path))

    # Garde-fou anti-derive : le terme juridique du modele doit rester intact
    # et l'ancienne paraphrase « 99 ans » ne doit jamais reapparaitre.
    assert "pour une durée indéterminée." in text
    assert "99 ans" not in text


def test_autorisation_domiciliation_opening_agrees_masculine(tmp_path: Path) -> None:
    # Le modele source fige l'ouverture au feminin (« Je soussignée »). Pour un
    # signataire masculin, la couche genre doit l'accorder en « Je soussigné ».
    text = _docx_text(_generate(tmp_path, Gender.MASCULIN))

    assert "Je soussigné Monsieur Jean Durand autorise la domiciliation" in text
    assert "Je soussignée Monsieur Jean Durand" not in text


def test_autorisation_domiciliation_opening_agrees_feminine(tmp_path: Path) -> None:
    # Pour une signataire feminine, l'ouverture figee « Je soussignée » du modele
    # est conservee telle quelle (accord deja correct).
    text = _docx_text(_generate(tmp_path, Gender.FEMININ))

    assert "Je soussignée Madame Marie Durand autorise la domiciliation" in text


def test_autorisation_domiciliation_ignores_free_address_for_wording(
    tmp_path: Path,
) -> None:
    adresse = "Bâtiment B, 4 impasse des Tests, Marseille 13002"

    text = _docx_text(_generate(tmp_path, adresse_domiciliation_affichee=adresse))

    # L'adresse libre de domiciliation n'est pas injectee : le modele s'appuie
    # sur l'adresse complete du siege/cabinet de la societe.
    assert adresse not in text
    assert "dans les locaux du cabinet au 80 avenue Marceau, 75008 Paris," in text


def test_autorisation_domiciliation_uses_company_seat_city(
    tmp_path: Path,
) -> None:
    text = _docx_text(_generate(tmp_path))

    # Le modele utilise l'adresse complete du siege ([num_voie_siege] [voie_siege],
    # [cp_siege] [ville_siege]), pas l'adresse de domiciliation libre.
    assert "15 rue du Libre, Lyon 69002" not in text
    assert "80 avenue Marceau, 75008 Paris" in text


def test_autorisation_domiciliation_does_not_use_signature_image(tmp_path: Path) -> None:
    missing_image = tmp_path / "signature_absente.png"

    output_path = _generate(tmp_path, image_optionnelle=missing_image)

    assert len(Document(output_path).inline_shapes) == 0


def test_autorisation_domiciliation_uses_signature_paragraphs_without_table(
    tmp_path: Path,
) -> None:
    document = Document(_generate(tmp_path))

    # Le modele source porte un unique tableau (l'en-tete de titre encadre).
    assert len(document.tables) == 1
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    assert "Fait à Paris" in paragraphs
    assert "Le 12 mai 2026" in paragraphs
    assert "Monsieur Jean Durand" in paragraphs
