from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    BailContext,
    BailParty,
    CessionAcquereur,
    CessionBanque,
    CessionCabinet,
    CessionContext,
    CessionDestinataire,
    CessionFinancement,
    CessionVendeur,
    Company,
    DocumentContext,
    DocumentGenerationContext,
    DocumentSignataire,
    DossierOptions,
    Person,
    Signature,
)
from sydel_doc_engine.generators.lot_03.appel_fond_sel import AppelFondSelGenerator
from sydel_doc_engine.generators.lot_03.avenant_contrat_bail import (
    AvenantContratBailGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _context(
    structure: str = "SELARL",
    *,
    cession: bool = True,
    type_cabinet: str = "dentaire",
    societe_en_cours_immatriculation: bool = True,
) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(cession=cession),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        societe=Company(
            denomination="SELARL CABINET MARTIN",
            ville_rcs="Paris",
            siege=Address(adresse_affichee="12 avenue des Ternes, 75017 Paris"),
        ),
        bail=BailContext(
            date_avenant=date(2026, 5, 14),
            date_signature_origine=date(2021, 9, 1),
            societe_en_cours_immatriculation=societe_en_cours_immatriculation,
            bailleur_accepte_changement_locataire=True,
            bailleur=BailParty(
                civilite_affichage="Monsieur",
                prenom="Paul",
                nom="Leroy",
                profession="bailleur",
                date_naissance=date(1970, 1, 5),
                ville_naissance="Lyon",
                nationalite="française",
                adresse_affichee="8 rue Victor Hugo, 69002 Lyon",
            ),
            locataire=BailParty(
                civilite_affichage="Monsieur",
                civilite_courte="Docteur",
                prenom="Camille",
                nom="Martin",
                profession="chirurgien-dentiste",
                date_naissance=date(1984, 6, 20),
                ville_naissance="Paris",
                nationalite="française",
                adresse_affichee="4 rue du Bac, 75007 Paris",
            ),
        ),
        cession=CessionContext(
            type_cabinet=type_cabinet,
            financement=CessionFinancement(
                banque=CessionBanque(nom="BANQUE EXEMPLE"),
                destinataire=CessionDestinataire(
                    civilite_affichage="Monsieur",
                    prenom="Louis",
                    nom="Bernard",
                ),
                montant_deblocage="150 000",
            ),
            cabinet=CessionCabinet(
                denomination_ou_adresse_affichee="Cabinet dentaire des Ternes",
            ),
            vendeur=CessionVendeur(
                civilite_affichage="Monsieur",
                prenom="Camille",
                nom="Martin",
            ),
            acquereur=CessionAcquereur(
                denomination_societe="SELARL CABINET MARTIN",
            ),
        ),
        document=DocumentContext(
            nombre_exemplaires_lettres="quatre",
            signataire=DocumentSignataire(prenom="Camille", nom="Martin"),
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_no_source_placeholders(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and all(
        borders.find(qn(f"w:{edge}")) is not None
        for edge in ("top", "left", "bottom", "right")
    )


def test_avenant_contrat_bail_generates_source_wording_and_signature_table(
    tmp_path: Path,
) -> None:
    output_path = AvenantContratBailGenerator().generate(_context(), tmp_path)

    assert output_path == tmp_path / "avenant_contrat_bail.docx"
    text = _docx_text(output_path)
    assert "Avenant n°1 au bail du 14/05/2026" in text
    assert "ARTICLE 1 : changement de locataire" in text
    assert "les démarches seront finies" in text
    assert "Le Bailleur" in text
    assert "L’ancien locataire" in text
    assert text.count("Le nouveau locataire") == 1
    document = Document(output_path)
    assert abs(document.sections[0].top_margin - Cm(1.75)) < 300
    assert abs(document.sections[0].bottom_margin - Cm(0.5)) < 300
    party_marker = next(p for p in document.paragraphs if p.text.endswith("le Bailleur »"))
    assert party_marker.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    assert party_marker.runs[0].bold is True
    assert party_marker.runs[0].underline is True
    article = next(p for p in document.paragraphs if p.text.startswith("ARTICLE 1"))
    assert article.runs[0].underline is True
    assert _table_has_explicit_borders(document.tables[-1])
    _assert_no_source_placeholders(text)


def test_appel_fond_sel_generates_dentaire_request(tmp_path: Path) -> None:
    output_path = AppelFondSelGenerator().generate(_context(), tmp_path)

    assert output_path == tmp_path / "appel_fond_sel.docx"
    text = _docx_text(output_path)
    assert "BANQUE EXEMPLE" in text
    assert "Objet : demande de déblocage des fonds" in text
    assert "150 000" in text
    assert "cabinet dentaire exploité au Cabinet dentaire des Ternes" in text
    assert "Montant du fond" not in text
    document = Document(output_path)
    subject = next(p for p in document.paragraphs if p.text.startswith("Objet"))
    assert subject.runs[0].bold is True
    assert subject.runs[0].underline is True
    amount = next(p for p in document.paragraphs if p.text == "150 000")
    assert amount.alignment == WD_ALIGN_PARAGRAPH.CENTER
    signature = next(p for p in document.paragraphs if p.text == "Camille Martin")
    assert signature.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    # Retour UAT Rafael (DOC-008) : bloc banque + lieu/date aligne a DROITE.
    banque = next(p for p in document.paragraphs if p.text == "BANQUE EXEMPLE")
    assert banque.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    lieu_date = next(p for p in document.paragraphs if p.text.startswith("Paris, le"))
    assert lieu_date.alignment == WD_ALIGN_PARAGRAPH.RIGHT
    _assert_no_source_placeholders(text)


def test_appel_fond_sel_generates_medical_request(tmp_path: Path) -> None:
    # L'appel de fonds est un document commun « Si cession » : il doit etre genere
    # pour une cession MEDICALE, avec « cabinet medical » (et non plus « dentaire »).
    ctx = _context(type_cabinet="medical")

    output_path = AppelFondSelGenerator().generate(ctx, tmp_path)

    assert output_path == tmp_path / "appel_fond_sel.docx"
    text = _docx_text(output_path)
    assert "cabinet médical exploité au Cabinet dentaire des Ternes" in text
    assert "cabinet dentaire exploité" not in text
    _assert_no_source_placeholders(text)


def test_avenant_contrat_bail_blocks_already_registered_company(tmp_path: Path) -> None:
    ctx = _context(societe_en_cours_immatriculation=False)

    with pytest.raises(ValueError, match="societe_en_cours_immatriculation"):
        AvenantContratBailGenerator().generate(ctx, tmp_path)


def test_orchestrator_selects_bail_batch_for_selarl_dentaire() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id for document in orchestrator.select_documents_for_context(_context())
    }

    assert {"DOC-007", "DOC-008"}.issubset(selected_ids)


def test_orchestrator_selects_appel_fonds_for_selarl_medical() -> None:
    # L'appel de fonds (DOC-008) est commun a toute cession SELARL : il doit etre
    # selectionne pour une cession MEDICALE, pas seulement dentaire.
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(
            _context("SELARL", type_cabinet="medical"),
        )
    }

    assert {"DOC-007", "DOC-008"}.issubset(selected_ids)


@pytest.mark.parametrize("type_cabinet", ["dentaire", "medical"])
def test_orchestrator_excludes_appel_fonds_for_non_selarl(type_cabinet: str) -> None:
    # L'appel de fonds SEL reste borne a la structure SELARL : exclu pour SELAS.
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(
            _context("SELAS", type_cabinet=type_cabinet),
        )
    }

    assert "DOC-008" not in selected_ids


def test_orchestrator_excludes_bail_batch_when_cession_option_is_false() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(_context(cession=False))
    }

    assert "DOC-007" not in selected_ids
    assert "DOC-008" not in selected_ids
