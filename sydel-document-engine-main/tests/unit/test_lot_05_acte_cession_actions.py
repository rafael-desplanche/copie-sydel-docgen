from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    AssocieCible,
    CessionActions,
    DocumentContext,
    DocumentGenerationContext,
    DossierOptions,
    OperationSpfpl,
    Person,
    Signature,
    SocieteCible,
    SocieteSpfpl,
    SpfplConjoint,
    SpfplOrdre,
    SpfplPerson,
    SpfplRepresentant,
)
from sydel_doc_engine.generators.lot_05.acte_cession_actions_spfpl import (
    ActeCessionActionsSpfplGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _base_context() -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SPFPL cession",
        dossier_options=DossierOptions(cession=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(
            lieu="Paris",
            date=date(2026, 5, 15),
            nombre_exemplaires="trois",
            prestataire_signature_electronique="Yousign",
        ),
        operation_spfpl=OperationSpfpl(
            type="cession",
            nature_titres="actions",
            document_demande="acte_cession_actions",
        ),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            forme_sociale="SPFPLAS",
            capital_social="60 000",
            ville_rcs="Paris",
            numero_rcs="en cours",
            departement_inscription_ordre="Paris",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
            representant=SpfplRepresentant(
                civilite_affichage="Monsieur",
                civilite_courte="M.",
                prenom="Camille",
                nom="Martin",
                fonction="President",
            ),
        ),
        cedant=_cedant(),
        societe_cible=SocieteCible(
            denomination="SELAS CABINET MARTIN",
            forme_sociale="SELAS",
            forme_sociale_complete="societe d'exercice liberal par actions simplifiee",
            profession_reglementee="chirurgien-dentiste",
            profession_reglementee_pluriel="chirurgiens-dentistes",
            capital_social="10 000",
            nb_actions_total=100,
            valeur_nominale_action="100",
            valeur_nominale_action_lettres="cent euros",
            departement_inscription_ordre="Paris",
            presentation_dirigeants=(
                "Le Président est Docteur Camille Martin et les Directeurs Généraux "
                "sont Docteur Louise Bernard et Docteur Paul Durand."
            ),
            siege=Address(adresse_affichee="12 avenue des Ternes, 75017 Paris"),
            ville_rcs="Paris",
            numero_rcs="900 000 001",
        ),
        associes_cible=[
            AssocieCible(
                type="personne_morale",
                denomination="SELAS HOLDING EXEMPLE",
                nb_actions_avant=10,
                nb_actions_avant_lettres="dix",
            ),
            AssocieCible(
                civilite_affichage="Docteur",
                prenom="Louise",
                nom="Bernard",
                nb_actions_avant=30,
                nb_actions_avant_lettres="trente",
            ),
            AssocieCible(
                civilite_affichage="Docteur",
                prenom="Camille",
                nom="Martin",
                nb_actions_avant=60,
                nb_actions_avant_lettres="soixante",
                est_cedant=True,
            ),
        ],
        cession_actions=CessionActions(
            nb_actions=60,
            nb_actions_lettres="soixante",
            prix_total="60000",
            prix_total_lettres="soixante mille euros",
            prix_unitaire_action="1000",
            prix_unitaire_action_lettres="mille euros",
            modalites_paiement="credit_bancaire_comptant_cheque_banque",
            nombre_exemplaires_lettres="trois",
            agrement_unanime_confirme=True,
            pv_agrement_coherent=True,
            gap_applicable=True,
            representant_cessionnaire_confirme=True,
            titre_signature_cedant="Dr",
        ),
        document=DocumentContext(nombre_exemplaires_lettres="trois"),
    )


def _cedant() -> SpfplPerson:
    return SpfplPerson(
        civilite_affichage="Docteur",
        prenom="Camille",
        nom="Martin",
        genre=Gender.MASCULIN,
        profession="chirurgien-dentiste",
        profession_reglementee="chirurgien-dentiste",
        profession_reglementee_pluriel="chirurgiens-dentistes",
        date_naissance=date(1980, 1, 2),
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="francaise",
        situation_maritale="marie",
        regime_matrimonial="la communauté légale",
        conjoint=SpfplConjoint(
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Martin",
        ),
        adresse_personnelle=Address(adresse_affichee="5 rue Royale, 75008 Paris"),
        adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ordre=SpfplOrdre(
            professionnel="Ordre des chirurgiens-dentistes",
            departement="Paris",
            numero="12345",
            numero_rpps="10000000001",
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def test_acte_cession_actions_generates_source_vocabulary_and_clean_docx(
    tmp_path: Path,
) -> None:
    output_path = ActeCessionActionsSpfplGenerator().generate(_base_context(), tmp_path)

    text = _docx_text(output_path)

    assert output_path.name == "acte_cession_actions_spfpl.docx"
    assert "Cession d'actions" in text
    assert "OBJET DU CONTRAT : CESSION D'ACTIONS" in text
    assert "Actions Cédées" in text
    assert "Titres Cédés" in text
    assert "Remplir les conditions exigées par la loi pour détenir des actions de SELAS" in text
    assert "GARANTIE D'ACTIF ET DE PASSIF / GAP" in text
    assert "service Yousign" in text
    assert "Cession de parts" not in text
    assert "parts sociales" not in text
    assert "[" not in text
    assert "]" not in text


def test_acte_cession_actions_blocks_non_actions_context(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.operation_spfpl.nature_titres = "parts"

    with pytest.raises(ValueError, match="nature_titres"):
        ActeCessionActionsSpfplGenerator().generate(ctx, tmp_path)


def test_acte_cession_actions_blocks_unconfirmed_source_overlays(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.cession_actions.gap_applicable = False

    with pytest.raises(ValueError, match="gap_applicable"):
        ActeCessionActionsSpfplGenerator().generate(ctx, tmp_path)


def test_acte_cession_actions_blocks_incoherent_price(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.cession_actions.prix_total = "61000"

    with pytest.raises(ValueError, match="prix_total"):
        ActeCessionActionsSpfplGenerator().generate(ctx, tmp_path)


def test_orchestrator_selects_acte_actions_only_for_explicit_actions_context() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())
    selected = orchestrator.select_documents_for_context(_base_context())

    assert "DOC-029" in [document.doc_id for document in selected]

    ctx = _base_context()
    ctx.operation_spfpl.document_demande = "acte_cession_parts"

    selected_without_actions = orchestrator.select_documents_for_context(ctx)
    assert "DOC-029" not in [document.doc_id for document in selected_without_actions]
