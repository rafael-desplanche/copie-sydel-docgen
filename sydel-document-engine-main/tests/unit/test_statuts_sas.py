from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    CapitalSouscripteur,
    CapitalSouscription,
    CessionBanque,
    DepotFonds,
    DocumentGenerationContext,
    DossierOptions,
    ExerciceSocial,
    Person,
    Signature,
    SocieteSpfpl,
    SpfplConjoint,
    SpfplOrdre,
    SpfplPerson,
    StatutsPresident,
    StatutsSas,
)
from sydel_doc_engine.generators.lot_04.statuts_sas import StatutsSasGenerator
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _context() -> DocumentGenerationContext:
    actionnaire = SpfplPerson(
        civilite_affichage="Docteur",
        prenom="Camille",
        nom="Martin",
        genre=Gender.MASCULIN,
        profession="medecin",
        qualification_principale="Médecin cardiologue",
        date_naissance="2 janvier 1980",
        ville_naissance="Paris",
        departement_naissance="75",
        nationalite="française",
        situation_maritale="Marié",
        regime_matrimonial="la communauté légale",
        conjoint=SpfplConjoint(
            civilite_affichage="Madame",
            prenom="Alice",
            nom="Martin",
        ),
        adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
        ordre=SpfplOrdre(
            departement="Paris",
            numero="12345",
            numero_rpps="10000000001",
        ),
        nb_actions=120,
    )
    return DocumentGenerationContext(
        structure="SAS",
        dossier_options=DossierOptions(associe_unique=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        statuts_sas=StatutsSas(type="spfpl_medecins", profession="medecin"),
        societe_spfpl=SocieteSpfpl(
            denomination="SPFPL MARTIN",
            capital_social="12 000",
            capital_social_lettres="douze mille",
            nb_actions_total=120,
            nb_actions_total_lettres="cent vingt",
            valeur_nominale_action="100",
            valeur_nominale_action_lettres="cent euros",
            profession="medecin",
            siege=Address(adresse_affichee="10 rue de la Paix, 75002 Paris"),
        ),
        actionnaire_unique=actionnaire,
        president=StatutsPresident(
            ref_associe_index=0,
            civilite_affichage="Docteur",
            prenom="Camille",
            nom="Martin",
            adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
            duree_mandat="illimitee",
        ),
        depot_fonds=DepotFonds(
            banque=CessionBanque(nom="BANQUE EXEMPLE"),
            montant="12 000",
        ),
        exercice_social=ExerciceSocial(
            debut="1er janvier",
            fin="31 décembre",
            date_cloture_premier_exercice="31 décembre 2026",
        ),
        capital_souscription=CapitalSouscription(
            nb_actions_total=120,
            valeur_nominale_action="100",
            apports_numeraire_montant="12 000",
            souscripteurs=[
                CapitalSouscripteur(
                    civilite_affichage="Docteur",
                    prenom="Camille",
                    nom="Martin",
                    profession="medecin",
                    adresse_personnelle_affichee="5 rue Royale, 75008 Paris",
                    nb_actions=120,
                    qualite="actionnaire unique",
                )
            ],
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for section in document.sections:
        texts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text
    assert "OU\n" not in text


def test_statuts_sas_generates_spfpl_medecins_unique_shareholder_docx(
    tmp_path: Path,
) -> None:
    output_path = StatutsSasGenerator().generate(_context(), tmp_path)

    text = _docx_text(output_path)
    document = Document(output_path)
    article_1 = next(p for p in document.paragraphs if p.text == "ARTICLE 1 - FORME")
    acceptance = next(
        p for p in document.paragraphs if "Bon pour acceptation des fonctions" in p.text
    )

    assert output_path.name == "statuts_sas_spfpl_medecins.docx"
    assert "SPFPL MARTIN" in text
    assert "Société de Participations Financières de Profession Libérale de Médecins" in text
    assert "Par le Docteur Camille Martin 12 000" in text
    assert "Le Docteur Camille Martin 120 actions" in text
    assert "L’Associé Unique, Monsieur Camille Martin" in text
    assert "BANQUE EXEMPLE" in text
    assert "SPFPL MARTIN - Statuts constitutifs" in text
    assert any(run.underline for run in article_1.runs)
    assert any(run.italic for run in acceptance.runs)
    _assert_clean(text)


def test_statuts_sas_is_selected_only_for_confirmed_spfpl_medecins_context() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected = orchestrator.select_documents_for_context(_context())

    assert "DOC-015" in [document.doc_id for document in selected]

    ctx = _context()
    ctx.statuts_sas = None

    selected_without_statuts = orchestrator.select_documents_for_context(ctx)

    assert "DOC-015" not in [document.doc_id for document in selected_without_statuts]


def test_statuts_sas_blocks_multiple_souscripteurs(tmp_path: Path) -> None:
    ctx = _context()
    ctx.capital_souscription.souscripteurs.append(
        CapitalSouscripteur(
            civilite_affichage="Docteur",
            prenom="Louise",
            nom="Bernard",
            profession="medecin",
            adresse_personnelle_affichee="9 rue Bleue, 75009 Paris",
            nb_actions=1,
        )
    )

    with pytest.raises(ValueError, match="exactement un souscripteur"):
        StatutsSasGenerator().generate(ctx, tmp_path)


def test_statuts_sas_blocks_incoherent_capital(tmp_path: Path) -> None:
    ctx = _context()
    ctx.societe_spfpl.nb_actions_total = 121

    with pytest.raises(ValueError, match="coherents"):
        StatutsSasGenerator().generate(ctx, tmp_path)


def test_statuts_sas_blocks_president_distinct_from_unique_shareholder(
    tmp_path: Path,
) -> None:
    ctx = _context()
    ctx.president.nom = "Bernard"

    with pytest.raises(ValueError, match="president doit designer"):
        StatutsSasGenerator().generate(ctx, tmp_path)


def test_statuts_sas_blocks_non_married_wording_without_source_variant(
    tmp_path: Path,
) -> None:
    ctx = _context()
    ctx.actionnaire_unique.situation_maritale = "Célibataire"

    with pytest.raises(ValueError, match="phrase matrimoniale"):
        StatutsSasGenerator().generate(ctx, tmp_path)
