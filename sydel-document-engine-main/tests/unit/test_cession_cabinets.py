from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    CessionAccessibiliteCabinetDentaire,
    CessionAcquereur,
    CessionBailProfessionnel,
    CessionCabinet,
    CessionConjoint,
    CessionContext,
    CessionCreditVendeur,
    CessionExercice,
    CessionFinancement,
    CessionPrecedentProprietaire,
    CessionPret,
    CessionPrix,
    CessionRepresentant,
    CessionSalarie,
    CessionScm,
    CessionValidations,
    CessionVendeur,
    DocumentContext,
    DocumentGenerationContext,
    DossierOptions,
    Person,
    Signature,
)
from sydel_doc_engine.generators.lot_03.acte_cession_cabinet_dentaire import (
    ActeCessionCabinetDentaireGenerator,
)
from sydel_doc_engine.generators.lot_03.acte_cession_cabinet_medical import (
    ActeCessionCabinetMedicalGenerator,
)
from sydel_doc_engine.generators.lot_03.compromis_cession_cabinet_dentaire import (
    CompromisCessionCabinetDentaireGenerator,
)
from sydel_doc_engine.generators.lot_03.compromis_cession_cabinet_medical import (
    CompromisCessionCabinetMedicalGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _context(
    *,
    etape: str = "acte",
    type_cabinet: str = "medical",
    credit_vendeur: bool = False,
    validations: CessionValidations | None = None,
    salaries: list[CessionSalarie] | None = None,
    vendeur_genre: Gender = Gender.MASCULIN,
    representant_genre: Gender = Gender.FEMININ,
) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SELARL",
        dossier_options=DossierOptions(cession=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        cession=CessionContext(
            type_cabinet=type_cabinet,
            etape=etape,
            vendeur=CessionVendeur(
                civilite_affichage="Docteur",
                genre=vendeur_genre,
                prenom="Jean",
                nom="Durand",
                profession="chirurgien-dentiste" if type_cabinet == "dentaire" else "medecin",
                date_naissance=date(1975, 3, 10),
                ville_naissance="Lyon",
                departement_naissance="69",
                cp_naissance="69002",
                pays_naissance="France",
                nationalite="francaise",
                adresse_affichee="4 rue Victor Hugo, 69002 Lyon",
                adresse_exercice_affichee="10 rue du Cabinet, 75008 Paris",
                numero_siren="123 456 789",
                numero_ordre="ORD-123",
                numero_rpps="10101010101",
                ordre_departemental="Paris",
                situation_maritale="marie",
                regime_matrimonial="communaute reduite aux acquets",
                conjoint=CessionConjoint(
                    civilite_affichage="Madame",
                    prenom="Claire",
                    nom="Durand",
                ),
            ),
            acquereur=CessionAcquereur(
                denomination_societe="SELARL CABINET DURAND",
                forme_sociale="SELARL",
                capital_social="10 000",
                siege=Address(adresse_affichee="20 avenue de Wagram, 75017 Paris"),
                rcs_ville="Paris",
                numero_rcs="999 888 777",
                numero_siret="999 888 777 00012",
                date_immatriculation=date(2026, 1, 15),
                date_inscription_ordre=date(2026, 2, 1),
                representant=CessionRepresentant(
                    civilite_affichage="Docteur",
                    genre=representant_genre,
                    prenom="Alice",
                    nom="Moreau",
                    fonction="gerante",
                ),
            ),
            cabinet=CessionCabinet(
                nature_fonds_liberal="medecin generaliste",
                adresse_affichee="10 rue du Cabinet, 75008 Paris",
                adresse_locaux_affichee="10 rue du Cabinet, 75008 Paris",
                telephone="01 44 00 00 00",
                superficie_local="80 m2",
                date_origine_propriete=date(2020, 1, 1),
                annees_acquisition_patientele="2020",
                prix_origine_propriete="120 000",
                precedent_proprietaire=CessionPrecedentProprietaire(
                    civilite_affichage="Docteur",
                    prenom="Paul",
                    nom="Bernard",
                ),
            ),
            bail_professionnel=CessionBailProfessionnel(
                date_bail=date(2021, 1, 1),
                duree="six annees",
                date_debut=date(2021, 1, 1),
                date_fin=date(2027, 1, 1),
                date_reconduction_1=date(2027, 1, 1),
                date_reconduction_2=date(2033, 1, 1),
                loyer_mensuel="2 000 euros",
                activite_autorisee_affichee="activite medicale et paramedicale",
            ),
            exercices=[
                CessionExercice(periode="2023", chiffre_affaires="210 000", resultat="80 000"),
                CessionExercice(periode="2024", chiffre_affaires="220 000", resultat="85 000"),
                CessionExercice(periode="2025", chiffre_affaires="230 000", resultat="90 000"),
            ],
            prix=CessionPrix(
                total="300 000",
                total_lettres="trois cent mille euros",
                elements_corporels="50 000",
                elements_corporels_lettres="cinquante mille euros",
                elements_incorporels="250 000",
                elements_incorporels_lettres="deux cent cinquante mille euros",
            ),
            financement=CessionFinancement(
                pret=CessionPret(montant="240 000", taux="4 %", duree="sept ans"),
                credit_vendeur=CessionCreditVendeur(
                    actif=credit_vendeur,
                    # Regle NotebookLM : unite du credit-vendeur = ANNEES (« Trois ans »).
                    # Le wording du modele rend « [duree_credit_vendeur] ans ».
                    duree="trois",
                    montant="60 000",
                    taux="3 %",
                    majoration_interet_retard="2 points",
                ),
            ),
            scm=CessionScm(actif=credit_vendeur, nb_parts_a_ceder="10"),
            salaries=salaries or [],
            accessibilite_cabinet_dentaire=CessionAccessibiliteCabinetDentaire(
                information_requise="Information d'accessibilite dentaire fournie par contexte.",
            ),
            date_limite_realisation=date(2026, 9, 30),
            validations=validations
            or CessionValidations(
                mentions_bail_medical_validees=True,
                origine_compromis_medical_validee=True,
                date_realisation_compromis_validee=True,
                ligne_contrats_travail_medical_supprimee=True,
                salaries_dentaire_deux_valides=True,
            ),
        ),
        document=DocumentContext(
            nombre_pages_lettres="vingt",
            nombre_exemplaires_lettres="quatre",
            annexes=["ETAT DES ELEMENTS CORPORELS CEDES", "COPIE 2035 AMORTISSEMENTS"],
        ),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_no_residual_tokens(text: str) -> None:
    # Le rendu par remplissage de template ne doit laisser aucun token [xxx] residuel.
    assert "[" not in text
    assert "]" not in text


@pytest.mark.parametrize(
    ("generator", "ctx", "filename", "expected_text"),
    [
        (
            ActeCessionCabinetMedicalGenerator(),
            _context(credit_vendeur=True),
            "acte_cession_cabinet_medical.docx",
            "Ordre des Médecins",
        ),
        (
            CompromisCessionCabinetMedicalGenerator(),
            _context(etape="compromis"),
            "compromis_cession_cabinet_medical.docx",
            "Ordre des Médecins",
        ),
        (
            ActeCessionCabinetDentaireGenerator(),
            _context(
                type_cabinet="dentaire",
                salaries=[
                    CessionSalarie(civilite_affichage="Madame", prenom="Lea", nom="Petit"),
                    CessionSalarie(civilite_affichage="Monsieur", prenom="Noe", nom="Robert"),
                ],
            ),
            "acte_cession_cabinet_dentaire.docx",
            "Petit",
        ),
        (
            CompromisCessionCabinetDentaireGenerator(),
            _context(etape="compromis", type_cabinet="dentaire"),
            "compromis_cession_cabinet_dentaire.docx",
            "Entre les soussignés",
        ),
    ],
)
def test_cession_cabinet_generators_render_docx(
    generator,
    ctx: DocumentGenerationContext,
    filename: str,
    expected_text: str,
    tmp_path: Path,
) -> None:
    output_path = generator.generate(ctx, tmp_path)

    assert output_path == tmp_path / filename
    text = _docx_text(output_path)
    # Texte juridique d'origine preserve (rendu fidele par template, non paraphrase).
    assert "Entre les soussignés" in text
    assert expected_text in text
    # Tokens du contexte injectes (vendeur + acquereur).
    assert "SELARL CABINET" in text
    assert "Durand" in text or "Martin" in text
    # Dates rendues en francais long, pas en ISO.
    assert "10 mars 1975" in text or "20 juin 1984" in text
    _assert_no_residual_tokens(text)


def test_acte_medical_blocks_without_medical_bail_validation(tmp_path: Path) -> None:
    ctx = _context(
        validations=CessionValidations(
            mentions_bail_medical_validees=False,
            ligne_contrats_travail_medical_supprimee=True,
        ),
    )

    with pytest.raises(ValueError, match="mentions_bail_medical_validees"):
        ActeCessionCabinetMedicalGenerator().generate(ctx, tmp_path)


def test_credit_vendeur_blocks_outside_medical_acte(tmp_path: Path) -> None:
    ctx = _context(etape="compromis", credit_vendeur=True)

    with pytest.raises(ValueError, match="credit_vendeur"):
        CompromisCessionCabinetMedicalGenerator().generate(ctx, tmp_path)


def test_credit_vendeur_duree_rendered_in_years(tmp_path: Path) -> None:
    # FIX 4 : l'unite de duree du credit-vendeur est l'annee (« ... ans »).
    ctx = _context(credit_vendeur=True)

    text = _docx_text(ActeCessionCabinetMedicalGenerator().generate(ctx, tmp_path))

    assert "dans un délai maximum de trois ans" in text
    assert "Au terme du délai de trois ans" in text
    assert "mois" not in text.split("crédit-vendeur à hauteur")[1].split("intérêt annuel")[0]
    _assert_no_residual_tokens(text)


def test_origine_propriete_describes_vendeur_created_by_default(tmp_path: Path) -> None:
    # FIX 1 : l'origine decrit le VENDEUR (cedant) ; defaut « cree ».
    for generator, etape in (
        (ActeCessionCabinetMedicalGenerator(), "acte"),
        (CompromisCessionCabinetMedicalGenerator(), "compromis"),
    ):
        ctx = _context(etape=etape, credit_vendeur=(etape == "acte"))
        text = _docx_text(generator.generate(ctx, tmp_path / etape))
        # Sujet = vendeur (Docteur Jean Durand), pas l'acquereur (Alice Moreau).
        assert "Docteur Jean Durand est propriétaire des éléments constitutifs du cabinet" in text
        assert "pour l’avoir régulièrement créé le 1 janvier 2020." in text
        assert "Alice Moreau est propriétaire" not in text
        _assert_no_residual_tokens(text)


def test_origine_propriete_purchased_describes_vendeur(tmp_path: Path) -> None:
    # FIX 1 : variante « achete » -> origine via achat anterieur du vendeur.
    ctx = _context(credit_vendeur=True)
    ctx.cession.cabinet.origine_propriete_mode = "achete"
    text = _docx_text(ActeCessionCabinetMedicalGenerator().generate(ctx, tmp_path))

    assert "Docteur Jean Durand est propriétaire des éléments constitutifs du cabinet" in text
    assert "pour les avoir régulièrement acquis auprès de Docteur Paul Bernard" in text
    # Wording source « au prix de <prix> euros » : la donnee porte le seul montant.
    assert "au prix de 120 000 euros." in text
    _assert_no_residual_tokens(text)


def test_origine_propriete_complex_case_blocks_without_validation(tmp_path: Path) -> None:
    # GARDE-FOU souplesse : une origine COMPLEXE (mode non standard) bloque tant
    # qu'un texte libre valide a la main n'est pas fourni.
    ctx = _context(credit_vendeur=True)
    ctx.cession.cabinet.origine_propriete_mode = "succession"

    with pytest.raises(ValueError, match="origine_propriete_mode"):
        ActeCessionCabinetMedicalGenerator().generate(ctx, tmp_path)


def test_origine_propriete_complex_case_renders_free_text_when_validated(tmp_path: Path) -> None:
    # Cas COMPLEXE valide : texte libre rendu tel quel (relecture humaine).
    ctx = _context(credit_vendeur=True)
    ctx.cession.cabinet.origine_propriete_mode = "succession"
    ctx.cession.cabinet.description_origine_propriete = (
        "Le vendeur a recueilli le cabinet par voie de succession de son père en 2015."
    )
    ctx.cession.validations.origine_propriete_complexe_validee = True

    text = _docx_text(ActeCessionCabinetMedicalGenerator().generate(ctx, tmp_path))

    assert "par voie de succession de son père en 2015." in text
    _assert_no_residual_tokens(text)


def test_medical_models_have_no_dentaire_leak(tmp_path: Path) -> None:
    # FIX 2 : aucune mention « dentaire » dans les cessions MEDICALES.
    acte = _docx_text(
        ActeCessionCabinetMedicalGenerator().generate(_context(credit_vendeur=True), tmp_path / "a")
    )
    compromis = _docx_text(
        CompromisCessionCabinetMedicalGenerator().generate(
            _context(etape="compromis"), tmp_path / "c"
        )
    )
    for text in (acte, compromis):
        lowered = text.lower()
        assert "dentaire" not in lowered
        assert "dentiste" not in lowered
        assert "stomatologue" not in lowered
        # Le wording medical propre est conserve.
        assert "cabinet médical" in lowered


def test_acte_dentaire_salaries_zero_renders_neant(tmp_path: Path) -> None:
    # Regle NotebookLM : 0 salarie -> convention systeme "Néant".
    ctx = _context(type_cabinet="dentaire", salaries=[])

    text = _docx_text(ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path))

    assert "Néant." in text
    assert "De reprendre les contrats de travail de" not in text
    _assert_no_residual_tokens(text)


def test_acte_dentaire_salaries_one_renders_single(tmp_path: Path) -> None:
    # Regle NotebookLM : 1 salarie -> liste a un element (poste optionnel rendu).
    ctx = _context(
        type_cabinet="dentaire",
        salaries=[
            CessionSalarie(
                civilite_affichage="Madame", prenom="Lea", nom="Petit", poste="assistante dentaire"
            )
        ],
    )

    text = _docx_text(ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path))

    assert "De reprendre les contrats de travail de Madame Lea Petit" in text
    assert "en qualité de assistante dentaire" in text
    assert " et de " not in text.split("contrats de travail de")[1].split(".")[0]
    _assert_no_residual_tokens(text)


def test_acte_dentaire_salaries_three_renders_full_list(tmp_path: Path) -> None:
    # Regle NotebookLM : N salaries -> liste complete (assouplissement 0/1/N).
    ctx = _context(
        type_cabinet="dentaire",
        salaries=[
            CessionSalarie(civilite_affichage="Madame", prenom="Lea", nom="Petit"),
            CessionSalarie(civilite_affichage="Monsieur", prenom="Noe", nom="Robert"),
            CessionSalarie(civilite_affichage="Madame", prenom="Ines", nom="Faure"),
        ],
    )

    text = _docx_text(ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path))

    assert "Madame Lea Petit, Monsieur Noe Robert et de Madame Ines Faure" in text
    _assert_no_residual_tokens(text)


def test_acte_dentaire_salary_requires_complete_identity(tmp_path: Path) -> None:
    # Garde-fou : un salarie liste doit avoir civilite/prenom/nom (identite complete).
    ctx = _context(
        type_cabinet="dentaire",
        salaries=[CessionSalarie(civilite_affichage="Madame", prenom="Lea")],
    )

    with pytest.raises(ValueError, match="salaries"):
        ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path)


_DENT_SALARIES = [
    CessionSalarie(civilite_affichage="Madame", prenom="Lea", nom="Petit"),
    CessionSalarie(civilite_affichage="Monsieur", prenom="Noe", nom="Robert"),
]


def test_acte_dentaire_vendeur_masculin_agreement(tmp_path: Path) -> None:
    # Modele dentaire fige au FEMININ -> vendeur masculin = retour au masculin.
    ctx = _context(
        type_cabinet="dentaire",
        salaries=_DENT_SALARIES,
        vendeur_genre=Gender.MASCULIN,
    )

    text = _docx_text(ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path))

    assert "chirurgien-dentiste, né le 10 mars 1975" in text
    assert "née le 10 mars 1975" not in text
    assert "Inscrit au tableau du Conseil départemental" in text
    assert "Inscrite au tableau du Conseil départemental" not in text
    # Forme de role invariante conservee.
    assert "le soussigné de première part" in text
    _assert_no_residual_tokens(text)


def test_acte_dentaire_vendeur_feminin_agreement(tmp_path: Path) -> None:
    # Modele dentaire deja feminin -> vendeur feminin = conserve tel quel.
    ctx = _context(
        type_cabinet="dentaire",
        salaries=_DENT_SALARIES,
        vendeur_genre=Gender.FEMININ,
    )

    text = _docx_text(ActeCessionCabinetDentaireGenerator().generate(ctx, tmp_path))

    assert "chirurgien-dentiste, née le 10 mars 1975" in text
    assert "Inscrite au tableau du Conseil départemental" in text
    _assert_no_residual_tokens(text)


def test_acte_dentaire_representant_gender_agreement(tmp_path: Path) -> None:
    # « domicilié(e) en cette qualité » est pilote par le genre du representant.
    masc = _docx_text(
        ActeCessionCabinetDentaireGenerator().generate(
            _context(
                type_cabinet="dentaire",
                salaries=_DENT_SALARIES,
                representant_genre=Gender.MASCULIN,
            ),
            tmp_path / "masc",
        )
    )
    fem = _docx_text(
        ActeCessionCabinetDentaireGenerator().generate(
            _context(
                type_cabinet="dentaire",
                salaries=_DENT_SALARIES,
                representant_genre=Gender.FEMININ,
            ),
            tmp_path / "fem",
        )
    )

    assert "domicilié en cette qualité audit siège" in masc
    assert "domiciliée en cette qualité" not in masc
    assert "domiciliée en cette qualité audit siège" in fem
    # « Représentée » accorde la societe (toujours feminin) : jamais touche.
    assert "Représentée par" in masc
    assert "Représentée par" in fem


def test_compromis_medical_vendeur_feminin_agreement(tmp_path: Path) -> None:
    # Modele medical fige au MASCULIN -> vendeur feminin = accord au feminin.
    ctx = _context(
        etape="compromis",
        type_cabinet="medical",
        vendeur_genre=Gender.FEMININ,
    )

    text = _docx_text(CompromisCessionCabinetMedicalGenerator().generate(ctx, tmp_path))

    assert "medecin, née le 10 mars 1975" in text
    assert "inscrite au tableau du Conseil départemental" in text
    # GARDE-FOU : le « Tableau » de la societe (Ordre) ne doit jamais devenir feminin parasite.
    assert "au Tableau de l’Ordre des Médecins" in text
    _assert_no_residual_tokens(text)


def test_compromis_medical_vendeur_masculin_agreement(tmp_path: Path) -> None:
    ctx = _context(
        etape="compromis",
        type_cabinet="medical",
        vendeur_genre=Gender.MASCULIN,
    )

    text = _docx_text(CompromisCessionCabinetMedicalGenerator().generate(ctx, tmp_path))

    assert "medecin, né le 10 mars 1975" in text
    assert "née le 10 mars 1975" not in text
    assert "inscrit au tableau du Conseil départemental" in text
    _assert_no_residual_tokens(text)


def test_acte_medical_keeps_inclusive_birth_form(tmp_path: Path) -> None:
    # L'acte medical fige la forme inclusive « né(e) le » : jamais accordee,
    # quel que soit le genre du vendeur (aucune paire ne la matche).
    masc = _docx_text(
        ActeCessionCabinetMedicalGenerator().generate(
            _context(credit_vendeur=True, vendeur_genre=Gender.MASCULIN),
            tmp_path / "masc",
        )
    )
    fem = _docx_text(
        ActeCessionCabinetMedicalGenerator().generate(
            _context(credit_vendeur=True, vendeur_genre=Gender.FEMININ),
            tmp_path / "fem",
        )
    )

    assert "né(e) le 10 mars 1975" in masc
    assert "né(e) le 10 mars 1975" in fem


def test_acte_medical_renders_conjoint_prenom_and_nom(tmp_path: Path) -> None:
    # Retour UAT Rafael (DOC-009) : la clause de situation maritale doit afficher
    # le PRENOM ET le NOM du conjoint (« marié(e) à Prenom Nom, sous le régime... »).
    text = _docx_text(
        ActeCessionCabinetMedicalGenerator().generate(
            _context(credit_vendeur=True),
            tmp_path,
        )
    )

    # Conjoint = Claire Durand, regime = communaute reduite aux acquets (cf. _context).
    assert "marie à Claire Durand, sous le régime de communaute reduite aux acquets" in text
    _assert_no_residual_tokens(text)


def test_orchestrator_selects_only_requested_cession_cabinet_document() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(
            _context(etape="compromis", type_cabinet="dentaire"),
        )
    }

    assert "DOC-012" in selected_ids
    assert "DOC-009" not in selected_ids
    assert "DOC-010" not in selected_ids
    assert "DOC-011" not in selected_ids
