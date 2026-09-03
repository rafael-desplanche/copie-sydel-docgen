from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Associe,
    Company,
    DocumentGenerationContext,
    DossierOptions,
    FraisCommunsContext,
    LocauxContext,
    PacteAssociesScmContext,
    PartieFraisCommuns,
    Person,
    PraticienScm,
    ReglementInterieurScmContext,
    ScmRepresentant,
    ScmSatellitesOptions,
    ScmSocietePartie,
    Signature,
)
from sydel_doc_engine.generators.lot_05.contrat_frais_communs import (
    ContratFraisCommunsGenerator,
)
from sydel_doc_engine.generators.lot_05.liste_depenses_communes_scm import (
    ListeDepensesCommunesScmGenerator,
)
from sydel_doc_engine.generators.lot_05.pacte_associes_scm import PacteAssociesScmGenerator
from sydel_doc_engine.generators.lot_05.reglement_interieur_scm import (
    ReglementInterieurScmGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text
    ]
    return "\n".join(paragraphs + cells)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _party(index: int, *, forme_juridique: str = "SELARL") -> PartieFraisCommuns:
    return PartieFraisCommuns(
        societe=ScmSocietePartie(
            denomination=f"SEL DOCTEUR {index}",
            forme_juridique=forme_juridique,
            capital_social="1 000 euros",
            siege=Address(adresse_affichee=f"{index} rue des Soins, 7500{index} Paris"),
            ville_rcs="Paris",
            numero_rcs=f"900 000 00{index}",
        ),
        representant=ScmRepresentant(
            civilite_affichage="Monsieur" if index == 1 else "Madame",
            prenom="Jean" if index == 1 else "Alice",
            nom="Durand" if index == 1 else "Martin",
            identite_affichee="Jean Durand" if index == 1 else "Alice Martin",
            titre_affichage="Docteur",
            fonction="gerant",
        ),
    )


def _base_context() -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SCM",
        dossier_options=DossierOptions(scm_satellites=True),
        scm_satellites=ScmSatellitesOptions(
            pacte_associes=True,
            liste_depenses_communes=True,
            contrat_frais_communs=True,
            reglement_interieur=True,
        ),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Durand",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 15)),
        societe=Company(
            denomination="SCM SANTE PARIS",
            forme_juridique="SCM",
            capital_social="2 000 euros",
            siege=Address(adresse_affichee="10 avenue de la Sante, 75014 Paris"),
            ville_rcs="Paris",
            numero_rcs="900 111 222",
            nb_parts_total=200,
        ),
        associes=[
            Associe(
                genre=Gender.MASCULIN,
                civilite_affichage="Monsieur",
                prenom="Jean",
                nom="Durand",
                nb_parts=100,
            ),
            Associe(
                genre=Gender.FEMININ,
                civilite_affichage="Madame",
                prenom="Alice",
                nom="Martin",
                nb_parts=100,
            ),
        ],
        pacte_associes=PacteAssociesScmContext(ville_tribunal="Paris"),
        frais_communs=FraisCommunsContext(date_effet_contrat="1er juin 2026"),
        reglement_interieur=ReglementInterieurScmContext(
            seuil_depense_commune="1 500 euros",
            annee_reference_charges="2026",
            date_fin_gestion_administrative="31 decembre 2026",
            date_attribution_responsabilites="1er janvier",
        ),
        parties_frais_communs=[_party(1), _party(2)],
        praticiens=[
            PraticienScm(identite_affichee="Jean Durand", telephone="01 00 00 00 01"),
            PraticienScm(identite_affichee="Alice Martin", telephone="01 00 00 00 02"),
        ],
        locaux=LocauxContext(adresse_affichee="10 avenue de la Sante, 75014 Paris"),
    )


def test_scm_satellite_generators_create_clean_docx(tmp_path: Path) -> None:
    ctx = _base_context()

    outputs = [
        PacteAssociesScmGenerator().generate(ctx, tmp_path),
        ListeDepensesCommunesScmGenerator().generate(ctx, tmp_path),
        ContratFraisCommunsGenerator().generate(ctx, tmp_path),
        ReglementInterieurScmGenerator().generate(ctx, tmp_path),
    ]
    texts = {path.name: _docx_text(path) for path in outputs}

    assert {path.name for path in outputs} == {
        "pacte_associes_scm.docx",
        "liste_depenses_communes_scm.docx",
        "contrat_frais_communs.docx",
        "reglement_interieur_scm.docx",
    }
    assert "PACTE D" in texts["pacte_associes_scm.docx"]
    assert "SCM SANTE PARIS" in texts["pacte_associes_scm.docx"]
    assert "Monsieur Jean Durand" in texts["pacte_associes_scm.docx"]
    assert "CONTRAT D'EXERCICE PROFESSIONNEL" in texts["contrat_frais_communs.docx"]
    assert "1er juin 2026" in texts["contrat_frais_communs.docx"]
    assert "DENOMINATION DE LA DEPENSE" in texts["liste_depenses_communes_scm.docx"]
    assert "Frais de prothèse" in texts["liste_depenses_communes_scm.docx"]
    assert "Jean Durand" in texts["liste_depenses_communes_scm.docx"]
    assert "REGLEMENT INTERIEUR DE LA SOCIETE CIVILE DE MOYENS" in texts[
        "reglement_interieur_scm.docx"
    ]
    assert "1 500 euros" in texts["reglement_interieur_scm.docx"]
    for text in texts.values():
        _assert_clean(text)


def test_pacte_associes_scm_requires_two_associes(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.associes.append(
        Associe(
            genre=Gender.MASCULIN,
            civilite_affichage="Monsieur",
            prenom="Paul",
            nom="Bernard",
            nb_parts=1,
        )
    )

    with pytest.raises(ValueError, match="exactement deux associes"):
        PacteAssociesScmGenerator().generate(ctx, tmp_path)


def test_contrat_frais_communs_requires_satellite_flag(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.scm_satellites = ScmSatellitesOptions(contrat_frais_communs=False)

    with pytest.raises(ValueError, match="contrat_frais_communs"):
        ContratFraisCommunsGenerator().generate(ctx, tmp_path)


def test_liste_depenses_communes_requires_two_associes(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.associes.pop()

    with pytest.raises(ValueError, match="exactement deux associes"):
        ListeDepensesCommunesScmGenerator().generate(ctx, tmp_path)


def test_liste_depenses_communes_requires_satellite_flag(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.scm_satellites = ScmSatellitesOptions(liste_depenses_communes=False)

    with pytest.raises(ValueError, match="liste_depenses_communes"):
        ListeDepensesCommunesScmGenerator().generate(ctx, tmp_path)


def test_reglement_interieur_rejects_different_party_forms(tmp_path: Path) -> None:
    ctx = _base_context()
    ctx.parties_frais_communs[1] = _party(2, forme_juridique="SELAS")

    with pytest.raises(ValueError, match="meme forme_juridique"):
        ReglementInterieurScmGenerator().generate(ctx, tmp_path)


def test_orchestrator_selects_only_enabled_scm_satellites() -> None:
    ctx = _base_context()
    ctx.scm_satellites = ScmSatellitesOptions(
        pacte_associes=True,
        liste_depenses_communes=True,
        contrat_frais_communs=False,
        reglement_interieur=True,
    )
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected = orchestrator.select_documents_for_context(ctx)

    selected_ids = [document.doc_id for document in selected]
    assert "DOC-026" in selected_ids
    assert "DOC-027" not in selected_ids
    assert "DOC-028" in selected_ids
    assert "DOC-030" in selected_ids
