from __future__ import annotations

from datetime import date
from pathlib import Path
from unicodedata import normalize

from docx import Document
from streamlit.testing.v1 import AppTest

from sydel_doc_engine.front_app.data_entry import build_clean_data_entry
from sydel_doc_engine.front_app.dossier_selection import dossier_type_by_label
from sydel_doc_engine.front_app.generation import build_clean_generation_plan
from sydel_doc_engine.front_app.legacy_boundary import legacy_boundary_items
from sydel_doc_engine.front_app.routing import clean_front_routes
from sydel_doc_engine.front_app.selarl_slice import (
    PROFESSION_DENTISTE,
    PROFESSION_MEDECIN,
    build_generation_context,
    generate_selarl_dossier,
    selected_selarl_document_codes,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog
from sydel_doc_engine.scenarios.selarl import (
    build_selarl_scenario,
    cession_fixture_for_profession,
    scm_cession_fixture,
)


def test_clean_front_routes_are_minimal() -> None:
    assert [route.label for route in clean_front_routes()] == [
        "Type de dossier",
        "Donnees a saisir",
        "Generation",
    ]


def test_clean_front_selarl_slice_is_generable_for_medecin() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = _valid_selarl_input(PROFESSION_MEDECIN)

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is True
    assert plan.status == "ready"
    assert plan.target_engine_adapter == "front_app.selarl_slice"
    assert plan.document_codes == (
        "DOC-001",
        "DOC-002",
        "DOC-003",
        "DOC-004",
        "DOC-034",
        "DOC-017",
    )
    assert "DOC-006" not in plan.document_codes


def test_clean_front_selarl_slice_switches_statuts_for_dentiste() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = _valid_selarl_input(PROFESSION_DENTISTE)

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is True
    assert "DOC-016" in plan.document_codes
    assert "DOC-017" not in plan.document_codes


def test_clean_front_selarl_slice_adds_regime_batch_only_for_regime() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = _valid_selarl_input(PROFESSION_MEDECIN, regime_communautaire=True)

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is True
    assert "DOC-005" in plan.document_codes
    assert "DOC-006" in plan.document_codes
    assert any(
        row.doc_code == "DOC-006" and row.status == "generable"
        for row in plan.document_rows
    )


def test_clean_front_selarl_medecin_regime_derives_conjoint_only_when_active() -> None:
    standard_ctx = build_generation_context(_valid_selarl_input(PROFESSION_MEDECIN))
    regime_ctx = build_generation_context(
        _valid_selarl_input(PROFESSION_MEDECIN, regime_communautaire=True)
    )

    assert standard_ctx.conjoint is None
    assert standard_ctx.regime_communautaire is None
    assert regime_ctx.conjoint is not None
    assert regime_ctx.conjoint.prenom == "Claire"
    assert regime_ctx.conjoint.nom == "Martin"
    assert regime_ctx.conjoint.adresse_perso is not None
    assert regime_ctx.conjoint.adresse_perso.adresse_affichee == "10 rue Test, 75001 Paris"
    assert regime_ctx.statuts_sel is not None
    assert regime_ctx.statuts_sel.overlay == "selarl_medecin"
    assert regime_ctx.regime_communautaire is not None
    assert regime_ctx.regime_communautaire.date_courrier_avertissement == date.today()
    assert regime_ctx.regime_communautaire.renonciation is not None
    assert (
        regime_ctx.regime_communautaire.renonciation.nombre_exemplaires_lettres
        == "quatre"
    )


def test_clean_front_selarl_regime_does_not_require_conjoint_address() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    kwargs = _valid_selarl_kwargs(PROFESSION_MEDECIN, regime_communautaire=True)
    data_entry = build_clean_data_entry(dossier_type, **kwargs)

    plan = build_clean_generation_plan(dossier_type, data_entry)
    ctx = build_generation_context(data_entry)

    assert plan.can_generate is True
    assert ctx.conjoint is not None
    assert ctx.conjoint.adresse_perso is not None
    assert ctx.conjoint.adresse_perso.adresse_affichee == "10 rue Test, 75001 Paris"


def test_clean_front_selarl_medecin_separation_de_biens_generates_statuts(
    tmp_path: Path,
) -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = _valid_selarl_input(
        PROFESSION_MEDECIN,
        married_separation=True,
    )

    plan = build_clean_generation_plan(dossier_type, data_entry)
    ctx = build_generation_context(data_entry)
    generated = generate_selarl_dossier(data_entry, tmp_path / "selarl-medecin-separation")
    statuts_path = next(path for path in generated.docx_paths if path.name.startswith("statuts"))
    statuts_text = _docx_text(statuts_path)

    assert plan.can_generate is True
    assert "DOC-005" not in plan.document_codes
    assert "DOC-006" not in plan.document_codes
    assert ctx.conjoint is not None
    assert ctx.conjoint.adresse_perso is None
    assert ctx.associes[0].conjoint is not None
    assert (
        "marié sous le régime de la séparation de biens avec Madame Claire Martin"
        in statuts_text
    )


def test_clean_front_selarl_medecin_separation_de_biens_blocks_without_conjoint() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    kwargs = _valid_selarl_kwargs(PROFESSION_MEDECIN, married_separation=True)
    kwargs.update(
        {
            "conjoint_civilite": "",
            "conjoint_prenom": "",
            "conjoint_nom": "",
        }
    )
    data_entry = build_clean_data_entry(dossier_type, **kwargs)

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is False
    assert any("conjoint" in blocker.casefold() for blocker in plan.blockers)


def test_clean_front_selarl_regime_ui_never_exposes_conjoint_address_fields() -> None:
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=120)

    app.checkbox(key="selarl_regime_communautaire").set_value(True)
    app.run(timeout=120)

    conjoint_address_labels = [
        widget.label
        for widget in app.text_input
        if "conjoint" in widget.label.casefold()
        and "adresse" in widget.label.casefold()
    ]
    conjoint_address_keys = [
        str(widget.key)
        for widget in app.text_input
        if "conjoint_adresse" in str(widget.key)
    ]

    assert conjoint_address_labels == []
    assert conjoint_address_keys == []


def test_clean_front_selarl_slice_blocks_out_of_scope_cases() -> None:
    # La cession est desormais SUPPORTEE quand les donnees cession sont fournies
    # (cession_context). Demander la cession (flag) sans donnees reste bloque.
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = build_clean_data_entry(
        dossier_type,
        **{
            **_valid_selarl_kwargs(PROFESSION_MEDECIN),
            "cession": True,
        },
    )

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is False
    assert "Cession demandee mais donnees cession manquantes." in plan.blockers


def test_clean_front_selarl_cession_cabinet_medical_generates_acte(tmp_path: Path) -> None:
    # Cession avec donnees (scenario fige) -> l'acte de cession cabinet medical est generable.
    data = build_selarl_scenario("selarl_medecin_cession_cabinet_medical")
    dossier_type = dossier_type_by_label("SELARL creation V1")

    plan = build_clean_generation_plan(dossier_type, data)

    assert plan.can_generate is True
    assert "DOC-009" in plan.document_codes
    assert "DOC-007" in plan.document_codes
    # Appel de fonds (DOC-008) = document commun « Si cession », present aussi en medical.
    assert "DOC-008" in plan.document_codes

    result = generate_selarl_dossier(data, tmp_path)
    names = {path.name for path in result.docx_paths}
    assert "acte_cession_cabinet_medical.docx" in names
    assert "avenant_contrat_bail.docx" in names
    assert "appel_fond_sel.docx" in names


def test_clean_front_selarl_cession_cabinet_dentaire_generates_full_pack(tmp_path: Path) -> None:
    # Cession dentaire : acte (DOC-011) + avenant bail (DOC-007) + appel de fonds (DOC-008).
    data = build_selarl_scenario("selarl_dentiste_cession_cabinet_dentaire")
    dossier_type = dossier_type_by_label("SELARL creation V1")

    plan = build_clean_generation_plan(dossier_type, data)

    assert plan.can_generate is True
    assert {"DOC-011", "DOC-008", "DOC-007"}.issubset(set(plan.document_codes))

    result = generate_selarl_dossier(data, tmp_path)
    names = {path.name for path in result.docx_paths}
    assert "acte_cession_cabinet_dentaire.docx" in names
    assert "appel_fond_sel.docx" in names
    assert "avenant_contrat_bail.docx" in names


def test_clean_front_selarl_cession_scm_generates_scm_docs(tmp_path: Path) -> None:
    # Cession de parts de SCM : PV AGE (DOC-031) + courrier SDE (DOC-032) + acte (DOC-033).
    data = build_selarl_scenario("selarl_dentiste_cession_scm")
    dossier_type = dossier_type_by_label("SELARL creation V1")

    plan = build_clean_generation_plan(dossier_type, data)

    assert plan.can_generate is True
    assert {"DOC-031", "DOC-032", "DOC-033"}.issubset(set(plan.document_codes))

    result = generate_selarl_dossier(data, tmp_path)
    names = {path.name for path in result.docx_paths}
    assert "pv_age_cession_parts_scm.docx" in names
    assert "courrier_sde_cession_scm.docx" in names
    assert "acte_cession_parts_scm.docx" in names


def test_cession_context_medical_selects_acte_bail_appel_fonds() -> None:
    # (a) cession_context medical + bail -> DOC-009 (acte medical) + DOC-007 (bail)
    # + DOC-008 (appel de fonds).
    cession, bail = cession_fixture_for_profession(PROFESSION_MEDECIN)
    data_entry = _valid_selarl_input(
        PROFESSION_MEDECIN,
        cession=True,
        cession_context=cession,
        bail_context=bail,
    )

    codes = selected_selarl_document_codes(data_entry)

    assert "DOC-009" in codes
    assert "DOC-007" in codes
    assert "DOC-008" in codes
    assert "DOC-011" not in codes


def test_cession_context_dentaire_selects_acte_and_appel_fonds() -> None:
    # (a) cession_context dentaire -> DOC-011 (acte dentaire) + DOC-008 (appel de fonds).
    cession, bail = cession_fixture_for_profession(PROFESSION_DENTISTE)
    data_entry = _valid_selarl_input(
        PROFESSION_DENTISTE,
        cession=True,
        cession_context=cession,
        bail_context=bail,
    )

    codes = selected_selarl_document_codes(data_entry)

    assert "DOC-011" in codes
    assert "DOC-008" in codes
    assert "DOC-009" not in codes


def test_scm_cession_context_selects_scm_docs() -> None:
    # (c) scm_cession_context -> DOC-031 / DOC-032 / DOC-033.
    data_entry = _valid_selarl_input(
        PROFESSION_DENTISTE,
        scm=True,
        scm_cession_context=scm_cession_fixture(),
    )

    codes = selected_selarl_document_codes(data_entry)

    assert {"DOC-031", "DOC-032", "DOC-033"}.issubset(set(codes))


def test_scm_flag_without_data_is_blocked() -> None:
    # Le garde-fou reste : SCM coche sans donnees -> bloque (pas de generation muette).
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = build_clean_data_entry(
        dossier_type,
        **{**_valid_selarl_kwargs(PROFESSION_MEDECIN), "scm": True},
    )

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is False
    assert any("SCM" in blocker for blocker in plan.blockers)


def test_clean_front_ui_prefill_generates_cession_medical_without_residual_tokens(
    tmp_path: Path,
    monkeypatch,
) -> None:
    # (b) Chemin UI complet : profession medecin, bouton de test (active la cession +
    # prereremplit selarl_cession_*), generation -> acte cession medical + appel de fonds
    # + bail, sans aucun token [xxx] residuel.
    _assert_ui_prefill_cession_generates(
        tmp_path,
        monkeypatch,
        profession_label="Medecin",
        expected_doc="acte_cession_cabinet_medical.docx",
    )


def test_clean_front_ui_prefill_generates_cession_dentaire_without_residual_tokens(
    tmp_path: Path,
    monkeypatch,
) -> None:
    # (b) Chemin UI complet pour le cabinet dentaire (acte dentaire + salaries).
    _assert_ui_prefill_cession_generates(
        tmp_path,
        monkeypatch,
        profession_label="Chirurgien-dentiste",
        expected_doc="acte_cession_cabinet_dentaire.docx",
    )


def _assert_ui_prefill_cession_generates(
    tmp_path: Path,
    monkeypatch,
    *,
    profession_label: str,
    expected_doc: str,
) -> None:
    from sydel_doc_engine.front_app import shell

    monkeypatch.setattr(shell, "ARTIFACTS_DIR", tmp_path / "ui-cession")
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=180)
    app.selectbox(key="selarl_profession").set_value(profession_label)
    app = app.run(timeout=180)

    app.button(key="clean_generate_test_data").click()
    app = app.run(timeout=180)

    assert app.checkbox(key="selarl_cession").value is True
    assert app.checkbox(key="selarl_scm").value is True
    assert not any("Blocage" in item.value for item in app.caption)
    assert app.button(key="clean_generate_dossier").disabled is False

    app.button(key="clean_generate_dossier").click()
    app = app.run(timeout=180)

    download_labels = [item.label for item in app.get("download_button")]
    assert f"Telecharger {expected_doc}" in download_labels
    assert "Telecharger appel_fond_sel.docx" in download_labels
    assert "Telecharger avenant_contrat_bail.docx" in download_labels
    assert "Telecharger pv_age_cession_parts_scm.docx" in download_labels

    generated = app.session_state["clean_generated_dossier"]
    combined_text = "\n".join(
        _docx_text(Path(path)) for path in generated["docx_paths"]
    )
    assert "[" not in combined_text
    assert "]" not in combined_text


def test_clean_front_ui_creation_only_unchanged(
    tmp_path: Path,
    monkeypatch,
) -> None:
    # (d) Sans cession ni SCM : la creation seule genere toujours les 6 documents,
    # inchangee par le cablage du sous-formulaire.
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = _valid_selarl_input(PROFESSION_MEDECIN)

    plan = build_clean_generation_plan(dossier_type, data_entry)
    generated = generate_selarl_dossier(data_entry, tmp_path / "creation-only")

    assert plan.can_generate is True
    assert plan.document_codes == (
        "DOC-001",
        "DOC-002",
        "DOC-003",
        "DOC-004",
        "DOC-034",
        "DOC-017",
    )
    assert data_entry.cession_context is None
    assert data_entry.scm_cession_context is None
    assert len(generated.docx_paths) == 6


def test_clean_front_cession_form_returns_none_without_flag() -> None:
    # Le sous-formulaire ne rend rien et renvoie (None, None) quand la cession
    # n'est pas demandee -> aucun expander parasite dans le wizard de base.
    from sydel_doc_engine.front_app import shell

    cession_context, bail_context = shell._render_cession_form(False, PROFESSION_MEDECIN)
    scm_context = shell._render_scm_cession_form(False)

    assert cession_context is None
    assert bail_context is None
    assert scm_context is None


def test_clean_front_selarl_cession_compromis_generates(tmp_path: Path) -> None:
    # Compromis de cession : médical (DOC-010) et dentaire (DOC-012), même moteur que l'acte.
    dossier_type = dossier_type_by_label("SELARL creation V1")
    for scenario, expected_doc, filename in (
        (
            "selarl_medecin_cession_compromis_medical",
            "DOC-010",
            "compromis_cession_cabinet_medical.docx",
        ),
        (
            "selarl_dentiste_cession_compromis_dentaire",
            "DOC-012",
            "compromis_cession_cabinet_dentaire.docx",
        ),
    ):
        data = build_selarl_scenario(scenario)
        plan = build_clean_generation_plan(dossier_type, data)
        assert plan.can_generate is True
        assert expected_doc in plan.document_codes
        result = generate_selarl_dossier(data, tmp_path / scenario)
        names = {path.name for path in result.docx_paths}
        assert filename in names


def test_clean_front_selarl_context_selects_only_expected_engine_docs() -> None:
    ctx = build_generation_context(_valid_selarl_input(PROFESSION_DENTISTE))
    selected = DocumentOrchestrator(build_seed_catalog()).select_documents_for_context(ctx)
    selected_codes = {document.doc_id for document in selected}

    assert {"DOC-001", "DOC-002", "DOC-003", "DOC-004", "DOC-034", "DOC-016"}.issubset(
        selected_codes
    )
    assert "DOC-017" not in selected_codes
    assert "DOC-005" not in selected_codes
    assert "DOC-006" not in selected_codes


def test_clean_front_selarl_context_derives_hidden_ux_values() -> None:
    ctx = build_generation_context(_valid_selarl_input(PROFESSION_MEDECIN))

    assert ctx.personne_signataire.genre.value == "masculin"
    assert ctx.personne_signataire.date_naissance == date(1984, 4, 12)
    assert ctx.personne_signataire.ville_naissance == "Paris"
    assert ctx.personne_signataire.ville_naissance_article_au is False
    assert ctx.societe.capital_social_lettres == "mille"
    assert ctx.capital is not None
    assert ctx.capital.valeur_nominale_titre == "10"
    assert ctx.capital.nombre_titres_total_lettres == "cent"
    assert ctx.capital.valeur_nominale_titre_lettres == "dix"
    assert ctx.reunion is not None
    assert ctx.reunion.date_lettres == "vingt-six mai deux mille vingt-six"
    assert ctx.reunion.heure is None
    assert ctx.reunion.president is not None
    assert ctx.reunion.president.civilite_president_seance == "Monsieur"
    assert ctx.reunion.president.prenom_president_seance == "Jean"
    assert ctx.reunion.president.nom_personne_seance == "Martin"
    assert ctx.signature.prestataire_signature_electronique == "Yousign"
    assert ctx.gerance is not None
    assert ctx.gerance.seuil_achat_materiel == "5000"
    assert ctx.gerance.seuil_emprunt == "10000"
    assert ctx.mandataire is not None
    assert ctx.mandataire.cabinet == "SYDEL"
    assert ctx.exercice_social is not None
    assert ctx.exercice_social.lieux[0].adresse_affichee == "20 avenue du Siege, 75002 Paris"
    assert ctx.associes[0].nb_parts == 100


def test_clean_front_selarl_accepts_french_date_strings_outside_streamlit_range() -> None:
    dossier_type = dossier_type_by_label("SELARL creation V1")
    data_entry = build_clean_data_entry(
        dossier_type,
        **{
            **_valid_selarl_kwargs(PROFESSION_MEDECIN),
            "date_naissance": "31/12/1974",
            "signature_date": "27/05/2026",
            "decision_date": "27/05/2026",
        },
    )

    plan = build_clean_generation_plan(dossier_type, data_entry)

    assert plan.can_generate is True
    assert data_entry.date_naissance == date(1974, 12, 31)
    assert data_entry.signature_date == date(2026, 5, 27)
    assert data_entry.decision_date == date(2026, 5, 27)


def test_clean_front_selarl_maps_birth_city_article_au() -> None:
    ctx = build_generation_context(
        _valid_selarl_input(
            PROFESSION_MEDECIN,
            ville_naissance="Bourget",
            ville_naissance_article_au=True,
        )
    )

    assert ctx.personne_signataire.ville_naissance == "Bourget"
    assert ctx.personne_signataire.ville_naissance_article_au is True


def test_clean_front_selarl_generation_smoke(tmp_path: Path) -> None:
    generated = generate_selarl_dossier(
        _valid_selarl_input(PROFESSION_MEDECIN),
        tmp_path / "selarl-medecin",
    )

    assert len(generated.docx_paths) == 6
    assert generated.zip_path.exists()
    assert {path.name for path in generated.docx_paths} >= {
        "declaration_non_condamnation.docx",
        "autorisation_domiciliation.docx",
        "procuration.docx",
        "pv_nomination_gerant.docx",
        "demande_inscription_ordre.docx",
        "statuts_selarl_medecin.docx",
    }
    combined_text = "\n".join(_docx_text(path) for path in generated.docx_paths)
    assert "SELARL SELARL" not in combined_text
    assert "Société d’exercice libéral à responsabilité limitée de médecin" in combined_text
    assert "Conseil départemental de l'Ordre des médecins de 75" in combined_text
    assert "Au capital de 1 000 euros" in combined_text
    assert "Au capital de 1000" not in combined_text
    assert " medecin" not in combined_text


def test_clean_front_selarl_medecin_regime_communautaire_generation_smoke(
    tmp_path: Path,
) -> None:
    generated = generate_selarl_dossier(
        _valid_selarl_input(PROFESSION_MEDECIN, regime_communautaire=True),
        tmp_path / "selarl-medecin-regime-communautaire",
    )

    names = {path.name for path in generated.docx_paths}
    assert len(generated.docx_paths) == 8
    assert generated.zip_path.exists()
    assert names == {
        "declaration_non_condamnation.docx",
        "autorisation_domiciliation.docx",
        "procuration.docx",
        "pv_nomination_gerant.docx",
        "demande_inscription_ordre.docx",
        "statuts_selarl_medecin.docx",
        "lettre_renonciation_associe.docx",
        "lettre_avertissement_conjoint.docx",
    }
    assert "statuts_selarl_chirurgien_dentiste.docx" not in names

    combined_text = "\n".join(_docx_text(path) for path in generated.docx_paths)
    ascii_text = _ascii_text(combined_text)
    assert "[" not in combined_text
    assert "]" not in combined_text
    assert "SELARL SELARL" not in combined_text
    assert "RCS PARIS 788 531 432" not in combined_text
    assert "0153814303" not in combined_text
    assert "Société d’exercice libéral à responsabilité limitée de médecin" in combined_text
    assert "Au capital de 1 000 €" in combined_text
    assert "Au capital de 1000" not in combined_text
    assert " medecin" not in combined_text
    assert f"Par courrier en date du {date.today():%d/%m/%Y}" in combined_text
    assert "euros dependant de notre communaute." in ascii_text
    assert "regime de communaute" not in ascii_text
    assert "Madame Martin" in combined_text
    assert "10 rue Test" in combined_text
    assert "30 rue Conjoint" not in combined_text


def test_clean_front_legacy_boundary_is_explicit() -> None:
    items = legacy_boundary_items()
    decisions = {item.decision for item in items}
    components = {item.component for item in items}

    assert {"reused", "legacy_reference", "ignored_by_clean_front", "remove_later"}.issubset(
        decisions
    )
    assert "src/sydel_doc_engine/front_data/" in components
    assert "src/sydel_doc_engine/app/streamlit_app.py" in components


def test_clean_front_entrypoint_does_not_import_legacy_screens() -> None:
    source = "\n".join(
        Path(path).read_text(encoding="utf-8")
        for path in (
            "src/sydel_doc_engine/front_app/app.py",
            "src/sydel_doc_engine/front_app/shell.py",
            "src/sydel_doc_engine/front_app/selarl_slice.py",
        )
    )

    assert "business_wizard" not in source
    assert "single_document_mode" not in source
    assert "streamlit_app" not in source
    assert "front_internal_tool" not in source
    assert "Technique / diagnostic" not in source
    assert "Document unitaire" not in source
    assert "Debug interne" not in source


def test_clean_front_streamlit_surface_is_not_legacy() -> None:
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=120)

    assert [item.value for item in app.subheader] == [
        "Type de dossier",
        "Donnees a saisir",
        "Generation",
    ]
    assert app.selectbox(key="clean_dossier_type").label == "Type de dossier"
    assert app.selectbox(key="clean_dossier_type").value == "SELARL creation V1"
    assert app.selectbox(key="selarl_profession").label == "Profession"
    assert not any(str(widget.key) == "selarl_case_mode" for widget in app.selectbox)
    assert app.button(key="clean_generate_test_data").label == "Generer des donnees de test"
    assert app.button(key="clean_generate_dossier").disabled is True
    assert app.button(key="selarl_signature_date_today").label == "Aujourd'hui"
    assert len(app.radio) == 0
    assert len(app.table) == 0
    assert len(app.expander) == 0
    assert not any(item.label == "Outils internes" for item in app.checkbox)

    visible_labels = {
        *[item.label for item in app.text_input],
        *[item.label for item in app.number_input],
        *[item.label for item in app.selectbox],
        *[item.label for item in app.date_input],
    }
    assert len(app.date_input) == 0
    assert "Genre" not in visible_labels
    assert "Titre affichage" not in visible_labels
    assert "Capital social en lettres" not in visible_labels
    assert "Nombre de parts en lettres" not in visible_labels
    assert "Valeur nominale d'une part (€)" not in visible_labels
    assert "Valeur nominale en lettres" not in visible_labels
    assert "Date reunion en lettres" not in visible_labels
    assert "Heure de decision" not in visible_labels
    assert "Regime matrimonial" not in visible_labels
    assert "Prestataire signature electronique" not in visible_labels
    assert "Seuil achat materiel" not in visible_labels
    assert "Seuil emprunt" not in visible_labels
    assert "Lieu exercice" not in visible_labels
    assert "Civilite mandataire" not in visible_labels
    assert "Civilite conjoint" not in visible_labels
    assert app.selectbox(key="selarl_nationalite_choice").label == "Nationalite"
    assert "Portugaise" in app.selectbox(key="selarl_nationalite_choice").options
    assert app.selectbox(key="selarl_situation_maritale").label == "Situation matrimoniale"
    assert (
        app.checkbox(key="selarl_regime_communautaire").label
        == "Documents regime de la communaute"
    )
    assert (
        app.text_input(key="selarl_departement_ordre").label
        == "Departement d'inscription a l'ordre"
    )


def test_clean_front_streamlit_no_longer_exposes_multi_associes_case() -> None:
    # SELARL = unipersonnelle (decision Gad 2026-06-04) : aucun selecteur de cas multi,
    # aucun champ de sous-formulaire multi-associes ne doit subsister dans le wizard.
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=120)

    assert not any(str(widget.key) == "selarl_case_mode" for widget in app.selectbox)
    multi_keys = [
        str(widget.key)
        for group in (app.number_input, app.selectbox, app.text_input)
        for widget in group
        if "doc004" in str(widget.key)
    ]
    assert multi_keys == []


def test_clean_front_streamlit_random_data_button_prefills_generable_case(
    tmp_path: Path,
    monkeypatch,
) -> None:
    from sydel_doc_engine.front_app import shell

    monkeypatch.setattr(shell, "ARTIFACTS_DIR", tmp_path / "random-data")
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=120)

    app.button(key="clean_generate_test_data").click()
    app = app.run(timeout=120)

    assert app.text_input(key="selarl_dossier_reference").value.startswith("TEST-SELARL-")
    assert app.number_input(key="selarl_capital_social").value > 0
    assert app.number_input(key="selarl_nb_parts_total").value > 0
    assert app.button(key="clean_generate_dossier").disabled is False
    assert not any("Blocage" in item.value for item in app.caption)

    app.button(key="clean_generate_dossier").click()
    app = app.run(timeout=120)

    assert any(
        item.label == "Telecharger le dossier ZIP"
        for item in app.get("download_button")
    )


def test_clean_front_streamlit_generation_exposes_download_buttons(
    tmp_path: Path,
    monkeypatch,
) -> None:
    from sydel_doc_engine.front_app import shell

    monkeypatch.setattr(shell, "ARTIFACTS_DIR", tmp_path / "streamlit-downloads")
    app = AppTest.from_file("src/sydel_doc_engine/front_app/app.py").run(timeout=120)
    _fill_valid_streamlit_selarl_form(app)
    app = app.run(timeout=120)

    assert app.button(key="clean_generate_dossier").disabled is False

    app.button(key="clean_generate_dossier").click()
    app = app.run(timeout=120)

    download_labels = [item.label for item in app.get("download_button")]
    assert download_labels == [
        "Telecharger le dossier ZIP",
        "Telecharger declaration_non_condamnation.docx",
        "Telecharger autorisation_domiciliation.docx",
        "Telecharger procuration.docx",
        "Telecharger pv_nomination_gerant.docx",
        "Telecharger demande_inscription_ordre.docx",
        "Telecharger statuts_selarl_medecin.docx",
    ]


def _valid_selarl_input(
    profession: str,
    *,
    regime_communautaire: bool = False,
    married_separation: bool = False,
    **overrides: object,
):
    dossier_type = dossier_type_by_label("SELARL creation V1")
    values = _valid_selarl_kwargs(
        profession,
        regime_communautaire=regime_communautaire,
        married_separation=married_separation,
    )
    values.update(overrides)
    return build_clean_data_entry(
        dossier_type,
        **values,
    )


def _valid_selarl_kwargs(
    profession: str,
    *,
    regime_communautaire: bool = False,
    married_separation: bool = False,
) -> dict[str, object]:
    is_married = regime_communautaire or married_separation
    return {
        "dossier_reference": "B-SELARL-001",
        "profession": profession,
        "dossier_unipersonnel": True,
        "regime_communautaire": regime_communautaire,
        "civilite": "Monsieur",
        "prenom": "Jean",
        "nom": "Martin",
        "date_naissance": date(1984, 4, 12),
        "ville_naissance": "Paris",
        "departement_naissance": "75",
        "nationalite": "française",
        "nom_pere": "Pierre Martin",
        "nom_mere": "Anne Martin",
        "adresse_num_voie": "10",
        "adresse_voie": "rue Test",
        "adresse_cp": "75001",
        "adresse_ville": "Paris",
        "situation_maritale": "marie" if is_married else "celibataire",
        "regime_matrimonial": (
            "regime de communaute"
            if regime_communautaire
            else "separation de biens"
            if married_separation
            else ""
        ),
        "numero_ordre": "ORD-123",
        "numero_rpps": "10000000001",
        "departement_ordre": "75",
        "denomination": "SELARL MARTIN",
        "capital_social": "1000",
        "nb_parts_total": 100,
        "valeur_nominale_part": "10",
        "siege_num_voie": "20",
        "siege_voie": "avenue du Siege",
        "siege_cp": "75002",
        "siege_ville": "Paris",
        "ville_rcs": "Paris",
        "ordre_adresse_ligne_1": "1 rue de l'Ordre",
        "ordre_cp": "75008",
        "ordre_ville": "Paris",
        "signature_lieu": "Paris",
        "signature_date": date(2026, 5, 26),
        "decision_date": date(2026, 5, 26),
        "depot_banque_nom": "Banque Test",
        "depot_banque_adresse": "30 boulevard Banque, 75009 Paris",
        "exercice_debut": "1er janvier",
        "exercice_fin": "31 decembre",
        "exercice_cloture_premier": "31 decembre 2026",
        **(
            {
                "conjoint_civilite": "Madame",
                "conjoint_prenom": "Claire",
                "conjoint_nom": "Martin",
            }
            if profession == PROFESSION_DENTISTE or regime_communautaire or married_separation
            else {}
        ),
    }


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _ascii_text(value: str) -> str:
    return normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")


def _fill_valid_streamlit_selarl_form(app: AppTest) -> None:
    values = {
        "selarl_dossier_reference": "B-SELARL-DOWNLOAD-TEST",
        "selarl_prenom": "Jean",
        "selarl_nom": "Martin",
        "selarl_date_naissance": "31/12/1974",
        "selarl_ville_naissance": "Paris",
        "selarl_departement_naissance": "75",
        "selarl_nom_pere": "Pierre Martin",
        "selarl_nom_mere": "Anne Martin",
        "selarl_numero_ordre": "ORD-123",
        "selarl_numero_rpps": "10000000001",
        "selarl_adresse_num_voie": "10",
        "selarl_adresse_voie": "rue Test",
        "selarl_adresse_cp": "75001",
        "selarl_adresse_ville": "Paris",
        "selarl_denomination": "SELARL MARTIN",
        "selarl_ville_rcs": "Paris",
        "selarl_siege_num_voie": "20",
        "selarl_siege_voie": "avenue du Siege",
        "selarl_siege_cp": "75002",
        "selarl_siege_ville": "Paris",
        "selarl_departement_ordre": "75",
        "selarl_ordre_adresse_ligne_1": "1 rue de l'Ordre",
        "selarl_ordre_cp": "75008",
        "selarl_ordre_ville": "Paris",
        "selarl_signature_lieu": "Paris",
        "selarl_signature_date": "27/05/2026",
        "selarl_decision_date": "27/05/2026",
        "selarl_depot_banque_nom": "Banque Test",
        "selarl_depot_banque_adresse": "30 boulevard Banque, 75009 Paris",
        "selarl_exercice_debut": "1er janvier",
        "selarl_exercice_fin": "31 decembre",
        "selarl_exercice_cloture_premier": "31 decembre 2026",
    }
    for key, value in values.items():
        app.text_input(key=key).set_value(value)
    app.checkbox(key="selarl_ville_naissance_article_au").set_value(False)
    app.number_input(key="selarl_capital_social").set_value(1000)
    app.number_input(key="selarl_nb_parts_total").set_value(100)
