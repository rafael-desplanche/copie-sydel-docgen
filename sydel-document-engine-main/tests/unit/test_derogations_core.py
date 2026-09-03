from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Company,
    CompanyInscriptionOrdre,
    Contact,
    DerogationContext,
    DerogationRole,
    DocumentGenerationContext,
    DossierOptions,
    Person,
    Signature,
    SiteDeclare,
    SiteExistant,
)
from sydel_doc_engine.generators.lot_03.demande_derogation_cumul_selarl_bnc import (
    DemandeDerogationCumulSelarlBncGenerator,
)
from sydel_doc_engine.generators.lot_03.formulaire_derogation_sites_sel import (
    FormulaireDerogationSitesSelGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _context(derogation_type: str = "multi_sites_sel") -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure="SELARL",
        dossier_options=DossierOptions(derogation=True),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Camille",
            nom="Martin",
            numero_inscription_ordre="ORD-123",
            qualification_principale="chirurgien-dentiste",
            contact=Contact(telephone="01 44 00 00 00", email="camille@example.test"),
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
        societe=Company(
            denomination="SELARL CABINET MARTIN",
            siege=Address(
                adresse_affichee="12 avenue des Ternes, 75017 Paris",
                cp="75017",
                ville="Paris",
            ),
            inscription_ordre=CompanyInscriptionOrdre(
                departement="Paris",
                ville="Paris",
                numero="SEL-456",
            ),
        ),
        derogation=DerogationContext(
            type=derogation_type,
            mode_rendu="formulaire_a_completer",
            representant_legal=DerogationRole(
                prenom="Camille",
                nom="Martin",
                fonction="gerant",
                contact=Contact(email="camille@example.test"),
            ),
            associe_exercant=DerogationRole(
                prenom="Camille",
                nom="Martin",
                qualification_principale="chirurgien-dentiste",
            ),
            sites_existants_present=True,
        ),
        site_declare=SiteDeclare(
            adresse_affichee="24 rue du Nouveau Site, 75016 Paris",
            date_debut_activite=date(2026, 9, 1),
        ),
        sites_existants=[
            SiteExistant(
                adresse_affichee="12 avenue des Ternes, 75017 Paris",
                date_debut_activite=date(2024, 1, 15),
                temps_hebdomadaire="5 demi-journees",
            )
        ],
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _assert_no_source_placeholders(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _table_has_explicit_borders(table) -> bool:
    borders = table._tbl.tblPr.find(qn("w:tblBorders"))
    return borders is not None and all(
        borders.find(qn(f"w:{edge}")) is not None
        for edge in ("top", "left", "bottom", "right")
    )


def test_formulaire_derogation_sites_sel_generates_prefilled_form(tmp_path: Path) -> None:
    output_path = FormulaireDerogationSitesSelGenerator().generate(_context(), tmp_path)

    assert output_path == tmp_path / "formulaire_derogation_sites_sel_formulaire_a_completer.docx"
    text = _docx_text(output_path)
    assert "Declaration prealable d'ouverture d'un site distinct" in text
    assert "SELARL CABINET MARTIN" in text
    assert "Mandat (gerant/president/...) : gerant" in text
    assert "Qualification : chirurgien-dentiste" in text
    assert "☒ OUI" in text
    assert "24 rue du Nouveau Site, 75016 Paris" in text
    document = Document(output_path)
    assert abs(document.sections[0].top_margin - Cm(2.0)) < 300
    section = next(p for p in document.paragraphs if p.text == "I - Identification du declarant")
    assert section.runs[0].underline is True
    _assert_no_source_placeholders(text)


def test_formulaire_derogation_sites_sel_blocks_without_role_mapping(
    tmp_path: Path,
) -> None:
    ctx = _context()
    ctx.derogation.representant_legal = None

    with pytest.raises(ValueError, match="representant_legal"):
        FormulaireDerogationSitesSelGenerator().generate(ctx, tmp_path)


def test_demande_derogation_cumul_selarl_bnc_generates_prefilled_form(
    tmp_path: Path,
) -> None:
    output_path = DemandeDerogationCumulSelarlBncGenerator().generate(
        _context("cumul_sel_bnc"),
        tmp_path,
    )

    assert output_path == (
        tmp_path / "demande_derogation_cumul_selarl_bnc_formulaire_a_completer.docx"
    )
    text = _docx_text(output_path)
    assert "Demande de cumul d'exercices en societe d'exercice liberal (SEL)" in text
    assert "Sous le numero : ORD-123" in text
    assert "Denomination sociale : SELARL CABINET MARTIN" in text
    assert "Je soussigne(e) Dr Camille Martincertifie :" in text
    assert "Fait le 14/05/2026" in text
    document = Document(output_path)
    assert abs(document.sections[0].top_margin - Cm(3.25)) < 500
    assert any("PIECES A JOINDRE" in table.cell(0, 0).text for table in document.tables)
    assert all(_table_has_explicit_borders(table) for table in document.tables)
    _assert_no_source_placeholders(text)


def test_cumul_selarl_bnc_blocks_outside_selarl(tmp_path: Path) -> None:
    ctx = _context("cumul_sel_bnc")
    ctx.structure = "SELAS"

    with pytest.raises(ValueError, match="SELARL"):
        DemandeDerogationCumulSelarlBncGenerator().generate(ctx, tmp_path)


def test_orchestrator_selects_only_requested_derogation_core_document() -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    selected_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(_context("cumul_sel_bnc"))
    }

    assert "DOC-014" in selected_ids
    assert "DOC-013" not in selected_ids
