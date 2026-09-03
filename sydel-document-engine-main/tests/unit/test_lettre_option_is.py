from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    CentreImpots,
    Company,
    DocumentGenerationContext,
    DossierOptions,
    Person,
    Signature,
    StatutsCivilsAssocie,
    StatutsCivilsContext,
    StatutsCivilsParts,
)
from sydel_doc_engine.generators.lot_05.lettre_option_is import LettreOptionIsGenerator


def _docx_text(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        if cell.text
    ]
    return "\n".join(paragraphs + cells)


def _assert_clean(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def _base_context(*, option_is: bool = True, structure: str = "SCI") -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(option_is=option_is),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Durand",
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 15)),
        impots=CentreImpots(
            service="Service des impots des entreprises",
            centre="SIE Paris Centre",
            adresse_ligne_1="6 rue Paganini",
            adresse_ligne_2="TSA 10000",
            cp="75020",
            ville="Paris",
        ),
        societe=Company(
            denomination="SCI EXEMPLE",
            forme_sociale="SCI",
            capital_social="1000",
            siren="123 456 789",
            siege=Address(
                num_voie="12",
                voie="avenue Victor Hugo",
                cp="75016",
                ville="Paris",
            ),
        ),
        statuts_civils=StatutsCivilsContext(
            type="sci" if structure == "SCI" else "sci_iris",
            capital_social="1000",
            nb_parts_total=100,
            associes=[
                StatutsCivilsAssocie(
                    type_personne="personne_physique",
                    civilite_affichage="Monsieur",
                    prenom="Jean",
                    nom="Durand",
                    adresse_personnelle=Address(
                        num_voie="1",
                        voie="rue Exemple",
                        cp="75000",
                        ville="Paris",
                    ),
                    parts=StatutsCivilsParts(nb=40, qualite_associe="associe"),
                ),
                StatutsCivilsAssocie(
                    type_personne="personne_morale",
                    denomination="SEL IRIS",
                    siege=Address(
                        num_voie="2",
                        voie="rue Pro",
                        cp="75000",
                        ville="Paris",
                    ),
                    parts=StatutsCivilsParts(nb=60),
                ),
            ],
        ),
    )


def test_lettre_option_is_generates_clean_docx(tmp_path: Path) -> None:
    output_path = LettreOptionIsGenerator().generate(_base_context(), tmp_path)
    text = _docx_text(output_path)

    assert output_path.name == "lettre_option_is.docx"
    assert "Service des impots des entreprises" in text
    assert "Demande d'option pour le régime de l'impôt sur les sociétés" in text
    assert "SCI EXEMPLE" in text
    assert "123 456 789" in text
    assert "Monsieur Jean Durand, demeurant 1 rue Exemple, 75000 Paris" in text
    assert "La société SEL IRIS, ayant son siège social au 2 rue Pro, 75000 Paris" in text
    assert "Le gérant" in text
    _assert_clean(text)


def test_lettre_option_is_requires_option_flag(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="option_is"):
        LettreOptionIsGenerator().generate(_base_context(option_is=False), tmp_path)


def test_lettre_option_is_rejects_non_sci_structure(tmp_path: Path) -> None:
    ctx = _base_context(structure="SELARL")

    with pytest.raises(ValueError, match="dossier.structure"):
        LettreOptionIsGenerator().generate(ctx, tmp_path)
