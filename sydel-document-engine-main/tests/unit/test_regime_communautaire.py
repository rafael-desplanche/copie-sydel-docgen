from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Apport,
    Associe,
    Company,
    DocumentGenerationContext,
    DossierOptions,
    Person,
    RegimeCommunautaire,
    RegimeCommunautaireAvertissement,
    RegimeCommunautaireRenonciation,
    Signature,
)
from sydel_doc_engine.generators.lot_02.lettre_avertissement_conjoint import (
    LettreAvertissementConjointGenerator,
)
from sydel_doc_engine.generators.lot_02.lettre_renonciation_associe import (
    LettreRenonciationAssocieGenerator,
)
from sydel_doc_engine.orchestrator.service import DocumentOrchestrator
from sydel_doc_engine.registry.catalog import build_seed_catalog


def _context(
    structure: str = "SELAS",
    *,
    regime_communautaire: bool = True,
    forme_sociale_abregee: str | None = "SELAS",
    qualite_renoncee: str | None = "actionnaire",
    date_courrier_avertissement: date | str | None = "14 mai 2026",
) -> DocumentGenerationContext:
    return DocumentGenerationContext(
        structure=structure,
        dossier_options=DossierOptions(regime_communautaire=regime_communautaire),
        personne_signataire=Person(
            genre=Gender.MASCULIN,
            civilite="Monsieur",
            prenom="Jean",
            nom="Durand",
            fonction_dirigeant="président",
            adresse_perso=Address(
                num_voie="12",
                voie="rue de l'Associe",
                cp="75001",
                ville="Paris",
            ),
        ),
        conjoint=Person(
            genre=Gender.FEMININ,
            civilite="Madame",
            prenom="Claire",
            nom="Durand",
            adresse_perso=Address(
                num_voie="24",
                voie="rue de la Paix",
                cp="75002",
                ville="Paris",
            ),
        ),
        societe=Company(
            forme_sociale="SELAS",
            forme_sociale_complete="société d'exercice libéral par actions simplifiée",
            forme_sociale_abregee=forme_sociale_abregee,
            denomination="RC SANTE",
            capital_social="1 000",
            siege=Address(
                num_voie="80",
                voie="avenue Marceau",
                cp="75008",
                ville="Paris",
            ),
        ),
        apport=Apport(montant="500", montant_lettres="cinq cents"),
        regime_communautaire=RegimeCommunautaire(
            avertissement=RegimeCommunautaireAvertissement(
                date_signature=date(2026, 5, 14),
            ),
            renonciation=RegimeCommunautaireRenonciation(
                lieu_signature="Paris",
                date_signature=date(2026, 5, 15),
                nombre_exemplaires_lettres="deux",
            ),
            date_courrier_avertissement=date_courrier_avertissement,
            regime_matrimonial="communauté",
            qualite_renoncee=qualite_renoncee,
        ),
        signature=Signature(lieu="Paris", date=date(2026, 5, 14)),
    )


def _docx_text(path: Path) -> str:
    document = Document(path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(text for text in texts if text)


def _matching_paragraphs(path: Path, text: str):
    return [paragraph for paragraph in Document(path).paragraphs if paragraph.text == text]


def _first_paragraph_starting_with(path: Path, prefix: str):
    for paragraph in Document(path).paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise AssertionError(f"Paragraphe introuvable : {prefix}")


def _assert_no_source_placeholders(text: str) -> None:
    assert "[" not in text
    assert "]" not in text


def test_regime_communautaire_selas_generates_both_documents(tmp_path: Path) -> None:
    ctx = _context("SELAS", forme_sociale_abregee="SELAS")

    renonciation = LettreRenonciationAssocieGenerator().generate(ctx, tmp_path)
    avertissement = LettreAvertissementConjointGenerator().generate(ctx, tmp_path)

    assert renonciation == tmp_path / "lettre_renonciation_associe.docx"
    assert avertissement == tmp_path / "lettre_avertissement_conjoint.docx"

    renonciation_text = _docx_text(renonciation)
    avertissement_text = _docx_text(avertissement)
    assert "Par courrier en date du 14 mai 2026" in renonciation_text
    assert "euros dépendant de notre communauté." in renonciation_text
    assert "euros dépendant de notre regime de communaute." not in renonciation_text
    assert "personnellement actionnaire de cette société" in renonciation_text
    assert "Fait pour servir et valoir ce que de droit." in renonciation_text
    assert "RCS" not in renonciation_text
    assert "à la SELAS RC SANTE" in avertissement_text
    assert "Le  14/05/2026" in avertissement_text
    renonciation_section = Document(renonciation).sections[0]
    assert abs(renonciation_section.left_margin - Cm(3.17)) < 300
    assert abs(renonciation_section.right_margin - Cm(3.17)) < 300
    assert _matching_paragraphs(renonciation, "À Paris")[0].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    assert "Le 15/05/2026" not in renonciation_text
    renonciation_subject = _matching_paragraphs(
        renonciation,
        "Objet : Lettre de renonciation à revendiquer la qualité d'associé",
    )[0]
    assert renonciation_subject.runs[0].bold is True
    assert renonciation_subject.runs[0].underline is True
    assert _matching_paragraphs(renonciation, "Claire Durand")[0].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    avertissement_section = Document(avertissement).sections[0]
    assert abs(avertissement_section.left_margin - Cm(3.17)) < 300
    assert abs(avertissement_section.right_margin - Cm(3.17)) < 300
    company_header = _matching_paragraphs(avertissement, "RC SANTE")[0]
    assert company_header.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert company_header.runs[0].bold is True
    assert _matching_paragraphs(avertissement, "Madame Durand")[0].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    assert _matching_paragraphs(avertissement, "Le  14/05/2026")[0].alignment == (
        WD_ALIGN_PARAGRAPH.RIGHT
    )
    avertissement_subject = _matching_paragraphs(
        avertissement,
        "Objet : Lettre d'avertissement au conjoint en cas d'apport d'un bien commun.",
    )[0]
    assert avertissement_subject.runs[0].bold is True
    assert avertissement_subject.runs[0].underline is True
    apporteur_quality = _matching_paragraphs(
        avertissement,
        "Agissant en qualité de futur président",
    )[0]
    assert apporteur_quality.runs[0].italic is True
    instruction = _first_paragraph_starting_with(avertissement, "(Faire précéder")
    assert instruction.runs[0].italic is True
    _assert_no_source_placeholders(renonciation_text)
    _assert_no_source_placeholders(avertissement_text)


def test_regime_communautaire_selarl_uses_societe_mention_without_abregee(
    tmp_path: Path,
) -> None:
    ctx = _context("SELARL", forme_sociale_abregee=None, qualite_renoncee="associé")

    text = _docx_text(LettreAvertissementConjointGenerator().generate(ctx, tmp_path))

    assert "à la Société RC SANTE" in text
    assert "à la SELAS RC SANTE" not in text
    _assert_no_source_placeholders(text)


def test_selarl_avertissement_uses_written_form_and_associe_address(
    tmp_path: Path,
) -> None:
    ctx = _context("SELARL", forme_sociale_abregee="SELARL", qualite_renoncee="associé")
    ctx.associes = [
        Associe(
            genre=Gender.MASCULIN,
            civilite_affichage="Monsieur",
            prenom="Jean",
            nom="Durand",
            nb_parts=100,
            profession_reglementee="médecin",
        )
    ]

    text = _docx_text(LettreAvertissementConjointGenerator().generate(ctx, tmp_path))

    assert "Société d’exercice libéral à responsabilité limitée de médecin" in text
    assert "12 rue de l'Associe" in text
    assert "75001 Paris" in text
    assert "24 rue de la Paix" not in text
    assert "75002 Paris" not in text
    assert "Fait en quatre exemplaires" in text
    assert "Fait en trois exemplaires" not in text


@pytest.mark.parametrize("structure", ["SPFPL cession", "SPFPL apport"])
def test_regime_communautaire_spfpl_structures_use_abregee_overlay(
    tmp_path: Path,
    structure: str,
) -> None:
    ctx = _context(structure, forme_sociale_abregee="SPFPL")

    text = _docx_text(LettreAvertissementConjointGenerator().generate(ctx, tmp_path))

    assert "à la SPFPL RC SANTE" in text
    _assert_no_source_placeholders(text)


def test_renonciation_falls_back_to_avertissement_date_when_generated_as_batch(
    tmp_path: Path,
) -> None:
    ctx = _context(date_courrier_avertissement=None)

    text = _docx_text(LettreRenonciationAssocieGenerator().generate(ctx, tmp_path))

    assert "Par courrier en date du 14/05/2026" in text


def test_regime_communautaire_blocks_when_batch_option_is_false(tmp_path: Path) -> None:
    ctx = _context(regime_communautaire=False)

    with pytest.raises(ValueError, match="regime_communautaire"):
        LettreAvertissementConjointGenerator().generate(ctx, tmp_path)


def test_renonciation_blocks_when_qualite_renoncee_is_missing(tmp_path: Path) -> None:
    ctx = _context(qualite_renoncee=None)

    with pytest.raises(ValueError, match="qualite_renoncee"):
        LettreRenonciationAssocieGenerator().generate(ctx, tmp_path)


def test_selas_avertissement_blocks_when_abregee_is_missing(tmp_path: Path) -> None:
    ctx = _context("SELAS", forme_sociale_abregee=None)

    with pytest.raises(ValueError, match="forme_sociale_abregee"):
        LettreAvertissementConjointGenerator().generate(ctx, tmp_path)


def test_orchestrator_generates_regime_communautaire_batch_only_when_enabled(
    tmp_path: Path,
) -> None:
    orchestrator = DocumentOrchestrator(build_seed_catalog())

    enabled_ids = {
        document.doc_id for document in orchestrator.select_documents_for_context(_context("SELAS"))
    }
    disabled_ids = {
        document.doc_id
        for document in orchestrator.select_documents_for_context(
            _context("SELAS", regime_communautaire=False),
        )
    }

    assert {"DOC-005", "DOC-006"}.issubset(enabled_ids)
    assert "DOC-005" not in disabled_ids
    assert "DOC-006" not in disabled_ids
