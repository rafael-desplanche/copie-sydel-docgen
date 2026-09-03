from __future__ import annotations

from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    CapitalSouscription,
    DocumentGenerationContext,
    SocieteSpfpl,
    SpfplConjoint,
    SpfplOrdre,
    SpfplPerson,
    StatutsPresident,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_centered_block,
    add_hyphen_list_item,
    add_paragraph,
    add_statuts_annex_heading,
    add_statuts_article_heading,
    add_statuts_body_paragraph,
    add_statuts_hanging_list_item,
    add_statuts_signature_block,
    new_document,
)

DOCUMENT_CODE = "CODE-STATUTS-SAS-001"
OUTPUT_FILENAME = "statuts_sas_spfpl_medecins.docx"
STATUTS_SAS_TYPE = "spfpl_medecins"
STATUTS_SAS_PROFESSION = "medecin"


class StatutsSasGenerator:
    """Generateur from-scratch des statuts SAS / SPFPL medecins V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        data = _ResolvedStatutsSas.from_context(ctx)
        document = new_document()
        footer = document.sections[0].footer.paragraphs[0]
        footer.text = f"{data.denomination} - Statuts constitutifs"

        add_centered_block(
            document,
            [
                (data.denomination, True, False),
                (
                    "Société de Participations Financières de Profession Libérale de "
                    "Médecins par actions simplifiée",
                    False,
                    False,
                ),
                f"Au capital de {data.capital_social}",
                f"Siège social : {data.adresse_siege}",
            ],
        )
        add_paragraph(document, "STATUTS", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

        _add_initial_shareholder_block(document, data)
        _add_articles(document, data)
        _add_signature_and_annex(document, data)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


class _ResolvedStatutsSas:
    def __init__(
        self,
        *,
        denomination: str,
        capital_social: str,
        capital_social_lettres: str,
        nb_actions_total: int,
        nb_actions_total_lettres: str,
        valeur_nominale_action: str,
        valeur_nominale_action_lettres: str,
        adresse_siege: str,
        actionnaire: SpfplPerson,
        president: StatutsPresident,
        banque_nom: str,
        exercice_debut: str,
        exercice_fin: str,
        exercice_cloture_1: str,
        signature_lieu: str,
    ) -> None:
        self.denomination = denomination
        self.capital_social = capital_social
        self.capital_social_lettres = capital_social_lettres
        self.nb_actions_total = nb_actions_total
        self.nb_actions_total_lettres = nb_actions_total_lettres
        self.valeur_nominale_action = valeur_nominale_action
        self.valeur_nominale_action_lettres = valeur_nominale_action_lettres
        self.adresse_siege = adresse_siege
        self.actionnaire = actionnaire
        self.president = president
        self.banque_nom = banque_nom
        self.exercice_debut = exercice_debut
        self.exercice_fin = exercice_fin
        self.exercice_cloture_1 = exercice_cloture_1
        self.signature_lieu = signature_lieu

    @classmethod
    def from_context(cls, ctx: DocumentGenerationContext) -> _ResolvedStatutsSas:
        _validate_sas_scope(ctx)
        societe = _required_societe_spfpl(ctx)
        actionnaire = _required_actionnaire(ctx)
        president = _required_president(ctx)
        capital = _required_capital_souscription(ctx)
        _validate_actionnaire_unique(ctx, actionnaire, president, capital)
        _validate_capital(societe, actionnaire, capital)
        _validate_marital_sentence(actionnaire)

        if ctx.depot_fonds is None or ctx.depot_fonds.banque is None:
            raise ValueError(f"depot_fonds.banque est obligatoire pour {DOCUMENT_CODE}.")
        if ctx.exercice_social is None:
            raise ValueError(f"exercice_social est obligatoire pour {DOCUMENT_CODE}.")

        depot_montant = ctx.depot_fonds.montant
        capital_social = _required_text(societe.capital_social, "societe_spfpl.capital_social")
        if depot_montant is not None and depot_montant.strip() != capital_social:
            raise ValueError(
                "depot_fonds.montant doit etre coherent avec "
                f"societe_spfpl.capital_social pour {DOCUMENT_CODE}."
            )

        return cls(
            denomination=_required_text(societe.denomination, "societe_spfpl.denomination"),
            capital_social=capital_social,
            capital_social_lettres=_required_text(
                societe.capital_social_lettres,
                "societe_spfpl.capital_social_lettres",
            ),
            nb_actions_total=_required_int(
                societe.nb_actions_total,
                "societe_spfpl.nb_actions_total",
            ),
            nb_actions_total_lettres=_required_text(
                societe.nb_actions_total_lettres,
                "societe_spfpl.nb_actions_total_lettres",
            ),
            valeur_nominale_action=_required_text(
                societe.valeur_nominale_action,
                "societe_spfpl.valeur_nominale_action",
            ),
            valeur_nominale_action_lettres=_required_text(
                societe.valeur_nominale_action_lettres,
                "societe_spfpl.valeur_nominale_action_lettres",
            ),
            adresse_siege=_address_display(societe),
            actionnaire=actionnaire,
            president=president,
            banque_nom=_required_text(ctx.depot_fonds.banque.nom, "depot_fonds.banque.nom"),
            exercice_debut=_required_text(
                ctx.exercice_social.debut,
                "exercice_social.debut",
            ),
            exercice_fin=_required_text(ctx.exercice_social.fin, "exercice_social.fin"),
            exercice_cloture_1=_required_text(
                ctx.exercice_social.date_cloture_premier_exercice,
                "exercice_social.date_cloture_premier_exercice",
            ),
            signature_lieu=ctx.signature.lieu,
        )


def _add_initial_shareholder_block(document, data: _ResolvedStatutsSas) -> None:
    actionnaire = data.actionnaire
    conjoint = _required_conjoint(actionnaire)
    ordre = _required_ordre(actionnaire)
    add_paragraph(document, "Le soussigné :")
    add_hyphen_list_item(document, _person_name(actionnaire, "actionnaire_unique"))
    add_paragraph(
        document,
        f"{_qualification(actionnaire)} de profession",
    )
    departement_naissance = _required_text(
        actionnaire.departement_naissance,
        "actionnaire_unique.departement_naissance",
    )
    add_paragraph(
        document,
        "Né le "
        f"{_format_display_date(actionnaire.date_naissance, 'actionnaire_unique.date_naissance')} "
        f"à {_required_text(actionnaire.ville_naissance, 'actionnaire_unique.ville_naissance')} "
        f"({departement_naissance})",
    )
    add_paragraph(
        document,
        f"Demeurant {_person_address(actionnaire, 'actionnaire_unique')}",
    )
    situation_maritale = _required_text(
        actionnaire.situation_maritale,
        "actionnaire_unique.situation_maritale",
    )
    regime_matrimonial = _required_text(
        actionnaire.regime_matrimonial,
        "actionnaire_unique.regime_matrimonial",
    )
    add_paragraph(
        document,
        f"{situation_maritale} sous le régime de {regime_matrimonial} "
        f"avec {_conjoint_name(conjoint)}",
    )
    nationalite = _required_text(actionnaire.nationalite, "actionnaire_unique.nationalite")
    add_paragraph(
        document,
        f"De nationalité {nationalite}",
    )
    add_paragraph(
        document,
        "Inscrit au tableau de l’Ordre des Médecins de "
        f"{_required_text(ordre.departement, 'actionnaire_unique.ordre.departement')} "
        "sous le numéro national "
        f"{_required_text(ordre.numero, 'actionnaire_unique.ordre.numero')} "
        "et sous le numéro RPPS "
        f"{_required_text(ordre.numero_rpps, 'actionnaire_unique.ordre.numero_rpps')}.",
    )
    add_paragraph(document, "Ci après dénommé l’ « Associé Unique », ou l’ « Actionnaire Unique »")
    add_paragraph(
        document,
        "Ou pris avec d’éventuels autres associés les « Associés », voire séparément "
        "l’ « Associé »,",
    )
    add_paragraph(
        document,
        "A établi ainsi qu’il suit les statuts de la société de participations financières "
        "de profession libérale de médecins par actions simplifiée unipersonnelle qu’il a "
        "décidé d’instituer (ci-après dénommé la « Société »), telle que précisé ci-après "
        "sous la condition suspensive de son incription au Tableau de l’Ordre des médecins.",
    )


def _add_articles(document, data: _ResolvedStatutsSas) -> None:
    _article(document, "ARTICLE 1 - FORME")
    _paragraphs(
        document,
        [
            "Il est formé une Société de Participations Financières de Professions Libérales "
            "de Médecins constituée sous forme de société par actions simplifiée "
            "unipersonnelle régie par les présents statuts , ainsi que par les lois et "
            "règlements en vigueur, notamment :",
            "Par l’ordonnance n°2023-77 du 8 février 2023 relative à l'exercice en société "
            "des professions libérales réglementées,",
            "Par les dispositions du Code de commerce applicables aux sociétés commerciales "
            "et plus particulièrement aux sociétés par actions simplifiée,",
            "Par le Code de la Santé publique,",
            "Par le Code de déontologie médicale codifiée aux articles R.4127-1 et suivants "
            "du Code de la Santé publique.",
            " Elle fonctionne sous la même forme avec un ou plusieurs associés.",
        ],
    )
    _article(document, "ARTICLE 2 - OBJET")
    _paragraphs(
        document,
        [
            "La Société a pour objet :",
            "- la détention de parts ou d’actions de Sociétés d’Exercice Libéral ayant pour "
            "objet la profession de Médecin;",
            "- toutes activités accessoires en relation directe avec leur objet et sous "
            "réserve d’être destinées exclusivement aux sociétés ou groupements dans "
            "lesquelles elle détient des participations dans le respect des dispositions "
            "législatives et réglementaires en vigueur;",
            "- et plus généralement, la réalisation de toutes opérations de quelque nature "
            "qu’elles soient, notamment financière, pouvant se rattacher, directement ou "
            "indirectement à l’objet spécifié ci-dessus pouvant en faciliter la réalisation.",
        ],
    )
    _article(document, "ARTICLE 3 - DENOMINATION")
    _paragraphs(
        document,
        [
            f"La dénomination de la Société est : {data.denomination}",
            "Dans tous les actes, lettres, facture, et autres documents émanant de la "
            "Société et destinés à des tiers, la dénomination sociale doit être précédée "
            "ou suivie immédiatement de la mention « Société de Participations Financières "
            "de Professions Libérales de Médecins par actions simplifiée » ou des initiales "
            "« SPFPL Médecins par actions simplifiée» de l'énonciation du montant du "
            "capital social, du lieu de son siège social et suivi du numéro d’immatriculation "
            "au Registre du commerce et des sociétés. Cette dénomination sociale pourra "
            "également être suivie du numéro d’inscription au Tableau de l’Ordre des Médecins.",
        ],
    )
    _article(document, "ARTICLE 4 - SIEGE SOCIAL")
    _paragraphs(
        document,
        [
            f"Le siège social est fixé au {data.adresse_siege}.",
            "Il ne peut être transféré, que ce soit dans le même département, dans un "
            "département limitrophe, ou dans tout autre lieu quel qu’il soit, que par décision "
            "d’un ou plusieurs associés représentant plus de la moitié des parts sociales ou "
            "par décision du gérant seul. Cette décision sera retranscrite par un procès-verbal.",
        ],
    )
    _article(document, "ARTICLE 5 - DUREE")
    _paragraphs(
        document,
        [
            "La durée de la Société est fixée à QUATRE VINGT DIX NEUF (99) ans à compter "
            "de la date de son immatriculation au Registre du Commerce et des Sociétés, "
            "sauf dissolution anticipée ou prorogation.",
            "Les décisions de prorogation de la durée de la Société ou de dissolution "
            "anticipée sont prises par l'associé unique ou par décision collective "
            "extraordinaire des associés.",
        ],
    )
    _article(document, "ARTICLE 6 - APPORTS")
    _paragraphs(
        document,
        [
            "Il a été apporté en numéraire :",
            "Par le Docteur "
            f"{data.actionnaire.prenom} {data.actionnaire.nom} {data.capital_social}",
            f"Soit au total la somme de {data.capital_social}",
            "Cette somme a été dès avant ce jour, déposée au crédit d'un compte ouvert au "
            f"nom de la Société dans les livres de la Banque {data.banque_nom}",
        ],
    )
    _article(document, "ARTICLE 7 - CAPITAL SOCIAL")
    _paragraphs(
        document,
        [
            f"Le capital social est fixé à {data.capital_social_lettres} "
            f"({data.capital_social}) euros, divisé en {data.nb_actions_total_lettres} "
            f"({data.nb_actions_total}) actions de {data.valeur_nominale_action_lettres} "
            f"({data.valeur_nominale_action}) chacune, entièrement libérées et attribuées "
            "comme suit :",
            f"Le Docteur {data.actionnaire.prenom} {data.actionnaire.nom} "
            f"{data.actionnaire.nb_actions} actions",
            f"Soit un total de {data.nb_actions_total} actions",
        ],
    )
    _fixed_articles_8_to_11(document)
    _article_12_president(document, data)
    _fixed_articles_13_to_16(document)
    _article(document, "ARTICLE 17 - EXERCICE SOCIAL - COMPTES SOCIAUX")
    _paragraphs(
        document,
        [
            "Chaque exercice social a une durée d'une année, qui commence le "
            f"{data.exercice_debut} et finit le {data.exercice_fin}.",
            "Par exception, le premier exercice commencera le jour de l'immatriculation de "
            "la Société au Registre du Commerce et des Sociétés et se terminera le "
            f"{data.exercice_cloture_1}.",
            "Les comptes annuels (bilan, compte de résultat et annexe), l'inventaire, le "
            "rapport de gestion et les rapports spéciaux de la gérance ainsi que, le cas "
            "échéant, les rapports du Commissaire aux Comptes sont établis conformément "
            "aux lois et règlements en vigueur.",
            "L'associé unique approuve les comptes annuels et décide l'affectation du "
            "résultat dans les six mois de la clôture de l'exercice social.",
            "En cas de pluralité d'associés, l'Assemblée des associés approuve les comptes "
            "annuels dans les six mois de la clôture de l'exercice social.",
        ],
    )
    _fixed_articles_18_to_27(document)


def _fixed_articles_8_to_11(document) -> None:
    _article(document, "ARTICLE 8 – QUALITE DES ASSOCIES")
    _paragraphs(
        document,
        [
            "Le capital et les droits de vote de la Société sont exclusivement détenus par "
            "les associés qui exercent la profession de médecin au sein des SEL dans "
            "lesquelles la Société détient des participations, et ce en application des "
            "dispositions de l’ordonnance n°2023-77 du 8 février 2023 qui exige que plus "
            "de la moitié du capital et des droits de vote soit détenue par des personnes "
            "exerçant la profession de médecin.",
            "Le complément du capital social peut être détenu cumulativement ou "
            "individuellement par :",
            "Pendant un délai de dix ans, des personnes physiques qui, ayant cessé toute "
            "activité professionnelle, ont exercé cette profession au sein de la société ;",
            "Les ayants droit des personnes physiques mentionnées ci-dessus pendant un "
            "délai de cinq ans suivant leur décès ;",
            "Des personnes exerçant l'une des professions libérales de santé constituant "
            "l'objet social de la société.",
        ],
    )
    _article(document, "ARTICLE 9 – AUGMENTATION ET REDUCTION DU CAPITAL")
    _paragraphs(
        document,
        [
            "Article 9.1 Augmentation du capital",
            "La société pourra, en vertu d’une décision collective des Associés ou de "
            "l’Associé unique, être augmenté en une ou plusieurs fois, par tous moyens "
            "et voies de droit, notamment par :",
            "La création d’actions nouvelles attribuées en représentation d’apports en "
            "nature ou en numéraire ;",
            "L’augmentation du montant du nominal de celles existant déjà, en cas "
            "d’incorporation au capital de bénéfices, reports à nouveau, primes d’émission "
            "ou réserves disponibles.",
            "L’augmentation de capital par la Société ne peut avoir pour effet de "
            "contrevenir aux stipulations de l’article 8 des présents statuts.",
            "Les attributions d’actions nouvelles, s’ils n’ont pas déjà la qualité d’associé, "
            "devront être agréés dans les conditions fixées par l’article 10.1 des statuts "
            "pour les cessions d’actions. Les attributaires doivent solliciter leur agrément "
            "au moment de leur souscription.",
            "Les actions nouvelles doivent être entièrement libérées et réparties dès leur "
            "création.",
            "Article 9.2 Réduction du capital",
            "Le capital social peut être réduit, en vertu d’une décision collective "
            "extraordinaire des associés ou de l’associé unique, par voie de réduction du "
            "nombre des actions ou de leur valeur nominale et suivant les modalités des "
            "dispositions légales et réglementaires en vigueur.",
        ],
    )
    _article(document, "ARTICLE 10 - CESSION ET TRANSMISSION DES ACTIONS")
    _paragraphs(document, _ARTICLE_10_PARAGRAPHS)
    _article(document, "ARTICLE 11 - COMPTES COURANTS")
    _paragraphs(
        document,
        [
            "Outre leurs apports, l'associé unique ou les associés pourront verser ou "
            "laisser à disposition de la Société toutes sommes dont elle pourrait avoir "
            "besoin. Ces sommes seront inscrites au crédit d'un compte ouvert au nom de "
            "l'associé.",
            "Les comptes courants ne doivent jamais être débiteurs et la Société a la "
            "faculté d'en rembourser tout ou partie, après avis donné par écrit un mois à "
            "l'avance, sauf stipulation contraire.",
        ],
    )


def _article_12_president(document, data: _ResolvedStatutsSas) -> None:
    _article(document, "ARTICLE 12 – PRESIDENT")
    _paragraphs(
        document,
        [
            "NOMINATION ET POUVOIRS",
            "La Société est administrée par un président, personne physique, Associé de la "
            "Société, exerçant la médecine et associé au sein des sociétés d’exercice "
            "libéral dans lesquelles la Société détient des actions.",
            "Le président est désigné par l'associé unique ou, en cas de pluralité "
            "d'associés, par les associés représentant plus de la moitié des actions.",
            "Dans les rapports avec les tiers, les pouvoirs du président sont les plus "
            "étendus pour agir en toute circonstance au nom de la Société, dans la limite "
            "de l’objet social et sous réserve des pouvoirs que la loi attribue expressément "
            "à l'associé unique ou aux associés.",
            "Dans les rapports avec les tiers, la société est engagée même par les actes du "
            "Président qui ne relèvent pas de I'objet social, à moins qu'elle ne prouve que "
            "les tiers savaient que I'acte dépassait cet obiet ou qu'ils ne pouvaíent "
            "l'ignorer compte tenu des circonstances, étant exclu que seule la publication "
            "des statuts suffise à constituer une preuve.",
            "Sur le plan interne, le président peut faire tous les actes de gestion "
            "conformes à l'intérêt de la Société.",
            "Le président est révocable par décision de l'associé unique ou, en cas de "
            "pluralité d'associés, par décision des associés représentant plus de la moitié "
            "des actions.",
            f"L’Associé Unique, Monsieur {data.president.prenom} {data.president.nom},",
            f"Demeurant {data.president.adresse_personnelle_affichee}",
            "est nommé président de la Société et ce pour une durée illimitée.",
            "REMUNERATION",
            "La rémunération du Président se divise en deux parties : une correspondant à "
            "la rémunération de ses fonctions techniques et une autre au titre de son mandat "
            "social dans la société. Le montant total de cette rémunération pourra être "
            "fixée chaque année dans le procès-verbal d'approbation des comptes. Les "
            "Associés, lors de cette assemblée générale, pourront également définir les "
            "modalités de prise en charge de cette rémunération ainsi que son traitement "
            "comptable et fiscal.",
            "Outre sa rémunération, le président sera remboursé, sur justificatifs, de ses "
            "frais de déplacement et de représentation.",
        ],
    )


def _fixed_articles_13_to_16(document) -> None:
    _paragraphs(document, _ARTICLES_13_TO_16_PARAGRAPHS)


def _fixed_articles_18_to_27(document) -> None:
    _paragraphs(document, _ARTICLES_18_TO_27_PARAGRAPHS)


def _add_signature_and_annex(document, data: _ResolvedStatutsSas) -> None:
    add_statuts_signature_block(
        document,
        [
            f"Fait à {data.signature_lieu}",
            "Le",
            f"{data.president.prenom} {data.president.nom}",
        ],
        mention_lines=[
            "Faire précéder de la mention",
            "« Bon pour acceptation des fonctions de Président »",
        ],
    )
    document.add_page_break()
    add_statuts_annex_heading(
        document,
        "ANNEXE",
        "ETAT DES ENGAGEMENTS PRIS AVANT\nLA CONSTITUTION DE LA SOCIETE",
    )
    add_paragraph(document, "Ouverture d'un compte bancaire auprès de la Banque.")


def _article(document, title: str) -> None:
    add_statuts_article_heading(document, title)


def _paragraphs(document, paragraphs: list[str]) -> None:
    for paragraph in paragraphs:
        if paragraph.startswith("- "):
            add_statuts_hanging_list_item(document, paragraph[2:])
            continue
        if paragraph.startswith("ARTICLE "):
            _article(document, paragraph)
        else:
            add_statuts_body_paragraph(document, paragraph)


def _validate_sas_scope(ctx: DocumentGenerationContext) -> None:
    if ctx.structure != "SAS":
        raise ValueError(f"dossier.structure doit etre SAS pour {DOCUMENT_CODE}.")
    if ctx.statuts_sas is None:
        raise ValueError(f"statuts_sas est obligatoire pour {DOCUMENT_CODE}.")
    statuts_type = _required_text(ctx.statuts_sas.type, "statuts_sas.type").lower()
    profession = _required_text(ctx.statuts_sas.profession, "statuts_sas.profession").lower()
    if statuts_type != STATUTS_SAS_TYPE:
        raise ValueError(f"statuts_sas.type doit etre {STATUTS_SAS_TYPE} pour {DOCUMENT_CODE}.")
    if _normalize_profession(profession) != STATUTS_SAS_PROFESSION:
        raise ValueError(
            f"statuts_sas.profession doit etre {STATUTS_SAS_PROFESSION} pour {DOCUMENT_CODE}."
        )


def _validate_actionnaire_unique(
    ctx: DocumentGenerationContext,
    actionnaire: SpfplPerson,
    president: StatutsPresident,
    capital: CapitalSouscription,
) -> None:
    if ctx.associes and len(ctx.associes) != 1:
        raise ValueError("les statuts SAS V1 sont limites a un actionnaire unique.")
    if len(capital.souscripteurs) != 1:
        raise ValueError(
            "capital_souscription.souscripteurs doit contenir exactement un "
            f"souscripteur pour {DOCUMENT_CODE}."
        )
    if president.ref_associe_index != 0:
        raise ValueError(f"president.ref_associe_index doit etre 0 pour {DOCUMENT_CODE}.")
    _required_text(president.prenom, "president.prenom")
    _required_text(president.nom, "president.nom")
    _required_text(president.adresse_personnelle_affichee, "president.adresse_personnelle_affichee")
    if not _same_text(president.prenom, actionnaire.prenom) or not _same_text(
        president.nom,
        actionnaire.nom,
    ):
        raise ValueError(
            "president doit designer la meme personne que actionnaire_unique "
            f"pour {DOCUMENT_CODE}."
        )


def _validate_capital(
    societe: SocieteSpfpl,
    actionnaire: SpfplPerson,
    capital: CapitalSouscription,
) -> None:
    societe_total = _required_int(societe.nb_actions_total, "societe_spfpl.nb_actions_total")
    actionnaire_actions = _required_int(actionnaire.nb_actions, "actionnaire_unique.nb_actions")
    capital_total = _required_int(
        capital.nb_actions_total,
        "capital_souscription.nb_actions_total",
    )
    souscripteur = capital.souscripteurs[0]
    souscripteur_actions = _required_int(
        souscripteur.nb_actions,
        "capital_souscription.souscripteurs[0].nb_actions",
    )
    if len({societe_total, actionnaire_actions, capital_total, souscripteur_actions}) != 1:
        raise ValueError(
            "societe_spfpl.nb_actions_total, actionnaire_unique.nb_actions et "
            "capital_souscription doivent etre coherents pour "
            f"{DOCUMENT_CODE}."
        )
    valeur_societe = _required_text(
        societe.valeur_nominale_action,
        "societe_spfpl.valeur_nominale_action",
    )
    valeur_capital = _required_text(
        capital.valeur_nominale_action,
        "capital_souscription.valeur_nominale_action",
    )
    if valeur_societe != valeur_capital:
        raise ValueError(
            "societe_spfpl.valeur_nominale_action doit correspondre a "
            "capital_souscription.valeur_nominale_action pour "
            f"{DOCUMENT_CODE}."
        )
    if not _same_text(souscripteur.prenom, actionnaire.prenom) or not _same_text(
        souscripteur.nom,
        actionnaire.nom,
    ):
        raise ValueError(
            "capital_souscription.souscripteurs[0] doit correspondre a "
            f"actionnaire_unique pour {DOCUMENT_CODE}."
        )


def _validate_marital_sentence(actionnaire: SpfplPerson) -> None:
    situation = _required_text(
        actionnaire.situation_maritale,
        "actionnaire_unique.situation_maritale",
    )
    if not _normalize_profession(situation).startswith("mari"):
        raise ValueError(
            "la phrase matrimoniale source n'est stabilisee que pour une situation "
            f"mariee pour {DOCUMENT_CODE}."
        )
    _required_text(actionnaire.regime_matrimonial, "actionnaire_unique.regime_matrimonial")
    _required_conjoint(actionnaire)


def _required_societe_spfpl(ctx: DocumentGenerationContext) -> SocieteSpfpl:
    if ctx.societe_spfpl is None:
        raise ValueError(f"societe_spfpl est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.societe_spfpl


def _required_actionnaire(ctx: DocumentGenerationContext) -> SpfplPerson:
    if ctx.actionnaire_unique is None:
        raise ValueError(f"actionnaire_unique est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.actionnaire_unique


def _required_president(ctx: DocumentGenerationContext) -> StatutsPresident:
    if ctx.president is None:
        raise ValueError(f"president est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.president


def _required_capital_souscription(ctx: DocumentGenerationContext) -> CapitalSouscription:
    if ctx.capital_souscription is None:
        raise ValueError(f"capital_souscription est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.capital_souscription


def _required_conjoint(person: SpfplPerson) -> SpfplConjoint:
    if person.conjoint is None:
        raise ValueError(f"actionnaire_unique.conjoint est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(person.conjoint.civilite_affichage, "actionnaire_unique.conjoint.civilite")
    _required_text(person.conjoint.prenom, "actionnaire_unique.conjoint.prenom")
    _required_text(person.conjoint.nom, "actionnaire_unique.conjoint.nom")
    return person.conjoint


def _required_ordre(person: SpfplPerson) -> SpfplOrdre:
    if person.ordre is None:
        raise ValueError(f"actionnaire_unique.ordre est obligatoire pour {DOCUMENT_CODE}.")
    return person.ordre


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def _required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _address_display(societe: SocieteSpfpl) -> str:
    if societe.siege is None:
        raise ValueError(f"societe_spfpl.siege est obligatoire pour {DOCUMENT_CODE}.")
    if societe.siege.adresse_affichee:
        return societe.siege.adresse_affichee.strip()
    return (
        f"{_required_text(societe.siege.num_voie, 'societe_spfpl.siege.num_voie')} "
        f"{_required_text(societe.siege.voie, 'societe_spfpl.siege.voie')}, "
        f"{_required_text(societe.siege.cp, 'societe_spfpl.siege.cp')} "
        f"{_required_text(societe.siege.ville, 'societe_spfpl.siege.ville')}"
    )


def _person_address(person: SpfplPerson, field_name: str) -> str:
    if person.adresse_personnelle_affichee:
        return person.adresse_personnelle_affichee.strip()
    if person.adresse_personnelle is None:
        raise ValueError(f"{field_name}.adresse_personnelle est obligatoire pour {DOCUMENT_CODE}.")
    if person.adresse_personnelle.adresse_affichee:
        return person.adresse_personnelle.adresse_affichee.strip()
    return (
        f"{_required_text(person.adresse_personnelle.num_voie, f'{field_name}.adresse.num_voie')} "
        f"{_required_text(person.adresse_personnelle.voie, f'{field_name}.adresse.voie')}, "
        f"{_required_text(person.adresse_personnelle.cp, f'{field_name}.adresse.cp')} "
        f"{_required_text(person.adresse_personnelle.ville, f'{field_name}.adresse.ville')}"
    )


def _person_name(person: SpfplPerson, field_name: str) -> str:
    return (
        f"{_required_text(person.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{_required_text(person.prenom, f'{field_name}.prenom')} "
        f"{_required_text(person.nom, f'{field_name}.nom')}"
    )


def _conjoint_name(conjoint: SpfplConjoint) -> str:
    return (
        f"{_required_text(conjoint.civilite_affichage, 'actionnaire_unique.conjoint.civilite')} "
        f"{_required_text(conjoint.prenom, 'actionnaire_unique.conjoint.prenom')} "
        f"{_required_text(conjoint.nom, 'actionnaire_unique.conjoint.nom')}"
    )


def _qualification(person: SpfplPerson) -> str:
    if person.qualification_principale:
        return person.qualification_principale.strip()
    return _required_text(person.profession, "actionnaire_unique.qualification_principale")


def _format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return _required_text(value, field_name)


def _same_text(left: str | None, right: str | None) -> bool:
    if left is None or right is None:
        return False
    return left.strip().casefold() == right.strip().casefold()


def _normalize_profession(value: str) -> str:
    return value.casefold().replace("é", "e").replace("è", "e").replace("ê", "e")


_ARTICLE_10_PARAGRAPHS = [
    "Les actions ne peuvent être transmises ou cédées qu’au profit d’une personne "
    "justifiant de l’une des qualités énoncées à l’article 8 des présents statuts. "
    "Cette réserve vaut pour tous les cas de transmission ou de cession.",
    "Article 10.1 Clause d’agrément",
    "Toutes les cessions d’actions, y compris entre Associés, entre conjoints et entre "
    "ascendants et descendants sont obligatoirement soumises à la présente clause "
    "d’agrément, sauf transmission pour cause de décès ou lorsque le capital est détenu "
    "par un Associé unique.",
    "L’agrément est donné par la majorité des trois quarts des actions de la Société, "
    "actions du cédant, et éventuelles actions du cessionnaire comprises, les votes "
    "correspondant à ces deux derniers types de actions étant automatiquement considérés "
    "comme en faveur de l’agrément, même si leurs titulaires ne sont pas en mesure "
    "d’exprimer leur vote.",
    "Pour obtenir cet agrément, l’Associé qui veut transmettre tout ou partie des actions "
    "qu’il détient doit notifier son projet à la Société et aux Associés en indiquant le "
    "nom, prénoms, profession et domicile du cessionnaire proposé, le nombre d’actions "
    "qu’il désire transmettre et, s’il s’agit d’une vente, le prix convenu. Cette "
    "notification doit être effectuée par acte extrajudiciaire ou par lettre recommandée "
    "avec demande d’avis de réception.",
    "Dans les quinze jours à compter de la notification du projet de cession, le président "
    "convoque l’assemblée des Associés pour qu’elle délibère sur le projet de cession ou "
    "consulte les Associés par écrit sur ledit projet.",
    "La décision des Associés n’a pas à être motivée. Elle est notifiée, dès son prononcé, "
    "au cédant, par le président, par lettre recommandée avec demande d’avis de réception.",
    "Si le cessionnaire proposé est agréé, la cession doit être régularisée dans le délai "
    "maximal de 90 jours à partir de la notification de la décision des Associés, à "
    "défaut de quoi, une nouvelle demande d’agrément serait nécessaire.",
    "Si la société n’a pas fait connaître sa décision dans le délai de 30 jours à compter "
    "de la notification du projet faite par le cédant, le consentement à la cession est "
    "réputé acquis.",
    "Si la société a refusé de consentir à la cession, le cédant peut, dans un délai de "
    "15 jours de la notification de refus qui lui est faite, signifier par lettre "
    "recommandée avec demande d’avis de réception, adressé à la Société, qu’il renonce "
    "à son projet de cession.",
    "A défaut de renonciation de la part du cédant, la Société est tenu dans un délais de "
    "trois mois à compter de son refus effectif ou réputé acquis, d’acquérir ou de faire "
    "acquérir les actions.",
    "Si, à l’expiration du délai imparti, la Société n’a pas racheté ou fait racheter les "
    "actions, l’Associé cédant peut réaliser la cession initialement prévue.",
    "La Société peut également, avec le consentement de l’Associé candidat cédant, décider, "
    "dans le même délai, de réduire son capital du montant de la valeur nominale des "
    "actions dont la cession est proposée, et racheter celles-ci.",
    "La procédure d’agrément peut être substituée par un acte unanime des Associés agréant "
    "la cession.",
    "À défaut d'accord entre les parties, le prix des actions est fixé dans les conditions "
    "prévues à l'article 1843-4 du Code civil.",
    "Article 10.2 Formalités",
    "Toute cession d’actions doit être constatée par un acte notarié ou sous seing privé.",
    "Pour être opposable à la Société, elle doit lui être signifiée par exploit d'huissier "
    "ou être acceptée par elle dans un acte notarié. La signification peut être remplacée "
    "par le dépôt d'un original de l'acte de cession au siège social contre remise par le "
    "président d'une attestation de ce dépôt.",
    "Pour être opposable aux tiers, elle doit en outre avoir été déposée au greffe, en "
    "annexe au Registre du commerce et des Sociétés.",
    "En cas de dissolution de la communauté de biens existant entre l'associé unique et "
    "son conjoint, ou en cas de décès de l’associé unique, la Société continue de plein "
    "droit, que dans la mesure où le conjoint s’étant vu attribués les actions ou l’un "
    "des héritiers a la qualité d’associé conformément à l’article 8 des présents statuts. "
    "La Société continue soit sous la forme d'une Société Unipersonnelle par actions "
    "simplifiée si la totalité des actions est attribuée à l'un des époux, soit sous la "
    "forme d'une Société par actions simplifiée pluripersonnelle si les actions sont "
    "partagées entre les époux.",
    "Article 10.3 - Transmission par décès",
    "En cas de décès de l’associé unique, la société ne continue avec ses héritiers que "
    "dans la mesure où un ou plusieurs d’entre eux ont la qualité d’associés conformément "
    "à l’article 8 des présents statuts.",
    "Les ayants droit d'un associé décédé ne pourront conserver les parts de la Société "
    "que pendant un délai de cinq ans à compter du décès.",
    "Lorsque, à l'expiration du délai de cinq ans à compter du décès de leur auteur, les "
    "héritiers et ayants droit n'ont pas cédé les parts qu'ils détiennent, la Société peut, "
    "nonobstant leur opposition, décider de réduire son capital et de les racheter.",
    "Les dispositions de l'alinéa qui précède ne sont pas applicables aux héritiers et "
    "ayants droit qui, au jour du décès de leur auteur, sont déjà membres de la Société "
    "ni à ceux qui acquièrent la qualité d'associé professionnel avant l'expiration du "
    "délai visé à cet alinéa.",
    "En cas de décès d'un ayant droit, d'un professionnel assimilé ou d'un associé externe, "
    "ses parts sont librement transmises au profit de toute personne qui est déjà membre "
    "de la Société.",
    "Tous autres héritiers ou ayants droit ne deviennent associés que s'ils reçoivent "
    "l'agrément de la majorité des trois quarts des actions de la Société.",
    "Article 10.4 Nullité des cessions",
    "Toutes les cessions ou transmissions d’actions effectuées en violation des "
    "dispositions de l’article 10 des présents statuts sont nulles.",
]

_ARTICLES_13_TO_16_PARAGRAPHS = [
    "ARTICLE 13 - DIRECTEURS GENERAUX",
    "ARTICLE 13-1 - DESIGNATION",
    "Le ou les Directeurs généraux sont des personnes physiques, choisis parmi des "
    "professionnels exerçants réalisant leur activité au sein de la société dans lesquelles "
    "la société de participations financières de professions libérales détient des "
    "participations.",
    "Sur la proposition du Président, la collectivité des associés peut nommer à la "
    "majorité des deux tiers des voix, un ou plusieurs directeurs généraux, portant le "
    "titre de directeur général ou de directeur général délégué, et investis, sauf "
    "disposition contraire inopposable aux tiers, des mêmes pouvoirs que le Président.",
    "Le Directeur Général peut être lié par un contrat de travail à condition que ce contrat "
    "corresponde à un emploi effectif.",
    "Le Directeur Général devra être désigné parmi les associés exerçant la profession au "
    "sein de la société faisant l’objet de la prise de participations.",
    "ARTICLE 13-2 - DUREE",
    "La durée des fonctions du Directeur Général est fixée dans la décision de nomination.",
    "Toutefois, en cas de cessation des fonctions du Président, le Directeur Général "
    "conserve ses fonctions jusqu'à la nomination du nouveau Président, sauf décision "
    "contraire des associés.",
    "Les fonctions de Directeur Général prennent fin soit par le décès, la démission, la "
    "révocation, l'expiration de son mandat, soit par l'ouverture d'une procédure de "
    "redressement ou de liquidation judiciaire.",
    "Le Directeur Général peut démissionner de son mandat sous réserve de respecter un "
    "préavis de 15 jours, lequel pourra être réduit lors de la consultation de la "
    "collectivité des actionnaires qui aura à statuer sur le remplacement du Président "
    "démissionnaire.",
    "ARTICLE 13-3 - REVOCATION",
    "Le Directeur Général peut être révoqué à tout moment, sans juste motif, par décision "
    "de la collectivité des associés, sur proposition du Président, prise à la majorité "
    "simple.",
    "Cette révocation n'ouvre droit à aucune indemnisation.",
    "En outre le Directeur Général sera révoqué de plein droit, sans indemnisation dans "
    "les cas suivants :",
    "- Interdiction de diriger, gérer, administrer, ou contrôler une entreprise ou personne "
    "morale, incapacité ou faillite personnelle du Directeur Général personne physique ;",
    "- Exclusion du Directeur Général Associé,",
    "- Interdiction définitive d'exercer la profession de médecin.",
    "ARTICLE 13-4 - REMUNERATION",
    "Le Directeur Général peut recevoir une rémunération qui se divise en deux parties : "
    "une correspondant à la rémunération de ses fonctions techniques et une autre au titre "
    "de son mandat social dans la société. Le montant total de cette rémunération pourra "
    "être fixée chaque année dans le procès-verbal d'approbation des comptes. Les "
    "Associés, lors de cette assemblée générale, pourront également définir les modalités "
    "de prise en charge de cette rémunération ainsi que son traitement comptable et fiscal.",
    "En outre, le Directeur Général est remboursé de ses frais de représentation et de "
    "déplacement sur justificatifs.",
    "ARTICLE 13-5 - POUVOIRS",
    "Le Directeur Général dispose des mêmes pouvoirs que le Président, sous réserve des "
    "limitations éventuellement fixées par la décision de nomination ou décision ultérieure.",
    "Le Directeur Général ne dispose pas du pouvoir de représenter la société à l'égard "
    "des tiers.",
    "ARTICLE 14 - CONVENTIONS ENTRE LA SOCIETE ET SES DIRIGEANTS",
    "Les conventions qui interviennent directement ou par personne interposée entre la "
    "société et l'un de ses dirigeants ou l’un de ses associés disposant d’une fraction "
    "des droits de vote suppérieure à 10% sont mentionnées au registre des décisions des "
    "associés sont soumises à l'approbation de la collectivité des associés, conformément "
    "aux dispositions de l’article L.227-10 du Code de commerce, sur rapport du "
    "commissaire aux comptes, ou à défaut du président.",
    "Par dérogation aux dispositions du premier alinéa, lorsque la société ne comprend "
    "qu’un seul associé, il est seulement fait mention au registre des décisions des "
    "conventions intervenues directement ou par personne interposées entre la société et "
    "son dirigeant.",
    "S'il n'existe pas de Commissaire aux Comptes, les conventions conclues par le dirigeant "
    "non associé sont soumises à l'approbation préalable de l'associé unique ou de "
    "l'assemblée des associés.",
    "A peine de nullité du contrat, il est interdit aux dirigeants ou associés autres que "
    "les personnes morales de contracter sous quelque forme que ce soit, des emprunts "
    "auprès de la Société, de se faire consentir par elle un découvert, en compte courant "
    "ou autrement, ainsi que de faire cautionner ou avaliser par elle leurs engagements "
    "envers les tiers. Cette interdiction s'applique aux conjoint, ascendants et descendants "
    "des dirigeants ou associés ainsi qu'à toute personne interposée et aux représentants "
    "légaux des personnes morales associées.",
    "Ces conventions devront êtres communiquées par les intéressés au conseil départemental "
    "concerné dans le mois suivant leur conclusion, conformément à l’article L.4113-9 du "
    "code de la Santé Publique.",
    "ARTICLE 15 - DECISIONS D’ACTIONNAIRES",
    "L'associé unique exerce les pouvoirs dévolus à l'assemblée des actionnaires. Il ne "
    "peut déléguer ses pouvoirs. Ses décisions sont constatées par des procès-verbaux "
    "signés par lui et répertoriés dans un registre coté et paraphé comme les registres "
    "d'assemblées.",
    "En cas de pluralité d'associés, les décisions collectives sont prises, au choix de la "
    "gérance, en assemblée ou par consultation écrite des associés. Elles peuvent aussi "
    "résulter du consentement de tous les associés exprimé dans un acte.",
    "Toutefois, la réunion d'une assemblée est obligatoire pour statuer sur l'approbation "
    "annuelle des comptes ou sur demande d'un ou plusieurs associés détenant la moitié des "
    "actions ou détenant, s'ils représentent au moins le quart des associés, le quart des "
    "actions.",
    "Les Assemblées Générales sont convoquées et délibèrent dans les conditions et avec "
    "les effets fixés par les lois et règlements en vigueur.",
    "En cas de consultation écrite, la gérance adresse à chaque associé, par lettre "
    "recommandée, le texte des résolutions proposées ainsi que les documents nécessaires "
    "à l'information des associés.",
    "Les associés disposent d'un délai de quinze jours à compter de la date de réception "
    "du projet de résolutions pour transmettre leur vote à la gérance par lettre recommandée. "
    "Tout associé n'ayant pas répondu dans le délai ci-dessus est considéré comme s'étant "
    "abstenu.",
    "Chaque associé a le droit de participer aux décisions collectives et dispose d'un "
    "nombre de voix égal à celui des actions qu'il possède. Un associé peut se faire "
    "représenter par son conjoint à moins que la Société ne comprenne que les deux époux. "
    "Sauf si les associés sont au nombre de deux, un associé peut se faire représenter "
    "par un autre associé.",
    "Si une ou plusieurs actions sont grevées d'usufruit, le droit de vote appartient à "
    "l’associé exerçant la profession de médecin, y compris pour les décisions concernant "
    "l'affectation des résultats.",
    "ARTICLE 16 - COMMISSAIRES AUX COMPTES",
    "Un ou plusieurs Commissaires aux Comptes titulaires et suppléants peuvent ou doivent "
    "être désignés dans les conditions prévues par l'article L. 227-9-1 du Code de commerce.",
    "Ils sont nommés pour une durée de six exercices et exercent leurs fonctions dans les "
    "conditions et avec les effets prévus par les dispositions législatives et réglementaires "
    "en vigueur.",
]

_ARTICLES_18_TO_27_PARAGRAPHS = [
    "ARTICLE 18 - AFFECTATION ET REPARTITION DES BENEFICES",
    "Le bénéfice distribuable est constitué par le bénéfice de l'exercice, diminué des "
    "pertes antérieures et des sommes portées en réserve en application de la loi et des "
    "statuts, et augmenté du report bénéficiaire.",
    "En cas de pluralité d'associés, l'Assemblée des associés détermine la part attribuée "
    "à chacun des associés. L'associé unique ou l'Assemblée des associés détermine les "
    "modalités de mise en paiement des dividendes, qui doit intervenir dans un délai de "
    "neuf mois après la clôture de l'exercice, sauf prolongation par décision de justice.",
    "De même, l'associé unique ou l'Assemblée Générale peut décider la distribution de "
    "sommes prélevées sur les réserves disponibles en indiquant expressément les postes "
    "de réserves sur lesquels les prélèvements ont été effectués.",
    "Toutefois, les dividendes sont prélevés par priorité sur le bénéfice distribuable de "
    "l'exercice.",
    "L'associé unique ou l'Assemblée Générale peut également décider d'affecter les sommes "
    "distribuables aux réserves et au report à nouveau, en totalité ou en partie.",
    "Aucune distribution ne peut être faite lorsque les capitaux propres sont ou "
    "deviendraient à la suite de celle-ci inférieurs au montant du capital augmenté des "
    "réserves que la loi ne permet pas de distribuer.",
    "ARTICLE 19 - CAPITAUX PROPRES INFERIEURS A LA MOITIE DU CAPITAL SOCIAL",
    "Si, du fait de pertes constatées dans les documents comptables, les capitaux propres "
    "de la Société deviennent inférieurs à la moitié du capital social, l'associé unique "
    "ou, en cas de pluralité d'associés, l'Assemblée statuant à la majorité requise pour "
    "la modification des statuts doit, dans les quatre mois qui suivent l'approbation des "
    "comptes ayant fait apparaître cette perte, décider, s'il y a lieu à dissolution "
    "anticipée de la Société.",
    "Si la dissolution n'est pas prononcée, le capital doit être, sous réserve des "
    "dispositions légales relatives au capital minimum dans les sociétés par actions "
    "simplifiée et, dans le délai fixé par la loi, réduit d'un montant égal à celui des "
    "pertes qui n'ont pu être imputées sur les réserves si dans ce délai les capitaux "
    "propres ne sont pas redevenus au moins égaux à la moitié du capital social.",
    "En cas d'inobservation de ces prescriptions, tout intéressé peut demander en justice "
    "la dissolution de la Société. Il en est de même si l'Assemblée n'a pu délibérer "
    "valablement.",
    "Toutefois, le Tribunal ne peut prononcer la dissolution si, au jour où il statue sur "
    "le fond, la régularisation a eu lieu.",
    "ARTICLE 20 – EXCLUSION",
    "Tout associé peut être exclu de la Société dans les cas suivants :",
    "Pour tout associé personne morale, en cas de modification de son contrôle au sens de "
    "l'article",
    "L. 233-3 du Code de commerce ;",
    "Pour tout associé, personne physique ou morale, en cas de :",
    "Mise en redressement judiciaire ;",
    "Exercice d’une activité concurrente à celle de la SEL ;",
    "Violation d'une clause statutaire causant un préjudice pour la Société ;",
    "Violation de toute autre règle légale, réglementaire ou contractuelle sanctionnée",
    "expressément par l’exclusion ;",
    "Cessation de toute activité au sein de SEL ;",
    "Exclusion de la SEL ;",
    "Violation du règlement intérieur de la SEL.",
    "L'exclusion est décidée en assemblée générale à la majorité de deux tiers des droits "
    "de vote des associés de la Société présents ou représentés lors de la consultation ; "
    "l'associé dont l'exclusion est proposée participe au vote et ses actions sont prises "
    "en compte pour le calcul de cette majorité.",
    "Si l'exclusion n'est pas décidée, l'associé conserve ses droits et obligations "
    "attachés à la qualité d’associé. Il reste propriétaire de ses actions.",
    "Aucune décision d'exclusion ne peut être prise si l'associé intéressé n'a pas été "
    "régulièrement convoqué à l'assemblée générale, quinze jours au moins avant la date "
    "prévue, par lettre recommandée avec accusé de réception, et, s'il n'a pas été mis "
    "à même de présenter sa défense sur les faits précis qui lui sont reprochés.",
    "Les actions de l'associé exclu sont acquises par la Société qui doit ensuite réduire "
    "son capital social.",
    "Le transfert de propriété des actions et le paiement du prix devront intervenir au "
    "plus tard à l’expiration d’un délai de six mois suivant le jour de la décision "
    "collective ayant prononcé l’exclusion de l’associé concerné, à première demande de "
    "la Société. Pour le cas où le cédant serait défaillant, le(s) Cessionnaire(s) "
    "pourrai(en)t séquestrer auprès de tout établissement bancaire ou financier acceptant "
    "cette mission ou, à défaut, la Société, le prix de cession. Dans ce cas, la simple "
    "remise à la Société des copies de la notification de la mise en jeu de la présente "
    "clause et d’une copie de la convention de séquestre vaudrait ordre de mouvement et "
    "obligerait la Société à passer les écritures qui en résulteraient dans le registre "
    "des mouvements de titres de la Société.",
    "ARTICLE 21 - DISSOLUTION - LIQUIDATION",
    "La Société est dissoute à l'arrivée du terme statutaire de sa durée, sauf prorogation "
    "régulière, ou s'il survient une cause de dissolution prévue par la loi.",
    "Si la Société ne comprend qu'un seul associé personne morale, la dissolution pour "
    "quelque cause que ce soit entraîne la transmission universelle du patrimoine à "
    "l'associé unique personne morale, sans qu'il y ait lieu à liquidation.",
    "Les créanciers de la Société peuvent faire opposition à la dissolution dans le délai "
    "de trente jours à compter de la publication de celle-ci. Le Tribunal judiciaire saisi "
    "de l'opposition peut soit la rejeter, soit ordonner le paiement des créances, soit "
    "ordonner la constitution de garanties si la Société en offre et si elles sont jugées "
    "suffisantes. La transmission à l'associé unique personne morale du patrimoine de la "
    "Société et la disparition de la personnalité morale de celle-ci n'interviennent qu'à "
    "l'issue du délai d'opposition ou, le cas échéant, lorsque l'opposition a été rejetée "
    "en première instance ou que le remboursement des créances a été effectué ou les "
    "garanties constituées.",
    "Si la Société comprend au moins deux associés ou un seul associé personne physique, "
    "la dissolution, pour quelque cause que ce soit, entraîne sa liquidation. Cette "
    "liquidation est effectuée dans les conditions et selon les modalités prévues par les "
    "dispositions légales et réglementaires en vigueur au moment de son ouverture.",
    "La liquidation est faite par le président alors en fonction à moins qu'une décision "
    "collective ne désigne un autre liquidateur.",
    "Le ou les liquidateurs ont les pouvoirs les plus étendus pour réaliser l'actif même "
    "à l'amiable et acquitter le passif. Il peut être autorisé par les associés à continuer "
    "les affaires en cours ou à en engager de nouvelles pour les besoins de la liquidation.",
    "ARTICLE 22 - TRANSFORMATION DE LA SOCIETE",
    "La transformation de la Société en une société commerciale d'une autre forme (SAS – SA)"
    "peut être décidée par les associés statuant aux conditions de majorité et selon les "
    "modalités requises par la loi.",
    "ARTICLE 23 - CONTESTATIONS",
    "Tous les litiges ou différends relatifs notamment à la validité, l’interprétation, "
    "l’exécution ou la résolution des présents statuts seront soumis avant tout recours à "
    "une conciliation confiée au Conseil Départemental de l’Ordre des Médecins.",
    "Toutes les contestations qui pourraient surgir pendant la durée de la Société ou lors "
    "de sa liquidation entre la Société et l'associé unique ou entre la Société et les "
    "associés ou entre les associés eux-mêmes, relativement aux affaires sociales ou à "
    "l'exécution des présents statuts, seront soumises aux tribunaux compétents.",
    "ARTICLE 24– CONDITION SUSPENSIVE",
    "La constitution de la présente Société est réalisée sous la condition suspensive de "
    "son inscription au tableau de l’Ordre des Médecins. Cette inscription emportera "
    "automatiquement la levée de la présente condition suspensive .",
    "ARTICLE 25– ORDRE PROFESSIONNEL",
    "Les sociétés de participations financières doivent êlre inscrites sut la liste ou au "
    "Tableau du ou des Ordres professionnels concernés. Une fois par an, la Société "
    "adresse à I'Ordre professionnel dont elle relève un état de la composition de son "
    "capital social.",
    "La Société fait connaitre au Président du Conseil de l'Ordre compétent, dans un délai "
    "de tente (30) jours, à compter de la date à laquelle il se produit, tout changement "
    "dans la situation déclarée, avec les pièces justificatives idoines.",
    "Et une fois par an, la société de participations financières de professions libérales "
    "adresse au Président du Conseil de l'Ordre compétent un état de la composition de son "
    "capital social et des droits de vote afférents, ainsi qu'une version à jour de ses "
    "statuts.",
    "Sont également adressées par les associés de la société, dans les conditions prévues "
    "au premier alinéa, les conventions contenant des clauses portant sur l'organisation "
    "et les pouvoirs des organes de direction, d'administration ou de surveillance ayant "
    "fait l'objet d'une modification au cours de l'exercice écoulé.",
    "ARTICLE 26– FRAIS",
    "Tous les frais, droits et honoraires résultant des présents statuts seront portés au "
    "compte des frais généraux du premier exercice social.",
    "ARTICLE 27 - JOUISSANCE DE LA PERSONNALITE MORALE - POUVOIRS",
    "Est annexé aux présents statuts un état énumérant les actes accomplis avant leur "
    "signature pour le compte de la Société en formation avec l’indication pour chacun "
    "d’eux, des engagements qui en résulteront pour le compte de la Société.",
    "L'immatriculation de la Société au Registre du Commerce et des Sociétés emportera de "
    "plein droit reprise par la Société de ces actes et engagements qui seront réputés "
    "avoir été faits et souscrits dès l'origine par la Société. La société ne jouira de "
    "la personnalité morale qu'à compter de cette ìmmatriculation au registre du commerce "
    "et des sociétés.",
    "Enfin, le Président de la Société, agira au nom de la Société en formation, jusqu’à "
    "son immatriculation au Registre du Commerce et des Sociétés, notamment pour effectuer "
    "les formalités de publicité relatives à la constitution de la Société.",
]
