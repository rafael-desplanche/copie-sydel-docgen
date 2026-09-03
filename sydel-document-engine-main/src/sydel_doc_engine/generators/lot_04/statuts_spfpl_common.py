from __future__ import annotations

from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    ApportTitres,
    CapitalSouscription,
    DocumentGenerationContext,
    SocieteCible,
    SocieteSpfpl,
    SpfplPerson,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_paragraph,
    add_statuts_article_heading,
    add_statuts_body_paragraph,
    add_statuts_hanging_list_item,
    add_statuts_part_heading,
    add_statuts_signature_block,
    new_document,
)

DOCUMENT_CODE = "CODE-STATUTS-SPFPL-001"
SPFPL_CESSION_STRUCTURE = "SPFPL cession"
SPFPL_APPORT_STRUCTURE = "SPFPL apport"
OPERATION_CESSION = "cession"
OPERATION_APPORT = "apport"


def required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return required_text(value, field_name)


def required_societe_spfpl(ctx: DocumentGenerationContext) -> SocieteSpfpl:
    if ctx.societe_spfpl is None:
        raise ValueError(f"societe_spfpl est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.societe_spfpl


def required_capital_souscription(ctx: DocumentGenerationContext) -> CapitalSouscription:
    if ctx.capital_souscription is None:
        raise ValueError(f"capital_souscription est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.capital_souscription


def required_apport_titres(ctx: DocumentGenerationContext) -> ApportTitres:
    if ctx.apport_titres is None:
        raise ValueError(f"apport_titres est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.apport_titres


def required_societe_cible(ctx: DocumentGenerationContext) -> SocieteCible:
    if ctx.societe_cible is None:
        raise ValueError(f"societe_cible est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.societe_cible


def required_actionnaire_unique(ctx: DocumentGenerationContext) -> SpfplPerson:
    if ctx.actionnaire_unique is not None:
        return ctx.actionnaire_unique
    if ctx.operation_spfpl and ctx.operation_spfpl.type == OPERATION_CESSION and ctx.cedant:
        return ctx.cedant
    if ctx.operation_spfpl and ctx.operation_spfpl.type == OPERATION_APPORT and ctx.apporteur:
        return ctx.apporteur
    raise ValueError(f"actionnaire_unique est obligatoire pour {DOCUMENT_CODE}.")


def validate_common_statuts_context(
    ctx: DocumentGenerationContext,
    *,
    structure: str,
    operation: str,
) -> None:
    if ctx.structure != structure:
        raise ValueError(f"dossier.structure doit etre {structure} pour {DOCUMENT_CODE}.")
    if ctx.operation_spfpl is None:
        raise ValueError(f"operation_spfpl est obligatoire pour {DOCUMENT_CODE}.")
    operation_type = required_text(ctx.operation_spfpl.type, "operation_spfpl.type").lower()
    if operation_type != operation:
        raise ValueError(f"operation_spfpl.type doit etre {operation} pour {DOCUMENT_CODE}.")
    if ctx.dossier_options is None:
        raise ValueError(f"dossier.options est obligatoire pour {DOCUMENT_CODE}.")
    if operation == OPERATION_CESSION and not ctx.dossier_options.cession:
        raise ValueError(f"dossier.options.cession doit etre vrai pour {DOCUMENT_CODE}.")
    if operation == OPERATION_APPORT and not ctx.dossier_options.apport:
        raise ValueError(f"dossier.options.apport doit etre vrai pour {DOCUMENT_CODE}.")
    if operation == OPERATION_CESSION and ctx.dossier_options.apport:
        raise ValueError(f"un seul overlay SPFPL peut etre rendu pour {DOCUMENT_CODE}.")
    if operation == OPERATION_APPORT and ctx.dossier_options.cession:
        raise ValueError(f"un seul overlay SPFPL peut etre rendu pour {DOCUMENT_CODE}.")
    if len(ctx.associes) > 1:
        raise ValueError(
            f"les statuts SPFPL multi-associes sont bloques en V1 pour {DOCUMENT_CODE}."
        )
    if (
        ctx.capital_souscription is not None
        and len(ctx.capital_souscription.souscripteurs) > 1
    ):
        raise ValueError(
            f"les statuts SPFPL multi-associes sont bloques en V1 pour {DOCUMENT_CODE}."
        )


def company_siege_display(societe: SocieteSpfpl | SocieteCible, field_name: str) -> str:
    if societe.siege is None:
        raise ValueError(f"{field_name}.siege est obligatoire pour {DOCUMENT_CODE}.")
    if societe.siege.adresse_affichee:
        return societe.siege.adresse_affichee.strip()
    return (
        f"{required_text(societe.siege.num_voie, f'{field_name}.siege.num_voie')} "
        f"{required_text(societe.siege.voie, f'{field_name}.siege.voie')}, "
        f"{required_text(societe.siege.cp, f'{field_name}.siege.cp')} "
        f"{required_text(societe.siege.ville, f'{field_name}.siege.ville')}"
    )


def person_address_display(person: SpfplPerson, field_name: str) -> str:
    if person.adresse_personnelle_affichee:
        return person.adresse_personnelle_affichee.strip()
    if person.adresse_personnelle is None:
        raise ValueError(f"{field_name}.adresse_personnelle est obligatoire pour {DOCUMENT_CODE}.")
    address = person.adresse_personnelle
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{required_text(address.num_voie, f'{field_name}.adresse_personnelle.num_voie')} "
        f"{required_text(address.voie, f'{field_name}.adresse_personnelle.voie')}, "
        f"{required_text(address.cp, f'{field_name}.adresse_personnelle.cp')} "
        f"{required_text(address.ville, f'{field_name}.adresse_personnelle.ville')}"
    )


def render_statuts_docx(
    blocks: tuple[str, ...],
    replacements: dict[str, str],
    output_path: Path,
) -> Path:
    docx = new_document()
    index = 0
    while index < len(blocks):
        block = blocks[index]
        text = replace_placeholders(block, replacements)
        if text == "STATUTS" or _is_major_heading(text):
            if _is_major_heading(text):
                add_statuts_part_heading(docx, text, mode="boxed")
            else:
                add_paragraph(docx, text, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        elif text.startswith("ARTICLE "):
            add_statuts_article_heading(docx, text, underline=False)
        elif text.startswith("Fait à ") or text.startswith("Fait a "):
            signature_lines, mention_lines, index = _collect_signature_lines(
                blocks,
                replacements,
                index,
            )
            add_statuts_signature_block(docx, signature_lines, mention_lines=mention_lines)
            continue
        elif text.startswith("- "):
            add_statuts_hanging_list_item(docx, text[2:])
        elif _looks_like_numbered_list_item(text):
            add_statuts_hanging_list_item(docx, text, marker=None)
        else:
            add_statuts_body_paragraph(docx, text)
        index += 1

    full_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    if "[" in full_text or "]" in full_text:
        raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(output_path)
    return output_path


def replace_placeholders(text: str, replacements: dict[str, str]) -> str:
    rendered = text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered


def founder_common_replacements(founder: SpfplPerson, field_name: str) -> dict[str, str]:
    ordre = founder.ordre
    if ordre is None:
        raise ValueError(f"{field_name}.ordre est obligatoire pour {DOCUMENT_CODE}.")
    return {
        "[civilite]": required_text(
            founder.civilite_affichage,
            f"{field_name}.civilite_affichage",
        ),
        "[prenom]": required_text(founder.prenom, f"{field_name}.prenom"),
        "[prenoms]": required_text(founder.prenoms or founder.prenom, f"{field_name}.prenoms"),
        "[nom]": required_text(founder.nom, f"{field_name}.nom"),
        "[profession]": required_text(founder.profession, f"{field_name}.profession"),
        "[date_naissance]": format_display_date(
            founder.date_naissance,
            f"{field_name}.date_naissance",
        ),
        "[ville_naissance]": required_text(
            founder.ville_naissance,
            f"{field_name}.ville_naissance",
        ),
        "[departement_naissance]": required_text(
            founder.departement_naissance,
            f"{field_name}.departement_naissance",
        ),
        "[adresse_personnelle]": person_address_display(founder, field_name),
        "[situation_maritale]": required_text(
            founder.situation_maritale,
            f"{field_name}.situation_maritale",
        ),
        "[nationalite]": required_text(founder.nationalite, f"{field_name}.nationalite"),
        "[numero_ordre]": required_text(ordre.numero, f"{field_name}.ordre.numero"),
        "[numero_rpps]": required_text(ordre.numero_rpps, f"{field_name}.ordre.numero_rpps"),
    }


def _is_major_heading(text: str) -> bool:
    headings = {
        "DECISIONS DES ACTIONNAIRES",
        "RESULTATS SOCIAUX",
        "TRANSFORMATION DE LA SOCIETE",
        "DISSOLUTION – LIQUIDATION",
        "CONTESTATIONS",
        "CONSTITUTION DE LA SOCIETE",
        "ANNEXE 1",
        "ETAT DES ENGAGEMENTS PRIS AVANT",
        "LA CONSTITUTION DE LA SOCIETE",
    }
    return text in headings


def _collect_signature_lines(
    blocks: tuple[str, ...],
    replacements: dict[str, str],
    start_index: int,
) -> tuple[list[str], list[str], int]:
    signature_lines: list[str] = []
    mention_lines: list[str] = []
    index = start_index
    while index < len(blocks):
        rendered = replace_placeholders(blocks[index], replacements)
        if index > start_index and (rendered.startswith("ANNEXE") or _is_major_heading(rendered)):
            break
        if "Bon pour acceptation" in rendered:
            mention_lines.append(rendered)
        else:
            signature_lines.append(rendered)
        index += 1
    return signature_lines, mention_lines, index


def _looks_like_numbered_list_item(text: str) -> bool:
    return len(text) > 2 and text[0].isdigit() and text[1] in {"°", "."}
