from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    DocumentGenerationContext,
    StatutsCivilsAssocie,
    StatutsCivilsContext,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_paragraph,
    add_statuts_article_heading,
    add_statuts_body_paragraph,
    add_statuts_hanging_list_item,
    add_statuts_matrix_table,
    add_statuts_part_heading,
    add_statuts_signature_block,
    add_statuts_signature_grid,
    add_statuts_title_box,
    new_document,
)

DOCUMENT_CODE = "CODE-STATUTS-CIVILS-CORE-001"
MAX_ASSOCIES = 6


@dataclass(frozen=True)
class StatutsCivilTemplate:
    source_name_contains: tuple[str, ...]
    output_filename: str
    expected_structure: str
    expected_type: str
    associate_slice: tuple[int, int]
    apport_slice: tuple[int, int]
    capital_slice: tuple[int, int]
    signature_slice: tuple[int, int] | None = None
    append_signatures_after: int | None = None


SCS_TEMPLATE = StatutsCivilTemplate(
    source_name_contains=("SCS",),
    output_filename="statuts_scs.docx",
    expected_structure="SCS",
    expected_type="scs",
    associate_slice=(14, 18),
    apport_slice=(43, 58),
    capital_slice=(63, 76),
    append_signatures_after=253,
)

SCI_TEMPLATE = StatutsCivilTemplate(
    source_name_contains=("SCI",),
    output_filename="statuts_sci.docx",
    expected_structure="SCI",
    expected_type="sci",
    associate_slice=(25, 44),
    apport_slice=(97, 111),
    capital_slice=(120, 131),
    signature_slice=(612, 623),
)

SCI_IRIS_TEMPLATE = StatutsCivilTemplate(
    source_name_contains=("SCI", "IRIS"),
    output_filename="statuts_sci_iris.docx",
    expected_structure="SCI IRIS",
    expected_type="sci_iris",
    associate_slice=(24, 42),
    apport_slice=(95, 109),
    capital_slice=(120, 131),
    signature_slice=(629, 636),
)


def generate_statuts_civil_docx(
    ctx: DocumentGenerationContext,
    output_dir: Path,
    template: StatutsCivilTemplate,
) -> Path:
    data = _ResolvedStatutsCivil.from_context(ctx, template)
    source = _source_path(template)
    source_doc = Document(source)
    output_doc = new_document()
    output_doc.sections[0].footer.paragraphs[0].text = f"{data.denomination} - Statuts constitutifs"

    replacements = data.common_replacements()
    skip_until = -1
    for index, paragraph in enumerate(source_doc.paragraphs):
        if index < skip_until:
            continue
        if index == template.associate_slice[0]:
            _add_associate_block(output_doc, data)
            skip_until = template.associate_slice[1]
            continue
        if index == template.apport_slice[0]:
            _add_apport_block(output_doc, data)
            skip_until = template.apport_slice[1]
            continue
        if index == template.capital_slice[0]:
            _add_capital_block(output_doc, data)
            skip_until = template.capital_slice[1]
            continue
        if template.signature_slice is not None and index == template.signature_slice[0]:
            _add_signature_block(output_doc, data)
            skip_until = template.signature_slice[1]
            continue

        text = paragraph.text.strip()
        if not text:
            continue
        rendered = _replace_placeholders(text, replacements)
        _add_rendered_paragraph(output_doc, rendered)
        if template.expected_type == "sci_iris" and index == 561:
            _add_resultat_groupes_block(output_doc, data)
        if (
            template.append_signatures_after is not None
            and index == template.append_signatures_after
        ):
            _add_signature_block(output_doc, data)

    full_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
    if "[" in full_text or "]" in full_text:
        raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / template.output_filename
    output_doc.save(output_path)
    return output_path


class _ResolvedStatutsCivil:
    def __init__(
        self,
        *,
        template: StatutsCivilTemplate,
        statuts: StatutsCivilsContext,
        denomination: str,
        forme_sociale: str,
        adresse_siege: str,
        siege_num_voie: str,
        siege_voie: str,
        siege_cp: str,
        siege_ville: str,
        ville_rcs: str,
        signature_lieu: str,
        signature_date: str,
        associes: list[StatutsCivilsAssocie],
    ) -> None:
        self.template = template
        self.statuts = statuts
        self.denomination = denomination
        self.forme_sociale = forme_sociale
        self.adresse_siege = adresse_siege
        self.siege_num_voie = siege_num_voie
        self.siege_voie = siege_voie
        self.siege_cp = siege_cp
        self.siege_ville = siege_ville
        self.ville_rcs = ville_rcs
        self.signature_lieu = signature_lieu
        self.signature_date = signature_date
        self.associes = associes

    @classmethod
    def from_context(
        cls,
        ctx: DocumentGenerationContext,
        template: StatutsCivilTemplate,
    ) -> _ResolvedStatutsCivil:
        if ctx.structure != template.expected_structure:
            raise ValueError(
                f"dossier.structure doit etre {template.expected_structure} pour {DOCUMENT_CODE}."
            )
        if ctx.statuts_civils is None:
            raise ValueError(f"statuts_civils est obligatoire pour {DOCUMENT_CODE}.")
        statuts_type = _required_text(ctx.statuts_civils.type, "statuts_civils.type").lower()
        if statuts_type != template.expected_type:
            raise ValueError(
                f"statuts_civils.type doit etre {template.expected_type} pour {DOCUMENT_CODE}."
            )
        if ctx.societe is None:
            raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
        if ctx.societe.siege is None:
            raise ValueError(f"societe.siege est obligatoire pour {DOCUMENT_CODE}.")
        associes = list(ctx.statuts_civils.associes)
        _validate_associes(associes, template)
        _validate_capital_totals(ctx.statuts_civils, associes)
        if template.expected_type == "scs":
            _validate_scs(ctx.statuts_civils, associes)
        if template.expected_type == "sci":
            _validate_sci(associes)
        if template.expected_type == "sci_iris":
            _validate_sci_iris(ctx.statuts_civils, associes)
        _validate_template_fields(ctx.statuts_civils, template)

        return cls(
            template=template,
            statuts=ctx.statuts_civils,
            denomination=_required_text(ctx.societe.denomination, "societe.denomination"),
            forme_sociale=_required_text(
                ctx.statuts_civils.forme_sociale or ctx.societe.forme_sociale,
                "statuts_civils.forme_sociale",
            ),
            adresse_siege=_address_display(ctx.societe.siege, "societe.siege"),
            siege_num_voie=_required_text(ctx.societe.siege.num_voie, "societe.siege.num_voie"),
            siege_voie=_required_text(ctx.societe.siege.voie, "societe.siege.voie"),
            siege_cp=_required_text(ctx.societe.siege.cp, "societe.siege.cp"),
            siege_ville=_required_text(ctx.societe.siege.ville, "societe.siege.ville"),
            ville_rcs=_required_text(ctx.societe.ville_rcs, "societe.ville_rcs"),
            signature_lieu=ctx.signature.lieu,
            signature_date=_format_display_date(ctx.signature.date, "signature.date"),
            associes=associes,
        )

    def common_replacements(self) -> dict[str, str]:
        statuts = self.statuts
        depot = statuts.capital_depot
        return {
            "[denomination_societe]": self.denomination,
            "[forme_sociale]": self.forme_sociale,
            "[mention_capital_variable]": _text_or_empty(statuts.mention_capital_variable),
            "[capital_social]": _required_text(
                statuts.capital_social,
                "statuts_civils.capital_social",
            ),
            "[capital_lettres]": _required_text(
                statuts.capital_social_lettres,
                "statuts_civils.capital_social_lettres",
            ),
            "[capital_autorise]": _text_or_empty(statuts.capital_autorise),
            "[capital_autorise_lettres]": _text_or_empty(statuts.capital_autorise_lettres),
            "[capital_social_maximal]": _text_or_empty(statuts.capital_maximal),
            "[capital_social_maximal_lettres]": _text_or_empty(statuts.capital_maximal_lettres),
            "[nb_parts]": str(
                _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
            ),
            "[nb_parts_total]": str(
                _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
            ),
            "[nb_parts_lettres]": _required_text(
                statuts.nb_parts_total_lettres,
                "statuts_civils.nb_parts_total_lettres",
            ),
            "[nb_parts_total_lettres]": _required_text(
                statuts.nb_parts_total_lettres,
                "statuts_civils.nb_parts_total_lettres",
            ),
            "[valeur_nominale_part]": _required_text(
                statuts.valeur_nominale_part,
                "statuts_civils.valeur_nominale_part",
            ),
            "[valeur_nominale_part_lettres]": _text_or_empty(statuts.valeur_nominale_part_lettres),
            "[plage_parts_total]": _text_or_empty(statuts.plage_parts_totale),
            "[parts_debut]": str(_first_part_number(self.associes)),
            "[parts_fin]": str(_last_part_number(self.associes)),
            "[adresse_siege]": self.adresse_siege,
            "[num_voie_siege]": self.siege_num_voie,
            "[voie_siege]": self.siege_voie,
            "[cp_siege]": self.siege_cp,
            "[ville_siege]": self.siege_ville,
            "[ville_rcs]": self.ville_rcs,
            "[duree_societe]": _text_or_empty(statuts.duree_societe),
            "[nom_banque]": _required_text(
                depot.banque_nom if depot else None,
                "statuts_civils.capital_depot.banque_nom",
            ),
            "[adresse_banque]": _required_text(
                depot.banque_adresse if depot else None,
                "statuts_civils.capital_depot.banque_adresse",
            ),
            "[date_cloture_exercice_1]": _required_text(
                statuts.date_cloture_premier_exercice,
                "statuts_civils.date_cloture_premier_exercice",
            ),
            "[lieu_signature]": _required_text(
                self.signature_lieu,
                "signature.lieu",
            ),
            "[date_signature]": self.signature_date,
            "[nombre_exemplaires_lettres]": _text_or_empty(statuts.nombre_exemplaires_lettres),
            "[denomination_cabinet_mandataire]": _text_or_empty(
                statuts.denomination_cabinet_mandataire
            ),
        }


def _add_associate_block(document, data: _ResolvedStatutsCivil) -> None:
    for associe in data.associes:
        if _is_morale(associe):
            _add_morale_identity(document, associe)
        else:
            _add_physical_identity(document, associe)


def _add_apport_block(document, data: _ResolvedStatutsCivil) -> None:
    if data.template.expected_type == "scs":
        add_paragraph(
            document, "Le capital social est constitue par les apports en numeraires suivants :"
        )
        add_paragraph(document, "Associes commandites :", bold=True)
        for associe in _associes_by_role(data.associes, "commandite"):
            _add_apport_line(document, associe)
        total_commandites = _required_text(
            data.statuts.total_apports_commandites,
            "statuts_civils.total_apports_commandites",
        )
        add_paragraph(
            document,
            f"Le montant total verse par le commandite est de {total_commandites}.",
        )
        add_paragraph(document, "Associes commanditaires :", bold=True)
        for associe in _associes_by_role(data.associes, "commanditaire"):
            _add_apport_line(document, associe, commanditaire=True)
    else:
        for associe in data.associes:
            _add_apport_line(document, associe)
    capital_social = _required_text(data.statuts.capital_social, "statuts_civils.capital_social")
    add_paragraph(
        document,
        f"SOIT AU TOTAL {capital_social} euros",
    )
    depot = data.statuts.capital_depot
    banque_nom = _required_text(
        depot.banque_nom if depot else None,
        "statuts_civils.capital_depot.banque_nom",
    )
    banque_adresse = _required_text(
        depot.banque_adresse if depot else None,
        "statuts_civils.capital_depot.banque_adresse",
    )
    add_paragraph(
        document,
        "Les associes declarent et reconnaissent que la somme liberee, d'un montant de "
        f"{capital_social} euros, "
        "a ete deposee integralement et avant ce jour, au credit d'un compte ouvert, "
        "au nom de la societe en formation, a la banque "
        f"{banque_nom}, {banque_adresse}.",
    )


def _add_capital_block(document, data: _ResolvedStatutsCivil) -> None:
    for associe in data.associes:
        parts = _required_parts(associe)
        add_paragraph(document, _signature_label(associe))
        if data.template.expected_type == "sci_iris":
            add_paragraph(
                document,
                "A concurrence de "
                f"{_required_text(parts.nb_lettres, 'associes[].parts.nb_lettres')} parts, "
                f"ci {parts.nb} parts Numerotees de "
                f"{_required_int(parts.debut, 'associes[].parts.debut')} a "
                f"{_required_int(parts.fin, 'associes[].parts.fin')}.",
            )
        else:
            qualite = f", {parts.qualite_associe}" if parts.qualite_associe else ""
            add_paragraph(document, f"- {_signature_label(associe)}{qualite},")
            add_paragraph(
                document,
                "Proprietaire de "
                f"{_required_text(parts.nb_lettres, 'associes[].parts.nb_lettres')} parts sociales "
                f"{parts.nb} parts sociales",
            )
            if parts.plage_affichee:
                add_paragraph(document, f"Numerotees de {parts.plage_affichee}")
    add_paragraph(
        document,
        "SOIT AU TOTAL "
        f"{_required_int(data.statuts.nb_parts_total, 'statuts_civils.nb_parts_total')} parts",
    )


def _add_signature_block(document, data: _ResolvedStatutsCivil) -> None:
    if data.template.signature_slice is not None:
        add_paragraph(document, f"A {data.signature_lieu}, le {data.signature_date}")
    signers = [_signature_label(a) for a in data.associes if a.est_signataire]
    if data.template.expected_type == "scs":
        add_statuts_signature_grid(document, signers, mention="Lu et approuve")
        return
    for signer in signers:
        add_statuts_signature_block(
            document,
            [signer],
            bold=True,
            underline=True,
        )


def _add_resultat_groupes_block(document, data: _ResolvedStatutsCivil) -> None:
    rows = []
    for group in data.statuts.resultat_groupes_parts:
        parts_debut = _required_int(
            group.parts_debut,
            "statuts_civils.resultat_groupes_parts[].parts_debut",
        )
        parts_fin = _required_int(
            group.parts_fin,
            "statuts_civils.resultat_groupes_parts[].parts_fin",
        )
        quote_part = _required_text(
            group.quote_part_resultat_exceptionnel,
            "statuts_civils.resultat_groupes_parts[].quote_part_resultat_exceptionnel",
        )
        rows.append((f"Parts {parts_debut} a {parts_fin}", quote_part))
    if data.statuts.resultat_quote_part_exceptionnel_total:
        rows.append(("Total", data.statuts.resultat_quote_part_exceptionnel_total))
    add_statuts_matrix_table(
        document,
        ("Groupe de parts", "Quote-part de résultat exceptionnel"),
        rows,
    )


def _add_apport_line(
    document,
    associe: StatutsCivilsAssocie,
    *,
    commanditaire: bool = False,
) -> None:
    apport = associe.apport
    if apport is None:
        raise ValueError(f"associes[].apport est obligatoire pour {DOCUMENT_CODE}.")
    montant = (apport.montant_commanditaire or apport.montant) if commanditaire else apport.montant
    montant_lettres = (
        (apport.montant_commanditaire_lettres or apport.montant_lettres)
        if commanditaire
        else apport.montant_lettres
    )
    add_paragraph(document, f"- {_signature_label(associe)} apporte,")
    add_paragraph(
        document,
        f"la somme de {_required_text(montant_lettres, 'associes[].apport.montant_lettres')}, "
        f"{_required_text(montant, 'associes[].apport.montant')}",
    )


def _add_physical_identity(document, associe: StatutsCivilsAssocie) -> None:
    gender = associe.genre or Gender.MASCULIN
    born = "Nee" if gender == Gender.FEMININ else "Ne"
    add_paragraph(document, _signature_label(associe))
    add_paragraph(
        document,
        f"{born} le {_format_display_date(associe.date_naissance, 'associes[].date_naissance')} "
        f"a {_required_text(associe.ville_naissance, 'associes[].ville_naissance')} "
        f"({_required_text(associe.departement_naissance, 'associes[].departement_naissance')})",
    )
    add_paragraph(
        document,
        f"De nationalite {_required_text(associe.nationalite, 'associes[].nationalite')}",
    )
    add_paragraph(
        document,
        _required_text(associe.situation_maritale, "associes[].situation_maritale"),
    )
    add_paragraph(document, f"Demeurant {_person_address(associe)}")


def _add_morale_identity(document, associe: StatutsCivilsAssocie) -> None:
    add_paragraph(document, _signature_label(associe))
    add_paragraph(
        document,
        f"{_required_text(associe.forme_juridique, 'associes[].forme_juridique')} "
        f"au capital de {_required_text(associe.capital_social, 'associes[].capital_social')}, "
        f"ayant son siege {_address_display(associe.siege, 'associes[].siege')}, "
        f"immatriculee au RCS de {_required_text(associe.ville_rcs, 'associes[].ville_rcs')} "
        f"sous le numero {_required_text(associe.numero_rcs, 'associes[].numero_rcs')}.",
    )
    if associe.representant is None:
        raise ValueError(
            f"associes[].representant est obligatoire pour une personne morale {DOCUMENT_CODE}."
        )
    add_paragraph(
        document,
        "Representee par "
        f"{_required_text(associe.representant.civilite_affichage, 'representant.civilite')} "
        f"{_required_text(associe.representant.prenom, 'associes[].representant.prenom')} "
        f"{_required_text(associe.representant.nom, 'associes[].representant.nom')}, "
        f"{_required_text(associe.representant.fonction, 'associes[].representant.fonction')}.",
    )


def _validate_associes(
    associes: list[StatutsCivilsAssocie],
    template: StatutsCivilTemplate,
) -> None:
    if not associes:
        raise ValueError(f"au moins un associe est obligatoire pour {DOCUMENT_CODE}.")
    if len(associes) > MAX_ASSOCIES:
        raise ValueError(f"les statuts civils sont limites a 6 associes pour {DOCUMENT_CODE}.")
    for associe in associes:
        _required_parts(associe)
        if _is_morale(associe) and template.expected_type == "sci":
            raise ValueError(
                "les associes personnes morales SCI sont hors source observee V1 "
                f"pour {DOCUMENT_CODE}."
            )


def _validate_capital_totals(
    statuts: StatutsCivilsContext,
    associes: list[StatutsCivilsAssocie],
) -> None:
    total_parts = sum(_required_int(_required_parts(a).nb, "associes[].parts.nb") for a in associes)
    expected_parts = _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
    if total_parts != expected_parts:
        raise ValueError(
            "la somme des parts doit correspondre a statuts_civils.nb_parts_total "
            f"pour {DOCUMENT_CODE}."
        )
    total_apports = sum(_amount_to_int(_required_apport(a).montant) for a in associes)
    expected_capital = _amount_to_int(statuts.capital_social)
    if total_apports != expected_capital:
        raise ValueError(
            "la somme des apports doit correspondre a statuts_civils.capital_social "
            f"pour {DOCUMENT_CODE}."
        )


def _validate_scs(statuts: StatutsCivilsContext, associes: list[StatutsCivilsAssocie]) -> None:
    if not _associes_by_role(associes, "commandite"):
        raise ValueError(f"au moins un associe commandite est obligatoire pour {DOCUMENT_CODE}.")
    if not _associes_by_role(associes, "commanditaire"):
        raise ValueError(f"au moins un associe commanditaire est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(statuts.total_apports_commandites, "statuts_civils.total_apports_commandites")
    _required_text(statuts.capital_maximal, "statuts_civils.capital_maximal")
    _required_text(statuts.capital_maximal_lettres, "statuts_civils.capital_maximal_lettres")


def _validate_sci(associes: list[StatutsCivilsAssocie]) -> None:
    for associe in associes:
        if _is_morale(associe):
            raise ValueError(f"les personnes morales SCI sont bloquees en V1 pour {DOCUMENT_CODE}.")


def _validate_sci_iris(
    statuts: StatutsCivilsContext,
    associes: list[StatutsCivilsAssocie],
) -> None:
    if not any(_is_morale(associe) for associe in associes):
        raise ValueError(
            f"SCI IRIS requiert l'associe personne morale source en V1 pour {DOCUMENT_CODE}."
        )
    if not statuts.resultat_groupes_parts:
        raise ValueError(
            f"statuts_civils.resultat_groupes_parts est obligatoire pour {DOCUMENT_CODE}."
        )
    for group in statuts.resultat_groupes_parts:
        _required_int(group.parts_debut, "statuts_civils.resultat_groupes_parts[].parts_debut")
        _required_int(group.parts_fin, "statuts_civils.resultat_groupes_parts[].parts_fin")
        _required_text(
            group.quote_part_resultat_exceptionnel,
            "statuts_civils.resultat_groupes_parts[].quote_part_resultat_exceptionnel",
        )


def _validate_template_fields(
    statuts: StatutsCivilsContext,
    template: StatutsCivilTemplate,
) -> None:
    _required_text(statuts.capital_social, "statuts_civils.capital_social")
    _required_text(statuts.capital_social_lettres, "statuts_civils.capital_social_lettres")
    _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
    _required_text(statuts.nb_parts_total_lettres, "statuts_civils.nb_parts_total_lettres")
    _required_text(statuts.valeur_nominale_part, "statuts_civils.valeur_nominale_part")
    _required_text(
        statuts.date_cloture_premier_exercice, "statuts_civils.date_cloture_premier_exercice"
    )
    if template.expected_type in {"sci", "sci_iris"}:
        _required_text(statuts.mention_capital_variable, "statuts_civils.mention_capital_variable")
        _required_text(statuts.capital_autorise, "statuts_civils.capital_autorise")
        _required_text(statuts.capital_autorise_lettres, "statuts_civils.capital_autorise_lettres")
    if template.expected_type == "scs":
        _required_text(
            statuts.valeur_nominale_part_lettres, "statuts_civils.valeur_nominale_part_lettres"
        )
        _required_text(statuts.plage_parts_totale, "statuts_civils.plage_parts_totale")
        _required_text(statuts.duree_societe, "statuts_civils.duree_societe")
        _required_text(
            statuts.nombre_exemplaires_lettres, "statuts_civils.nombre_exemplaires_lettres"
        )
        _required_text(
            statuts.denomination_cabinet_mandataire,
            "statuts_civils.denomination_cabinet_mandataire",
        )


def _source_path(template: StatutsCivilTemplate) -> Path:
    lot_dir = Path("project/source_documents/lot_04")
    candidates = []
    for path in lot_dir.glob("*.docx"):
        name = path.name.upper()
        if all(token.upper() in name for token in template.source_name_contains):
            if template.expected_type == "sci" and "IRIS" in name:
                continue
            candidates.append(path)
    if len(candidates) != 1:
        raise ValueError(f"source DOCX introuvable ou ambigue pour {DOCUMENT_CODE}: {candidates}")
    return candidates[0]


def _add_rendered_paragraph(document, text: str) -> None:
    if text == "STATUTS":
        add_statuts_title_box(document, text)
    elif text.startswith("TITRE "):
        add_statuts_part_heading(document, text)
    elif text.startswith("ARTICLE ") or text.startswith("Article "):
        add_statuts_article_heading(document, text, left_indent_cm=0.25)
    elif text.startswith("- "):
        add_statuts_hanging_list_item(document, text[2:])
    else:
        add_statuts_body_paragraph(document, text)


def _replace_placeholders(text: str, replacements: dict[str, str]) -> str:
    rendered = text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not str(value).strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return str(value).strip()


def _text_or_empty(value: str | None) -> str:
    if value is None:
        return ""
    return str(value).strip()


def _required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _required_parts(associe: StatutsCivilsAssocie):
    if associe.parts is None:
        raise ValueError(f"associes[].parts est obligatoire pour {DOCUMENT_CODE}.")
    _required_int(associe.parts.nb, "associes[].parts.nb")
    _required_text(associe.parts.nb_lettres, "associes[].parts.nb_lettres")
    return associe.parts


def _required_apport(associe: StatutsCivilsAssocie):
    if associe.apport is None:
        raise ValueError(f"associes[].apport est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(associe.apport.montant, "associes[].apport.montant")
    _required_text(associe.apport.montant_lettres, "associes[].apport.montant_lettres")
    return associe.apport


def _format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return _required_text(value, field_name)


def _address_display(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{_required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{_required_text(address.voie, f'{field_name}.voie')} - "
        f"{_required_text(address.cp, f'{field_name}.cp')} "
        f"{_required_text(address.ville, f'{field_name}.ville')}"
    )


def _person_address(associe: StatutsCivilsAssocie) -> str:
    if associe.adresse_personnelle_affichee:
        return associe.adresse_personnelle_affichee.strip()
    return _address_display(associe.adresse_personnelle, "associes[].adresse_personnelle")


def _signature_label(associe: StatutsCivilsAssocie) -> str:
    if _is_morale(associe):
        denomination = _required_text(associe.denomination, "associes[].denomination")
        if associe.representant is None:
            return denomination
        return (
            f"{denomination}, representee par "
            f"{_required_text(associe.representant.civilite_affichage, 'representant.civilite')} "
            f"{_required_text(associe.representant.prenom, 'associes[].representant.prenom')} "
            f"{_required_text(associe.representant.nom, 'associes[].representant.nom')}"
        )
    prenoms = associe.prenoms or associe.prenom
    return (
        f"{_required_text(associe.civilite_affichage, 'associes[].civilite_affichage')} "
        f"{_required_text(prenoms, 'associes[].prenoms')} "
        f"{_required_text(associe.nom, 'associes[].nom')}"
    )


def _is_morale(associe: StatutsCivilsAssocie) -> bool:
    return associe.type_personne == "personne_morale"


def _associes_by_role(
    associes: list[StatutsCivilsAssocie],
    role: str,
) -> list[StatutsCivilsAssocie]:
    return [a for a in associes if (a.role_statutaire or "").lower() == role]


def _first_part_number(associes: list[StatutsCivilsAssocie]) -> int:
    values = [a.parts.debut for a in associes if a.parts and a.parts.debut is not None]
    return min(values) if values else 1


def _last_part_number(associes: list[StatutsCivilsAssocie]) -> int:
    values = [a.parts.fin for a in associes if a.parts and a.parts.fin is not None]
    if values:
        return max(values)
    return sum(_required_int(_required_parts(a).nb, "associes[].parts.nb") for a in associes)


def _amount_to_int(value: str | None) -> int:
    text = _required_text(value, "montant")
    normalized = (
        text.replace(" ", "")
        .replace("\u00a0", "")
        .replace("euros", "")
        .replace("euro", "")
        .replace("EUR", "")
        .replace("€", "")
        .strip()
    )
    if not normalized.isdigit():
        raise ValueError(f"montant numerique attendu pour {DOCUMENT_CODE}: {value}")
    return int(normalized)
