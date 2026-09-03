from __future__ import annotations

from datetime import date
from pathlib import Path
from unicodedata import normalize

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Associe,
    Company,
    DocumentGenerationContext,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_paragraph,
    add_statuts_annex_heading,
    add_statuts_article_heading,
    add_statuts_body_paragraph,
    add_statuts_hanging_list_item,
    add_statuts_signature_block,
    add_statuts_title_box,
    new_document,
)
from sydel_doc_engine.utils.grammar import apply_gender_pairs

DOCUMENT_CODE = "CODE-STATUTS-SEL-001"
STRUCTURE_SELARL = "SELARL"
STRUCTURE_SELAS = "SELAS"
OVERLAY_SELARL_DENTISTE = "selarl_dentiste"
OVERLAY_SELARL_MEDECIN = "selarl_medecin"
OVERLAY_SELAS_MEDECIN = "selas_medecin"


def required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return required_text(value, field_name)


def validate_sel_context(
    ctx: DocumentGenerationContext,
    *,
    expected_structure: str,
    expected_overlay: str,
) -> None:
    if ctx.structure != expected_structure:
        raise ValueError(f"dossier.structure doit etre {expected_structure} pour {DOCUMENT_CODE}.")
    if ctx.statuts_sel is not None and ctx.statuts_sel.overlay != expected_overlay:
        raise ValueError(f"statuts_sel.overlay doit etre {expected_overlay} pour {DOCUMENT_CODE}.")
    if len(ctx.associes) != 1:
        raise ValueError(
            f"les statuts SEL multi-associes sont bloques en V1 pour {DOCUMENT_CODE}."
        )
    if ctx.capital_souscription and len(ctx.capital_souscription.souscripteurs) > 1:
        raise ValueError(
            f"les statuts SEL multi-associes sont bloques en V1 pour {DOCUMENT_CODE}."
        )
    if ctx.dirigeant_nomine is not None and not _dirigeant_is_unique_associe(ctx):
        raise ValueError(
            "la signature du dirigeant non associe reste manuelle en V1 "
            f"pour {DOCUMENT_CODE}."
        )


def validate_selas_second_lieu(ctx: DocumentGenerationContext) -> bool:
    if ctx.exercice_social is None or len(ctx.exercice_social.lieux) < 2:
        return False
    second_lieu = ctx.exercice_social.lieux[1]
    has_name = bool(second_lieu.nom and second_lieu.nom.strip())
    has_address = bool(second_lieu.adresse_affichee and second_lieu.adresse_affichee.strip())
    if has_name != has_address:
        raise ValueError(
            "exercice_social.lieux[1].nom et adresse_affichee doivent etre fournis "
            f"ensemble pour {DOCUMENT_CODE}."
        )
    return has_name and has_address


def common_replacements(
    ctx: DocumentGenerationContext,
    *,
    title_type: str,
) -> dict[str, str]:
    company = required_company(ctx)
    associate = required_associe_unique(ctx)
    replacements = {
        "[denomination_societe]": required_text(company.denomination, "societe.denomination"),
        "[adresse_siege]": address_display(company.siege, "societe.siege"),
        "[capital_social]": capital_amount(ctx, company),
        "[capital_lettres]": capital_amount_letters(ctx, company),
        "[forme_sociale_complete]": required_text(
            company.forme_sociale_complete or company.forme_sociale_libelle_long,
            "societe.forme_sociale_complete",
        ),
        "[civilite]": required_text(
            associate.civilite_affichage,
            "associes[0].civilite_affichage",
        ),
        "[prenom]": required_text(associate.prenom, "associes[0].prenom"),
        "[nom]": required_text(associate.nom, "associes[0].nom"),
        "[PRENOM]": required_text(associate.prenom, "associes[0].prenom"),
        "[NOM]": required_text(associate.nom, "associes[0].nom"),
        "[profession]": required_text(associate.profession, "associes[0].profession"),
        "[profession_reglementee]": required_text(
            associate.profession_reglementee,
            "associes[0].profession_reglementee",
        ),
        "[profession_reglementee_pluriel]": required_text(
            associate.profession_reglementee_pluriel,
            "associes[0].profession_reglementee_pluriel",
        ),
        "[date_naissance]": format_display_date(
            associate.date_naissance,
            "associes[0].date_naissance",
        ),
        "[ville_naissance]": required_text(
            associate.ville_naissance,
            "associes[0].ville_naissance",
        ),
        "[departement_naissance]": required_text(
            associate.departement_naissance,
            "associes[0].departement_naissance",
        ),
        "[nationalite]": required_text(associate.nationalite, "associes[0].nationalite"),
        "[adresse_personnelle]": person_address_display(associate),
        "[situation_maritale]": marital_status_display(associate),
        "[situation_matrimoniale_statuts]": statuts_sel_matrimonial_clause(associate),
        "[regime_matrimonial]": matrimonial_regime_display(associate),
        "[qualite_associe_article_8]": article_8_associate_label(ctx, associate),
        "[nb_parts_total]": str(capital_titles_total(ctx)),
        "[nb_parts_total_lettres]": capital_titles_total_letters(ctx),
        "[nb_actions]": str(capital_titles_total(ctx)),
        "[valeur_nominale_part]": capital_title_value(ctx, title_type),
        "[valeur_nominale_action]": capital_title_value(ctx, title_type),
        "[montant_apport]": apport_amount(ctx, associate),
        "[montant_apport_lettres]": apport_amount_letters(ctx, associate),
        "[apport_personne_1]": apport_amount(ctx, associate),
        "[apport_lettres_personne_1]": apport_amount_letters(ctx, associate),
        "[lieu_signature]": required_text(ctx.signature.lieu, "signature.lieu"),
        "[date_signature]": format_display_date(ctx.signature.date, "signature.date"),
    }
    _validate_unique_associate_capital(ctx, associate)
    return replacements


def add_conjoint_replacements(
    replacements: dict[str, str],
    associate: Associe,
) -> None:
    if associate.conjoint is None:
        raise ValueError(f"associes[0].conjoint est obligatoire pour {DOCUMENT_CODE}.")
    replacements.update(
        {
            "[civilite_conjoint]": required_text(
                associate.conjoint.civilite_affichage,
                "associes[0].conjoint.civilite_affichage",
            ),
            "[prenom_conjoint]": required_text(
                associate.conjoint.prenom,
                "associes[0].conjoint.prenom",
            ),
            "[nom_conjoint]": required_text(
                associate.conjoint.nom,
                "associes[0].conjoint.nom",
            ),
        }
    )


def marital_status_display(associate: Associe) -> str:
    value = required_text(
        associate.situation_maritale,
        "associes[0].situation_maritale",
    )
    normalized = _normalized_text(value)
    if normalized in {"marie", "mariee"}:
        return "mariée" if associate.genre == Gender.FEMININ else "marié"
    return value


def matrimonial_regime_display(associate: Associe) -> str:
    value = required_text(
        associate.regime_matrimonial,
        "associes[0].regime_matrimonial",
    )
    normalized = _normalized_text(value)
    if "communaute" in normalized and "legale" in normalized:
        return "la communauté légale"
    if "communaute" in normalized:
        return "la communauté"
    for prefix in ("sous le régime de ", "sous le regime de ", "régime de ", "regime de "):
        if value.lower().startswith(prefix):
            return value[len(prefix) :].strip()
    return value


def statuts_sel_matrimonial_clause(associate: Associe) -> str:
    status = marital_status_display(associate)
    normalized_status = _normalized_text(status)
    if normalized_status not in {"marie", "mariee"}:
        return status
    conjoint = associate.conjoint
    if conjoint is None:
        raise ValueError(f"associes[0].conjoint est obligatoire pour {DOCUMENT_CODE}.")
    conjoint_label = " ".join(
        (
            required_text(
                conjoint.civilite_affichage,
                "associes[0].conjoint.civilite_affichage",
            ),
            required_text(conjoint.prenom, "associes[0].conjoint.prenom"),
            required_text(conjoint.nom, "associes[0].conjoint.nom"),
        )
    )
    return (
        f"{status} sous le régime de {statuts_sel_matrimonial_regime(associate)} "
        f"avec {conjoint_label}"
    )


def statuts_sel_matrimonial_regime(associate: Associe) -> str:
    value = required_text(
        associate.regime_matrimonial,
        "associes[0].regime_matrimonial",
    )
    normalized = _normalized_text(value)
    if "separation" in normalized and "bien" in normalized:
        return "la séparation de biens"
    if "communaute" in normalized:
        return "la communauté"
    return matrimonial_regime_display(associate)


def article_8_associate_label(
    ctx: DocumentGenerationContext,
    associate: Associe,
) -> str:
    if len(ctx.associes) == 1:
        return "associée unique" if associate.genre == Gender.FEMININ else "associé unique"
    if all(other.genre == Gender.FEMININ for other in ctx.associes):
        return "associées"
    return "associés"


def _normalized_text(value: str) -> str:
    ascii_text = normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    return " ".join(ascii_text.lower().split())


def add_ordre_replacements(
    replacements: dict[str, str],
    associate: Associe,
) -> None:
    if associate.ordre is None:
        raise ValueError(f"associes[0].ordre est obligatoire pour {DOCUMENT_CODE}.")
    replacements.update(
        {
            "[numero_ordre]": required_text(
                associate.ordre.numero,
                "associes[0].ordre.numero",
            ),
            "[numero_rpps]": required_text(
                associate.ordre.numero_rpps,
                "associes[0].ordre.numero_rpps",
            ),
            "[ordre_departemental]": required_text(
                associate.ordre.departement,
                "associes[0].ordre.departement",
            ),
            "[ville_ordre]": required_text(
                associate.ordre.ville or associate.ordre.departement,
                "associes[0].ordre.ville",
            ),
            "[ordre_professionnel]": required_text(
                associate.ordre.professionnel,
                "associes[0].ordre.professionnel",
            ),
        }
    )


def add_depot_replacements(
    replacements: dict[str, str],
    ctx: DocumentGenerationContext,
    *,
    require_address: bool,
) -> None:
    if ctx.depot_fonds is None or ctx.depot_fonds.banque is None:
        raise ValueError(f"depot_fonds.banque est obligatoire pour {DOCUMENT_CODE}.")
    replacements["[nom_banque]"] = required_text(
        ctx.depot_fonds.banque.nom,
        "depot_fonds.banque.nom",
    )
    if require_address:
        replacements["[adresse_banque]"] = required_text(
            ctx.depot_fonds.banque.adresse_affichee,
            "depot_fonds.banque.adresse_affichee",
        )


def add_exercice_replacements(
    replacements: dict[str, str],
    ctx: DocumentGenerationContext,
    *,
    require_debut_fin: bool,
    require_lieu: bool,
) -> None:
    if ctx.exercice_social is None:
        raise ValueError(f"exercice_social est obligatoire pour {DOCUMENT_CODE}.")
    if require_debut_fin:
        replacements.update(
            {
                "[debut_exercice]": required_text(
                    ctx.exercice_social.debut,
                    "exercice_social.debut",
                ),
                "[fin_exercice]": required_text(
                    ctx.exercice_social.fin,
                    "exercice_social.fin",
                ),
            }
        )
    replacements["[date_cloture_exercice_1]"] = required_text(
        ctx.exercice_social.date_cloture_premier_exercice,
        "exercice_social.date_cloture_premier_exercice",
    )
    if require_lieu:
        replacements["[adresse_lieu_exercice]"] = first_lieu_exercice(ctx)


def render_statuts_sel_docx(
    blocks: tuple[str, ...],
    replacements: dict[str, str],
    output_path: Path,
    *,
    associate: Associe,
    skip_personne_2_line: bool = False,
    render_selas_second_lieu: bool = False,
    title_box_bordered: bool = True,
    annex_page_break: bool = False,
) -> Path:
    docx = new_document()
    signature_mode = False
    for index, block in enumerate(blocks):
        if skip_personne_2_line and "[civilite_personne_2]" in block:
            continue
        if "[nom_lieu_exercice_2]" in block and not render_selas_second_lieu:
            continue
        if index == 4:
            add_statuts_title_box(docx, "STATUTS", bordered=title_box_bordered)
        text = replace_placeholders(block, replacements)
        text = apply_gender_variants(text, associate)
        if _is_heading(text):
            if text.startswith("ANNEXE"):
                signature_mode = False
            if text.startswith("ANNEXE"):
                if annex_page_break:
                    docx.add_page_break()
                add_statuts_annex_heading(docx, text)
            else:
                add_paragraph(docx, text, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        elif text.startswith("ARTICLE "):
            add_statuts_article_heading(docx, text)
        elif text.startswith("Fait à ") or text.startswith("Fait a "):
            signature_mode = True
            add_statuts_signature_block(docx, [text])
        elif signature_mode and (
            "Faire précéder" in text
            or "Faire prÃ©cÃ©der" in text
            or text.startswith("«")
            or text.startswith("Â«")
        ):
            add_statuts_signature_block(docx, [], mention_lines=[text])
        elif signature_mode:
            add_statuts_signature_block(docx, [text])
        elif text.startswith("-") or text.startswith("-\t"):
            add_statuts_hanging_list_item(docx, text.lstrip("-\t "))
        else:
            add_statuts_body_paragraph(docx, text)

    full_text = "\n".join(paragraph.text for paragraph in docx.paragraphs)
    if "[" in full_text or "]" in full_text:
        raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(output_path)
    return output_path


def replace_placeholders(text: str, replacements: dict[str, str]) -> str:
    rendered = text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered


# Paires d'accord en genre des statuts SEL, pilotees par le genre de l'associe.
# Chaque paire = (forme_masculin, forme_feminin), chaine EXACTE telle que figee
# dans les blocs sources (cf. *_BLOCKS de statuts_sel_exercice_templates.py).
# JAMAIS de regex de terminaison : uniquement ces chaines litterales ancrees.
#  - "LE SOUSSIGNE\xa0:" -> "LA SOUSSIGNEE\xa0:" : l'entete figee au masculin doit
#    s'accorder pour une associee (insecable avant les deux-points conserve).
#  - "ne le " -> "nee le " : la date de naissance dans la ligne d'identification.
# Les variantes mojibake (nÃ©) couvrent un eventuel rendu Word mal encode.
_STATUTS_GENDER_PAIRS: list[tuple[str, str]] = [
    ("LE SOUSSIGNE\xa0:", "LA SOUSSIGNÉE\xa0:"),
    (", né le ", ", née le "),
    ("né le ", "née le "),
    (", nÃ© le ", ", nÃ©e le "),
    ("nÃ© le ", "nÃ©e le "),
]


def apply_gender_variants(text: str, associate: Associe) -> str:
    """Accorde l'entete et la ligne de naissance des statuts selon l'associe.

    Remplace l'ancienne logique unidirectionnelle (masc->fem sur « né le »
    seulement) par un accord BIDIRECTIONNEL pilote par `associate.genre` via
    `grammar.apply_gender_pairs`. Pour un homme, l'entete masculine « LE
    SOUSSIGNE » et « né le » des blocs sont laissees telles quelles ; pour une
    femme, elles deviennent « LA SOUSSIGNÉE » et « née le ».
    """
    return apply_gender_pairs(text, associate.genre, _STATUTS_GENDER_PAIRS)


def required_company(ctx: DocumentGenerationContext) -> Company:
    if ctx.societe is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.societe


def required_associe_unique(ctx: DocumentGenerationContext) -> Associe:
    if len(ctx.associes) != 1:
        raise ValueError(
            f"les statuts SEL multi-associes sont bloques en V1 pour {DOCUMENT_CODE}."
        )
    return ctx.associes[0]


def address_display(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{required_text(address.voie, f'{field_name}.voie')}, "
        f"{required_text(address.cp, f'{field_name}.cp')} "
        f"{required_text(address.ville, f'{field_name}.ville')}"
    )


def person_address_display(associate: Associe) -> str:
    if associate.adresse_personnelle_affichee:
        return associate.adresse_personnelle_affichee.strip()
    return address_display(associate.adresse_personnelle, "associes[0].adresse_personnelle")


def first_lieu_exercice(ctx: DocumentGenerationContext) -> str:
    if ctx.exercice_social is None or not ctx.exercice_social.lieux:
        raise ValueError(f"exercice_social.lieux[0] est obligatoire pour {DOCUMENT_CODE}.")
    return required_text(
        ctx.exercice_social.lieux[0].adresse_affichee,
        "exercice_social.lieux[0].adresse_affichee",
    )


def capital_amount(ctx: DocumentGenerationContext, company: Company) -> str:
    return required_text(
        company.capital_social or company.capital or (ctx.capital.montant if ctx.capital else None),
        "capital.montant",
    )


def capital_amount_letters(ctx: DocumentGenerationContext, company: Company) -> str:
    return required_text(
        company.capital_social_lettres or (ctx.capital.montant_lettres if ctx.capital else None),
        "capital.montant_lettres",
    )


def capital_titles_total(ctx: DocumentGenerationContext) -> int:
    if ctx.capital is None:
        raise ValueError(f"capital est obligatoire pour {DOCUMENT_CODE}.")
    return required_int(
        ctx.capital.nombre_titres_total or ctx.capital.nb_parts_total,
        "capital.nombre_titres_total",
    )


def capital_titles_total_letters(ctx: DocumentGenerationContext) -> str:
    if ctx.capital is None:
        raise ValueError(f"capital est obligatoire pour {DOCUMENT_CODE}.")
    return required_text(
        ctx.capital.nombre_titres_total_lettres,
        "capital.nombre_titres_total_lettres",
    )


def capital_title_value(ctx: DocumentGenerationContext, title_type: str) -> str:
    if ctx.capital is None:
        raise ValueError(f"capital est obligatoire pour {DOCUMENT_CODE}.")
    field_name = (
        "capital.valeur_nominale_part"
        if title_type == "parts_sociales"
        else "capital.valeur_nominale_titre"
    )
    return required_text(
        ctx.capital.valeur_nominale_titre or ctx.capital.valeur_nominale_part,
        field_name,
    )


def apport_amount(ctx: DocumentGenerationContext, associate: Associe) -> str:
    return required_text(
        associate.apport_numeraire or (ctx.apport.montant if ctx.apport else None),
        "associes[0].apport_numeraire",
    )


def apport_amount_letters(ctx: DocumentGenerationContext, associate: Associe) -> str:
    return required_text(
        associate.apport_numeraire_lettres
        or (ctx.apport.montant_lettres if ctx.apport else None),
        "associes[0].apport_numeraire_lettres",
    )


def _validate_unique_associate_capital(
    ctx: DocumentGenerationContext,
    associate: Associe,
) -> None:
    total = capital_titles_total(ctx)
    if associate.nb_parts != total:
        raise ValueError(
            "associes[0].nb_parts doit etre coherent avec "
            f"capital.nombre_titres_total pour {DOCUMENT_CODE}."
        )


def _dirigeant_is_unique_associe(ctx: DocumentGenerationContext) -> bool:
    if ctx.dirigeant_nomine is None or len(ctx.associes) != 1:
        return True
    associate = ctx.associes[0]
    if ctx.dirigeant_nomine.ref_associe_index == 0:
        return True
    return (
        ctx.dirigeant_nomine.prenom == associate.prenom
        and ctx.dirigeant_nomine.nom == associate.nom
    )


def _is_heading(text: str) -> bool:
    if text == "STATUTS":
        return True
    return text in {
        "DECISIONS DES ACTIONNAIRES",
        "RESULTATS SOCIAUX",
        "TRANSFORMATION DE LA SOCIETE",
        "DISSOLUTION â€“ LIQUIDATION",
        "CONTESTATIONS",
        "CONSTITUTION DE LA SOCIETE",
        "ANNEXE",
        "ANNEXE 1",
    }
