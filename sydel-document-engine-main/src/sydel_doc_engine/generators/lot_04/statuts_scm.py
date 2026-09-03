from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    DocumentGenerationContext,
    StatutsCivilsAssocie,
    StatutsCivilsContext,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_paragraph,
    add_statuts_article_heading,
    add_statuts_body_paragraph,
    add_statuts_part_heading,
    add_statuts_signature_block,
    add_statuts_title_box,
    new_document,
)

DOCUMENT_CODE = "CODE-STATUTS-SCM-001"
MAX_ASSOCIES = 6
SOURCE_PATH = Path("project/source_documents/lot_04/Statuts SCM.docx")
OUTPUT_FILENAME = "statuts_scm.docx"

ASSOCIATE_SLICE = (25, 42)
APPORT_SLICE = (81, 90)
CAPITAL_SLICE = (95, 100)
SIGNATURE_SLICE = (262, 269)


class StatutsScmGenerator:
    """Generateur from-scratch des statuts SCM V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        data = _ResolvedStatutsScm.from_context(ctx)
        source_doc = Document(SOURCE_PATH)
        output_doc = new_document()
        output_doc.sections[0].footer.paragraphs[0].text = (
            f"{data.denomination} - Statuts constitutifs"
        )

        replacements = data.replacements()
        skip_until = -1
        for index, paragraph in enumerate(source_doc.paragraphs):
            if index < skip_until:
                continue
            if index == ASSOCIATE_SLICE[0]:
                _add_associate_block(output_doc, data)
                skip_until = ASSOCIATE_SLICE[1]
                continue
            if index == APPORT_SLICE[0]:
                _add_apport_block(output_doc, data)
                skip_until = APPORT_SLICE[1]
                continue
            if index == CAPITAL_SLICE[0]:
                _add_capital_block(output_doc, data)
                skip_until = CAPITAL_SLICE[1]
                continue
            if index == SIGNATURE_SLICE[0]:
                _add_signature_block(output_doc, data)
                skip_until = SIGNATURE_SLICE[1]
                continue

            text = paragraph.text.strip()
            if not text:
                continue
            _add_rendered_paragraph(output_doc, _replace_placeholders(text, replacements))

        full_text = "\n".join(paragraph.text for paragraph in output_doc.paragraphs)
        if "[" in full_text or "]" in full_text:
            raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        output_doc.save(output_path)
        return output_path


class _ResolvedStatutsScm:
    def __init__(
        self,
        *,
        statuts: StatutsCivilsContext,
        denomination: str,
        denomination_courte: str,
        forme_sociale: str,
        siege_num_voie: str,
        siege_voie: str,
        siege_cp: str,
        siege_ville: str,
        signature_lieu: str,
        signature_date: str,
        associes: list[StatutsCivilsAssocie],
    ) -> None:
        self.statuts = statuts
        self.denomination = denomination
        self.denomination_courte = denomination_courte
        self.forme_sociale = forme_sociale
        self.siege_num_voie = siege_num_voie
        self.siege_voie = siege_voie
        self.siege_cp = siege_cp
        self.siege_ville = siege_ville
        self.signature_lieu = signature_lieu
        self.signature_date = signature_date
        self.associes = associes

    @classmethod
    def from_context(cls, ctx: DocumentGenerationContext) -> _ResolvedStatutsScm:
        if ctx.structure != "SCM":
            raise ValueError(f"dossier.structure doit etre SCM pour {DOCUMENT_CODE}.")
        if ctx.societe is None:
            raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
        if ctx.societe.siege is None:
            raise ValueError(f"societe.siege est obligatoire pour {DOCUMENT_CODE}.")
        if ctx.statuts_civils is None:
            raise ValueError(f"statuts_civils est obligatoire pour {DOCUMENT_CODE}.")
        statuts_type = _required_text(ctx.statuts_civils.type, "statuts_civils.type").lower()
        if statuts_type != "scm":
            raise ValueError(f"statuts_civils.type doit etre scm pour {DOCUMENT_CODE}.")

        associes = list(ctx.statuts_civils.associes)
        _validate_associes(associes)
        _validate_capital_totals(ctx.statuts_civils, associes)
        _validate_statuts_fields(ctx.statuts_civils)
        _validate_signatures(associes)

        return cls(
            statuts=ctx.statuts_civils,
            denomination=_required_text(ctx.societe.denomination, "societe.denomination"),
            denomination_courte=_required_text(
                ctx.societe.denomination_courte,
                "societe.denomination_courte",
            ),
            forme_sociale=_required_text(
                ctx.statuts_civils.forme_sociale or ctx.societe.forme_sociale,
                "statuts_civils.forme_sociale",
            ),
            siege_num_voie=_required_text(ctx.societe.siege.num_voie, "societe.siege.num_voie"),
            siege_voie=_required_text(ctx.societe.siege.voie, "societe.siege.voie"),
            siege_cp=_required_text(ctx.societe.siege.cp, "societe.siege.cp"),
            siege_ville=_required_text(ctx.societe.siege.ville, "societe.siege.ville"),
            signature_lieu=ctx.signature.lieu,
            signature_date=_format_display_date(ctx.signature.date, "signature.date"),
            associes=associes,
        )

    def replacements(self) -> dict[str, str]:
        depot = self.statuts.capital_depot
        return {
            "[denomination_societe]": self.denomination,
            "[denomination_societe_courte]": self.denomination_courte,
            "[forme_sociale]": self.forme_sociale,
            "[capital_social]": _required_text(
                self.statuts.capital_social,
                "statuts_civils.capital_social",
            ),
            "[capital_lettres]": _required_text(
                self.statuts.capital_social_lettres,
                "statuts_civils.capital_social_lettres",
            ),
            "[nb_parts]": str(
                _required_int(self.statuts.nb_parts_total, "statuts_civils.nb_parts_total")
            ),
            "[valeur_nominale_part]": _required_text(
                self.statuts.valeur_nominale_part,
                "statuts_civils.valeur_nominale_part",
            ),
            "[num_voie_siege]": self.siege_num_voie,
            "[voie_siege]": self.siege_voie,
            "[cp_siege]": self.siege_cp,
            "[ville_siege]": self.siege_ville,
            "[nom_banque]": _required_text(
                depot.banque_nom if depot else None,
                "statuts_civils.capital_depot.banque_nom",
            ),
            "[adresse_banque]": _required_text(
                depot.banque_adresse if depot else None,
                "statuts_civils.capital_depot.banque_adresse",
            ),
            "[lieu_signature]": self.signature_lieu,
        }


def _add_associate_block(document, data: _ResolvedStatutsScm) -> None:
    for index, associe in enumerate(data.associes):
        if index:
            add_paragraph(document, "ET", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        if _is_morale(associe):
            _add_morale_identity(document, associe)
        else:
            _add_physical_identity(document, associe)


def _add_apport_block(document, data: _ResolvedStatutsScm) -> None:
    for associe in data.associes:
        apport = _required_apport(associe)
        add_paragraph(
            document,
            f"{_signature_label(associe)} apporte à la Société la somme de "
            f"{_required_text(apport.montant_lettres, 'associes[].apport.montant_lettres')}",
        )
        add_paragraph(
            document,
            f"ci- {_required_text(apport.montant, 'associes[].apport.montant')}.",
        )
    capital_lettres = _required_text(
        data.statuts.capital_social_lettres,
        "statuts_civils.capital_social_lettres",
    )
    capital_social = _required_text(
        data.statuts.capital_social,
        "statuts_civils.capital_social",
    )
    add_paragraph(
        document,
        f"Total des apports {capital_lettres} ({capital_social})",
    )
    depot = data.statuts.capital_depot
    banque_nom = _required_text(
        depot.banque_nom if depot else None,
        "statuts_civils.capital_depot.banque_nom",
    )
    banque_adresse = _required_text(
        depot.banque_adresse if depot else None,
        "statuts_civils.capital_depot.banque_adresse",
    )
    add_paragraph(
        document,
        "Cette somme a été intégralement déposée par les associés conformément à la loi, "
        "au crédit d’un compte ouvert au nom de la société en formation dans les livres "
        f"de la banque {banque_nom} {banque_adresse}.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_capital_block(document, data: _ResolvedStatutsScm) -> None:
    for associe in data.associes:
        parts = _required_parts(associe)
        add_paragraph(document, f"{_signature_label(associe)} {parts.nb} parts")
    add_paragraph(
        document,
        "Total du nombre de parts composant le capital social : "
        f"{_required_int(data.statuts.nb_parts_total, 'statuts_civils.nb_parts_total')} parts",
    )


def _add_signature_block(document, data: _ResolvedStatutsScm) -> None:
    add_paragraph(document, f"Fait à {data.signature_lieu},")
    add_paragraph(document, f"Le {data.signature_date}")
    for associe in [a for a in data.associes if a.est_signataire]:
        add_statuts_signature_block(
            document,
            [_signature_label(associe)],
            mention_lines=["« Lu et approuvé »"],
        )
        add_paragraph(document, "", alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after_pt=18)


def _add_morale_identity(document, associe: StatutsCivilsAssocie) -> None:
    capital_social = _required_text(associe.capital_social, "associes[].capital_social")
    add_paragraph(document, _required_text(associe.denomination, "associes[].denomination"))
    add_paragraph(
        document,
        f"{_required_text(associe.forme_juridique, 'associes[].forme_juridique')} de "
        f"{_required_text(associe.profession, 'associes[].profession')}",
    )
    add_paragraph(
        document,
        f"Au capital social de {capital_social}",
    )
    add_paragraph(
        document,
        f"Siège social : {_address_display(associe.siege, 'associes[].siege')}",
    )
    add_paragraph(
        document,
        f"{_required_text(associe.numero_rcs, 'associes[].numero_rcs')} R.C.S. de "
        f"{_required_text(associe.ville_rcs, 'associes[].ville_rcs')}",
    )
    representant = _required_representant(associe)
    add_paragraph(
        document,
        "Représentée par "
        f"{_required_text(representant.civilite_affichage, 'associes[].representant.civilite')} "
        f"{_required_text(representant.prenom, 'associes[].representant.prenom')} "
        f"{_required_text(representant.nom, 'associes[].representant.nom')}, "
        "en sa qualité de "
        f"{_required_text(representant.fonction, 'associes[].representant.fonction')}",
    )


def _add_physical_identity(document, associe: StatutsCivilsAssocie) -> None:
    gender = associe.genre or Gender.MASCULIN
    born = "Née" if gender == Gender.FEMININ else "Né"
    add_paragraph(document, _signature_label(associe))
    add_paragraph(
        document,
        f"{_required_text(associe.profession, 'associes[].profession')} de profession",
    )
    add_paragraph(
        document,
        f"{born} le {_format_display_date(associe.date_naissance, 'associes[].date_naissance')} "
        f"à {_required_text(associe.ville_naissance, 'associes[].ville_naissance')}",
    )
    add_paragraph(document, f"Demeurant {_person_address(associe)}")
    add_paragraph(
        document,
        f"De nationalité {_required_text(associe.nationalite, 'associes[].nationalite')}",
    )
    add_paragraph(
        document,
        _required_text(associe.situation_maritale, "associes[].situation_maritale"),
    )


def _validate_associes(associes: list[StatutsCivilsAssocie]) -> None:
    if not associes:
        raise ValueError(f"au moins un associe est obligatoire pour {DOCUMENT_CODE}.")
    if len(associes) > MAX_ASSOCIES:
        raise ValueError(f"les statuts SCM sont limites a 6 associes pour {DOCUMENT_CODE}.")
    for associe in associes:
        _required_apport(associe)
        _required_parts(associe)
        _required_text(associe.profession, "associes[].profession")
        if _is_morale(associe):
            _required_text(associe.denomination, "associes[].denomination")
            _required_text(associe.forme_juridique, "associes[].forme_juridique")
            _required_text(associe.capital_social, "associes[].capital_social")
            _address_display(associe.siege, "associes[].siege")
            _required_text(associe.numero_rcs, "associes[].numero_rcs")
            _required_text(associe.ville_rcs, "associes[].ville_rcs")
            _required_representant(associe)
        else:
            _required_text(associe.civilite_affichage, "associes[].civilite_affichage")
            _required_text(associe.prenom or associe.prenoms, "associes[].prenom")
            _required_text(associe.nom, "associes[].nom")
            _format_display_date(associe.date_naissance, "associes[].date_naissance")
            _required_text(associe.ville_naissance, "associes[].ville_naissance")
            _person_address(associe)
            _required_text(associe.nationalite, "associes[].nationalite")
            _required_text(associe.situation_maritale, "associes[].situation_maritale")


def _validate_capital_totals(
    statuts: StatutsCivilsContext,
    associes: list[StatutsCivilsAssocie],
) -> None:
    total_parts = sum(_required_int(_required_parts(a).nb, "associes[].parts.nb") for a in associes)
    expected_parts = _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
    if total_parts != expected_parts:
        raise ValueError(
            "la somme des parts doit correspondre a statuts_civils.nb_parts_total "
            f"pour {DOCUMENT_CODE}."
        )

    total_apports = sum(_amount_to_int(_required_apport(a).montant) for a in associes)
    expected_capital = _amount_to_int(statuts.capital_social)
    if total_apports != expected_capital:
        raise ValueError(
            "la somme des apports doit correspondre a statuts_civils.capital_social "
            f"pour {DOCUMENT_CODE}."
        )


def _validate_statuts_fields(statuts: StatutsCivilsContext) -> None:
    _required_text(statuts.forme_sociale, "statuts_civils.forme_sociale")
    _required_text(statuts.capital_social, "statuts_civils.capital_social")
    _required_text(statuts.capital_social_lettres, "statuts_civils.capital_social_lettres")
    _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
    _required_text(statuts.valeur_nominale_part, "statuts_civils.valeur_nominale_part")
    depot = statuts.capital_depot
    _required_text(depot.banque_nom if depot else None, "statuts_civils.capital_depot.banque_nom")
    _required_text(
        depot.banque_adresse if depot else None,
        "statuts_civils.capital_depot.banque_adresse",
    )


def _validate_signatures(associes: list[StatutsCivilsAssocie]) -> None:
    signataires = [associe for associe in associes if associe.est_signataire]
    if len(signataires) != len(associes):
        raise ValueError(
            f"les signatures SCM doivent correspondre aux associes pour {DOCUMENT_CODE}."
        )
    for associe in signataires:
        _signature_label(associe)


def _required_apport(associe: StatutsCivilsAssocie):
    if associe.apport is None:
        raise ValueError(f"associes[].apport est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(associe.apport.montant, "associes[].apport.montant")
    _required_text(associe.apport.montant_lettres, "associes[].apport.montant_lettres")
    return associe.apport


def _required_parts(associe: StatutsCivilsAssocie):
    if associe.parts is None:
        raise ValueError(f"associes[].parts est obligatoire pour {DOCUMENT_CODE}.")
    _required_int(associe.parts.nb, "associes[].parts.nb")
    return associe.parts


def _required_representant(associe: StatutsCivilsAssocie):
    if associe.representant is None:
        raise ValueError(
            f"associes[].representant est obligatoire pour une personne morale {DOCUMENT_CODE}."
        )
    _required_text(associe.representant.civilite_affichage, "associes[].representant.civilite")
    _required_text(associe.representant.prenom, "associes[].representant.prenom")
    _required_text(associe.representant.nom, "associes[].representant.nom")
    _required_text(associe.representant.fonction, "associes[].representant.fonction")
    return associe.representant


def _signature_label(associe: StatutsCivilsAssocie) -> str:
    if _is_morale(associe):
        denomination = _required_text(associe.denomination, "associes[].denomination")
        representant = _required_representant(associe)
        civilite = _required_text(
            representant.civilite_affichage,
            "associes[].representant.civilite",
        )
        return (
            f"{denomination}, représentée par "
            f"{civilite} "
            f"{_required_text(representant.prenom, 'associes[].representant.prenom')} "
            f"{_required_text(representant.nom, 'associes[].representant.nom')}"
        )
    prenoms = associe.prenoms or associe.prenom
    return (
        f"{_required_text(associe.civilite_affichage, 'associes[].civilite_affichage')} "
        f"{_required_text(prenoms, 'associes[].prenoms')} "
        f"{_required_text(associe.nom, 'associes[].nom')}"
    )


def _is_morale(associe: StatutsCivilsAssocie) -> bool:
    return associe.type_personne == "personne_morale"


def _person_address(associe: StatutsCivilsAssocie) -> str:
    if associe.adresse_personnelle_affichee:
        return associe.adresse_personnelle_affichee.strip()
    return _address_display(associe.adresse_personnelle, "associes[].adresse_personnelle")


def _address_display(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{_required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{_required_text(address.voie, f'{field_name}.voie')} "
        f"{_required_text(address.cp, f'{field_name}.cp')} "
        f"{_required_text(address.ville, f'{field_name}.ville')}"
    )


def _format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return _required_text(value, field_name)


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not str(value).strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return str(value).strip()


def _required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _amount_to_int(value: str | None) -> int:
    text = _required_text(value, "montant")
    normalized = (
        text.replace(" ", "")
        .replace("\u00a0", "")
        .replace("euros", "")
        .replace("euro", "")
        .replace("EUR", "")
        .replace("€", "")
        .strip()
    )
    if not normalized.isdigit():
        raise ValueError(f"montant numerique attendu pour {DOCUMENT_CODE}: {value}")
    return int(normalized)


def _add_rendered_paragraph(document, text: str) -> None:
    if text == "STATUTS":
        add_statuts_title_box(document, text)
    elif text.startswith("TITRE "):
        add_statuts_part_heading(document, text)
    elif text.startswith("Article "):
        add_statuts_article_heading(document, text, left_indent_cm=0.25)
    else:
        add_statuts_body_paragraph(document, text)


def _replace_placeholders(text: str, replacements: dict[str, str]) -> str:
    rendered = text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered
