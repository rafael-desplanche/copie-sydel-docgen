from __future__ import annotations

from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import Address, Company, DocumentGenerationContext
from sydel_doc_engine.rendering.docx_builder import (
    add_centered_block,
    add_framed_title,
    add_paragraph,
    add_signature_block,
    add_spacer,
    new_document,
)
from sydel_doc_engine.utils.grammar import subject_line

OUTPUT_FILENAME = "procuration.docx"

MANDATAIRE_NOM = "SYDEL"
MANDATAIRE_ADRESSE = "80 avenue Marceau, 75008 PARIS"

MANDATE_PARAGRAPH_1 = (
    "De pour moi et en mon nom faire tous dépôts, immatriculations, modifications, radiations "
    "et de recevoir le registre des bénéficiaires effectifs, concernant mon entreprise auprès "
    "des registres."
)
MANDATE_PARAGRAPH_2 = (
    "En conséquence, faire toutes déclarations et démarches, produire toutes pièces "
    "justificatives, "
    "effectuer tout dépôt de pièces, signer tous documents, requêtes et documents utiles, élire "
    "domicile, substituer en totalité ou en partie, et en général faire tout ce qui sera "
    "nécessaire."
)
MANDATE_PARAGRAPH_3 = "L’exécution de ce mandat vaudra décharge au mandataire."
LEGAL_EFFECT_PARAGRAPH = "Fait pour servir et valoir ce que de droit."


class ProcurationGenerator:
    """Générateur cible du DOC-003."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        person = ctx.personne_signataire
        company = _required_company(ctx.societe)
        personal_address = _required_address(
            person.adresse_perso,
            "personne_signataire.adresse_perso",
        )
        company_address = _required_address(company.siege, "societe.siege")

        civilite = _required_text(person.civilite, "personne_signataire.civilite")
        prenom = _required_text(person.prenom, "personne_signataire.prenom")
        nom = _required_text(person.nom, "personne_signataire.nom")
        fonction_dirigeant = _required_text(
            person.fonction_dirigeant,
            "personne_signataire.fonction_dirigeant",
        )
        forme_sociale = _required_text(company.forme_sociale, "societe.forme_sociale")
        denomination_societe = _required_text(company.denomination, "societe.denomination")
        company_designation = _company_designation(company, forme_sociale, denomination_societe)
        lieu_signature = _required_text(ctx.signature.lieu, "signature.lieu")

        document = new_document()
        _add_title(document)
        _add_paragraph(
            document,
            (
                f"{subject_line(person.genre)} {civilite} {prenom} {nom}, demeurant au "
                f"{personal_address}, agissant en qualité de {fonction_dirigeant} de la "
                f"{company_designation}, dont le siège est situé "
                f"{company_address}"
            ),
        )
        _add_paragraph(document, "Donne par les présentes pouvoir à :")
        _add_mandataire_block(document)
        for text in (MANDATE_PARAGRAPH_1, MANDATE_PARAGRAPH_2, MANDATE_PARAGRAPH_3):
            _add_paragraph(document, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(document, LEGAL_EFFECT_PARAGRAPH)
        add_spacer(document, space_after_pt=6)
        _add_final_block(
            document,
            lieu_signature=lieu_signature,
            date_signature=_format_date(ctx.signature.date),
            signatory_name=f"{prenom} {nom}",
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError("societe est obligatoire pour DOC-003.")
    return company


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour DOC-003.")
    return value.strip()


def _required_address(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour DOC-003.")
    num_voie = _required_text(address.num_voie, f"{field_name}.num_voie")
    voie = _required_text(address.voie, f"{field_name}.voie")
    ville = _required_text(address.ville, f"{field_name}.ville")
    cp = _required_text(address.cp, f"{field_name}.cp")
    return f"{num_voie} {voie}, {cp} {ville}"


def _company_designation(company: Company, forme: str, denomination: str) -> str:
    if _denomination_starts_with_form(denomination, company, forme):
        return denomination
    return f"{forme} {denomination}"


def _denomination_starts_with_form(
    denomination: str,
    company: Company,
    forme: str,
) -> bool:
    normalized_denomination = _normalize_for_prefix(denomination)
    candidates = [
        forme,
        company.forme_sociale_abregee,
        company.forme_sociale_affichage,
        company.forme_juridique,
    ]
    return any(
        normalized_denomination.startswith(_normalize_for_prefix(candidate) + " ")
        or normalized_denomination == _normalize_for_prefix(candidate)
        for candidate in candidates
        if candidate and _normalize_for_prefix(candidate)
    )


def _normalize_for_prefix(value: str) -> str:
    return " ".join(value.casefold().replace("’", "'").split())


def _format_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def _add_title(document) -> None:
    add_framed_title(document, ["Procuration"])


def _add_paragraph(
    document,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    add_paragraph(document, text, alignment=alignment)


def _add_mandataire_block(document) -> None:
    add_centered_block(
        document,
        [
            (MANDATAIRE_NOM, True, False),
            (MANDATAIRE_ADRESSE, False, True),
        ],
        space_after_pt=0,
    )


def _add_final_block(
    document,
    *,
    lieu_signature: str,
    date_signature: str,
    signatory_name: str,
) -> None:
    add_signature_block(
        document,
        [f"Fait à {lieu_signature}", f"Le {date_signature}", signatory_name],
    )
