from __future__ import annotations

from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import Address, DocumentGenerationContext
from sydel_doc_engine.rendering.docx_builder import (
    add_framed_title,
    add_legal_reminder,
    add_paragraph,
    add_signature_block,
    new_document,
)
from sydel_doc_engine.utils.grammar import birth_label, filiation_label, subject_line

OUTPUT_FILENAME = "declaration_non_condamnation.docx"

DECLARATION_TEXT = (
    "Déclare sur l’honneur, conformément aux dispositions de l’article A.123-51 du Code de "
    "commerce, n’avoir fait l’objet d’aucune condamnation pénale ni de sanction civile ou "
    "administrative de nature à m’interdire – soit d’exercer une activité commerciale – soit de "
    "gérer, d’administrer ou de diriger une personne morale."
)

RAPPEL_TITLE_SUFFIX = " : Article L123-5 du code de commerce"
RAPPEL_PARAGRAPH_1 = (
    "Le fait de donner, de mauvaise foi, des indications inexactes ou incomplètes en vue d’une "
    "immatriculation, d’une radiation ou d’une mention complémentaire ou rectificative au registre "
    "du commerce et des sociétés est puni d’une amende de 4500 euros et d’un emprisonnement de "
    "six mois."
)
RAPPEL_PARAGRAPH_2 = (
    "Les dispositions des deuxième et troisième alinéas de l’article L.123-4 sont applicables "
    "dans les cas prévus au présent article."
)


class DeclarationNonCondamnationGenerator:
    """Générateur cible du DOC-001."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        person = ctx.personne_signataire
        address = person.adresse_perso
        if address is None:
            raise ValueError("personne_signataire.adresse_perso est obligatoire pour DOC-001.")

        civilite = _required_text(person.civilite, "personne_signataire.civilite")
        prenom = _required_text(person.prenom, "personne_signataire.prenom")
        nom = _required_text(person.nom, "personne_signataire.nom")
        date_naissance = _required_date(
            person.date_naissance,
            "personne_signataire.date_naissance",
        )
        ville_naissance = _required_text(
            person.ville_naissance,
            "personne_signataire.ville_naissance",
        )
        nationalite = _required_text(person.nationalite, "personne_signataire.nationalite")
        nom_pere = _required_text(person.nom_pere, "personne_signataire.nom_pere")
        nom_mere = _required_text(person.nom_mere, "personne_signataire.nom_mere")
        lieu_signature = _required_text(ctx.signature.lieu, "signature.lieu")
        adresse_perso = _compose_required_address(address)

        document = new_document()
        _add_title(document)
        _add_identity_block(
            document,
            subject=f"{subject_line(person.genre)} {civilite} {prenom} {nom}",
            birth=f"{birth_label(person.genre)} {date_naissance} "
            f"{_birth_city_prefix(person)} {ville_naissance}.",
            address=f"demeurant au {adresse_perso}",
            nationality=f"de nationalité {nationalite}",
            filiation_father=f"{filiation_label(person.genre)} {nom_pere}",
            filiation_mother=f"et de Madame {nom_mere}",
        )
        _add_paragraph(
            document,
            DECLARATION_TEXT,
            space_before=10,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        _add_signature_block(
            document,
            lieu_signature=lieu_signature,
            date_signature=_format_date(ctx.signature.date),
            image_path=ctx.signature.image_optionnelle,
        )
        _add_legal_reminder(document)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour DOC-001.")
    return value.strip()


def _required_date(value: date | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour DOC-001.")
    return _format_date(value)


def _compose_required_address(address: Address) -> str:
    num_voie = _required_text(address.num_voie, "personne_signataire.adresse_perso.num_voie")
    voie = _required_text(address.voie, "personne_signataire.adresse_perso.voie")
    cp = _required_text(address.cp, "personne_signataire.adresse_perso.cp")
    ville = _required_text(address.ville, "personne_signataire.adresse_perso.ville")
    return f"{num_voie} {voie}, {cp} {ville}"


def _birth_city_prefix(person) -> str:
    return "au" if person.ville_naissance_article_au else "\u00e0"


def _format_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def _add_title(document) -> None:
    add_framed_title(
        document,
        [
            "DECLARATION DE NON CONDAMNATION",
            "EN APPLICATION DE L’ARTICLE A.123-51 du Code de Commerce",
        ],
    )


def _add_identity_block(
    document,
    *,
    subject: str,
    birth: str,
    address: str,
    nationality: str,
    filiation_father: str,
    filiation_mother: str,
) -> None:
    for line, bold in (
        (subject, True),
        (birth, False),
        (address, False),
        (nationality, False),
        (filiation_father, False),
        (filiation_mother, False),
    ):
        _add_paragraph(document, line, bold=bold)


def _add_paragraph(
    document,
    text: str,
    *,
    space_before: int = 0,
    bold: bool = False,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    add_paragraph(
        document,
        text,
        space_before_pt=space_before,
        bold=bold,
        alignment=alignment,
    )


def _add_signature_block(
    document,
    *,
    lieu_signature: str,
    date_signature: str,
    image_path: Path | None,
) -> None:
    add_signature_block(
        document,
        [f"Fait à {lieu_signature}", f"Le {date_signature}"],
        image_path=image_path,
    )


def _add_legal_reminder(document) -> None:
    add_legal_reminder(
        document,
        title="Rappel",
        title_suffix=RAPPEL_TITLE_SUFFIX,
        paragraphs=[RAPPEL_PARAGRAPH_1, RAPPEL_PARAGRAPH_2],
    )
