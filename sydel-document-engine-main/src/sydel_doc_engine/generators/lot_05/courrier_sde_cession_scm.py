from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.scm_cession_common import (
    add_body_paragraph,
    format_display_date,
    required_text,
    save_clean_document,
    validate_courrier_sde_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_header_logo,
    add_letter_place_date,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "courrier_sde_cession_scm.docx"


class CourrierSdeCessionScmGenerator:
    """Generateur from-scratch du courrier SDE cession SCM V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        scm_cession = validate_courrier_sde_context(ctx)
        enregistrement = scm_cession.enregistrement
        signataire = scm_cession.signataire_sde
        if enregistrement is None or signataire is None:
            raise ValueError("scm_cession.enregistrement et signataire_sde sont obligatoires.")

        document = new_document()
        # Logo SYDEL en header, aligne a GAUCHE (retour UAT Rafael DOC-032).
        add_header_logo(document, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        if ctx.structure == "SELAS":
            for line in [
                required_text(enregistrement.service, "scm_cession.enregistrement.service"),
                required_text(
                    enregistrement.centre_finances_publiques,
                    "scm_cession.enregistrement.centre_finances_publiques",
                ),
                required_text(
                    enregistrement.adresse_service,
                    "scm_cession.enregistrement.adresse_service",
                ),
                required_text(
                    enregistrement.cp_ville_service,
                    "scm_cession.enregistrement.cp_ville_service",
                ),
            ]:
                add_paragraph(document, line, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        add_letter_place_date(
            document,
            f"{ctx.signature.lieu}, le {format_display_date(ctx.signature.date, 'signature.date')}",
        )
        # Objet en gras + souligne (retour UAT Rafael DOC-032).
        add_paragraph(
            document,
            "Objet : Enregistrement actes de cession des parts de la société SCM",
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            bold=True,
            underline=True,
        )
        add_body_paragraph(document, "Madame, Monsieur,")
        exemplaires = (
            required_text(
                enregistrement.nombre_exemplaires,
                "scm_cession.enregistrement.nombre_exemplaires",
            )
            if ctx.structure == "SELAS"
            else "4"
        )
        add_body_paragraph(
            document,
            (
                "Je vous prie de bien vouloir trouver sous ce pli "
                f"{exemplaires} exemplaires de l'acte de cession pour les enregistrer."
            ),
        )
        add_body_paragraph(
            document,
            (
                "Vous trouverez également un chèque de "
                f"{_montant_droits(enregistrement)} "
                "euros "
                "correspondants aux droits d'enregistrements."
            ),
        )
        add_body_paragraph(
            document,
            (
                "Merci de bien vouloir me retourner les originaux chez Sydel. "
                "A cet effet, vous trouverez une enveloppe de retour timbrée."
            ),
        )
        add_body_paragraph(
            document,
            "Je vous prie d'agréer, Madame, Monsieur, mes salutations distinguées.",
        )
        # Signataire aligne a DROITE (retour UAT Rafael DOC-032).
        add_paragraph(
            document,
            (
                f"{required_text(signataire.prenom, 'scm_cession.signataire_sde.prenom')} "
                f"{required_text(signataire.nom, 'scm_cession.signataire_sde.nom')}"
            ),
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        return save_clean_document(document, output_dir, OUTPUT_FILENAME)


def _montant_droits(enregistrement) -> str:
    return required_text(
        enregistrement.montant_droits,
        "scm_cession.enregistrement.montant_droits",
    )
