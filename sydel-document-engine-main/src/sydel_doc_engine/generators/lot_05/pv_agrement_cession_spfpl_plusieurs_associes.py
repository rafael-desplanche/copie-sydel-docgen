from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.pv_agrement_common import (
    add_article_7_bis,
    add_ordre_du_jour,
    add_pouvoirs_resolution,
    add_pv_title,
    add_resolution_agrement,
    add_societe_cible_header,
    reunion_intro_lines,
)
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    associe_signature_name,
    person_display,
    presence_lines,
    required_int,
    required_societe_cible,
    required_text,
    validate_associe_unique,
    validate_cession_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_hyphen_list_item,
    add_paragraph,
    add_signature_lines,
    new_document,
)

OUTPUT_FILENAME = "pv_agrement_cession_spfpl_plusieurs_associes.docx"


class PvAgrementCessionSpfplPlusieursAssociesGenerator:
    """Generateur from-scratch du PV d'agrement SPFPL avec plusieurs associes."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_cession_context(ctx)
        validate_associe_unique(ctx, expected=False)
        societe_cible = required_societe_cible(ctx)

        docx = new_document()
        add_societe_cible_header(docx, ctx)
        add_pv_title(docx, "L'ASSEMBLEE GENERALE EXTRAORDINAIRE", ctx)
        for line in reunion_intro_lines(ctx):
            add_paragraph(docx, line)
        add_paragraph(
            docx,
            (
                "Les associes de la Societe "
                f"{required_text(societe_cible.denomination, 'societe_cible.denomination')}, "
                "au capital de "
                f"{required_text(societe_cible.capital_social, 'societe_cible.capital_social')} "
                "euros, compose de "
                f"{required_int(societe_cible.nb_parts_total, 'societe_cible.nb_parts_total')} "
                "parts, se sont reunis sur convocation reguliere de la gerance au siege "
                "de la Societe."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(docx, "Sont presents ou representes :")
        for line in presence_lines(ctx):
            add_hyphen_list_item(docx, line)
        add_paragraph(
            docx,
            (
                "Les associes presents ou representes disposent ensemble de la totalite "
                "des parts formant le capital de la societe. L'assemblee est habilitee "
                "a prendre les decisions extraordinaires."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        _add_president_sentence(docx, ctx)
        _add_depot_documents(docx, ctx)
        add_ordre_du_jour(docx)
        add_resolution_agrement(docx, ctx, subject="L'assemblee generale")
        add_article_7_bis(docx, ctx, subject="L'assemblee generale")
        add_pouvoirs_resolution(docx, subject="L'assemblee generale")
        add_paragraph(
            docx,
            (
                "De tout ce que dessus, il a ete dresse le present proces-verbal qui a "
                "ete signe apres lecture par tous les associes."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_signature_lines(
            docx,
            [
                associe_signature_name(associe, f"associes_cible[{index}]")
                for index, associe in enumerate(ctx.associes_cible)
                if associe.est_present_ou_represente
            ],
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _add_president_sentence(docx, ctx: DocumentGenerationContext) -> None:
    if ctx.reunion is None or ctx.reunion.president is None:
        raise ValueError("reunion.president est obligatoire pour CODE-SPFPL-AGR-INFO-001.")
    president = ctx.reunion.president
    add_paragraph(
        docx,
        (
            f"{person_display(president, 'reunion.president')} preside la seance en qualite "
            f"de {required_text(president.qualite, 'reunion.president.qualite')}."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_depot_documents(docx, ctx: DocumentGenerationContext) -> None:
    societe_spfpl_name = required_text(
        ctx.societe_spfpl.denomination if ctx.societe_spfpl else None,
        "societe_spfpl.denomination",
    )
    cedant_name = required_text(ctx.cedant.prenom if ctx.cedant else None, "cedant.prenom")
    cedant_name += f" {required_text(ctx.cedant.nom if ctx.cedant else None, 'cedant.nom')}"
    add_paragraph(
        docx,
        "Le President depose et met a la disposition des associes les documents suivants :",
    )
    for item in [
        "Les copies des convocations des associes ;",
        (
            f"Projet du contrat de cession des parts sociales detenues par {cedant_name} "
            f"au profit de la {societe_spfpl_name};"
        ),
        "Le rapport de la gerance ;",
        "Le texte des resolutions proposees.",
    ]:
        add_hyphen_list_item(docx, item)
    add_paragraph(
        docx,
        (
            "Le President declare que tous les documents prevus par la reglementation et "
            "les statuts ont bien ete adresses aux associes avec la convocation."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        docx,
        (
            "Ils ont ete tenus a leur disposition au siege social pendant le delai de "
            "quinze jours ayant precede l'assemblee."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        docx,
        (
            "L'assemblee lui donne acte de ses declarations et reconnait la validite de "
            "la convocation."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(docx, "Puis le President rappelle l'ordre du jour :")
