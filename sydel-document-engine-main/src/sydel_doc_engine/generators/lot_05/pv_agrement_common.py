from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    capital_after_lines,
    company_siege_display,
    format_display_date,
    person_display,
    required_cedant,
    required_cession_parts,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_centered_block,
    add_framed_title,
    add_hyphen_list_item,
    add_paragraph,
)


def add_societe_cible_header(docx, ctx: DocumentGenerationContext) -> None:
    societe_cible = required_societe_cible(ctx)
    siege = societe_cible.siege
    if siege is None:
        raise ValueError("societe_cible.siege est obligatoire pour CODE-SPFPL-AGR-INFO-001.")
    add_centered_block(
        docx,
        [
            required_text(societe_cible.denomination, "societe_cible.denomination"),
            required_text(societe_cible.forme_sociale, "societe_cible.forme_sociale"),
            f"Au capital de {_capital_social(societe_cible)} euros",
            (
                "Siege social : "
                f"{required_text(siege.num_voie, 'societe_cible.siege.num_voie')} "
                f"{required_text(siege.voie, 'societe_cible.siege.voie')}, "
                f"{required_text(siege.cp, 'societe_cible.siege.cp')} "
                f"{required_text(siege.ville, 'societe_cible.siege.ville')}"
            ),
            (
                "Immatriculee au RCS de "
                f"{required_text(societe_cible.ville_rcs, 'societe_cible.ville_rcs')} "
                "sous le n "
                f"{required_text(societe_cible.numero_rcs, 'societe_cible.numero_rcs')}"
            ),
        ],
    )


def add_pv_title(docx, middle_line: str, ctx: DocumentGenerationContext) -> None:
    decision_date = ctx.decision.date if ctx.decision else None
    add_framed_title(
        docx,
        [
            "PROCES-VERBAL DE",
            middle_line,
            f"DU {format_display_date(decision_date, 'decision.date')}",
        ],
    )


def reunion_intro_lines(ctx: DocumentGenerationContext) -> tuple[str, str]:
    reunion = ctx.reunion
    if reunion is None:
        raise ValueError("reunion est obligatoire pour CODE-SPFPL-AGR-INFO-001.")
    return (
        f"L'an {required_text(reunion.annee_lettres, 'reunion.annee_lettres')},",
        (
            f"Le {required_text(reunion.date_lettres, 'reunion.date_lettres')}, "
            f"a {required_text(reunion.heure, 'reunion.heure')},"
        ),
    )


def add_ordre_du_jour(docx) -> None:
    add_paragraph(docx, "Agrement d'un nouvel associe, la SPFPL ;")
    add_paragraph(docx, "Modification correlative des statuts ;")
    add_paragraph(docx, "Pouvoirs pour l'accomplissement des formalites.")
    add_paragraph(docx, "Des lors, il est decide de ce qui suit :")


def add_resolution_agrement(
    docx,
    ctx: DocumentGenerationContext,
    *,
    subject: str,
) -> None:
    societe_spfpl = required_societe_spfpl(ctx)
    societe_cible = required_societe_cible(ctx)
    cedant = required_cedant(ctx)
    cession_parts = required_cession_parts(ctx)
    add_paragraph(docx, "PREMIERE RESOLUTION", bold=True, space_before_pt=10)
    add_paragraph(
        docx,
        (
            f"{subject} autorise la cession par {person_display(cedant, 'cedant')} de "
            f"{required_int(cession_parts.nb_parts, 'cession_parts.nb_parts')} parts sociales "
            "qu'il detient de la "
            f"{required_text(societe_cible.denomination, 'societe_cible.denomination')}, a la "
            f"{required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')}, "
            f"numerotees de {_plage_parts(cession_parts)} "
            "inclus a compter de ce jour."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        docx,
        (
            f"Par consequent, {subject.lower()} agree la societe "
            f"{required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')} "
            "en qualite de nouvelle associee a compter de ce jour."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def add_article_7_bis(docx, ctx: DocumentGenerationContext, *, subject: str) -> None:
    societe_cible = required_societe_cible(ctx)
    add_paragraph(docx, "DEUXIEME RESOLUTION", bold=True, space_before_pt=10)
    add_paragraph(
        docx,
        (
            f"En consequence de la premiere resolution, {subject.lower()} decide, que "
            "l'article 7 bis des statuts sera modifie comme suit, a compter de ce jour :"
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(docx, "« Article 7 bis - Capital social")
    add_paragraph(
        docx,
        (
            "Le capital social de la Societe est fixe a "
            f"{required_text(societe_cible.capital_social, 'societe_cible.capital_social')} euros "
            f"({_capital_social_lettres(societe_cible)}) "
            "et est divise en "
            f"{required_int(societe_cible.nb_parts_total, 'societe_cible.nb_parts_total')} "
            "parts sociales d'un montant de "
            f"{_valeur_nominale_part(societe_cible)} "
            "euros chacune de nominal, entierement liberees, attribuees aux Associes de "
            "la maniere suivante :"
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    for line in capital_after_lines(ctx):
        add_hyphen_list_item(docx, line)
    add_paragraph(docx, "»")
    add_paragraph(docx, "Le reste de l'article est inchange.")


def add_pouvoirs_resolution(docx, *, subject: str) -> None:
    add_paragraph(docx, "TROISIEME RESOLUTION", bold=True, space_before_pt=10)
    add_paragraph(
        docx,
        (
            f"{subject} donne tous pouvoirs au porteur de copies ou d'extraits du present "
            "proces-verbal pour remplir toutes formalites de droit."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def add_societe_cible_context_sentence(docx, ctx: DocumentGenerationContext, text: str) -> None:
    societe_cible = required_societe_cible(ctx)
    add_paragraph(
        docx,
        text.format(
            denomination=required_text(societe_cible.denomination, "societe_cible.denomination"),
            capital=required_text(societe_cible.capital_social, "societe_cible.capital_social"),
            nb_parts=required_int(societe_cible.nb_parts_total, "societe_cible.nb_parts_total"),
            siege=company_siege_display(societe_cible, "societe_cible"),
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _capital_social(societe_cible) -> str:
    return required_text(societe_cible.capital_social, "societe_cible.capital_social")


def _plage_parts(cession_parts) -> str:
    return required_text(cession_parts.plage_parts, "cession_parts.plage_parts")


def _capital_social_lettres(societe_cible) -> str:
    return required_text(
        societe_cible.capital_social_lettres,
        "societe_cible.capital_social_lettres",
    )


def _valeur_nominale_part(societe_cible) -> str:
    return required_text(
        societe_cible.valeur_nominale_part,
        "societe_cible.valeur_nominale_part",
    )
