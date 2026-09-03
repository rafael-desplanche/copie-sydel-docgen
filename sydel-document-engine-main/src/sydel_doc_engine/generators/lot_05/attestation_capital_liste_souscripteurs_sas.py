from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    CapitalSouscripteur,
    DocumentGenerationContext,
)
from sydel_doc_engine.generators.lot_05.sas_satellites_common import (
    DOCUMENT_CODE,
    address_display,
    format_display_date,
    person_name,
    required_actionnaire_unique,
    required_capital_souscription,
    required_int,
    required_president,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_capital_consistency,
    validate_sas_satellite_scope,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_company_identity_block,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "attestation_capital_liste_souscripteurs_sas.docx"


class AttestationCapitalListeSouscripteursSasGenerator:
    """Generateur from-scratch de l'attestation capital / souscripteurs SAS V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        data = _ResolvedAttestationCapitalSas.from_context(ctx)
        document = new_document()

        add_company_identity_block(
            document,
            [
                data.denomination,
                f"Société par actions simplifiée au capital de {data.capital_social} euros",
                "Société de Participations Financières de Profession Libérale de "
                f"{data.profession_societe}",
                f"Siège social : {data.adresse_siege}",
            ],
        )
        add_paragraph(document, "ATTESTATION", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(
            document,
            "Liste des souscripteurs",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        add_paragraph(
            document,
            f"{data.president_nom} {data.profession_actionnaire}, demeurant "
            f"{data.adresse_actionnaire}, atteste que le capital de la société "
            f"{data.denomination} est réparti de la manière suivante :",
        )
        add_paragraph(document, f"Capital social : {data.capital_social} €")
        add_paragraph(
            document,
            f"Nombre d'actions : {data.nb_actions_total} actions d'un montant "
            f"d'{data.valeur_nominale_action} euro chacune",
        )
        add_paragraph(
            document,
            f"Répartition : {data.nb_actions_souscripteur} actions attribuées au Dr "
            f"{data.actionnaire_signature}, actionnaire unique",
        )
        add_paragraph(document, "Apports en nature :", bold=True)
        add_paragraph(
            document,
            f"{data.actionnaire_nom} fait apport de {data.nb_parts_apportees} parts de la "
            f"{data.societe_cible_forme} dénommée {data.societe_cible_denomination} "
            f"ayant son siège {data.societe_cible_siege}, immatriculée au RCS de "
            f"{data.societe_cible_ville_rcs} sous le numéro {data.societe_cible_numero_rcs} "
            f"pour une valeur de {data.apports_nature_montant} €",
        )
        add_paragraph(
            document,
            f"Total des apports en nature {data.apports_nature_montant} €",
        )
        add_paragraph(document, f"Apports en numéraire : {data.apports_numeraire_montant}")
        add_paragraph(
            document,
            f"Le Docteur {data.president_nom} a fait la totalité des apports en nature.",
        )
        add_paragraph(
            document,
            "Le présent état qui constate la souscription d'actions de la société "
            f"{data.denomination}, ainsi que l'apport de la somme de "
            f"{data.apports_nature_montant} euros correspondant à la totalité du nominal "
            "desdites actions, est certifié exact, sincère et véritable par le Président, "
            f"{data.president_nom}.",
        )
        add_paragraph(document, f"Fait à {data.lieu_signature}")
        add_paragraph(document, f"Le {data.date_signature}")
        add_paragraph(document, data.president_nom)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


class _ResolvedAttestationCapitalSas:
    def __init__(
        self,
        *,
        denomination: str,
        capital_social: str,
        profession_societe: str,
        adresse_siege: str,
        president_nom: str,
        actionnaire_nom: str,
        actionnaire_signature: str,
        profession_actionnaire: str,
        adresse_actionnaire: str,
        nb_actions_total: int,
        valeur_nominale_action: str,
        nb_actions_souscripteur: int,
        nb_parts_apportees: int,
        societe_cible_forme: str,
        societe_cible_denomination: str,
        societe_cible_siege: str,
        societe_cible_ville_rcs: str,
        societe_cible_numero_rcs: str,
        apports_nature_montant: str,
        apports_numeraire_montant: str,
        lieu_signature: str,
        date_signature: str,
    ) -> None:
        self.denomination = denomination
        self.capital_social = capital_social
        self.profession_societe = profession_societe
        self.adresse_siege = adresse_siege
        self.president_nom = president_nom
        self.actionnaire_nom = actionnaire_nom
        self.actionnaire_signature = actionnaire_signature
        self.profession_actionnaire = profession_actionnaire
        self.adresse_actionnaire = adresse_actionnaire
        self.nb_actions_total = nb_actions_total
        self.valeur_nominale_action = valeur_nominale_action
        self.nb_actions_souscripteur = nb_actions_souscripteur
        self.nb_parts_apportees = nb_parts_apportees
        self.societe_cible_forme = societe_cible_forme
        self.societe_cible_denomination = societe_cible_denomination
        self.societe_cible_siege = societe_cible_siege
        self.societe_cible_ville_rcs = societe_cible_ville_rcs
        self.societe_cible_numero_rcs = societe_cible_numero_rcs
        self.apports_nature_montant = apports_nature_montant
        self.apports_numeraire_montant = apports_numeraire_montant
        self.lieu_signature = lieu_signature
        self.date_signature = date_signature

    @classmethod
    def from_context(
        cls,
        ctx: DocumentGenerationContext,
    ) -> _ResolvedAttestationCapitalSas:
        validate_sas_satellite_scope(ctx, require_apport=True)
        societe = required_societe_spfpl(ctx)
        actionnaire = required_actionnaire_unique(ctx)
        president = required_president(ctx)
        capital = required_capital_souscription(ctx)
        souscripteur = _unique_souscripteur(capital.souscripteurs)
        societe_cible = required_societe_cible(ctx)
        if ctx.apport_titres is None:
            raise ValueError(f"apport_titres est obligatoire pour {DOCUMENT_CODE}.")
        validate_capital_consistency(societe, capital)
        _validate_souscripteur_matches_context(souscripteur, actionnaire, capital)

        return cls(
            denomination=required_text(societe.denomination, "societe_spfpl.denomination"),
            capital_social=required_text(societe.capital_social, "societe_spfpl.capital_social"),
            profession_societe=required_text(societe.profession, "societe_spfpl.profession"),
            adresse_siege=address_display(societe.siege, "societe_spfpl.siege"),
            president_nom=person_name(president, "president"),
            actionnaire_nom=person_name(actionnaire, "actionnaire_unique"),
            actionnaire_signature=(
                f"{required_text(actionnaire.prenom, 'actionnaire_unique.prenom')} "
                f"{required_text(actionnaire.nom, 'actionnaire_unique.nom')}"
            ),
            profession_actionnaire=required_text(
                actionnaire.profession,
                "actionnaire_unique.profession",
            ),
            adresse_actionnaire=required_text(
                actionnaire.adresse_personnelle_affichee,
                "actionnaire_unique.adresse_personnelle_affichee",
            ),
            nb_actions_total=required_int(
                capital.nb_actions_total,
                "capital_souscription.nb_actions_total",
            ),
            valeur_nominale_action=required_text(
                capital.valeur_nominale_action,
                "capital_souscription.valeur_nominale_action",
            ),
            nb_actions_souscripteur=required_int(
                souscripteur.nb_actions,
                "capital_souscription.souscripteurs[0].nb_actions",
            ),
            nb_parts_apportees=required_int(ctx.apport_titres.nb_parts, "apport_titres.nb_parts"),
            societe_cible_forme=required_text(
                societe_cible.forme_sociale,
                "societe_cible.forme_sociale",
            ),
            societe_cible_denomination=required_text(
                societe_cible.denomination,
                "societe_cible.denomination",
            ),
            societe_cible_siege=address_display(societe_cible.siege, "societe_cible.siege"),
            societe_cible_ville_rcs=required_text(
                societe_cible.ville_rcs,
                "societe_cible.ville_rcs",
            ),
            societe_cible_numero_rcs=required_text(
                societe_cible.numero_rcs,
                "societe_cible.numero_rcs",
            ),
            apports_nature_montant=required_text(
                capital.apports_nature_montant,
                "capital_souscription.apports_nature_montant",
            ),
            apports_numeraire_montant=required_text(
                capital.apports_numeraire_montant,
                "capital_souscription.apports_numeraire_montant",
            ),
            lieu_signature=ctx.signature.lieu,
            date_signature=format_display_date(ctx.signature.date, "signature.date"),
        )


def _unique_souscripteur(
    souscripteurs: list[CapitalSouscripteur],
) -> CapitalSouscripteur:
    if len(souscripteurs) != 1:
        raise ValueError(
            "capital_souscription.souscripteurs doit contenir exactement un "
            f"souscripteur pour {DOCUMENT_CODE}."
        )
    return souscripteurs[0]


def _validate_souscripteur_matches_context(
    souscripteur: CapitalSouscripteur,
    actionnaire,
    capital,
) -> None:
    souscripteur_actions = required_int(
        souscripteur.nb_actions,
        "capital_souscription.souscripteurs[0].nb_actions",
    )
    total_actions = required_int(
        capital.nb_actions_total,
        "capital_souscription.nb_actions_total",
    )
    if souscripteur_actions != total_actions:
        raise ValueError(
            "Le souscripteur unique doit souscrire la totalite des actions pour "
            f"{DOCUMENT_CODE}."
        )
    for field_name in ("civilite_affichage", "prenom", "nom"):
        souscripteur_value = getattr(souscripteur, field_name)
        actionnaire_value = getattr(actionnaire, field_name)
        if required_text(
            souscripteur_value,
            f"capital_souscription.souscripteurs[0].{field_name}",
        ) != required_text(actionnaire_value, f"actionnaire_unique.{field_name}"):
            raise ValueError(
                "capital_souscription.souscripteurs[0] doit correspondre a "
                f"actionnaire_unique pour {DOCUMENT_CODE}."
            )
