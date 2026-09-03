# ruff: noqa: E501
from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    Address,
    DocumentGenerationContext,
    ScmCessionAssocie,
    ScmCessionCedant,
    ScmCessionContext,
    ScmCessionCreditVendeur,
    ScmCessionEnregistrement,
    ScmCessionPrix,
    ScmCessionSignataire,
    ScmCessionSociete,
)
from sydel_doc_engine.rendering.docx_builder import add_paragraph

DOCUMENT_CODE = "FINAL-SCM-CESSION-WAVE-001"
SUPPORTED_STRUCTURES = {"SELARL": "selarl", "SELAS": "selas"}


def validate_scm_cession_enabled(ctx: DocumentGenerationContext) -> ScmCessionContext:
    if ctx.structure not in SUPPORTED_STRUCTURES:
        raise ValueError(f"dossier.structure doit etre SELARL ou SELAS pour {DOCUMENT_CODE}.")
    if ctx.dossier_options is None or not ctx.dossier_options.scm_cession:
        raise ValueError(f"dossier.options.scm_cession doit etre vrai pour {DOCUMENT_CODE}.")
    if ctx.scm_cession is None:
        raise ValueError(f"scm_cession est obligatoire pour {DOCUMENT_CODE}.")
    expected = SUPPORTED_STRUCTURES[ctx.structure]
    if ctx.scm_cession.variante_structure is not None:
        actual = ctx.scm_cession.variante_structure.strip().lower()
        if actual != expected:
            raise ValueError(
                "scm_cession.variante_structure doit correspondre a dossier.structure "
                f"pour {DOCUMENT_CODE}."
            )
    return ctx.scm_cession


def validate_pv_context(ctx: DocumentGenerationContext) -> ScmCessionContext:
    scm_cession = validate_scm_cession_enabled(ctx)
    scm_cedee = required_scm_cedee(scm_cession)
    required_cessionnaire(scm_cession)
    required_agrement(scm_cession, ctx.structure)
    required_associes(
        scm_cession.associes_presents,
        "scm_cession.associes_presents",
        expected_count=3,
        expected_total=scm_cedee.nb_parts_total,
        require_plage=False,
    )
    required_associes(
        scm_cession.associes_apres_cession,
        "scm_cession.associes_apres_cession",
        expected_count=4,
        expected_total=scm_cedee.nb_parts_total,
        require_plage=True,
    )
    if not scm_cession.signataires_pv:
        raise ValueError(f"scm_cession.signataires_pv est obligatoire pour {DOCUMENT_CODE}.")
    return scm_cession


def validate_courrier_sde_context(ctx: DocumentGenerationContext) -> ScmCessionContext:
    scm_cession = validate_scm_cession_enabled(ctx)
    enregistrement = required_enregistrement(scm_cession)
    required_text(enregistrement.montant_droits, "scm_cession.enregistrement.montant_droits")
    signataire = required_signataire_sde(scm_cession)
    required_text(signataire.prenom, "scm_cession.signataire_sde.prenom")
    required_text(signataire.nom, "scm_cession.signataire_sde.nom")
    required_text(ctx.signature.lieu, "signature.lieu")
    format_display_date(ctx.signature.date, "signature.date")
    if ctx.structure == "SELAS":
        required_text(enregistrement.service, "scm_cession.enregistrement.service")
        required_text(
            enregistrement.centre_finances_publiques,
            "scm_cession.enregistrement.centre_finances_publiques",
        )
        required_text(
            enregistrement.adresse_service,
            "scm_cession.enregistrement.adresse_service",
        )
        required_text(
            enregistrement.cp_ville_service,
            "scm_cession.enregistrement.cp_ville_service",
        )
        required_text(
            enregistrement.nombre_exemplaires,
            "scm_cession.enregistrement.nombre_exemplaires",
        )
    return scm_cession


def validate_acte_context(ctx: DocumentGenerationContext) -> ScmCessionContext:
    scm_cession = validate_scm_cession_enabled(ctx)
    scm_cedee = required_scm_cedee(scm_cession)
    cessionnaire = required_cessionnaire(scm_cession)
    cedant = required_cedant(scm_cession)
    required_associes(
        scm_cession.associes_avant_cession,
        "scm_cession.associes_avant_cession",
        expected_count=3,
        expected_total=scm_cedee.nb_parts_total,
        require_plage=False,
    )
    required_parts_cedees(scm_cession)
    required_prix(scm_cession)
    validate_credit_vendeur(scm_cession.credit_vendeur, ctx.structure)
    required_text(scm_cession.nombre_exemplaires_lettres, "scm_cession.nombre_exemplaires_lettres")
    required_text(ctx.signature.lieu, "signature.lieu")
    if ctx.structure == "SELAS":
        required_text(
            cedant.profession_reglementee_pluriel,
            "scm_cession.cedant.profession_reglementee_pluriel",
        )
        required_text(
            scm_cedee.forme_juridique,
            "scm_cession.scm_cedee.forme_juridique",
        )
        required_text(
            cessionnaire.forme_juridique,
            "scm_cession.cessionnaire.forme_juridique",
        )
        required_text(
            cessionnaire.representant.fonction if cessionnaire.representant else None,
            "scm_cession.cessionnaire.representant.fonction",
        )
        required_text(
            scm_cession.prestataire_signature_electronique
            or ctx.signature.prestataire_signature_electronique,
            "scm_cession.prestataire_signature_electronique",
        )
    if not scm_cession.representant_cessionnaire_confirme:
        raise ValueError(
            "scm_cession.representant_cessionnaire_confirme doit etre vrai "
            f"pour {DOCUMENT_CODE}."
        )
    _validate_representant_matches_cedant(cedant, cessionnaire)
    return scm_cession


def required_scm_cedee(scm_cession: ScmCessionContext) -> ScmCessionSociete:
    if scm_cession.scm_cedee is None:
        raise ValueError(f"scm_cession.scm_cedee est obligatoire pour {DOCUMENT_CODE}.")
    societe = scm_cession.scm_cedee
    required_text(societe.denomination, "scm_cession.scm_cedee.denomination")
    required_text(societe.capital_social, "scm_cession.scm_cedee.capital_social")
    address_display(societe.siege, "scm_cession.scm_cedee.siege")
    required_text(societe.ville_rcs, "scm_cession.scm_cedee.ville_rcs")
    required_text(societe.numero_rcs, "scm_cession.scm_cedee.numero_rcs")
    required_int(societe.nb_parts_total, "scm_cession.scm_cedee.nb_parts_total")
    required_text(
        societe.valeur_nominale_part,
        "scm_cession.scm_cedee.valeur_nominale_part",
    )
    required_text(societe.plage_parts_total, "scm_cession.scm_cedee.plage_parts_total")
    return societe


def required_cessionnaire(scm_cession: ScmCessionContext) -> ScmCessionSociete:
    if scm_cession.cessionnaire is None:
        raise ValueError(f"scm_cession.cessionnaire est obligatoire pour {DOCUMENT_CODE}.")
    societe = scm_cession.cessionnaire
    required_text(societe.denomination, "scm_cession.cessionnaire.denomination")
    required_text(societe.capital_social, "scm_cession.cessionnaire.capital_social")
    address_display(societe.siege, "scm_cession.cessionnaire.siege")
    required_text(societe.ville_rcs, "scm_cession.cessionnaire.ville_rcs")
    if societe.representant is None:
        raise ValueError(
            f"scm_cession.cessionnaire.representant est obligatoire pour {DOCUMENT_CODE}."
        )
    required_text(
        societe.representant.civilite_affichage,
        "scm_cession.cessionnaire.representant.civilite_affichage",
    )
    required_text(
        societe.representant.civilite_courte,
        "scm_cession.cessionnaire.representant.civilite_courte",
    )
    required_text(societe.representant.prenom, "scm_cession.cessionnaire.representant.prenom")
    required_text(societe.representant.nom, "scm_cession.cessionnaire.representant.nom")
    return societe


def required_cedant(scm_cession: ScmCessionContext) -> ScmCessionCedant:
    if scm_cession.cedant is None:
        raise ValueError(f"scm_cession.cedant est obligatoire pour {DOCUMENT_CODE}.")
    cedant = scm_cession.cedant
    for field_name, value in {
        "civilite_affichage": cedant.civilite_affichage,
        "prenom": cedant.prenom,
        "nom": cedant.nom,
        "profession": cedant.profession,
        "ville_naissance": cedant.ville_naissance,
        "departement_naissance": cedant.departement_naissance,
        "nationalite": cedant.nationalite,
        "adresse_affichee": cedant.adresse_affichee,
        "situation_maritale": cedant.situation_maritale,
        "numero_rpps": cedant.numero_rpps,
    }.items():
        required_text(value, f"scm_cession.cedant.{field_name}")
    format_display_date(cedant.date_naissance, "scm_cession.cedant.date_naissance")
    if cedant.ordre is None:
        raise ValueError(f"scm_cession.cedant.ordre est obligatoire pour {DOCUMENT_CODE}.")
    required_text(cedant.ordre.departemental, "scm_cession.cedant.ordre.departemental")
    required_text(cedant.ordre.numero, "scm_cession.cedant.ordre.numero")
    if cedant.conjoint is None:
        raise ValueError(f"scm_cession.cedant.conjoint est obligatoire pour {DOCUMENT_CODE}.")
    required_text(
        cedant.conjoint.civilite_affichage,
        "scm_cession.cedant.conjoint.civilite_affichage",
    )
    required_text(cedant.conjoint.prenom, "scm_cession.cedant.conjoint.prenom")
    required_text(cedant.conjoint.nom, "scm_cession.cedant.conjoint.nom")
    return cedant


def required_agrement(
    scm_cession: ScmCessionContext,
    structure: str | None,
):
    if scm_cession.agrement is None:
        raise ValueError(f"scm_cession.agrement est obligatoire pour {DOCUMENT_CODE}.")
    agrement = scm_cession.agrement
    format_display_date(agrement.date_pv, "scm_cession.agrement.date_pv")
    required_text(agrement.date_pv_lettres, "scm_cession.agrement.date_pv_lettres")
    if structure == "SELAS":
        required_text(agrement.delai_mois, "scm_cession.agrement.delai_mois")
        required_text(agrement.date_limite, "scm_cession.agrement.date_limite")
    return agrement


def required_parts_cedees(scm_cession: ScmCessionContext):
    if scm_cession.parts_cedees is None:
        raise ValueError(f"scm_cession.parts_cedees est obligatoire pour {DOCUMENT_CODE}.")
    required_int(scm_cession.parts_cedees.nb, "scm_cession.parts_cedees.nb")
    required_text(scm_cession.parts_cedees.plage, "scm_cession.parts_cedees.plage")
    return scm_cession.parts_cedees


def required_prix(scm_cession: ScmCessionContext) -> ScmCessionPrix:
    if scm_cession.prix is None:
        raise ValueError(f"scm_cession.prix est obligatoire pour {DOCUMENT_CODE}.")
    prix = scm_cession.prix
    required_text(prix.unitaire, "scm_cession.prix.unitaire")
    required_text(prix.unitaire_lettres, "scm_cession.prix.unitaire_lettres")
    required_text(prix.global_, "scm_cession.prix.global")
    required_text(prix.global_lettres, "scm_cession.prix.global_lettres")
    return prix


def required_enregistrement(scm_cession: ScmCessionContext) -> ScmCessionEnregistrement:
    if scm_cession.enregistrement is None:
        raise ValueError(f"scm_cession.enregistrement est obligatoire pour {DOCUMENT_CODE}.")
    return scm_cession.enregistrement


def required_signataire_sde(scm_cession: ScmCessionContext) -> ScmCessionSignataire:
    if scm_cession.signataire_sde is None:
        raise ValueError(f"scm_cession.signataire_sde est obligatoire pour {DOCUMENT_CODE}.")
    return scm_cession.signataire_sde


def required_associes(
    associes: list[ScmCessionAssocie],
    field_name: str,
    *,
    expected_count: int,
    expected_total: int | None,
    require_plage: bool,
) -> list[ScmCessionAssocie]:
    if len(associes) != expected_count:
        raise ValueError(
            f"{field_name} doit contenir exactement {expected_count} associes "
            f"pour {DOCUMENT_CODE}."
        )
    total = 0
    for index, associe in enumerate(associes):
        prefix = f"{field_name}[{index}]"
        associe_display(associe, prefix)
        if associe.parts is None:
            raise ValueError(f"{prefix}.parts est obligatoire pour {DOCUMENT_CODE}.")
        total += required_int(associe.parts.nb, f"{prefix}.parts.nb")
        if require_plage:
            required_text(associe.parts.plage, f"{prefix}.parts.plage")
    if expected_total is not None and total != expected_total:
        raise ValueError(f"{field_name} doit totaliser scm_cedee.nb_parts_total.")
    return associes


def validate_credit_vendeur(
    credit_vendeur: ScmCessionCreditVendeur | None,
    structure: str | None,
) -> None:
    if credit_vendeur is None or not credit_vendeur.actif:
        return
    required_text(credit_vendeur.montant, "scm_cession.credit_vendeur.montant")
    required_text(credit_vendeur.duree, "scm_cession.credit_vendeur.duree")
    required_text(credit_vendeur.taux, "scm_cession.credit_vendeur.taux")
    if structure == "SELAS":
        required_text(
            credit_vendeur.majoration_interet_retard,
            "scm_cession.credit_vendeur.majoration_interet_retard",
        )


def associe_display(associe: ScmCessionAssocie, field_name: str) -> str:
    if associe.type_personne == "personne_morale":
        return required_text(associe.denomination, f"{field_name}.denomination")
    return (
        f"{required_text(associe.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{required_text(associe.prenom, f'{field_name}.prenom')} "
        f"{required_text(associe.nom, f'{field_name}.nom')}"
    )


def associe_signature(associe: ScmCessionAssocie, field_name: str) -> str:
    if associe.type_personne == "personne_morale":
        return required_text(associe.denomination, f"{field_name}.denomination")
    return (
        f"{required_text(associe.prenom, f'{field_name}.prenom')} "
        f"{required_text(associe.nom, f'{field_name}.nom')}"
    )


def person_signature(civilite: str | None, prenom: str | None, nom: str | None, prefix: str) -> str:
    return (
        f"{required_text(civilite, f'{prefix}.civilite')} "
        f"{required_text(prenom, f'{prefix}.prenom')} "
        f"{required_text(nom, f'{prefix}.nom')}"
    )


def address_display(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{required_text(address.voie, f'{field_name}.voie')}, "
        f"{required_text(address.cp, f'{field_name}.cp')} "
        f"{required_text(address.ville, f'{field_name}.ville')}"
    )


def format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return required_text(value, field_name)


def required_text(value: str | None, field_name: str) -> str:
    if value is None or not str(value).strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return str(value).strip()


def required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return int(value)


def cessionnaire_forme(ctx: DocumentGenerationContext, cessionnaire: ScmCessionSociete) -> str:
    if ctx.structure == "SELARL":
        return "SELARL"
    return required_text(cessionnaire.forme_juridique, "scm_cession.cessionnaire.forme_juridique")


def cessionnaire_representant_fonction(
    ctx: DocumentGenerationContext,
    cessionnaire: ScmCessionSociete,
) -> str:
    if ctx.structure == "SELARL":
        return "gérant"
    return required_text(
        cessionnaire.representant.fonction if cessionnaire.representant else None,
        "scm_cession.cessionnaire.representant.fonction",
    )


def acte_signature_prestataire(ctx: DocumentGenerationContext, scm_cession: ScmCessionContext) -> str:
    if ctx.structure == "SELARL":
        return "Yousign"
    return required_text(
        scm_cession.prestataire_signature_electronique
        or ctx.signature.prestataire_signature_electronique,
        "scm_cession.prestataire_signature_electronique",
    )


def scm_cedee_forme(ctx: DocumentGenerationContext, scm_cedee: ScmCessionSociete) -> str:
    if ctx.structure == "SELARL":
        return "Société Civile de Moyens"
    return required_text(scm_cedee.forme_juridique, "scm_cession.scm_cedee.forme_juridique")


def scm_cedee_address_for_acte(
    ctx: DocumentGenerationContext,
    scm_cedee: ScmCessionSociete,
    cessionnaire: ScmCessionSociete,
) -> str:
    if ctx.structure == "SELARL":
        return address_display(cessionnaire.siege, "scm_cession.cessionnaire.siege")
    return address_display(scm_cedee.siege, "scm_cession.scm_cedee.siege")


def cedant_display(cedant: ScmCessionCedant) -> str:
    return (
        f"{required_text(cedant.civilite_affichage, 'scm_cession.cedant.civilite_affichage')} "
        f"{required_text(cedant.prenom, 'scm_cession.cedant.prenom')} "
        f"{required_text(cedant.nom, 'scm_cession.cedant.nom')}"
    )


def conjoint_display(cedant: ScmCessionCedant) -> str:
    if cedant.conjoint is None:
        raise ValueError(f"scm_cession.cedant.conjoint est obligatoire pour {DOCUMENT_CODE}.")
    return (
        f"{required_text(cedant.conjoint.civilite_affichage, 'scm_cession.cedant.conjoint.civilite_affichage')} "
        f"{required_text(cedant.conjoint.prenom, 'scm_cession.cedant.conjoint.prenom')} "
        f"{required_text(cedant.conjoint.nom, 'scm_cession.cedant.conjoint.nom')}"
    )


def add_body_paragraph(document: Any, text: str, *, bold: bool = False) -> None:
    add_paragraph(
        document,
        text,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=bold,
    )


def add_heading(document: Any, text: str) -> None:
    add_paragraph(document, text, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


def save_clean_document(document: Any, output_dir: Path, output_filename: str) -> Path:
    full_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    if "[" in full_text or "]" in full_text:
        raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")
    if "Ajouter en cas de CV" in full_text:
        raise ValueError(f"instruction credit-vendeur residuelle dans le rendu {DOCUMENT_CODE}.")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / output_filename
    document.save(output_path)
    return output_path


def _validate_representant_matches_cedant(
    cedant: ScmCessionCedant,
    cessionnaire: ScmCessionSociete,
) -> None:
    representant = cessionnaire.representant
    if representant is None:
        raise ValueError(
            f"scm_cession.cessionnaire.representant est obligatoire pour {DOCUMENT_CODE}."
        )
    if (
        required_text(representant.prenom, "scm_cession.cessionnaire.representant.prenom")
        != required_text(cedant.prenom, "scm_cession.cedant.prenom")
        or required_text(representant.nom, "scm_cession.cessionnaire.representant.nom")
        != required_text(cedant.nom, "scm_cession.cedant.nom")
    ):
        raise ValueError(
            "le representant de la SEL cessionnaire doit correspondre au cedant "
            f"pour le wording source V1 {DOCUMENT_CODE}."
        )
