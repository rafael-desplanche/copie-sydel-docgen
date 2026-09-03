from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    company_siege_display,
    format_display_date,
    person_address_display,
    person_signature,
    professional_entity_presentation,
    required_apport_titres,
    required_apporteur,
    required_commissaire_aux_apports,
    required_evaluateur_apport,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_apport_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_hyphen_list_item,
    add_paragraph,
    add_signature_lines,
    new_document,
)

OUTPUT_FILENAME = "contrat_apport_spfpl.docx"


class ContratApportSpfplGenerator:
    """Generateur from-scratch du contrat d'apport SEL vers SPFPL."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_apport_context(ctx)
        apporteur = required_apporteur(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        apport_titres = required_apport_titres(ctx)
        evaluateur = required_evaluateur_apport(ctx)
        commissaire = required_commissaire_aux_apports(ctx)
        apporteur_profession = required_text(
            apporteur.profession_reglementee,
            "apporteur.profession_reglementee",
        )
        apporteur_departement_naissance = required_text(
            apporteur.departement_naissance,
            "apporteur.departement_naissance",
        )
        spfpl_name = required_text(societe_spfpl.denomination, "societe_spfpl.denomination")
        spfpl_capital = required_text(societe_spfpl.capital_social, "societe_spfpl.capital_social")
        spfpl_siege = company_siege_display(societe_spfpl, "societe_spfpl")
        dirigeant_fonction = required_text(
            societe_spfpl.dirigeant.fonction if societe_spfpl.dirigeant else None,
            "societe_spfpl.dirigeant.fonction",
        )
        cible_name = required_text(societe_cible.denomination, "societe_cible.denomination")
        cible_capital = required_text(societe_cible.capital_social, "societe_cible.capital_social")
        cible_siege = company_siege_display(societe_cible, "societe_cible")
        cible_numero_rcs = required_text(societe_cible.numero_rcs, "societe_cible.numero_rcs")
        valeur_par_titre_lettres = required_text(
            apport_titres.valeur_par_titre_lettres,
            "apport_titres.valeur_par_titre_lettres",
        )
        valeur_globale_lettres = required_text(
            apport_titres.valeur_globale_lettres,
            "apport_titres.valeur_globale_lettres",
        )
        nb_actions_lettres = required_text(
            apport_titres.nb_actions_attribuees_lettres,
            "apport_titres.nb_actions_attribuees_lettres",
        )
        valeur_action_lettres = required_text(
            apport_titres.valeur_nominale_action_lettres,
            "apport_titres.valeur_nominale_action_lettres",
        )

        docx = new_document()
        add_paragraph(docx, "Contrat d'apport", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, "Entre les soussignes :", bold=True)
        add_hyphen_list_item(docx, f"{_civilite_nom(apporteur, 'apporteur')}")
        add_paragraph(
            docx,
            f"{apporteur_profession} de profession",
        )
        add_paragraph(
            docx,
            "Ne le "
            f"{format_display_date(apporteur.date_naissance, 'apporteur.date_naissance')} "
            f"a {required_text(apporteur.ville_naissance, 'apporteur.ville_naissance')} "
            f"({apporteur_departement_naissance})",
        )
        add_paragraph(docx, f"Demeurant {person_address_display(apporteur, 'apporteur')}")
        add_paragraph(
            docx,
            f"{required_text(apporteur.situation_maritale, 'apporteur.situation_maritale')} "
            f"avec {_conjoint_nom(apporteur)}",
        )
        add_paragraph(
            docx,
            f"De nationalite {required_text(apporteur.nationalite, 'apporteur.nationalite')}",
        )
        add_paragraph(docx, _ordre_apporteur(apporteur))
        add_paragraph(docx, 'Ci-apres designe "l\'apporteur" ou le soussigne de premiere part')

        add_paragraph(
            docx,
            required_text(societe_spfpl.denomination, "societe_spfpl.denomination"),
            bold=True,
            space_before_pt=10,
        )
        add_paragraph(
            docx,
            f"{required_text(societe_spfpl.forme_sociale, 'societe_spfpl.forme_sociale')} "
            f"au capital de {spfpl_capital} euros",
        )
        add_paragraph(
            docx,
            f"Societe de {required_text(societe_spfpl.activite, 'societe_spfpl.activite')}",
        )
        add_paragraph(docx, f"Siege social : {spfpl_siege}")
        add_paragraph(
            docx,
            "En cours d'immatriculation au RCS de "
            f"{required_text(societe_spfpl.ville_rcs, 'societe_spfpl.ville_rcs')}",
        )
        add_paragraph(
            docx,
            "Representee par son "
            f"{dirigeant_fonction}, "
            f"{_civilite_nom(apporteur, 'apporteur')}, domicilie en cette qualite audit siege.",
        )
        add_paragraph(docx, 'Ci-apres designee "la societe beneficiaire"')

        add_paragraph(docx, "Il a precedemment ete expose ce qui suit :", bold=True)
        add_paragraph(
            docx,
            "Les Parties ont decide que "
            f"{_civilite_nom(apporteur, 'apporteur')} apporte a la "
            f"{required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')} "
            f"{required_int(apport_titres.nb_parts, 'apport_titres.nb_parts')} "
            f"{required_text(apport_titres.nature_titres, 'apport_titres.nature_titres')} "
            f"de la {cible_name}, "
            f"{required_text(societe_cible.forme_sociale, 'societe_cible.forme_sociale')}, "
            f"au capital de {cible_capital} euros "
            f"dont le siege social est situe au {cible_siege}, "
            "immatriculee au RCS de "
            f"{required_text(societe_cible.ville_rcs, 'societe_cible.ville_rcs')} "
            f"sous le numero {cible_numero_rcs}.",
        )

        add_paragraph(docx, "Les biens apportes", bold=True)
        add_paragraph(
            docx,
            f"{_civilite_nom(apporteur, 'apporteur')}, soussigne de premiere part, "
            "apporte a la societe beneficiaire, sous les garanties ordinaires et de "
            "droit, la pleine propriete de "
            f"{required_int(apport_titres.nb_parts, 'apport_titres.nb_parts')} "
            f"{required_text(apport_titres.nature_titres, 'apport_titres.nature_titres')} "
            f"de la Societe {cible_name}.",
        )
        add_paragraph(
            docx,
            "L'apport est indivisible et porte obligatoirement sur la pleine et "
            "entiere propriete de "
            f"{required_int(apport_titres.nb_parts, 'apport_titres.nb_parts')} "
            "parts sociales de la Societe Apportee numerotees de "
            f"{required_text(apport_titres.plage_parts, 'apport_titres.plage_parts')}.",
        )

        add_paragraph(docx, "L'evaluation de l'apport", bold=True)
        add_paragraph(
            docx,
            "Le montant de l'apport est estime a "
            f"{valeur_par_titre_lettres} euros "
            f"({_valeur_par_titre(apport_titres)} euros) "
            "par part, soit le prix global de "
            f"{valeur_globale_lettres} euros "
            f"({_valeur_globale(apport_titres)} euros).",
        )
        add_paragraph(
            docx,
            "L'evaluation a ete effectuee par "
            f"{professional_entity_presentation(evaluateur, 'evaluateur_apport')}.",
        )
        add_paragraph(
            docx,
            "Dans ce cadre, "
            f"{professional_entity_presentation(commissaire, 'commissaire_aux_apports')}, "
            "en qualite de commissaire aux apports.",
        )

        add_paragraph(docx, "La remuneration de l'apport", bold=True)
        add_paragraph(
            docx,
            "En contrepartie de l'apport, il est attribue a l'apporteur "
            f"{nb_actions_lettres} "
            f"({_nb_actions_attribuees(apport_titres)}) "
            "actions nouvelles d'une valeur nominale de "
            f"{valeur_action_lettres} "
            "euro chacune.",
        )
        add_paragraph(docx, "Conditions suspensives", bold=True)
        add_hyphen_list_item(
            docx,
            "Inscription de la societe "
            f"{spfpl_name} au tableau de l'{_ordre_professionnel(apporteur)}.",
        )
        add_hyphen_list_item(
            docx,
            "Immatriculation de la societe au Registre du Commerce et des Societes de "
            f"{required_text(societe_spfpl.ville_rcs, 'societe_spfpl.ville_rcs')}.",
        )
        add_paragraph(docx, "Convention sur la preuve - signature electronique", bold=True)
        add_paragraph(
            docx,
            "Les Parties consentent expressement la faculte de proceder a la signature "
            "du present acte par le systeme de signature electronique.",
        )
        add_paragraph(
            docx,
            f"Fait a {ctx.signature.lieu} en {_nombre_exemplaires(ctx)} exemplaires",
        )
        add_paragraph(docx, f"Le {ctx.signature.date.strftime('%d/%m/%Y')}")
        add_signature_lines(
            docx,
            [
                f"{person_signature(apporteur, 'apporteur')} - En qualite d'apporteur",
                f"{spfpl_name} - En qualite de beneficiaire",
            ],
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _civilite_nom(person, field_name: str) -> str:
    return (
        f"{required_text(person.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{required_text(person.prenom, f'{field_name}.prenom')} "
        f"{required_text(person.nom, f'{field_name}.nom')}"
    )


def _conjoint_nom(person) -> str:
    if person.conjoint is None:
        raise ValueError("apporteur.conjoint est obligatoire.")
    return required_text(person.conjoint.nom, "apporteur.conjoint.nom")


def _ordre_apporteur(person) -> str:
    if person.ordre is None:
        raise ValueError("apporteur.ordre est obligatoire.")
    ordre_professionnel = required_text(
        person.ordre.professionnel,
        "apporteur.ordre.professionnel",
    )
    return (
        f"Inscrit au tableau de l'{ordre_professionnel} "
        f"de {required_text(person.ordre.departement, 'apporteur.ordre.departement')} "
        f"sous le n {required_text(person.ordre.numero, 'apporteur.ordre.numero')} "
        "et sous le numero RPPS "
        f"{required_text(person.ordre.numero_rpps, 'apporteur.ordre.numero_rpps')}"
    )


def _nombre_exemplaires(ctx: DocumentGenerationContext) -> str:
    if ctx.document and ctx.document.nombre_exemplaires_lettres:
        return ctx.document.nombre_exemplaires_lettres
    if ctx.signature.nombre_exemplaires:
        return ctx.signature.nombre_exemplaires
    raise ValueError("document.nombre_exemplaires_lettres est obligatoire.")


def _valeur_par_titre(apport_titres) -> str:
    return required_text(apport_titres.valeur_par_titre, "apport_titres.valeur_par_titre")


def _valeur_globale(apport_titres) -> str:
    return required_text(apport_titres.valeur_globale, "apport_titres.valeur_globale")


def _nb_actions_attribuees(apport_titres) -> int:
    return required_int(
        apport_titres.nb_actions_attribuees,
        "apport_titres.nb_actions_attribuees",
    )


def _ordre_professionnel(person) -> str:
    if person.ordre is None:
        raise ValueError("apporteur.ordre est obligatoire.")
    return required_text(person.ordre.professionnel, "apporteur.ordre.professionnel")
