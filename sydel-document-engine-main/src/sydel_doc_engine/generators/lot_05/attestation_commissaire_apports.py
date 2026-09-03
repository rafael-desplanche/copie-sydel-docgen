from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    company_siege_display,
    format_display_date,
    person_address_display,
    person_short_identity,
    professional_entity_presentation,
    required_apport_titres,
    required_apporteur,
    required_commissaire_aux_apports,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_apport_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_hyphen_list_item,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "attestation_commissaire_apports.docx"


class AttestationCommissaireApportsGenerator:
    """Generateur from-scratch de la designation du commissaire aux apports."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_apport_context(ctx)
        apporteur = required_apporteur(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        apport_titres = required_apport_titres(ctx)
        commissaire = required_commissaire_aux_apports(ctx)
        apporteur_departement_naissance = required_text(
            apporteur.departement_naissance,
            "apporteur.departement_naissance",
        )
        apporteur_profession = required_text(
            apporteur.profession_reglementee,
            "apporteur.profession_reglementee",
        )
        cible_forme = required_text(societe_cible.forme_sociale, "societe_cible.forme_sociale")
        cible_name = required_text(societe_cible.denomination, "societe_cible.denomination")
        cible_numero_rcs = required_text(societe_cible.numero_rcs, "societe_cible.numero_rcs")

        docx = new_document()
        add_paragraph(docx, person_signature_header(apporteur))
        add_paragraph(docx, person_address_display(apporteur, "apporteur"))
        add_paragraph(
            docx,
            "Acte de designation d'un commissaire aux apports",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_before_pt=10,
        )
        add_paragraph(
            docx,
            "Le soussigne, "
            f"{person_short_identity(apporteur, 'apporteur')}, "
            f"ne le {format_display_date(apporteur.date_naissance, 'apporteur.date_naissance')} "
            f"a {required_text(apporteur.ville_naissance, 'apporteur.ville_naissance')} "
            f"({apporteur_departement_naissance}), "
            f"{apporteur_profession}, "
            f"de nationalite {required_text(apporteur.nationalite, 'apporteur.nationalite')}, "
            f"demeurant {person_address_display(apporteur, 'apporteur')}, "
            f"{required_text(apporteur.situation_maritale, 'apporteur.situation_maritale')} "
            f"avec {_conjoint_nom(apporteur)}",
        )
        add_paragraph(
            docx,
            "seul futur associe de la societe "
            f"{required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')} "
            f"{required_text(societe_spfpl.forme_sociale, 'societe_spfpl.forme_sociale')} "
            f"de {required_text(societe_spfpl.profession, 'societe_spfpl.profession')} "
            "en cours de formation,",
        )
        add_paragraph(docx, "a prealablement expose et rappele ce qui suit :")
        add_paragraph(
            docx,
            "Le soussigne a decide de constituer une societe de "
            f"{required_text(societe_spfpl.activite, 'societe_spfpl.activite')} "
            "moyennant l'apport suivant :",
        )
        add_hyphen_list_item(
            docx,
            f"{required_int(apport_titres.nb_parts, 'apport_titres.nb_parts')} "
            f"parts de la {cible_forme} "
            f"denommee \"{cible_name}\", "
            f"ayant son siege {company_siege_display(societe_cible, 'societe_cible')}, "
            "immatriculee au RCS de "
            f"{required_text(societe_cible.ville_rcs, 'societe_cible.ville_rcs')} "
            f"sous le numero {cible_numero_rcs}.",
        )
        add_paragraph(docx, "Il a ete convenu ce qui suit :")
        add_paragraph(
            docx,
            "Aux fins de realisation de cet apport en nature a ladite societe, "
            "le soussigne nomme :",
        )
        add_paragraph(
            docx,
            professional_entity_presentation(commissaire, "commissaire_aux_apports")
            + ", en qualite de commissaire aux apports.",
        )
        add_paragraph(
            docx,
            "A l'effet d'etablir sous sa responsabilite un rapport sur la valeur "
            "dudit apport en nature, lequel sera annexe aux statuts de la societe "
            "conformement a l'article L. 223-9 du Code de commerce.",
        )
        add_paragraph(docx, f"Fait a {ctx.signature.lieu}")
        add_paragraph(docx, f"Le {ctx.signature.date.strftime('%d/%m/%Y')}")
        add_paragraph(docx, person_signature_header(apporteur), space_before_pt=12)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def person_signature_header(person) -> str:
    return (
        f"{required_text(person.prenom, 'apporteur.prenom')} "
        f"{required_text(person.nom, 'apporteur.nom')}"
    )


def _conjoint_nom(person) -> str:
    if person.conjoint is None:
        raise ValueError("apporteur.conjoint est obligatoire.")
    return required_text(person.conjoint.nom, "apporteur.conjoint.nom")
