from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from sydel_doc_engine.domain.models import (
    Address,
    CentreImpots,
    Company,
    DocumentGenerationContext,
    StatutsCivilsAssocie,
    StatutsCivilsContext,
)
from sydel_doc_engine.rendering.docx_builder import (
    LETTER_WIDE_STYLE_PROFILE,
    add_letter_place_date,
    add_paragraph,
    add_right_indented_block,
    add_spacer,
    add_subject_heading,
    new_document,
)

OUTPUT_FILENAME = "lettre_option_is.docx"
DOCUMENT_CODE = "CODE-OPTION-IS-001"
SUPPORTED_STRUCTURES = {"SCI", "SCI IRIS"}


class LettreOptionIsGenerator:
    """Generateur from-scratch de la lettre d'option IS V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        _validate_context(ctx)
        company = _required_company(ctx.societe)
        tax_office = _required_tax_office(ctx.impots)
        statuts = _required_statuts_civils(ctx.statuts_civils)

        document = new_document(LETTER_WIDE_STYLE_PROFILE)
        _add_tax_office_block(document, tax_office)
        _add_place_date_and_subject(document, ctx.signature.lieu, ctx.signature.date)
        _add_body_intro(document)
        _add_identification_table(document, company, statuts)
        _add_body_close(document)
        _add_signature(document)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _validate_context(ctx: DocumentGenerationContext) -> None:
    if ctx.structure not in SUPPORTED_STRUCTURES:
        supported = ", ".join(sorted(SUPPORTED_STRUCTURES))
        raise ValueError(f"dossier.structure doit etre dans [{supported}] pour {DOCUMENT_CODE}.")
    if ctx.dossier_options is None or not ctx.dossier_options.option_is:
        raise ValueError(f"dossier.options.option_is doit etre vrai pour {DOCUMENT_CODE}.")
    statuts = _required_statuts_civils(ctx.statuts_civils)
    expected_type = "sci_iris" if ctx.structure == "SCI IRIS" else "sci"
    if statuts.type != expected_type:
        raise ValueError(f"statuts_civils.type doit etre {expected_type} pour {DOCUMENT_CODE}.")


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return company


def _required_tax_office(tax_office: CentreImpots | None) -> CentreImpots:
    if tax_office is None:
        raise ValueError(f"impots est obligatoire pour {DOCUMENT_CODE}.")
    return tax_office


def _required_statuts_civils(statuts: StatutsCivilsContext | None) -> StatutsCivilsContext:
    if statuts is None:
        raise ValueError(f"statuts_civils est obligatoire pour {DOCUMENT_CODE}.")
    if not statuts.associes:
        raise ValueError(f"statuts_civils.associes est obligatoire pour {DOCUMENT_CODE}.")
    return statuts


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def _required_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _format_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def _add_tax_office_block(document: Any, tax_office: CentreImpots) -> None:
    add_right_indented_block(
        document,
        [
            _required_text(tax_office.service, "impots.service"),
            _required_text(tax_office.centre, "impots.centre"),
            _required_text(tax_office.adresse_ligne_1, "impots.adresse_ligne_1"),
            _required_text(tax_office.adresse_ligne_2, "impots.adresse_ligne_2"),
            (
                f"{_required_text(tax_office.cp, 'impots.cp')} "
                f"{_required_text(tax_office.ville, 'impots.ville')}"
            ),
        ],
        left_indent_cm=8.4,
        space_after_pt=2,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )
    add_spacer(document, space_after_pt=16)


def _add_place_date_and_subject(document: Any, lieu: str, signature_date: date) -> None:
    add_letter_place_date(
        document,
        f"Fait à {lieu}, le {_format_date(signature_date)}",
        space_after_pt=12,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )
    add_subject_heading(
        document,
        "Objet : Demande d'option pour le régime de l'impôt sur les sociétés",
        space_after_pt=12,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )


def _add_body_intro(document: Any) -> None:
    add_paragraph(document, "Madame, Monsieur,", style_profile=LETTER_WIDE_STYLE_PROFILE)
    add_paragraph(
        document,
        (
            "Nous vous informons que la société civile dont vous trouverez la description "
            "ci-après opte pour le régime de l'Impôt sur les Sociétés, et souhaite que cette "
            "option produise ses effets à compter de l'exercice ouvert dès l'immatriculation."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )
    add_paragraph(
        document,
        "Cette décision est prise en accord avec les associés participant.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )
    add_paragraph(
        document,
        (
            "État d'identification de la société et liste des associés au 1er jour du premier "
            "exercice d'option :"
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )


def _add_identification_table(
    document: Any,
    company: Company,
    statuts: StatutsCivilsContext,
) -> None:
    _validate_capital_distribution(statuts)
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.columns[0].width = Cm(6)
    table.columns[1].width = Cm(10)

    capital = _required_text(
        company.capital_social or statuts.capital_social,
        "societe.capital_social",
    )
    denomination = _required_text(company.denomination, "societe.denomination")
    _add_table_row(table, "Dénomination", denomination)
    _add_table_row(table, "Adresse (siège ou principal établissement)", _company_address(company))
    _add_table_row(table, "SIREN", _required_text(company.siren, "societe.siren"))
    label = (
        "Nom, prénom et adresse des différents associés de la société, "
        f"et répartition du capital de {capital} €"
    )
    for index, associe in enumerate(statuts.associes):
        _add_table_row(table, label, _associe_table_text(associe, index))

    add_spacer(document, space_after_pt=12)


def _add_table_row(table: Any, label: str, value: str) -> None:
    row_cells = table.add_row().cells
    row_cells[0].text = label
    row_cells[1].text = value
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)


def _validate_capital_distribution(statuts: StatutsCivilsContext) -> None:
    total = _required_int(statuts.nb_parts_total, "statuts_civils.nb_parts_total")
    associes_total = 0
    for index, associe in enumerate(statuts.associes):
        if associe.parts is None:
            raise ValueError(f"statuts_civils.associes[{index}].parts est obligatoire.")
        associes_total += _required_int(
            associe.parts.nb,
            f"statuts_civils.associes[{index}].parts.nb",
        )
    if associes_total != total:
        raise ValueError(
            "La repartition des parts doit correspondre a statuts_civils.nb_parts_total "
            f"pour {DOCUMENT_CODE}."
        )


def _company_address(company: Company) -> str:
    if company.siege is None:
        raise ValueError(f"societe.siege est obligatoire pour {DOCUMENT_CODE}.")
    return _address_display(company.siege, "societe.siege")


def _address_display(address: Address, field_name: str) -> str:
    if address.adresse_affichee:
        return address.adresse_affichee.strip()
    return (
        f"{_required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{_required_text(address.voie, f'{field_name}.voie')}, "
        f"{_required_text(address.cp, f'{field_name}.cp')} "
        f"{_required_text(address.ville, f'{field_name}.ville')}"
    )


def _associe_table_text(associe: StatutsCivilsAssocie, index: int) -> str:
    field_name = f"statuts_civils.associes[{index}]"
    if associe.parts is None:
        raise ValueError(f"{field_name}.parts est obligatoire pour {DOCUMENT_CODE}.")
    nb_parts = _required_int(associe.parts.nb, f"{field_name}.parts.nb")
    if associe.type_personne == "personne_morale":
        return _associe_morale_table_text(associe, field_name, nb_parts)
    return _associe_physique_table_text(associe, field_name, nb_parts)


def _associe_physique_table_text(
    associe: StatutsCivilsAssocie,
    field_name: str,
    nb_parts: int,
) -> str:
    address = associe.adresse_personnelle
    if associe.adresse_personnelle_affichee:
        address_display = associe.adresse_personnelle_affichee.strip()
    elif address is not None:
        address_display = _address_display(address, f"{field_name}.adresse_personnelle")
    else:
        raise ValueError(f"{field_name}.adresse_personnelle est obligatoire pour {DOCUMENT_CODE}.")
    qualite = _required_text(
        associe.parts.qualite_associe or associe.role_statutaire,
        f"{field_name}.parts.qualite_associe",
    )
    return (
        f"{_required_text(associe.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{_required_text(associe.prenom, f'{field_name}.prenom')} "
        f"{_required_text(associe.nom, f'{field_name}.nom')}, demeurant {address_display}, "
        f"{qualite}, détenant {nb_parts} parts."
    )


def _associe_morale_table_text(
    associe: StatutsCivilsAssocie,
    field_name: str,
    nb_parts: int,
) -> str:
    if associe.siege is None:
        raise ValueError(f"{field_name}.siege est obligatoire pour {DOCUMENT_CODE}.")
    return (
        f"La société {_required_text(associe.denomination, f'{field_name}.denomination')}, "
        f"ayant son siège social au {_address_display(associe.siege, f'{field_name}.siege')}, "
        f"détenant {nb_parts} parts."
    )


def _add_body_close(document: Any) -> None:
    add_paragraph(
        document,
        (
            "Nous vous remercions de bien vouloir nous confirmer que cette demande est validée, "
            "et vous prions de recevoir, Madame, Monsieur, l'assurance de notre parfaite "
            "considération."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )


def _add_signature(document: Any) -> None:
    add_spacer(document, space_after_pt=12)
    add_paragraph(
        document,
        "Le gérant",
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        style_profile=LETTER_WIDE_STYLE_PROFILE,
    )
