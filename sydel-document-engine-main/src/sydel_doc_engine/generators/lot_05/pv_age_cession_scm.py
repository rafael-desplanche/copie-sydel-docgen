# ruff: noqa: E501
from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext, ScmCessionAssocie
from sydel_doc_engine.generators.lot_05.scm_cession_common import (
    add_body_paragraph,
    add_heading,
    address_display,
    associe_display,
    format_display_date,
    required_text,
    save_clean_document,
    validate_pv_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_framed_title,
    add_hyphen_list_item,
    add_paragraph,
    add_signature_table,
    new_document,
)

OUTPUT_FILENAME = "pv_age_cession_parts_scm.docx"


class PvAgeCessionScmGenerator:
    """Generateur from-scratch du PV AGE de cession de parts SCM V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        scm_cession = validate_pv_context(ctx)
        scm_cedee = scm_cession.scm_cedee
        cessionnaire = scm_cession.cessionnaire
        agrement = scm_cession.agrement
        if scm_cedee is None or cessionnaire is None or agrement is None:
            raise ValueError("scm_cession est incomplet pour le PV AGE cession SCM.")

        document = new_document()
        add_paragraph(
            document,
            required_text(scm_cedee.denomination, "scm_cession.scm_cedee.denomination"),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_after_pt=2,
        )
        add_paragraph(
            document,
            "Société civile de moyens",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after_pt=2,
        )
        add_paragraph(
            document,
            f"Au capital de {required_text(scm_cedee.capital_social, 'scm_cession.scm_cedee.capital_social')} €",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after_pt=2,
        )
        add_paragraph(
            document,
            f"Siège social : {address_display(scm_cedee.siege, 'scm_cession.scm_cedee.siege')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            space_after_pt=2,
        )
        add_paragraph(
            document,
            (
                f"RCS de {required_text(scm_cedee.ville_rcs, 'scm_cession.scm_cedee.ville_rcs')} "
                f"sous le n° {required_text(scm_cedee.numero_rcs, 'scm_cession.scm_cedee.numero_rcs')}"
            ),
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_framed_title(
            document,
            [
                "PROCES-VERBAL DES DECISIONS",
                "DE L'ASSEMBLEE GENERALE EXTRAORDINAIRE",
                f"DU {format_display_date(agrement.date_pv, 'scm_cession.agrement.date_pv')}",
            ],
        )

        add_body_paragraph(
            document,
            f"L'an {required_text(agrement.date_pv_lettres, 'scm_cession.agrement.date_pv_lettres')}",
        )
        add_body_paragraph(
            document,
            (
                f"Les associés de la {required_text(scm_cedee.denomination, 'scm_cession.scm_cedee.denomination')}, "
                f"au capital de {required_text(scm_cedee.capital_social, 'scm_cession.scm_cedee.capital_social')} euros, "
                f"composé de {scm_cedee.nb_parts_total} parts de "
                f"{required_text(scm_cedee.valeur_nominale_part, 'scm_cession.scm_cedee.valeur_nominale_part')} euros chacune, "
                "se sont réunis sur convocation régulière de la gérance au siège de la Société."
            ),
        )
        add_body_paragraph(document, "Sont présents ou représentés :")
        for index, associe in enumerate(scm_cession.associes_presents, start=1):
            parts = associe.parts
            if parts is None:
                raise ValueError("associe present sans parts.")
            add_body_paragraph(
                document,
                f"{index}° {associe_display(associe, f'scm_cession.associes_presents[{index - 1}]')}, détenant {parts.nb} parts sociales",
            )
        add_body_paragraph(
            document,
            "Les associés présents ou représentés disposent ensemble la totalité des parts formant le capital de la société. L'assemblée est habilitée à prendre les décisions extraordinaires.",
        )
        president = scm_cession.associes_presents[2]
        add_body_paragraph(
            document,
            (
                f"{associe_display(president, 'scm_cession.associes_presents[2]')} "
                "préside la séance en qualité de gérant associé."
            ),
        )
        add_body_paragraph(
            document,
            "Le Président dépose et met à la disposition des associés les documents suivants :",
        )
        # Liste A (puces tiret) : documents deposes par le President (retour UAT Rafael).
        for item in [
            "Les copies des convocations des associés ;",
            "Un exemplaire du compromis de cession des parts sociales ;",
            "Le rapport de la gérance ;",
            "Le texte des résolutions proposées.",
        ]:
            add_hyphen_list_item(document, item, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for text in [
            "Le Président déclare que tous les documents prévus par la réglementation et les statuts ont bien été adressés aux associés avec la convocation.",
            "Ils ont été tenus à leur disposition au siège social pendant le délai de quinze jours ayant précédé l'assemblée.",
            "L'assemblée lui donne acte de ses déclarations et reconnaît la validité de la convocation.",
            "Puis le Président rappelle l'ordre du jour :",
        ]:
            add_body_paragraph(document, text)
        # Liste B (puces tiret) : ordre du jour (retour UAT Rafael).
        for item in [
            "Lecture du rapport de la gérance ;",
            f"Agrément d'un nouvel associé, la {required_text(cessionnaire.denomination, 'scm_cession.cessionnaire.denomination')} ;",
            "Modification corrélative des statuts.",
        ]:
            add_hyphen_list_item(document, item, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for text in [
            "Le président donne lecture aux associés du rapport de la gérance.",
            "Une discussion sans débat s'engage entre les associés.",
            "Plus personne ne demandant la parole, le Président met successivement aux voix les résolutions inscrites à l'ordre du jour.",
        ]:
            add_body_paragraph(document, text)

        add_heading(document, "PREMIERE RESOLUTION")
        add_body_paragraph(document, _agrement_resolution(ctx, cessionnaire.denomination or ""))
        add_body_paragraph(document, "Cette résolution est adoptée à l'unanimité")

        add_heading(document, "DEUXIEME RESOLUTION")
        add_body_paragraph(
            document,
            "L'assemblée générale, compte tenu de la résolution qui précède, et sous réserve de la réalisation définitive de la cession, décide, pour tenir compte de la nouvelle répartition du capital, de modifier l'article 7 des statuts qui sera rédigé ainsi :",
        )
        add_body_paragraph(
            document,
            (
                f"« A la suite de son évolution depuis la constitution de la Société, le capital social est fixé à "
                f"{required_text(scm_cedee.capital_social, 'scm_cession.scm_cedee.capital_social')} €. Il est divisé en "
                f"{scm_cedee.nb_parts_total} parts sociales, d'un montant nominal de "
                f"{required_text(scm_cedee.valeur_nominale_part, 'scm_cession.scm_cedee.valeur_nominale_part')} € chacune, "
                f"numérotées de {required_text(scm_cedee.plage_parts_total, 'scm_cession.scm_cedee.plage_parts_total')}, qui ont été attribuées aux associés tant en vertu des apports effectués lors de la constitution de la société, que des augmentations et réductions du capital et des cessions intervenues depuis la constitution de la société, à savoir :"
            ),
        )
        _add_repartition_apres_cession(document, scm_cession.associes_apres_cession)
        add_body_paragraph(
            document,
            f"Total égal au nombre de parts composant le capital social : {scm_cedee.nb_parts_total} parts ».",
        )
        add_body_paragraph(document, "Cette résolution est adoptée à l'unanimité.")

        add_heading(document, "TROISIEME RESOLUTION")
        add_body_paragraph(
            document,
            "L'assemblée générale confère tous pouvoirs au porteur d'une copie ou d'un extrait du présent procès-verbal afin d'accomplir toutes les formalités consécutives aux décisions prises.",
        )
        add_body_paragraph(document, "Cette résolution est adoptée à l'unanimité.")
        add_body_paragraph(
            document,
            "De tout ceci, il a été dressé le présent procès-verbal qui, après lecture, a été signé par la gérance, les associés présents.",
        )
        add_signature_table(document, _signature_rows(scm_cession.signataires_pv))

        return save_clean_document(document, output_dir, OUTPUT_FILENAME)


def _agrement_resolution(ctx: DocumentGenerationContext, cessionnaire_name: str) -> str:
    scm_cession = ctx.scm_cession
    if scm_cession is None or scm_cession.agrement is None:
        raise ValueError("scm_cession.agrement est obligatoire.")
    if ctx.structure == "SELAS":
        return (
            "L'assemblée générale, après avoir entendu lecture du rapport de la gérance et pris connaissance du projet de cession qui a été notifié à la société, décide d'agréer, comme nouvel associé la "
            f"{cessionnaire_name}, dans un délai de {required_text(scm_cession.agrement.delai_mois, 'scm_cession.agrement.delai_mois')} mois à compter de ce jour, soit jusqu'au {required_text(scm_cession.agrement.date_limite, 'scm_cession.agrement.date_limite')}."
        )
    return (
        "L'assemblée générale, après avoir entendu lecture du rapport de la gérance et pris connaissance du projet de cession qui a été notifié à la société, décide d'agréer, comme nouvel associé la "
        f"{cessionnaire_name}, à compter de ce jour."
    )


def _add_repartition_apres_cession(
    document,
    associes: list[ScmCessionAssocie],
) -> None:
    for index, associe in enumerate(associes):
        prefix = f"scm_cession.associes_apres_cession[{index}]"
        parts = associe.parts
        if parts is None:
            raise ValueError(f"{prefix}.parts est obligatoire.")
        add_body_paragraph(document, f"à {associe_display(associe, prefix)},")
        add_body_paragraph(document, f"à concurrence de {parts.nb} parts,")
        add_body_paragraph(
            document,
            f"numérotées de {required_text(parts.plage, f'{prefix}.parts.plage')},",
        )
        add_body_paragraph(document, f"ci                                    {parts.nb} parts")


def _signature_rows(signataires: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    current: list[str] = []
    for signataire in signataires:
        current.append(signataire)
        if len(current) == 2:
            rows.append(current)
            current = []
    if current:
        current.append("")
        rows.append(current)
    return rows
