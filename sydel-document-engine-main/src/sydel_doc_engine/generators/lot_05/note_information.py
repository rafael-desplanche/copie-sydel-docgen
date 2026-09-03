from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    OPERATION_APPORT,
    OPERATION_CESSION,
    capital_after_lines,
    company_siege_display,
    operation_party,
    person_signature,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_note_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_hyphen_list_item,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "note_information.docx"

OPERATION_PHRASES = {
    OPERATION_CESSION: "d'acquerir",
    OPERATION_APPORT: "de recevoir en apport en nature",
}

OPERATION_NOMS = {
    OPERATION_CESSION: "ladite cession",
    OPERATION_APPORT: "ledit apport",
}


class NoteInformationGenerator:
    """Generateur from-scratch de la note d'information SPFPL."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        operation_type = validate_note_context(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        party = operation_party(ctx)
        nb_titres = _operation_nb_titres(ctx)

        docx = new_document()
        add_paragraph(docx, "Note d'informations", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, "Constitution de la Societe", bold=True)
        add_paragraph(
            docx,
            required_text(societe_spfpl.denomination, "societe_spfpl.denomination"),
            bold=True,
        )
        add_paragraph(
            docx,
            (
                f"La {required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')}, "
                "en cours de constitution, dont le siege est situe "
                f"{company_siege_display(societe_spfpl, 'societe_spfpl')}, au capital de "
                f"{required_text(societe_spfpl.capital_social, 'societe_spfpl.capital_social')}, "
                f"prevoit {OPERATION_PHRASES[operation_type]}, des son immatriculation, "
                f"{nb_titres} parts de la "
                f"{required_text(societe_cible.denomination, 'societe_cible.denomination')}, "
                f"{required_text(societe_cible.forme_sociale, 'societe_cible.forme_sociale')} "
                f"de {_profession_reglementee(societe_cible)} "
                f"au capital de {_capital_social_cible(societe_cible)} "
                "divise en "
                f"{required_int(societe_cible.nb_parts_total, 'societe_cible.nb_parts_total')} "
                "parts, dont le siege social est situe "
                f"{company_siege_display(societe_cible, 'societe_cible')}, immatriculee au "
                f"RCS de {required_text(societe_cible.ville_rcs, 'societe_cible.ville_rcs')} "
                f"sous le numero {_numero_rcs_cible(societe_cible)}."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(
            docx,
            (
                f"Apres {OPERATION_NOMS[operation_type]}, le capital de la "
                f"{required_text(societe_cible.denomination, 'societe_cible.denomination')} "
                "sera decompose comme suit :"
            ),
        )
        for line in capital_after_lines(ctx):
            add_hyphen_list_item(docx, line)
        add_paragraph(docx, "________________________", space_before_pt=12)
        add_paragraph(docx, person_signature(party, _party_field_name(operation_type)))
        dirigeant = societe_spfpl.dirigeant
        fonction = required_text(
            dirigeant.fonction if dirigeant else None,
            "societe_spfpl.dirigeant.fonction",
        )
        add_paragraph(
            docx,
            (
                f"{fonction} de la "
                f"{required_text(societe_spfpl.denomination, 'societe_spfpl.denomination')}"
            ),
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _operation_nb_titres(ctx: DocumentGenerationContext) -> int:
    if ctx.operation_titres is not None and ctx.operation_titres.nb_titres is not None:
        return ctx.operation_titres.nb_titres
    if ctx.cession_parts is not None and ctx.cession_parts.nb_parts is not None:
        return ctx.cession_parts.nb_parts
    raise ValueError(
        "operation_titres.nb_titres ou cession_parts.nb_parts est obligatoire pour "
        "CODE-SPFPL-AGR-INFO-001."
    )


def _party_field_name(operation_type: str) -> str:
    if operation_type == OPERATION_CESSION:
        return "cedant"
    return "apporteur"


def _profession_reglementee(societe_cible) -> str:
    return required_text(
        societe_cible.profession_reglementee,
        "societe_cible.profession_reglementee",
    )


def _capital_social_cible(societe_cible) -> str:
    return required_text(
        societe_cible.capital_social,
        "societe_cible.capital_social",
    )


def _numero_rcs_cible(societe_cible) -> str:
    return required_text(
        societe_cible.numero_rcs,
        "societe_cible.numero_rcs",
    )
