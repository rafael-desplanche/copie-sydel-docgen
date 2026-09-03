from __future__ import annotations

from collections.abc import Mapping, Sequence
from datetime import date
from pathlib import Path
from typing import Any

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from sydel_doc_engine.domain.models import (
    Address,
    Associe,
    Company,
    DocumentGenerationContext,
    LocauxContext,
    PartieFraisCommuns,
    PraticienScm,
    ScmRepresentant,
    ScmSocietePartie,
)
from sydel_doc_engine.generators.lot_05.scm_satellites_templates import TemplateBlock
from sydel_doc_engine.rendering.docx_builder import add_paragraph, new_document

DOCUMENT_CODE = "CODE-SCM-SAT-DOCX-001"
SCM_STRUCTURE = "SCM"
TWO_PARTIES = 2


def generate_from_template(
    *,
    ctx: DocumentGenerationContext,
    output_dir: Path,
    output_filename: str,
    blocks: Sequence[TemplateBlock],
    replacements: Mapping[str, str],
) -> Path:
    document = new_document()
    for kind, payload in blocks:
        if kind == "p":
            _add_rendered_paragraph(document, _replace_placeholders(str(payload), replacements))
        elif kind == "table":
            _add_rendered_table(document, payload, replacements)
        else:
            raise ValueError(f"Bloc template inconnu pour {DOCUMENT_CODE}: {kind}")

    full_text = "\n".join(
        [paragraph.text for paragraph in document.paragraphs]
        + [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
    )
    if "[" in full_text or "]" in full_text:
        raise ValueError(f"placeholder source residuel dans le rendu {DOCUMENT_CODE}.")

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / output_filename
    document.save(output_path)
    return output_path


def validate_scm_satellite_enabled(ctx: DocumentGenerationContext, satellite_field: str) -> None:
    if ctx.structure != SCM_STRUCTURE:
        raise ValueError(f"dossier.structure doit etre SCM pour {DOCUMENT_CODE}.")
    if ctx.dossier_options is None or not ctx.dossier_options.scm_satellites:
        raise ValueError(f"dossier.options.scm_satellites doit etre vrai pour {DOCUMENT_CODE}.")
    if ctx.scm_satellites is None or not getattr(ctx.scm_satellites, satellite_field):
        raise ValueError(f"scm_satellites.{satellite_field} doit etre vrai pour {DOCUMENT_CODE}.")


def societe_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    company = _required_company(ctx.societe)
    return {
        "[denomination_societe]": _required_text(company.denomination, "societe.denomination"),
        "[forme_sociale]": _company_forme_juridique(company),
        "[capital_social]": _required_text(
            company.capital_social or company.capital,
            "societe.capital_social",
        ),
        "[adresse_siege]": _address_display(company.siege, "societe.siege"),
        "[ville_rcs]": _required_text(company.ville_rcs, "societe.ville_rcs"),
        "[numero_rcs]": _required_text(company.numero_rcs, "societe.numero_rcs"),
        "[nb_parts_sociales]": str(
            _required_value(company.nb_parts_total, "societe.nb_parts_total")
        ),
    }


def pacte_associes_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    associes = _required_two_associes(ctx.associes)
    if ctx.pacte_associes is None:
        raise ValueError(f"pacte_associes est obligatoire pour {DOCUMENT_CODE}.")
    replacements = societe_replacements(ctx)
    replacements.update(
        {
            "[civilite_personne_1]": associes[0].civilite_affichage,
            "[prenom_personne_1]": associes[0].prenom,
            "[nom_personne_1]": associes[0].nom,
            "[civilite_personne_2]": associes[1].civilite_affichage,
            "[prenom_personne_2]": associes[1].prenom,
            "[nom_personne_2]": associes[1].nom,
            "[ville_tribunal]": _required_text(
                ctx.pacte_associes.ville_tribunal,
                "pacte_associes.ville_tribunal",
            ),
            "[lieu_signature]": _required_text(ctx.signature.lieu, "signature.lieu"),
            "[date_signature]": _format_display_date(ctx.signature.date, "signature.date"),
        }
    )
    return replacements


def liste_depenses_communes_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    company = _required_company(ctx.societe)
    associes = _required_two_associes(ctx.associes)
    return {
        "[denomination_societe]": _required_text(company.denomination, "societe.denomination"),
        "[forme_sociale]": _company_forme_juridique(company),
        "[capital_social]": _required_text(
            company.capital_social or company.capital,
            "societe.capital_social",
        ),
        "[adresse_siege]": _address_display(company.siege, "societe.siege"),
        "[ville_rcs]": _required_text(company.ville_rcs, "societe.ville_rcs"),
        "[prenom_personne_1]": associes[0].prenom,
        "[nom_personne_1]": associes[0].nom,
        "[prenom_personne_2]": associes[1].prenom,
        "[nom_personne_2]": associes[1].nom,
    }


def contrat_frais_communs_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    parties = required_two_parties(ctx)
    locals_context = required_locaux(ctx.locaux)
    if ctx.frais_communs is None:
        raise ValueError(f"frais_communs est obligatoire pour {DOCUMENT_CODE}.")

    replacements = _party_replacements(parties[0], 1)
    replacements.update(_party_replacements(parties[1], 2))
    replacements.update(
        {
            "[adresse_locaux]": _required_text(
                locals_context.adresse_affichee,
                "locaux.adresse_affichee",
            ),
            "[date_effet_contrat]": _format_display_date(
                ctx.frais_communs.date_effet_contrat,
                "frais_communs.date_effet_contrat",
            ),
            "[lieu_signature]": _required_text(ctx.signature.lieu, "signature.lieu"),
            "[date_signature]": _format_display_date(ctx.signature.date, "signature.date"),
        }
    )
    return replacements


def reglement_interieur_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    company = _required_company(ctx.societe)
    parties = required_two_parties(ctx)
    practitioners = required_two_praticiens(ctx.praticiens)
    locals_context = required_locaux(ctx.locaux)
    if ctx.reglement_interieur is None:
        raise ValueError(f"reglement_interieur est obligatoire pour {DOCUMENT_CODE}.")

    _validate_same_forme_juridique(parties)
    replacements = {
        "[denomination_societe]": _required_text(company.denomination, "societe.denomination"),
        "[forme_sociale]": _required_text(
            parties[0].societe.forme_juridique if parties[0].societe else None,
            "parties_frais_communs[0].societe.forme_juridique",
        ),
    }
    replacements.update(_party_replacements(parties[0], 1, reglement=True))
    replacements.update(_party_replacements(parties[1], 2, reglement=True))
    replacements.update(
        {
            "[adresse_locaux]": _required_text(
                locals_context.adresse_affichee,
                "locaux.adresse_affichee",
            ),
            "[seuil_depense_commune]": _required_text(
                ctx.reglement_interieur.seuil_depense_commune,
                "reglement_interieur.seuil_depense_commune",
            ),
            "[annee_reference_charges]": _required_text(
                ctx.reglement_interieur.annee_reference_charges,
                "reglement_interieur.annee_reference_charges",
            ),
            "[date_fin_gestion_administrative]": _format_display_date(
                ctx.reglement_interieur.date_fin_gestion_administrative,
                "reglement_interieur.date_fin_gestion_administrative",
            ),
            "[date_attribution_responsabilites]": _format_display_date(
                ctx.reglement_interieur.date_attribution_responsabilites,
                "reglement_interieur.date_attribution_responsabilites",
            ),
            "[identite_praticien_1]": _required_text(
                practitioners[0].identite_affichee,
                "praticiens[0].identite_affichee",
            ),
            "[identite_praticien_2]": _required_text(
                practitioners[1].identite_affichee,
                "praticiens[1].identite_affichee",
            ),
            "[telephone_praticien_1]": _required_text(
                practitioners[0].telephone,
                "praticiens[0].telephone",
            ),
            "[telephone_praticien_2]": _required_text(
                practitioners[1].telephone,
                "praticiens[1].telephone",
            ),
            "[lieu_signature]": _required_text(ctx.signature.lieu, "signature.lieu"),
            "[date_signature]": _format_display_date(ctx.signature.date, "signature.date"),
        }
    )
    return replacements


def required_two_parties(ctx: DocumentGenerationContext) -> list[PartieFraisCommuns]:
    if len(ctx.parties_frais_communs) != TWO_PARTIES:
        raise ValueError(
            "parties_frais_communs doit contenir exactement deux parties "
            f"pour {DOCUMENT_CODE}."
        )
    for index, partie in enumerate(ctx.parties_frais_communs):
        _required_party(partie, index)
    return list(ctx.parties_frais_communs)


def required_two_praticiens(praticiens: Sequence[PraticienScm]) -> list[PraticienScm]:
    if len(praticiens) != TWO_PARTIES:
        raise ValueError(
            f"praticiens doit contenir exactement deux praticiens pour {DOCUMENT_CODE}."
        )
    return list(praticiens)


def required_locaux(locaux: LocauxContext | None) -> LocauxContext:
    if locaux is None:
        raise ValueError(f"locaux est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(locaux.adresse_affichee, "locaux.adresse_affichee")
    return locaux


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return company


def _required_two_associes(associes: Sequence[Associe]) -> list[Associe]:
    if len(associes) != TWO_PARTIES:
        raise ValueError(f"associes doit contenir exactement deux associes pour {DOCUMENT_CODE}.")
    return list(associes)


def _required_party(partie: PartieFraisCommuns, index: int) -> None:
    prefix = f"parties_frais_communs[{index}]"
    if partie.societe is None:
        raise ValueError(f"{prefix}.societe est obligatoire pour {DOCUMENT_CODE}.")
    if partie.representant is None:
        raise ValueError(f"{prefix}.representant est obligatoire pour {DOCUMENT_CODE}.")
    _required_societe_partie(partie.societe, f"{prefix}.societe")
    _required_representant(partie.representant, f"{prefix}.representant")


def _required_societe_partie(societe: ScmSocietePartie, prefix: str) -> None:
    _required_text(societe.denomination, f"{prefix}.denomination")
    _required_text(societe.forme_juridique, f"{prefix}.forme_juridique")
    _required_text(societe.capital_social, f"{prefix}.capital_social")
    _address_display(societe.siege, f"{prefix}.siege")
    _required_text(societe.ville_rcs, f"{prefix}.ville_rcs")
    _required_text(societe.numero_rcs, f"{prefix}.numero_rcs")


def _required_representant(representant: ScmRepresentant, prefix: str) -> None:
    _required_text(representant.civilite_affichage, f"{prefix}.civilite_affichage")
    _required_text(representant.prenom, f"{prefix}.prenom")
    _required_text(representant.nom, f"{prefix}.nom")
    _required_text(representant.fonction, f"{prefix}.fonction")


def _party_replacements(
    partie: PartieFraisCommuns,
    source_index: int,
    *,
    reglement: bool = False,
) -> dict[str, str]:
    index = source_index - 1
    prefix = f"parties_frais_communs[{index}]"
    societe = partie.societe
    representant = partie.representant
    if societe is None or representant is None:
        raise ValueError(f"{prefix} est incomplet pour {DOCUMENT_CODE}.")
    replacements = {
        f"[denomination_societe_{source_index}]": _required_text(
            societe.denomination,
            f"{prefix}.societe.denomination",
        ),
        f"[forme_sociale_societe_{source_index}]": _required_text(
            societe.forme_juridique,
            f"{prefix}.societe.forme_juridique",
        ),
        f"[capital_social_societe_{source_index}]": _required_text(
            societe.capital_social,
            f"{prefix}.societe.capital_social",
        ),
        f"[adresse_siege_societe_{source_index}]": _address_display(
            societe.siege,
            f"{prefix}.societe.siege",
        ),
        f"[ville_rcs_societe_{source_index}]": _required_text(
            societe.ville_rcs,
            f"{prefix}.societe.ville_rcs",
        ),
        f"[numero_rcs_societe_{source_index}]": _required_text(
            societe.numero_rcs,
            f"{prefix}.societe.numero_rcs",
        ),
        f"[civilite_representant_societe_{source_index}]": _required_text(
            representant.civilite_affichage,
            f"{prefix}.representant.civilite_affichage",
        ),
        f"[prenom_representant_societe_{source_index}]": _required_text(
            representant.prenom,
            f"{prefix}.representant.prenom",
        ),
        f"[nom_representant_societe_{source_index}]": _required_text(
            representant.nom,
            f"{prefix}.representant.nom",
        ),
        f"[fonction_representant_societe_{source_index}]": _required_text(
            representant.fonction,
            f"{prefix}.representant.fonction",
        ),
    }
    if reglement:
        replacements.update(
            {
                f"[titre_representant_societe_{source_index}]": _required_text(
                    representant.titre_affichage,
                    f"{prefix}.representant.titre_affichage",
                ),
                f"[identite_representant_societe_{source_index}]": _representant_identity(
                    representant,
                    f"{prefix}.representant",
                ),
            }
        )
    return replacements


def _validate_same_forme_juridique(parties: Sequence[PartieFraisCommuns]) -> None:
    formes = [
        _required_text(
            partie.societe.forme_juridique if partie.societe else None,
            f"parties_frais_communs[{index}].societe.forme_juridique",
        )
        for index, partie in enumerate(parties)
    ]
    if formes[0] != formes[1]:
        raise ValueError(
            "les deux parties du reglement interieur doivent avoir la meme forme_juridique "
            f"pour {DOCUMENT_CODE}."
        )


def _representant_identity(representant: ScmRepresentant, prefix: str) -> str:
    if representant.identite_affichee and representant.identite_affichee.strip():
        return representant.identite_affichee.strip()
    return (
        f"{_required_text(representant.prenom, f'{prefix}.prenom')} "
        f"{_required_text(representant.nom, f'{prefix}.nom')}"
    )


def _company_forme_juridique(company: Company) -> str:
    return _required_text(
        company.forme_juridique
        or company.forme_sociale
        or company.forme_sociale_complete
        or company.forme_sociale_libelle_long
        or company.forme_sociale_affichage,
        "societe.forme_juridique",
    )


def _address_display(address: Address | None, field_name: str) -> str:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if address.adresse_affichee and address.adresse_affichee.strip():
        return address.adresse_affichee.strip()
    return (
        f"{_required_text(address.num_voie, f'{field_name}.num_voie')} "
        f"{_required_text(address.voie, f'{field_name}.voie')}, "
        f"{_required_text(address.cp, f'{field_name}.cp')} "
        f"{_required_text(address.ville, f'{field_name}.ville')}"
    )


def _add_rendered_paragraph(document: Any, text: str) -> None:
    if not text.strip():
        return
    stripped = text.strip()
    upper = stripped.upper()
    if (
        stripped.startswith("TITRE ")
        or stripped.startswith("ARTICLE ")
        or stripped.startswith("Article ")
        or upper in {
            "PACTE D'ASSOCIES",
            "PACTE D’ASSOCIES",
            "CONTRAT D'EXERCICE PROFESSIONNEL",
            "A FRAIS COMMUNS",
            "REGLEMENT INTERIEUR DE LA SOCIETE CIVILE DE MOYENS",
        }
    ):
        add_paragraph(
            document,
            stripped,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_before_pt=8 if stripped.startswith(("ARTICLE ", "Article ")) else 0,
        )
    elif upper == stripped and len(stripped) <= 90:
        add_paragraph(document, stripped, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    else:
        add_paragraph(document, stripped, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _add_rendered_table(
    document: Any,
    rows_payload: str | list[list[str]],
    replacements: Mapping[str, str],
) -> None:
    if not isinstance(rows_payload, list):
        raise ValueError(f"Payload table invalide pour {DOCUMENT_CODE}.")
    column_count = max(len(row) for row in rows_payload)
    table = document.add_table(rows=0, cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_values in rows_payload:
        cells = table.add_row().cells
        for index in range(column_count):
            value = row_values[index] if index < len(row_values) else ""
            cells[index].text = _replace_placeholders(value, replacements)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)


def _replace_placeholders(text: str, replacements: Mapping[str, str]) -> str:
    rendered = text
    for placeholder, value in replacements.items():
        rendered = rendered.replace(placeholder, value)
    return rendered


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not str(value).strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return str(value).strip()


def _required_value(value: int | str | None, field_name: str) -> int | str:
    if value is None or not str(value).strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _format_display_date(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return _required_text(value, field_name)
