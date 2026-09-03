from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    capital_before_lines,
    company_siege_display,
    ordre_sentence,
    person_identity_sentence,
    person_signature,
    required_cedant,
    required_cession_parts,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_cession_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_hyphen_list_item,
    add_paragraph,
    add_signature_lines,
    new_document,
)

OUTPUT_FILENAME = "acte_cession_parts_spfpl.docx"


class ActeCessionPartsSpfplGenerator:
    """Generateur from-scratch de l'acte de cession de parts SPFPL."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_cession_context(ctx)
        cedant = required_cedant(ctx)
        cession_parts = required_cession_parts(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        representant = societe_spfpl.representant
        if representant is None:
            raise ValueError("societe_spfpl.representant est obligatoire.")
        spfpl_name = required_text(societe_spfpl.denomination, "societe_spfpl.denomination")
        spfpl_forme = required_text(societe_spfpl.forme_sociale, "societe_spfpl.forme_sociale")
        spfpl_capital = required_text(societe_spfpl.capital_social, "societe_spfpl.capital_social")
        spfpl_rcs = required_text(societe_spfpl.ville_rcs, "societe_spfpl.ville_rcs")
        spfpl_numero_rcs = required_text(societe_spfpl.numero_rcs, "societe_spfpl.numero_rcs")
        rep_civilite = required_text(
            representant.civilite_affichage,
            "societe_spfpl.representant.civilite_affichage",
        )
        rep_civilite_courte = required_text(
            representant.civilite_courte,
            "societe_spfpl.representant.civilite_courte",
        )
        rep_prenom = required_text(representant.prenom, "societe_spfpl.representant.prenom")
        rep_nom = required_text(representant.nom, "societe_spfpl.representant.nom")
        rep_fonction = required_text(
            representant.fonction,
            "societe_spfpl.representant.fonction",
        )
        cible_name = required_text(societe_cible.denomination, "societe_cible.denomination")
        cible_forme_complete = required_text(
            societe_cible.forme_sociale_complete,
            "societe_cible.forme_sociale_complete",
        )
        cible_capital = required_text(societe_cible.capital_social, "societe_cible.capital_social")
        cible_rcs = required_text(societe_cible.ville_rcs, "societe_cible.ville_rcs")
        cible_numero_rcs = required_text(societe_cible.numero_rcs, "societe_cible.numero_rcs")
        prix_unitaire_lettres = required_text(
            cession_parts.prix_unitaire_lettres,
            "cession_parts.prix_unitaire_lettres",
        )

        docx = new_document()
        add_paragraph(docx, "Cession de parts", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, "ENTRE", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, person_signature(cedant, "cedant"), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(docx, "ET", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(
            docx,
            f"La Societe {spfpl_name}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        add_paragraph(docx, "ENTRE LES SOUSSIGNES :", bold=True, space_before_pt=10)
        add_paragraph(docx, person_identity_sentence(cedant, "cedant"))
        add_paragraph(docx, ordre_sentence(cedant, "cedant"))
        add_paragraph(docx, 'Ci-apres denomme "LE CEDANT",')
        add_paragraph(docx, "D'une part,")
        add_paragraph(
            docx,
            f"La Societe {spfpl_name}",
        )
        add_paragraph(docx, spfpl_forme)
        add_paragraph(docx, f"Au capital de {spfpl_capital}")
        add_paragraph(
            docx,
            "Immatriculee au RCS de "
            f"{spfpl_rcs} sous le numero {spfpl_numero_rcs}",
        )
        add_paragraph(
            docx,
            f"Siege social : {company_siege_display(societe_spfpl, 'societe_spfpl')}",
        )
        add_paragraph(
            docx,
            "Representee aux presentes par "
            f"{rep_civilite} {rep_prenom} {rep_nom} en sa qualite de {rep_fonction} "
            "et ayant tout pouvoir a l'effet des presentes.",
        )
        add_paragraph(docx, 'Ci-apres denommee "LE CESSIONNAIRE",')
        add_paragraph(docx, "D'autre part,")

        add_paragraph(docx, "IL EST PREALABLEMENT EXPOSE CE QUI SUIT :", bold=True)
        add_paragraph(
            docx,
            "La Societe "
            f"{cible_name} est une {cible_forme_complete}, "
            f"au capital social de {cible_capital} "
            "divise en "
            f"{required_int(societe_cible.nb_parts_total, 'societe_cible.nb_parts_total')} "
            "parts sociales, immatriculee au Registre du Commerce et des Societes de "
            f"{cible_rcs} sous le numero {cible_numero_rcs}, "
            f"dont le siege est situe {company_siege_display(societe_cible, 'societe_cible')}.",
        )
        add_paragraph(docx, "Le capital social est reparti a ce jour comme suit :")
        for line in capital_before_lines(ctx):
            add_hyphen_list_item(docx, line)

        add_paragraph(docx, "CECI EXPOSE, IL EST CONVENU CE QUI SUIT :", bold=True)
        add_paragraph(docx, "OBJET DU CONTRAT : CESSION DE PARTS", bold=True)
        add_paragraph(
            docx,
            "Par les presentes, le Cedant cede ce jour, sous les garanties ordinaires "
            "de fait et de droit en la matiere, ainsi que celles consenties dans les "
            "presentes, a l'Acquereur, qui accepte, la pleine propriete de "
            f"{required_text(cession_parts.nb_parts_lettres, 'cession_parts.nb_parts_lettres')} "
            f"({required_int(cession_parts.nb_parts, 'cession_parts.nb_parts')}) parts "
            "qu'il detient.",
        )
        add_paragraph(docx, "PRIX ET MODALITES DE PAIEMENT DES PARTS", bold=True)
        add_paragraph(
            docx,
            "La cession a lieu moyennant le prix de "
            f"{prix_unitaire_lettres} "
            f"({required_text(cession_parts.prix_unitaire, 'cession_parts.prix_unitaire')}) "
            "euro par part cedee, soit un prix de "
            f"{required_text(cession_parts.prix_total, 'cession_parts.prix_total')} euros "
            f"({_cession_prix_total_lettres(cession_parts)}), "
            f"a payer par la Societe {spfpl_name}.",
        )
        add_paragraph(docx, "Le prix est paye ce jour par le moyen d'un cheque ou d'un virement.")
        add_paragraph(docx, "COMMUNICATION DU PRESENT CONTRAT AU CONSEIL DE L'ORDRE", bold=True)
        add_paragraph(
            docx,
            "Le present contrat sera, sans delai, communique au Conseil departemental "
            "de l'Ordre en vue de ses observations eventuelles.",
        )
        add_paragraph(docx, "FRAIS", bold=True)
        add_paragraph(
            docx,
            "Les frais et droits d'enregistrement afferents a la presente cession et "
            "tous les frais qui en sont la consequence seront supportes par le "
            "CESSIONNAIRE qui s'y oblige dans la mesure ou ces frais et droits se "
            "rattacheront a la cession d'action consentie.",
        )
        add_paragraph(docx, "CONVENTION SUR LA PREUVE - SIGNATURE ELECTRONIQUE", bold=True)
        add_paragraph(
            docx,
            "Les Parties consentent expressement la faculte de proceder a la signature "
            "du present acte par le systeme de signature electronique.",
        )

        add_paragraph(docx, f"Fait a {ctx.signature.lieu}")
        add_paragraph(docx, f"Le {ctx.signature.date.strftime('%d/%m/%Y')}")
        add_paragraph(
            docx,
            "En "
            f"{_nombre_exemplaires(cession_parts.nombre_exemplaires_lettres, ctx)} "
            "exemplaires originaux,",
        )
        add_signature_lines(
            docx,
            [
                f"Dr {person_signature(cedant, 'cedant')}",
                f"La societe {spfpl_name}",
                f"Representee par {rep_civilite_courte} {rep_prenom} {rep_nom}",
            ],
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _nombre_exemplaires(value: str | None, ctx: DocumentGenerationContext) -> str:
    if value:
        return value
    if ctx.document and ctx.document.nombre_exemplaires_lettres:
        return ctx.document.nombre_exemplaires_lettres
    raise ValueError("cession_parts.nombre_exemplaires_lettres est obligatoire.")


def _cession_prix_total_lettres(cession_parts) -> str:
    return required_text(
        cession_parts.prix_total_lettres,
        "cession_parts.prix_total_lettres",
    )
