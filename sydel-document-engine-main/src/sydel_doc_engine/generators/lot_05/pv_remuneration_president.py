from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.sas_satellites_common import (
    DOCUMENT_CODE,
    address_display,
    format_display_date,
    person_name,
    person_signature,
    personal_address_for_pv,
    required_actionnaire_unique,
    required_president,
    required_societe_spfpl,
    required_text,
    validate_sas_satellite_scope,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_company_identity_block,
    add_framed_title,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "pv_remuneration_president.docx"
REMUNERATION_TYPE_ABSENCE = "absence_remuneration"


class PvRemunerationPresidentGenerator:
    """Generateur from-scratch du PV remuneration president SAS V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        data = _ResolvedPvRemunerationPresident.from_context(ctx)
        document = new_document()

        add_company_identity_block(
            document,
            [
                data.denomination,
                data.forme_sociale,
                f"Au capital de {data.capital_social} euros",
                f"Siège social : {data.adresse_siege}",
                f"En cours d'immatriculation au RCS de {data.ville_rcs}",
            ],
        )
        add_framed_title(
            document,
            [
                "PROCES-VERBAL DES DECISIONS",
                "DE L'ASSOCIE UNIQUE",
                f"DU {data.date_signature}",
            ],
        )
        add_paragraph(document, data.actionnaire_nom)
        add_paragraph(document, f"Demeurant {data.adresse_actionnaire}.")
        add_paragraph(
            document,
            f"{data.qualite_associe} et {data.fonction_president} de la Société "
            f"{data.denomination} en cours de formation.",
        )
        add_paragraph(document, "a pris la décision suivante :")
        add_paragraph(document, f"Fixation de la rémunération du {data.fonction_president}")
        add_paragraph(document, "DECISION UNIQUE", bold=True)
        add_paragraph(
            document,
            f"{data.actionnaire_nom}, {data.qualite_associe}, décide qu'il ne percevra "
            "aucune rémunération au titre de son mandat de "
            f"{data.fonction_president}, à compter de son immatriculation, et ce, "
            f"jusqu'au {data.date_cloture_premier_exercice} inclus, date de la clôture "
            "du premier exercice social.",
        )
        add_paragraph(
            document,
            "Il pourra donc prétendre au remboursement sur justification de ses frais de "
            "représentation et de déplacement.",
        )
        add_paragraph(
            document,
            "De tout ce que dessus, l'associé unique a dressé et signé le présent "
            "procès-verbal.",
        )
        add_paragraph(document, f"Fait à {data.lieu_signature} en trois exemplaires")
        add_paragraph(document, "________________", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(document, data.signature_nom, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


class _ResolvedPvRemunerationPresident:
    def __init__(
        self,
        *,
        denomination: str,
        forme_sociale: str,
        capital_social: str,
        adresse_siege: str,
        ville_rcs: str,
        date_signature: str,
        actionnaire_nom: str,
        adresse_actionnaire: str,
        qualite_associe: str,
        fonction_president: str,
        date_cloture_premier_exercice: str,
        lieu_signature: str,
        signature_nom: str,
    ) -> None:
        self.denomination = denomination
        self.forme_sociale = forme_sociale
        self.capital_social = capital_social
        self.adresse_siege = adresse_siege
        self.ville_rcs = ville_rcs
        self.date_signature = date_signature
        self.actionnaire_nom = actionnaire_nom
        self.adresse_actionnaire = adresse_actionnaire
        self.qualite_associe = qualite_associe
        self.fonction_president = fonction_president
        self.date_cloture_premier_exercice = date_cloture_premier_exercice
        self.lieu_signature = lieu_signature
        self.signature_nom = signature_nom

    @classmethod
    def from_context(
        cls,
        ctx: DocumentGenerationContext,
    ) -> _ResolvedPvRemunerationPresident:
        validate_sas_satellite_scope(ctx)
        societe = required_societe_spfpl(ctx)
        actionnaire = required_actionnaire_unique(ctx)
        president = required_president(ctx)
        if ctx.exercice_social is None:
            raise ValueError(f"exercice_social est obligatoire pour {DOCUMENT_CODE}.")
        if ctx.remuneration_president is None:
            raise ValueError(
                f"remuneration_president est obligatoire pour {DOCUMENT_CODE}."
            )

        remuneration_type = required_text(
            ctx.remuneration_president.type,
            "remuneration_president.type",
        )
        if remuneration_type != REMUNERATION_TYPE_ABSENCE:
            raise ValueError(
                "remuneration_president.type doit valoir absence_remuneration pour "
                f"{DOCUMENT_CODE}."
            )
        date_cloture = required_text(
            ctx.exercice_social.date_cloture_premier_exercice,
            "exercice_social.date_cloture_premier_exercice",
        )
        date_fin_non_remuneree = required_text(
            ctx.remuneration_president.date_fin_non_remuneree,
            "remuneration_president.date_fin_non_remuneree",
        )
        if date_fin_non_remuneree != date_cloture:
            raise ValueError(
                "remuneration_president.date_fin_non_remuneree doit correspondre a "
                "exercice_social.date_cloture_premier_exercice pour "
                f"{DOCUMENT_CODE}."
            )

        return cls(
            denomination=required_text(societe.denomination, "societe_spfpl.denomination"),
            forme_sociale=required_text(societe.forme_sociale, "societe_spfpl.forme_sociale"),
            capital_social=required_text(societe.capital_social, "societe_spfpl.capital_social"),
            adresse_siege=address_display(societe.siege, "societe_spfpl.siege"),
            ville_rcs=required_text(societe.ville_rcs, "societe_spfpl.ville_rcs"),
            date_signature=format_display_date(ctx.signature.date, "signature.date"),
            actionnaire_nom=person_name(actionnaire, "actionnaire_unique"),
            adresse_actionnaire=personal_address_for_pv(actionnaire, "actionnaire_unique"),
            qualite_associe=required_text(
                actionnaire.qualite_associe,
                "actionnaire_unique.qualite_associe",
            ),
            fonction_president=required_text(president.fonction, "president.fonction"),
            date_cloture_premier_exercice=date_cloture,
            lieu_signature=ctx.signature.lieu,
            signature_nom=person_signature(actionnaire, "actionnaire_unique"),
        )
