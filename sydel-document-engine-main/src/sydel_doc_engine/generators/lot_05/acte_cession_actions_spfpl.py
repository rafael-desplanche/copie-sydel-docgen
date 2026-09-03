# ruff: noqa: E501
from __future__ import annotations

from decimal import Decimal, InvalidOperation
from pathlib import Path

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import CessionActions, DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    SPFPL_CESSION_STRUCTURE,
    associe_display_name,
    company_siege_display,
    format_display_date,
    person_address_display,
    person_display,
    required_cedant,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_cession_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_paragraph,
    add_signature_lines,
    new_document,
)

OUTPUT_FILENAME = "acte_cession_actions_spfpl.docx"
DOCUMENT_CODE = "CODE-ACTE-ACTIONS-001"
SOURCE_PAYMENT_MODE = "credit_bancaire_comptant_cheque_banque"


class ActeCessionActionsSpfplGenerator:
    """Generateur from-scratch de l'acte de cession d'actions SPFPL V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        _validate_actions_context(ctx)
        cedant = required_cedant(ctx)
        cession_actions = _required_cession_actions(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        representant = societe_spfpl.representant
        if representant is None:
            raise ValueError("societe_spfpl.representant est obligatoire.")

        spfpl_name = required_text(societe_spfpl.denomination, "societe_spfpl.denomination")
        spfpl_forme = required_text(societe_spfpl.forme_sociale, "societe_spfpl.forme_sociale")
        spfpl_capital = required_text(
            societe_spfpl.capital_social,
            "societe_spfpl.capital_social",
        )
        spfpl_rcs = required_text(societe_spfpl.ville_rcs, "societe_spfpl.ville_rcs")
        spfpl_numero_rcs = required_text(
            societe_spfpl.numero_rcs,
            "societe_spfpl.numero_rcs",
        )
        spfpl_ordre_departement = required_text(
            societe_spfpl.departement_inscription_ordre,
            "societe_spfpl.departement_inscription_ordre",
        )
        rep_civilite = required_text(
            representant.civilite_affichage,
            "societe_spfpl.representant.civilite_affichage",
        )
        rep_civilite_courte = required_text(
            representant.civilite_courte,
            "societe_spfpl.representant.civilite_courte",
        )
        rep_prenom = required_text(representant.prenom, "societe_spfpl.representant.prenom")
        rep_nom = required_text(representant.nom, "societe_spfpl.representant.nom")
        rep_fonction = required_text(
            representant.fonction,
            "societe_spfpl.representant.fonction",
        )
        cible_name = required_text(societe_cible.denomination, "societe_cible.denomination")
        cible_forme_complete = required_text(
            societe_cible.forme_sociale_complete,
            "societe_cible.forme_sociale_complete",
        )
        cible_capital = required_text(
            societe_cible.capital_social,
            "societe_cible.capital_social",
        )
        cible_actions_total = required_int(
            societe_cible.nb_actions_total,
            "societe_cible.nb_actions_total",
        )
        cible_valeur_action_lettres = required_text(
            societe_cible.valeur_nominale_action_lettres,
            "societe_cible.valeur_nominale_action_lettres",
        )
        cible_rcs = required_text(societe_cible.ville_rcs, "societe_cible.ville_rcs")
        cible_numero_rcs = required_text(societe_cible.numero_rcs, "societe_cible.numero_rcs")
        cible_ordre_departement = required_text(
            societe_cible.departement_inscription_ordre,
            "societe_cible.departement_inscription_ordre",
        )
        cible_profession = required_text(
            societe_cible.profession_reglementee,
            "societe_cible.profession_reglementee",
        )
        cible_profession_pluriel = required_text(
            societe_cible.profession_reglementee_pluriel,
            "societe_cible.profession_reglementee_pluriel",
        )
        presentation_dirigeants = required_text(
            societe_cible.presentation_dirigeants,
            "societe_cible.presentation_dirigeants",
        )
        nombre_exemplaires = _nombre_exemplaires(cession_actions, ctx)
        cedant_titre_signature = required_text(
            cession_actions.titre_signature_cedant,
            "cession_actions.titre_signature_cedant",
        )

        docx = new_document()
        _add_centered_heading(docx, "Cession d'actions")
        add_paragraph(docx, "ENTRE", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, person_display(cedant, "cedant"), alignment=WD_ALIGN_PARAGRAPH.CENTER)
        add_paragraph(docx, "ET", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(
            docx,
            f"La Société {spfpl_name}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        add_paragraph(docx, "ENTRE LES SOUSSIGNES :", bold=True, space_before_pt=10)
        add_paragraph(
            docx,
            (
                f"- {person_display(cedant, 'cedant')}, "
                f"{required_text(cedant.profession, 'cedant.profession')}, né le "
                f"{format_display_date(cedant.date_naissance, 'cedant.date_naissance')} "
                f"à {required_text(cedant.ville_naissance, 'cedant.ville_naissance')} "
                f"({required_text(cedant.departement_naissance, 'cedant.departement_naissance')}), "
                f"de nationalité {required_text(cedant.nationalite, 'cedant.nationalite')}, "
                f"demeurant {person_address_display(cedant, 'cedant')}, "
                f"{required_text(cedant.situation_maritale, 'cedant.situation_maritale')} "
                "sous le régime de "
                f"{required_text(cedant.regime_matrimonial, 'cedant.regime_matrimonial')} "
                f"avec {_conjoint_display(ctx)}, inscrit au tableau de l'Ordre des "
                f"{required_text(cedant.profession_reglementee_pluriel, 'cedant.profession_reglementee_pluriel')} "
                f"du {required_text(cedant.ordre.departement if cedant.ordre else None, 'cedant.ordre.departement')}, "
                "et sous le numéro RPPS "
                f"{required_text(cedant.ordre.numero_rpps if cedant.ordre else None, 'cedant.ordre.numero_rpps')}."
            ),
        )
        add_paragraph(docx, 'Ci-après dénommé "LE CEDANT",')
        add_paragraph(docx, "D'une part,")
        add_paragraph(docx, "ET", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(docx, f"- La Société {spfpl_name}")
        add_paragraph(
            docx,
            (
                f"{spfpl_forme} inscrite au tableau de l'Ordre des "
                f"{cible_profession_pluriel} du {spfpl_ordre_departement}."
            ),
        )
        add_paragraph(docx, f"Au capital de {spfpl_capital}")
        add_paragraph(
            docx,
            f"Immatriculée au RCS de {spfpl_rcs} sous le numéro {spfpl_numero_rcs}",
        )
        add_paragraph(docx, f"Siège social : {company_siege_display(societe_spfpl, 'societe_spfpl')}")
        add_paragraph(
            docx,
            (
                "Représentée aux présentes par "
                f"{rep_civilite} {rep_prenom} {rep_nom} en sa qualité de {rep_fonction} "
                "et ayant tout pouvoir à l'effet des présentes."
            ),
        )
        add_paragraph(docx, 'Ci-après dénommée "LE CESSIONNAIRE",')
        add_paragraph(docx, "D'autre part,")
        add_paragraph(
            docx,
            f"Ont procédé de la manière suivante à la cession des actions de la Société {cible_name}.",
        )
        add_paragraph(docx, 'Ci-après dénommée "LA SOCIETE",')

        add_paragraph(docx, "IL EST PREALABLEMENT EXPOSE CE QUI SUIT :", bold=True)
        add_paragraph(
            docx,
            (
                f"La Société {cible_name} est une {cible_forme_complete}, au capital social "
                f"de {cible_capital} divisé en {cible_actions_total} actions "
                f"d'{cible_valeur_action_lettres} de valeur nominale, entièrement libérées "
                f"dont le siège est situé au {company_siege_display(societe_cible, 'societe_cible')}."
            ),
        )
        add_paragraph(
            docx,
            (
                f"La Société {cible_name} est immatriculée au Registre du Commerce et des "
                f"Sociétés de {cible_rcs} sous le numéro {cible_numero_rcs} et inscrite au "
                f"tableau de l'Ordre Départemental des {cible_profession_pluriel} "
                f"du {cible_ordre_departement}."
            ),
        )
        add_paragraph(docx, f"Son objet social est l'exercice de la profession de {cible_profession}.")
        add_paragraph(docx, presentation_dirigeants)
        add_paragraph(docx, "Le capital social est réparti à ce jour comme suit :")
        _add_capital_table(docx, ctx)
        add_paragraph(
            docx,
            (
                f"Suite à sa création, la Société {spfpl_name}, s'est montrée intéressée "
                f"pour acquérir la quasi-totalité des actions de la Société {cible_name} "
                f"détenues par {person_display(cedant, 'cedant')}."
            ),
        )

        add_paragraph(docx, "ORIGINE DE PROPRIETE", bold=True)
        add_paragraph(
            docx,
            "Aux termes des statuts le capital social de la SOCIETE est actuellement détenu comme suit :",
        )
        for line in _origin_property_lines(ctx):
            add_paragraph(docx, line)
        add_paragraph(
            docx,
            (
                f"{person_display(cedant, 'cedant')}, le CEDANT, déclare qu'il est "
                "propriétaire des actions pour les avoir souscrites lors de la constitution "
                "de la SOCIETE."
            ),
        )

        add_paragraph(docx, "CECI EXPOSE, IL EST CONVENU CE QUI SUIT :", bold=True)
        add_paragraph(docx, "OBJET DU CONTRAT : CESSION D'ACTIONS", bold=True)
        add_paragraph(
            docx,
            (
                "Par les présentes, le Cédant cède ce jour, sous les garanties ordinaires "
                "de fait et de droit en la matière, ainsi que celles consenties dans les "
                "présentes, à l'Acquéreur, qui accepte, la pleine propriété de "
                f"{required_text(cession_actions.nb_actions_lettres, 'cession_actions.nb_actions_lettres')} "
                f"({required_int(cession_actions.nb_actions, 'cession_actions.nb_actions')}) "
                "actions qu'il détient, (ci-après, les \"Actions Cédées\" ou les "
                "\"Titres Cédés\"), ensemble avec tous les droits, titres et intérêts "
                "qui y sont attachés."
            ),
        )
        add_paragraph(
            docx,
            "Les Actions sont cédées intégralement libérées, libres de tout nantissement, "
            "sûretés, promesse ou autre empêchement ou restriction quelconque.",
        )
        _add_source_legal_blocks(docx, spfpl_name, cession_actions)
        add_paragraph(docx, f"Fait à {ctx.signature.lieu}")
        add_paragraph(docx, f"Le {format_display_date(ctx.signature.date, 'signature.date')}")
        add_paragraph(docx, f"En {nombre_exemplaires} exemplaires originaux,")
        add_signature_lines(
            docx,
            [
                f"{cedant_titre_signature} {required_text(cedant.prenom, 'cedant.prenom')} "
                f"{required_text(cedant.nom, 'cedant.nom')}",
                f"La société {spfpl_name}",
                f"Représentée par {rep_civilite_courte} {rep_prenom} {rep_nom}",
            ],
        )
        add_paragraph(docx, "Cadre réservé à l'administration")

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _validate_actions_context(ctx: DocumentGenerationContext) -> None:
    validate_cession_context(ctx)
    if ctx.structure != SPFPL_CESSION_STRUCTURE:
        raise ValueError(f"dossier.structure doit etre SPFPL cession pour {DOCUMENT_CODE}.")
    if ctx.operation_spfpl is None:
        raise ValueError(f"operation_spfpl est obligatoire pour {DOCUMENT_CODE}.")
    if (ctx.operation_spfpl.nature_titres or "").strip().lower() != "actions":
        raise ValueError(f"operation_spfpl.nature_titres doit etre actions pour {DOCUMENT_CODE}.")
    if (ctx.operation_spfpl.document_demande or "").strip().lower() != "acte_cession_actions":
        raise ValueError(
            f"operation_spfpl.document_demande doit etre acte_cession_actions pour {DOCUMENT_CODE}."
        )
    cedant = required_cedant(ctx)
    if cedant.genre != Gender.MASCULIN:
        raise ValueError("cedant.genre doit etre masculin pour le wording source V1.")
    cession_actions = _required_cession_actions(ctx)
    if cession_actions.modalites_paiement != SOURCE_PAYMENT_MODE:
        raise ValueError(
            "cession_actions.modalites_paiement doit confirmer le paiement source "
            f"{SOURCE_PAYMENT_MODE} pour {DOCUMENT_CODE}."
        )
    if not cession_actions.agrement_unanime_confirme:
        raise ValueError("cession_actions.agrement_unanime_confirme doit etre vrai.")
    if not cession_actions.pv_agrement_coherent:
        raise ValueError("cession_actions.pv_agrement_coherent doit etre vrai.")
    if not cession_actions.gap_applicable:
        raise ValueError("cession_actions.gap_applicable doit etre vrai.")
    if not cession_actions.representant_cessionnaire_confirme:
        raise ValueError("cession_actions.representant_cessionnaire_confirme doit etre vrai.")
    if (ctx.signature.prestataire_signature_electronique or "").strip().lower() != "yousign":
        raise ValueError("signature.prestataire_signature_electronique doit etre Yousign.")
    _validate_selas_dentiste_source_scope(ctx)
    _validate_capital_and_price(ctx, cession_actions)


def _validate_selas_dentiste_source_scope(ctx: DocumentGenerationContext) -> None:
    societe_cible = required_societe_cible(ctx)
    forme = (
        societe_cible.forme_sociale
        or societe_cible.forme_sociale_complete
        or ""
    ).lower()
    profession = (societe_cible.profession_reglementee or "").lower()
    if "selas" not in forme or profession != "chirurgien-dentiste":
        raise ValueError(
            "societe_cible doit etre une SELAS de chirurgien-dentiste pour le wording source V1."
        )


def _validate_capital_and_price(
    ctx: DocumentGenerationContext,
    cession_actions: CessionActions,
) -> None:
    societe_cible = required_societe_cible(ctx)
    total = required_int(societe_cible.nb_actions_total, "societe_cible.nb_actions_total")
    nb_cedees = required_int(cession_actions.nb_actions, "cession_actions.nb_actions")
    if not ctx.associes_cible:
        raise ValueError("associes_cible est obligatoire pour CODE-ACTE-ACTIONS-001.")

    total_before = 0
    cedant_before = 0
    for index, associe in enumerate(ctx.associes_cible):
        field_name = f"associes_cible[{index}]"
        nb_actions = required_int(associe.nb_actions_avant, f"{field_name}.nb_actions_avant")
        total_before += nb_actions
        if associe.est_cedant:
            cedant_before += nb_actions
    if total_before != total:
        raise ValueError("La repartition avant operation doit correspondre au total d'actions.")
    if cedant_before < nb_cedees:
        raise ValueError("Le cedant doit detenir au moins les actions cedees avant cession.")

    prix_unitaire = _parse_decimal(
        cession_actions.prix_unitaire_action,
        "cession_actions.prix_unitaire_action",
    )
    prix_total = _parse_decimal(cession_actions.prix_total, "cession_actions.prix_total")
    if prix_unitaire * Decimal(nb_cedees) != prix_total:
        raise ValueError("cession_actions.nb_actions * prix_unitaire_action doit egaler prix_total.")


def _required_cession_actions(ctx: DocumentGenerationContext) -> CessionActions:
    if ctx.cession_actions is None:
        raise ValueError(f"cession_actions est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.cession_actions


def _parse_decimal(value: str | None, field_name: str) -> Decimal:
    raw = required_text(value, field_name)
    normalized = raw.replace(" ", "").replace("\u00a0", "").replace(",", ".")
    try:
        return Decimal(normalized)
    except InvalidOperation as exc:
        raise ValueError(f"{field_name} doit etre numerique pour {DOCUMENT_CODE}.") from exc


def _conjoint_display(ctx: DocumentGenerationContext) -> str:
    cedant = required_cedant(ctx)
    if cedant.conjoint is None:
        raise ValueError("cedant.conjoint est obligatoire pour CODE-ACTE-ACTIONS-001.")
    return (
        f"{required_text(cedant.conjoint.civilite_affichage, 'cedant.conjoint.civilite_affichage')} "
        f"{required_text(cedant.conjoint.prenom, 'cedant.conjoint.prenom')} "
        f"{required_text(cedant.conjoint.nom, 'cedant.conjoint.nom')}"
    )


def _add_centered_heading(document, text: str) -> None:
    add_paragraph(
        document,
        "______________________________________________________________________________",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=0,
    )
    add_paragraph(document, text, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(
        document,
        "______________________________________________________________________________",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after_pt=0,
    )


def _add_capital_table(document, ctx: DocumentGenerationContext) -> None:
    societe_cible = required_societe_cible(ctx)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    header_cells[0].text = "Associés"
    header_cells[1].text = "Actions"
    for index, associe in enumerate(ctx.associes_cible):
        field_name = f"associes_cible[{index}]"
        row = table.add_row().cells
        row[0].text = associe_display_name(associe, field_name)
        row[1].text = str(required_int(associe.nb_actions_avant, f"{field_name}.nb_actions_avant"))
    total_row = table.add_row().cells
    total_row[0].text = "Total"
    total_row[1].text = str(required_int(societe_cible.nb_actions_total, "societe_cible.nb_actions_total"))


def _origin_property_lines(ctx: DocumentGenerationContext) -> list[str]:
    lines: list[str] = []
    for index, associe in enumerate(ctx.associes_cible):
        field_name = f"associes_cible[{index}]"
        lines.append(
            f"{associe_display_name(associe, field_name)}, détenant "
            f"{required_text(associe.nb_actions_avant_lettres, f'{field_name}.nb_actions_avant_lettres')} "
            f"({required_int(associe.nb_actions_avant, f'{field_name}.nb_actions_avant')}) actions ;"
        )
    return lines


def _add_source_legal_blocks(
    document,
    spfpl_name: str,
    cession_actions: CessionActions,
) -> None:
    add_paragraph(document, "NANTISSEMENT- PACTE D'ASSOCIES - AGREMENT", bold=True)
    add_paragraph(
        document,
        "Les Actions de la Société, objet de la cession, ne sont pas nanties ni données "
        "en garantie, et ne font l'objet d'aucune part ou poursuite de quelque nature "
        "que ce soit de la part d'un tiers.",
    )
    add_paragraph(
        document,
        "Conformément aux statuts de la Société, toute cession d'Actions de la Société "
        "au profit d'un tiers, nécessite l'agrément des Associés. Les Associés se sont "
        "réunis ce jour et ont agréé la cession à l'unanimité.",
    )
    add_paragraph(document, "PROPRIETE - JOUISSANCE", bold=True)
    add_paragraph(
        document,
        "L'Acquéreur sera propriétaire et aura la jouissance des Actions Cédées à compter "
        "de ce jour, (ci-après, la \"Date de Réalisation\"). À partir de cette date, "
        "l'Acquéreur sera subrogé dans tous les droits et obligations attachés aux "
        "Actions Cédées.",
    )
    add_paragraph(
        document,
        "Les dividendes afférents à l'exercice en cours au jour de la Date de Réalisation, "
        "dont la distribution pourrait être décidée et effectuée postérieurement au "
        "transfert de propriété des Actions seront répartis entre le Cessionnaire et le "
        "Cédant au prorata du temps pendant lequel, au cours dudit exercice, chacun "
        "d'entre eux aura été associé de la Société.",
    )
    add_paragraph(document, "PRIX ET MODALITES DE PAIEMENT DES ACTIONS", bold=True)
    add_paragraph(document, "4.1 Prix de cession", bold=True)
    add_paragraph(
        document,
        "La cession a lieu moyennant le prix global de "
        f"{required_text(cession_actions.prix_total_lettres, 'cession_actions.prix_total_lettres')} "
        f"({required_text(cession_actions.prix_total, 'cession_actions.prix_total')}) euros, "
        "soit un prix de "
        f"{required_text(cession_actions.prix_unitaire_action, 'cession_actions.prix_unitaire_action')} "
        "€ "
        f"({required_text(cession_actions.prix_unitaire_action_lettres, 'cession_actions.prix_unitaire_action_lettres')}) "
        f"par Action, à payer par la Société {spfpl_name}.",
    )
    add_paragraph(document, "4.2. Modalités de paiement du prix", bold=True)
    add_paragraph(
        document,
        f"Le prix est payé au moyen d'un crédit bancaire, comptant, ce jour, par la Société {spfpl_name}, "
        "par un chèque de banque, ce dont le CEDANT consent bonne et valable quittance.",
    )
    add_paragraph(document, "DONT QUITTANCE.")
    _add_declarations_and_gap(document)
    _add_general_and_formalities(document)


def _add_declarations_and_gap(document) -> None:
    add_paragraph(document, "DECLARATIONS DES PARTIES", bold=True)
    for text in [
        "Le Cédant et le Cessionnaire déclarent, chacun en ce qui le concerne :",
        "que leur état civil et leur existence juridique est conforme à celui ou celle indiqué en tête des présentes,",
        "qu'ils ne sont pas susceptibles d'être actuellement ou ultérieurement pour des faits existants à ce jour, l'objet de poursuites ou de mesures pouvant entraîner la confiscation totale ou partielle de leurs biens,",
        "qu'ils ont leur pleine capacité civile pour s'obliger dans le cadre des présentes et de leurs suites et, plus spécialement, qu'ils ne font pas présentement l'objet d'une procédure collective, ni ne sont susceptibles de l'être en raison de leurs professions et fonctions, ni ne sont en état de cessation des paiements ou déconfiture ;",
        "qu'il n'existe de leur chef aucune restriction d'ordre légal, judiciaire ou contractuel par suite de procédure de cessation des paiements, soumis à une procédure de sauvegarde ou de redressement ou liquidation judiciaire, et plus généralement aucun empêchement quelconque de nature à faire obstacle aux engagements de cession et d'acquisition objet des présentes",
        "Remplir les conditions exigées par la loi pour détenir des actions de SELAS de chirurgien-dentiste ;",
        "Qu'ils sont résidents français au sens de la réglementation des relations financières avec l'étranger.",
    ]:
        add_paragraph(document, text)
    add_paragraph(document, "GARANTIE D'ACTIF ET DE PASSIF / GAP", bold=True)
    for text in [
        "Le Cédant garantit le Cessionnaire contre toute diminution ou insuffisance d'actif, d'augmentation du passif ou révélation de passif nouveau, qui pourrait résulter notamment d'un redressement fiscal ou social, dont l'origine serait antérieure à la Date de Réalisation de la cession des Actions mais qui se révèlerait ultérieurement.",
        "Dès lors qu'apparaitra un évènement, un fait, un acte ou une omission révélant l'inexactitude et/ou la violation de l'une quelconque des Déclarations et Garanties et dont (i) l'existence ou la cause sera antérieure à la Date de Réalisation et (ii) l'existence ou les conséquences sont couvertes par la GAP (ci-après un « Fait Générateur »), le Cédant s'engage à indemniser le Cessionnaire de tout préjudice, perte ou passif effectivement subis par la Société ou par le Cessionnaire, résultant du Fait Générateur (le « Préjudice ») et à prendre en charge  l'intégralité des honoraires et frais raisonnablement engagés pour obtenir réparation de ce Préjudice.",
        "Il est expressément convenu qu'en cas de mise en jeu de la présente GAP, le Cédant ne pourra, en aucune manière et de quelque façon que ce soit, être exonéré en tout ou en partie de ses obligations d'indemnisation au titre des présentes par le fait que le Cessionnaire aurait pu avoir connaissance soit du fait des investigations qu'il a effectuées ou fait effectuer, soit des informations qui lui auraient été communiquées préalablement à la date des présentes, d'éléments d'information, de données ou de conclusions quels qu'ils soient, se rapportant à la réclamation concernée.",
        "Le montant de l'indemnisation sera rigoureusement égal au montant du Préjudice indemnisable. Le montant global des indemnisations dont le Cédant peut être redevable au titre du présent article sera toutefois plafonné au Prix de cession des Actions cédées.",
        "La présente garantie est consentie :",
        "s'agissant de toute réclamation en matière fiscale ou sociale fondée sur l'inexactitude d'une des Déclarations et Garanties : pour une durée de 6 (six) mois suivant l'expiration de la période de prescription applicable à toute action susceptible d'être engagée à l'encontre de la Société ou du Cessionnaire;",
        "pour une durée de 3 ans à compter de la Date de Réalisation de la cession.",
        "Les Parties s'engagent à faire leur affaire personnelle de la souscription d'une assurance contre les risques mentionnés dans la présente garantie.",
    ]:
        add_paragraph(document, text)


def _add_general_and_formalities(document) -> None:
    sections = [
        (
            "UNICITE DU CONTRAT",
            [
                "Le présent Contrat constitue l'entier accord des Parties sur les stipulations qui en sont l'objet. Les Parties n'ont contracté qu'en considération de chacune des stipulations du présent Contrat. En conséquence, ces stipulations ne sont pas susceptibles d'exécution séparée. Le préambule et les Annexes éventuelles au présent Contrat font partie intégrante de celui-ci.",
                "Le présent Contrat ne pourra faire l'objet d'un avenant ou être modifié que par accord écrit signé des Parties.",
            ],
        ),
        (
            "RENONCIATION",
            [
                "La renonciation par une Partie à une condition quelconque ou à faire valoir la violation d'une stipulation, d'un terme ou d'un engagement contenu dans le présent Contrat, dans un ou plusieurs cas, ne sera pas réputée ni interprétée comme une renonciation répétée ou persistante à cette condition ou à faire valoir la violation d'une autre stipulation, d'un autre terme ou engagement du présent Contrat.",
            ],
        ),
        (
            "NEGOCIATION ET EXECUTION DU CONTRAT",
            [
                "Les Parties reconnaissent que (i) le Contrat a été librement négocié de bonne foi entre elles, (ii) c'est en toute connaissance de cause, en conscience et de façon parfaitement éclairée, que les Parties concluent ledit Contrat dont elles déclarent comprendre et accepter chaque stipulation et condition ainsi que les obligations qui en découlent pour elles, et que (iii) le Contrat constitue un contrat de gré à gré au sens des dispositions de l'article 1110 du Code civil. Elles déclarent en particulier avoir bénéficié du temps de réflexion nécessaire et des conseils de leurs avocats avant la signature du Contrat.",
                "Chacune des Parties reconnaît que l'autre Partie lui a demandé les informations qui avaient pour elle une importance déterminante de son consentement à la conclusion du Contrat au sens des dispositions de l'article 1112-1 du Code civil.",
                "Chacune des Parties renonce irrévocablement, et de façon non équivoque, au bénéfice des dispositions de l'article 1195 du Code civil et s'interdit d'en solliciter l'application et/ou d'initier toute demande ou part judiciaire en résolution ou révision du Contrat et/ou ses suites sur le fondement des dispositions de l'article 1195 précité et/ou de l'équité visée par les dispositions de l'article 1194 du même Code, notamment en vue de demander la renégociation de ses termes et conditions, même en cas de changement de circonstances (quelles qu'elles soient, y compris matérielles ou légales) ou des conditions d'exécution (y compris financières) des présentes et/ou de leurs suites.",
                "Chacune des Parties reconnait et accepte expressément qu'en cas d'inexécution de ses obligations au titre du Contrat, toute autre Partie pourra en poursuivre l'exécution forcée en nature conformément aux dispositions de l'article 1221 du Code civil, sans préjudice des autres voies de recours prévues par les dispositions de l'article 1217 du même Code.",
            ],
        ),
        (
            "AUTONOMIE DES STIPULATIONS DE LA CONVENTION",
            [
                "Dans l'hypothèse où l'une quelconque des stipulations de la présente Convention serait considérée comme nulle, illégale, inopposable ou inapplicable d'une quelconque manière, en tout ou partie, les autres stipulations continueront à s'appliquer. Dans ce cas, les Parties devront, dans la mesure du possible, s'efforcer de substituer à cette stipulation une stipulation valide correspondant à l'esprit et à la finalité des présentes.",
            ],
        ),
        (
            "SIGNIFICATION DE LA CESSION",
            [
                "La présente cession de droits sociaux sera signifiée à la SOCIETE conformément aux dispositions de l'article 1690 du Code Civil.",
                "Toutefois, cette signification pourra être remplacée par le dépôt d'un original du présent acte au siège social contre remise par la Présidence d'une attestation de ce dépôt.",
            ],
        ),
        (
            "DECLARATION POUR L'ENREGISTREMENT",
            [
                "En vue de la perception des droits d'enregistrement, le CEDANT déclare que la Société est soumise à l'impôt sur les sociétés et que les actions cédées ne confèrent pas la jouissance de droits immobiliers.",
            ],
        ),
        ("POUVOIRS", ["Tous pouvoirs sont conférés au porteur d'originaux des présentes en vue de l'accomplissement de toutes formalités légales."]),
        (
            "COMMUNICATION DU PRESENT CONTRAT AU CONSEIL DE L'ORDRE",
            ["Le présent contrat sera, sans délai, communiqué au Conseil départemental de l'Ordre en vue de ses observations éventuelles."],
        ),
        (
            "FRAIS",
            ["Les frais et droits d'enregistrement afférents à la présente cession et tous les frais qui en sont la conséquence seront supportés par le CESSIONNAIRE qui s'y oblige dans la mesure où ces frais et droits se rattacheront à la cession d'action consentie."],
        ),
        (
            "AFFIRMATION DE SINCERITE",
            ["Conformément aux dispositions de l'article 1837 du Code Général des Impôts, les Parties déclarent sous les peines édictées par le Code Général des Impôts que le prix convenu entre elles est réel et qu'il n'est ni modifié, ni contredit par une contre-lettre."],
        ),
        (
            "LOI APPLICABLE - ATTRIBUTION DE JURIDICTION",
            ["La présente Cession est soumise au droit français et tous différends y afférents seront tranchés par les juridictions compétentes du siège social de la Société dont les actions font l'objet de la cession."],
        ),
        (
            "CONVENTION SUR LA PREUVE - SIGNATURE ELECTRONIQUE",
            [
                "Les Parties consentent expressément la faculté de procéder à la signature du présent acte par le système de signature électronique. Les Parties renoncent en conséquence expressément à signer et obtenir un quelconque acte original de ce dernier.",
                "Les Parties reconnaissent que le présent acte, tel que signé par voie électronique, constitue une preuve valable permettant d'apprécier les droits, les obligations et responsabilités des Parties et le consentement de leurs signataires.",
                "Le présent acte est signé par chacune des Parties dans le cadre du processus de signature électronique via le service Yousign.",
            ],
        ),
    ]
    for heading, paragraphs in sections:
        add_paragraph(document, heading, bold=True)
        for text in paragraphs:
            add_paragraph(document, text)


def _nombre_exemplaires(
    cession_actions: CessionActions,
    ctx: DocumentGenerationContext,
) -> str:
    if cession_actions.nombre_exemplaires_lettres:
        return cession_actions.nombre_exemplaires_lettres
    if ctx.document and ctx.document.nombre_exemplaires_lettres:
        return ctx.document.nombre_exemplaires_lettres
    if ctx.signature.nombre_exemplaires:
        return ctx.signature.nombre_exemplaires
    raise ValueError("cession_actions.nombre_exemplaires_lettres est obligatoire.")
