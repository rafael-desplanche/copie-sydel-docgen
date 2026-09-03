from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import CapitalSouscripteur, DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    company_siege_display,
    person_short_identity,
    required_apport_titres,
    required_apporteur,
    required_capital_souscription,
    required_int,
    required_societe_cible,
    required_societe_spfpl,
    required_text,
    validate_apport_context,
)
from sydel_doc_engine.rendering.docx_builder import add_paragraph, new_document

OUTPUT_FILENAME = "attestation_capital_liste_souscripteurs.docx"


class AttestationCapitalListeSouscripteursGenerator:
    """Generateur from-scratch de l'attestation capital SPFPL V1."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_apport_context(ctx)
        societe_spfpl = required_societe_spfpl(ctx)
        societe_cible = required_societe_cible(ctx)
        apporteur = required_apporteur(ctx)
        apport_titres = required_apport_titres(ctx)
        capital = required_capital_souscription(ctx)
        souscripteur = _unique_souscripteur(capital.souscripteurs)
        president = capital.president or souscripteur
        spfpl_name = required_text(societe_spfpl.denomination, "societe_spfpl.denomination")
        spfpl_capital = required_text(societe_spfpl.capital_social, "societe_spfpl.capital_social")
        souscripteur_prenom = required_text(souscripteur.prenom, _souscripteur_field("prenom"))
        souscripteur_nom = required_text(souscripteur.nom, _souscripteur_field("nom"))
        apport_nature = required_text(
            capital.apports_nature_montant,
            "capital_souscription.apports_nature_montant",
        )
        apports_numeraire = required_text(
            capital.apports_numeraire_montant,
            "capital_souscription.apports_numeraire_montant",
        )

        docx = new_document()
        add_paragraph(
            docx,
            spfpl_name,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        add_paragraph(
            docx,
            f"Societe par actions simplifiee au capital de {spfpl_capital} euros",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_paragraph(
            docx,
            "Societe de Participations Financieres de Profession Liberale de "
            f"{required_text(societe_spfpl.profession, 'societe_spfpl.profession')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_paragraph(
            docx,
            f"Siege social : {company_siege_display(societe_spfpl, 'societe_spfpl')}",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        add_paragraph(docx, "ATTESTATION", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        add_paragraph(
            docx,
            "Liste des souscripteurs",
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        add_paragraph(
            docx,
            f"{_souscripteur_identite(president, 'capital_souscription.president')}, "
            f"demeurant {_adresse(president, 'capital_souscription.president')}, "
            f"atteste que le capital de la societe {spfpl_name} "
            "est reparti de la maniere suivante :",
        )
        add_paragraph(docx, f"Capital social : {spfpl_capital} euros")
        add_paragraph(
            docx,
            "Nombre d'actions : "
            f"{required_int(capital.nb_actions_total, 'capital_souscription.nb_actions_total')} "
            f"actions d'un montant d'{_valeur_nominale_action(capital)} "
            "euro chacune",
        )
        add_paragraph(
            docx,
            "Repartition : "
            f"{required_int(souscripteur.nb_actions, _souscripteur_field('nb_actions'))} "
            f"actions attribuees au Dr {souscripteur_prenom} {souscripteur_nom}, "
            "actionnaire unique",
        )
        add_paragraph(docx, "Apports en nature :", bold=True)
        add_paragraph(
            docx,
            f"{person_short_identity(apporteur, 'apporteur')} fait apport de "
            f"{required_int(apport_titres.nb_parts, 'apport_titres.nb_parts')} parts de la "
            f"{required_text(societe_cible.forme_sociale, 'societe_cible.forme_sociale')} "
            f"denommee {required_text(societe_cible.denomination, 'societe_cible.denomination')} "
            f"ayant son siege {company_siege_display(societe_cible, 'societe_cible')}, "
            "immatriculee au RCS de "
            f"{required_text(societe_cible.ville_rcs, 'societe_cible.ville_rcs')} "
            f"sous le numero {required_text(societe_cible.numero_rcs, 'societe_cible.numero_rcs')} "
            f"pour une valeur de {apport_nature} euros.",
        )
        add_paragraph(docx, f"Total des apports en nature {apport_nature} euros")
        add_paragraph(docx, f"Apports en numeraire : {apports_numeraire}")
        add_paragraph(
            docx,
            "Le Docteur "
            f"{_souscripteur_nom(souscripteur)} a fait la totalite des apports en nature.",
        )
        add_paragraph(
            docx,
            f"Le present etat qui constate la souscription d'actions de la societe {spfpl_name}, "
            "ainsi que l'apport de la somme de "
            f"{apport_nature} euros correspondant a la totalite du nominal desdites actions, est "
            "certifie exact, sincere et veritable par le President, "
            f"{_souscripteur_identite(president, 'capital_souscription.president')}.",
        )
        add_paragraph(docx, f"Fait a {ctx.signature.lieu}")
        add_paragraph(docx, f"Le {ctx.signature.date.strftime('%d/%m/%Y')}")
        add_paragraph(docx, _souscripteur_identite(president, "capital_souscription.president"))

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _unique_souscripteur(souscripteurs: list[CapitalSouscripteur]) -> CapitalSouscripteur:
    if len(souscripteurs) != 1:
        raise ValueError(
            "capital_souscription.souscripteurs doit contenir exactement un "
            "souscripteur pour CODE-SPFPL-CORE-001."
        )
    return souscripteurs[0]


def _souscripteur_identite(souscripteur: CapitalSouscripteur, field_name: str) -> str:
    return (
        f"{required_text(souscripteur.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{required_text(souscripteur.prenom, f'{field_name}.prenom')} "
        f"{required_text(souscripteur.nom, f'{field_name}.nom')} "
        f"{required_text(souscripteur.profession, f'{field_name}.profession')}"
    )


def _souscripteur_nom(souscripteur: CapitalSouscripteur) -> str:
    civilite = required_text(
        souscripteur.civilite_affichage,
        _souscripteur_field("civilite_affichage"),
    )
    return (
        f"{civilite} "
        f"{required_text(souscripteur.prenom, _souscripteur_field('prenom'))} "
        f"{required_text(souscripteur.nom, _souscripteur_field('nom'))}"
    )


def _adresse(souscripteur: CapitalSouscripteur, field_name: str) -> str:
    return required_text(
        souscripteur.adresse_personnelle_affichee,
        f"{field_name}.adresse_personnelle_affichee",
    )


def _souscripteur_field(field_name: str) -> str:
    return f"capital_souscription.souscripteurs[0].{field_name}"


def _valeur_nominale_action(capital) -> str:
    return required_text(
        capital.valeur_nominale_action,
        "capital_souscription.valeur_nominale_action",
    )
