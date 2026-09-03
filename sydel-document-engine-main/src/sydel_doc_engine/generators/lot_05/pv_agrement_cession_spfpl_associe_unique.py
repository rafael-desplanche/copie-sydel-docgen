from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext
from sydel_doc_engine.generators.lot_05.pv_agrement_common import (
    add_article_7_bis,
    add_ordre_du_jour,
    add_pouvoirs_resolution,
    add_pv_title,
    add_resolution_agrement,
    add_societe_cible_header,
    reunion_intro_lines,
)
from sydel_doc_engine.generators.lot_05.spfpl_common import (
    person_display,
    person_signature,
    required_cedant,
    required_int,
    required_societe_cible,
    required_text,
    validate_associe_unique,
    validate_cession_context,
)
from sydel_doc_engine.rendering.docx_builder import add_paragraph, new_document

OUTPUT_FILENAME = "pv_agrement_cession_spfpl_associe_unique.docx"


class PvAgrementCessionSpfplAssocieUniqueGenerator:
    """Generateur from-scratch du PV d'agrement SPFPL en associe unique."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_cession_context(ctx)
        validate_associe_unique(ctx, expected=True)
        societe_cible = required_societe_cible(ctx)
        cedant = required_cedant(ctx)

        docx = new_document()
        add_societe_cible_header(docx, ctx)
        add_pv_title(docx, "L'ASSOCIE UNIQUE", ctx)
        for line in reunion_intro_lines(ctx):
            add_paragraph(docx, line)
        add_paragraph(
            docx,
            (
                f"{person_display(cedant, 'cedant')}, associe unique de la Societe "
                f"{required_text(societe_cible.denomination, 'societe_cible.denomination')}, "
                "au capital de "
                f"{required_text(societe_cible.capital_social, 'societe_cible.capital_social')} "
                "euros, compose de "
                f"{required_int(societe_cible.nb_parts_total, 'societe_cible.nb_parts_total')} "
                "parts, a pris les decisions suivantes :"
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_ordre_du_jour(docx)
        add_resolution_agrement(docx, ctx, subject="L'associe unique")
        add_article_7_bis(docx, ctx, subject="L'associe unique")
        add_pouvoirs_resolution(docx, subject="L'associe unique")
        add_paragraph(
            docx,
            (
                "De tout ce que dessus, il a ete dresse le present proces-verbal qui a "
                "ete signe apres lecture par l'associe unique."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(docx, person_signature(cedant, "cedant"), space_before_pt=12)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path
