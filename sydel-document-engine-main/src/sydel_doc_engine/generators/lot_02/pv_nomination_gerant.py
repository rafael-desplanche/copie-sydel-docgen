from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    Associe,
    BienImmobilier,
    CapitalContext,
    Company,
    DirigeantNomine,
    DocumentGenerationContext,
    Emprunt,
    ReunionPresident,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_centered_block,
    add_framed_title,
    add_hyphen_list_item,
    add_paragraph,
    add_signature_lines,
    add_spacer,
    new_document,
)

OUTPUT_FILENAME = "pv_nomination_gerant.docx"
DOCUMENT_CODE = "CODE-PV-001"

VOTE_FORMULA = "Cette résolution est adoptée à l’unanimité"
POWERS_TEXT = (
    "L’assemblée générale confère tous les pouvoirs au porteur d’un original à l’effet de procéder "
    "aux formalités d’enregistrement au greffe du Tribunal de Commerce de la Société."
)


class PvNominationGerantGenerator:
    """Générateur from-scratch du PV nomination gérant."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        company = _required_company(ctx.societe)
        capital = _required_capital(ctx.capital)
        dirigeant = _required_dirigeant(ctx.dirigeant_nomine)
        associes = _required_associes(ctx.associes)
        represented_associes = _represented_associes(associes)
        represented_parts = _validated_represented_parts(capital, represented_associes)
        emprunt = ctx.emprunt or Emprunt(actif=False)
        bien_immobilier = _required_bien_immobilier(ctx.bien_immobilier, emprunt)

        document = new_document()
        _add_company_header(document, company, associes)
        _add_title_and_meeting(document, ctx)
        _add_introduction(document, company, capital, associes)
        _add_associes_block(document, represented_associes, represented_parts)
        _add_order_of_business(document, ctx, dirigeant, emprunt, bien_immobilier)
        _add_nomination_decision(document, dirigeant)
        _add_borrowing_decision(document, emprunt, bien_immobilier)
        _add_powers_decision(document, emprunt)
        _add_closing_and_signatures(document, ctx, associes, dirigeant)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return company


def _required_capital(capital: CapitalContext | None) -> CapitalContext:
    if capital is None:
        raise ValueError(f"capital est obligatoire pour {DOCUMENT_CODE}.")
    return capital


def _required_dirigeant(dirigeant: DirigeantNomine | None) -> DirigeantNomine:
    if dirigeant is None:
        raise ValueError(f"dirigeant_nomine est obligatoire pour {DOCUMENT_CODE}.")
    return dirigeant


def _required_associes(associes: list[Associe]) -> list[Associe]:
    if not associes:
        raise ValueError(f"associes[] doit contenir au moins un associé pour {DOCUMENT_CODE}.")
    return associes


def _represented_associes(associes: list[Associe]) -> list[Associe]:
    represented = [associe for associe in associes if associe.est_present_ou_represente]
    if not represented:
        raise ValueError(
            "associes[] doit contenir au moins un associé présent ou représenté "
            f"pour {DOCUMENT_CODE}."
        )
    return represented


def _required_bien_immobilier(
    bien_immobilier: BienImmobilier | None,
    emprunt: Emprunt,
) -> BienImmobilier | None:
    if not emprunt.actif:
        return None
    if bien_immobilier is None:
        raise ValueError("bien_immobilier est obligatoire si emprunt.actif=true.")
    return bien_immobilier


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def _required_positive_int(value: int | None, field_name: str) -> int:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if value < 1:
        raise ValueError(f"{field_name} doit être supérieur ou égal à 1 pour {DOCUMENT_CODE}.")
    return value


def _required_display_value(value: date | str | None, field_name: str) -> str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, date):
        return value.strftime("%d/%m/%Y")
    return _required_text(value, field_name)


def _required_address(address: Address | None, field_name: str) -> Address:
    if address is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(address.num_voie, f"{field_name}.num_voie")
    _required_text(address.voie, f"{field_name}.voie")
    _required_text(address.cp, f"{field_name}.cp")
    _required_text(address.ville, f"{field_name}.ville")
    return address


def _address_inline(address: Address) -> str:
    num_voie = _required_text(address.num_voie, "adresse.num_voie")
    voie = _required_text(address.voie, "adresse.voie")
    cp = _required_text(address.cp, "adresse.cp")
    ville = _required_text(address.ville, "adresse.ville")
    return f"{num_voie} {voie}, {cp} {ville}"


def _address_no_comma(address: Address) -> str:
    num_voie = _required_text(address.num_voie, "adresse.num_voie")
    voie = _required_text(address.voie, "adresse.voie")
    cp = _required_text(address.cp, "adresse.cp")
    ville = _required_text(address.ville, "adresse.ville")
    return f"{num_voie} {voie} {cp} {ville}"


def _validated_represented_parts(
    capital: CapitalContext,
    represented_associes: list[Associe],
) -> int:
    nb_parts_total = _required_positive_int(capital.nb_parts_total, "capital.nb_parts_total")
    represented_parts = sum(associe.nb_parts for associe in represented_associes)
    if (
        capital.nb_parts_representees is not None
        and capital.nb_parts_representees != represented_parts
    ):
        raise ValueError(
            "capital.nb_parts_representees doit correspondre à la somme des parts des associés "
            f"présents ou représentés pour {DOCUMENT_CODE}."
        )
    if represented_parts != nb_parts_total:
        raise ValueError(
            "Les parts présentes ou représentées doivent correspondre à la totalité du capital "
            f"pour {DOCUMENT_CODE}."
        )
    return represented_parts


def _parts_label(nb_parts: int) -> str:
    return "part" if nb_parts == 1 else "parts"


def _nomination_agenda_label(fonction_affichage: str) -> str:
    normalized = fonction_affichage.strip().lower()
    if "gérant" in normalized or "gerant" in normalized:
        if normalized.endswith("s"):
            return "Nomination des premiers gérants"
        return "Nomination du gérant"
    if normalized.endswith("s"):
        return f"Nomination des {fonction_affichage}"
    return f"Nomination du {fonction_affichage}"


def _ne_label(genre: Gender) -> str:
    return "née" if genre == Gender.FEMININ else "né"


def _capital_social(company: Company) -> str:
    return _required_text(company.capital_social or company.capital, "societe.capital_social")


def _capital_social_header(company: Company) -> str:
    capital = _capital_social(company)
    if re.search(r"\b(?:euro|euros|eur|€)\s*$", capital, flags=re.IGNORECASE):
        return capital
    return f"{capital} euros"


def _forme_sociale_affichage(company: Company) -> str:
    return _required_text(
        company.forme_sociale_affichage or company.forme_sociale,
        "societe.forme_sociale_affichage",
    )


def _forme_sociale_header(company: Company, associes: list[Associe]) -> str:
    base = _known_forme_sociale_header(company) or _required_text(
        company.forme_sociale_complete
        or company.forme_sociale_libelle_long
        or company.forme_sociale_affichage
        or company.forme_sociale,
        "societe.forme_sociale_complete",
    )
    profession = _sel_profession_for_header(company, associes)
    if profession and not _normalized_contains_profession(base, profession):
        return f"{base} de {profession}"
    return base


def _known_forme_sociale_header(company: Company) -> str | None:
    acronym = (company.forme_sociale_abregee or company.forme_sociale or "").strip().upper()
    if acronym == "SELARL":
        return "Société d’exercice libéral à responsabilité limitée"
    if acronym == "SELAS":
        return "Société d’exercice libéral par actions simplifiée"
    return None


def _sel_profession_for_header(company: Company, associes: list[Associe]) -> str | None:
    acronym = (company.forme_sociale_abregee or company.forme_sociale or "").strip().upper()
    if acronym not in {"SELARL", "SELAS"}:
        return None
    for associe in associes:
        profession = (
            associe.profession_reglementee
            or associe.profession
            or associe.qualification_principale
        )
        if profession and profession.strip():
            return profession.strip()
    return None


def _normalized_contains_profession(base: str, profession: str) -> bool:
    normalized_base = _normalize_for_prefix(base)
    normalized_profession = _normalize_for_prefix(profession)
    return normalized_base.endswith(f" de {normalized_profession}")


def _capital_variable_mention(company: Company) -> str:
    if company.capital_variable is False:
        raise ValueError(
            "societe.capital_variable=false n'est pas couvert par la spec texte V1 "
            f"pour {DOCUMENT_CODE}."
        )
    return (
        company.capital_variable_mention
        if company.capital_variable_mention is not None
        else " à capital variable"
    )


def _capital_variable_formule_intro(company: Company) -> str:
    if company.capital_variable is False:
        raise ValueError(
            "societe.capital_variable=false n'est pas couvert par la spec texte V1 "
            f"pour {DOCUMENT_CODE}."
        )
    return (
        company.capital_variable_formule_intro
        if company.capital_variable_formule_intro is not None
        else "à capital variable"
    )


def _add_paragraph(
    document,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    space_before: int = 0,
    space_after: int = 6,
) -> None:
    add_paragraph(
        document,
        text,
        alignment=alignment,
        bold=bold,
        italic=italic,
        underline=underline,
        space_before_pt=space_before,
        space_after_pt=space_after,
    )


def _add_list_item(document, text: str) -> None:
    add_hyphen_list_item(document, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _add_decision_title(document, title: str) -> None:
    _add_paragraph(
        document,
        title,
        bold=True,
        underline=True,
        space_before=10,
        space_after=2,
    )


def _add_vote_formula(document) -> None:
    _add_paragraph(document, VOTE_FORMULA, italic=True)


def _add_company_header(document, company: Company, associes: list[Associe]) -> None:
    siege = _required_address(company.siege, "societe.siege")
    lines = [
        (_required_text(company.denomination, "societe.denomination"), True, False),
        _forme_sociale_header(company, associes),
        f"Au capital de {_capital_social_header(company)}",
        f"Siège social : {_address_no_comma(siege)}",
        "En cours d’immatriculation",
    ]
    add_centered_block(document, lines, space_after_pt=2)


def _add_title_and_meeting(document, ctx: DocumentGenerationContext) -> None:
    decision = ctx.decision
    reunion = ctx.reunion
    if decision is None:
        raise ValueError(f"decision est obligatoire pour {DOCUMENT_CODE}.")
    if reunion is None:
        raise ValueError(f"reunion est obligatoire pour {DOCUMENT_CODE}.")

    add_spacer(document)
    add_framed_title(
        document,
        [
            "PROCES-VERBAL DES DECISIONS",
            " DE L’ASSEMBLEE GENERALE",
            f" DU {_required_display_value(decision.date, 'decision.date')}",
        ],
    )
    _add_paragraph(document, f"Le {_required_text(reunion.date_lettres, 'reunion.date_lettres')}")


def _add_president_sentence(document, ctx: DocumentGenerationContext) -> None:
    if ctx.reunion is None or ctx.reunion.president is None:
        return
    president = ctx.reunion.president
    civilite = _president_civilite(president)
    prenom = _president_prenom(president)
    nom = _president_nom(president)
    if civilite is None and prenom is None and nom is None:
        return
    _add_paragraph(
        document,
        (
            f"{_required_text(civilite, 'reunion.president.civilite_president_seance')} "
            f"{_required_text(prenom, 'reunion.president.prenom_president_seance')} "
            f"{_required_text(nom, 'reunion.president.nom_personne_seance')} "
            "préside la séance."
        ),
    )


def _president_civilite(president: ReunionPresident) -> str | None:
    return president.civilite_president_seance or president.civilite_affichage


def _president_prenom(president: ReunionPresident) -> str | None:
    return president.prenom_president_seance or president.prenom


def _president_nom(president: ReunionPresident) -> str | None:
    return president.nom_personne_seance or president.nom


def _add_introduction(
    document,
    company: Company,
    capital: CapitalContext,
    associes: list[Associe],
) -> None:
    denomination = _required_text(company.denomination, "societe.denomination")
    company_designation = _company_designation_for_intro(company, denomination)
    nb_parts_total = _required_positive_int(capital.nb_parts_total, "capital.nb_parts_total")
    valeur_nominale = _required_text(
        capital.valeur_nominale_part,
        "capital.valeur_nominale_part",
    )
    common = (
        f"de la {company_designation}, au capital de {_capital_social(company)}, "
        f"composé de {nb_parts_total} parts de {valeur_nominale} euro chacune, "
    )
    text = f"Les associés {common}se sont réunis au siège social."
    _add_paragraph(document, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def _company_designation_for_intro(company: Company, denomination: str) -> str:
    forme = _forme_sociale_affichage(company)
    if _denomination_starts_with_form(denomination, company, forme):
        return denomination
    return f"{forme} {denomination}"


def _denomination_starts_with_form(
    denomination: str,
    company: Company,
    forme: str,
) -> bool:
    normalized_denomination = _normalize_for_prefix(denomination)
    candidates = [
        forme,
        company.forme_sociale_abregee,
    ]
    return any(
        normalized_denomination.startswith(_normalize_for_prefix(candidate) + " ")
        or normalized_denomination == _normalize_for_prefix(candidate)
        for candidate in candidates
        if candidate and _normalize_for_prefix(candidate)
    )


def _normalize_for_prefix(value: str) -> str:
    return " ".join(value.casefold().replace("’", "'").split())


def _add_associes_block(
    document,
    associes: list[Associe],
    represented_parts: int,
) -> None:
    _add_paragraph(document, "Sont présents ou représentés :")
    for associe in associes:
        _add_list_item(
            document,
            (
                f"{associe.civilite_affichage} {associe.prenom} {associe.nom}, "
                f"détenant {associe.nb_parts} {_parts_label(associe.nb_parts)},"
            ),
        )
    _add_paragraph(
        document,
        (
            "Les associés présents ou représentés disposent ensemble de la totalité des parts "
            "sociales. Cet ensemble est habilité à prendre des décisions."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_order_of_business(
    document,
    ctx: DocumentGenerationContext,
    dirigeant: DirigeantNomine,
    emprunt: Emprunt,
    bien_immobilier: BienImmobilier | None,
) -> None:
    fonction_affichage = _required_text(
        dirigeant.fonction_affichage,
        "dirigeant_nomine.fonction_affichage",
    )
    _add_president_sentence(document, ctx)
    _add_paragraph(document, "Le président rappelle l’ordre du jour :")
    _add_paragraph(document, f"· {_nomination_agenda_label(fonction_affichage)}")
    if emprunt.actif:
        bien_adresse = _address_inline(
            _required_address(
                bien_immobilier.adresse if bien_immobilier else None,
                "bien_immobilier.adresse",
            )
        )
        _add_paragraph(
            document,
            (
                "· Autorisation de contracter un emprunt pour l’achat d’un bien immobilier sis "
                f"{bien_adresse}"
            ),
        )
    _add_paragraph(document, "· Pouvoirs")


def _add_nomination_decision(document, dirigeant: DirigeantNomine) -> None:
    fonction_affichage = _required_text(
        dirigeant.fonction_affichage,
        "dirigeant_nomine.fonction_affichage",
    )
    address = _required_address(
        dirigeant.adresse_personnelle,
        "dirigeant_nomine.adresse_personnelle",
    )
    birth_date = _required_display_value(
        dirigeant.date_naissance,
        "dirigeant_nomine.date_naissance",
    )
    birth_city = _required_text(
        dirigeant.ville_naissance,
        "dirigeant_nomine.ville_naissance",
    )
    birth_department = _required_text(
        dirigeant.departement_naissance,
        "dirigeant_nomine.departement_naissance",
    )
    nationality = _required_text(
        dirigeant.nationalite,
        "dirigeant_nomine.nationalite",
    )
    _add_decision_title(document, "PREMIERE DECISION")
    _add_paragraph(
        document,
        (
            "L’assemblée générale décide de désigner en qualité de "
            f"{fonction_affichage} pour une durée indéterminée :"
        ),
    )
    _add_paragraph(
        document,
        (
            f"{dirigeant.civilite_affichage} {dirigeant.prenom} {dirigeant.nom}, "
            f"{_ne_label(dirigeant.genre)} le {birth_date} à {birth_city} "
            f"({birth_department}), de nationalité {nationality}, "
            f"demeurant au {_address_inline(address)}."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _add_vote_formula(document)


def _add_borrowing_decision(
    document,
    emprunt: Emprunt,
    bien_immobilier: BienImmobilier | None,
) -> None:
    if not emprunt.actif:
        return
    bien_adresse = _address_inline(
        _required_address(
            bien_immobilier.adresse if bien_immobilier else None,
            "bien_immobilier.adresse",
        )
    )
    montant = _required_text(emprunt.montant_max, "emprunt.montant_max")
    _add_decision_title(document, "DEUXIEME DECISION")
    _add_paragraph(
        document,
        (
            "L’assemblée générale décide de contracter un emprunt d’un montant "
            f"maximum de {montant} euros pour l’acquisition d’un bien immobilier sis "
            f"{bien_adresse}."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _add_vote_formula(document)


def _add_powers_decision(document, emprunt: Emprunt) -> None:
    title = "TROISIEME DECISION" if emprunt.actif else "DEUXIEME DECISION"
    _add_decision_title(document, title)
    _add_paragraph(document, POWERS_TEXT, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_vote_formula(document)


def _add_closing_and_signatures(
    document,
    ctx: DocumentGenerationContext,
    associes: list[Associe],
    dirigeant: DirigeantNomine,
) -> None:
    fonction_affichage = _required_text(
        dirigeant.fonction_affichage,
        "dirigeant_nomine.fonction_affichage",
    )
    lieu_signature = _required_text(ctx.signature.lieu, "signature.lieu")
    nombre_exemplaires = _required_text(
        ctx.signature.nombre_exemplaires,
        "signature.nombre_exemplaires",
    )
    _add_paragraph(
        document,
        f"Fait à {lieu_signature} en {nombre_exemplaires} exemplaires",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_spacer(document)
    add_signature_lines(
        document,
        [f"{associe.prenom} {associe.nom}" for associe in associes],
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    _add_paragraph(
        document,
        (
            "Faire précéder la signature de la mention « Bon pour acceptation des fonctions de "
            f"{fonction_affichage} »"
        ),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        italic=True,
    )
