from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import Address, Company, DocumentGenerationContext, Person
from sydel_doc_engine.generators.lot_02.regime_communautaire_common import (
    SELARL_STRUCTURE,
    city_line,
    company_capital_social,
    company_forme_sociale,
    company_forme_sociale_abregee,
    format_display_date,
    required_address,
    required_apport,
    required_company,
    required_regime_communautaire,
    required_text,
    street_line,
    validate_batch_enabled,
)
from sydel_doc_engine.rendering.docx_builder import (
    LETTER_WIDE_STYLE_PROFILE,
    add_company_identity_block,
    add_hyphen_list_item,
    add_italic_instruction,
    add_paragraph,
    add_right_aligned_lines,
    add_spacer,
    add_subject_heading,
    new_document,
)

OUTPUT_FILENAME = "lettre_avertissement_conjoint.docx"


class LettreAvertissementConjointGenerator:
    """Generateur from-scratch de la lettre d'avertissement au conjoint."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        structure = validate_batch_enabled(ctx)
        company = required_company(ctx.societe)
        apport = required_apport(ctx.apport)
        regime = required_regime_communautaire(ctx.regime_communautaire)
        if regime.avertissement is None:
            raise ValueError(
                "regime_communautaire.avertissement est obligatoire pour CODE-RC-001."
            )

        document = new_document(style_profile=LETTER_WIDE_STYLE_PROFILE)
        _add_company_block(document, company, ctx)
        add_spacer(document, space_after_pt=12)
        _add_conjoint_block(document, ctx)
        date_signature = format_display_date(
            regime.avertissement.date_signature,
            "regime_communautaire.avertissement.date_signature",
        )
        add_right_aligned_lines(
            document,
            [f"Le  {date_signature}"],
            space_after_pt=12,
        )
        add_subject_heading(
            document,
            "Objet : Lettre d'avertissement au conjoint en cas d'apport d'un bien commun.",
            space_after_pt=12,
        )
        add_paragraph(document, _conjoint_appel(ctx))
        add_paragraph(
            document,
            (
                "En application des dispositions de l'article 1832-2 alinéa 1 er du Code "
                "civil, je t’informe par la présente que j'ai l'intention de faire apport à "
                "une Société dont les caractéristiques sont décrites ci-après :"
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        _add_company_block(document, company, ctx)
        add_hyphen_list_item(
            document,
            (
                "d'une somme en numéraire de "
                f"{required_text(apport.montant_lettres, 'apport.montant_lettres')} "
                f"({required_text(apport.montant, 'apport.montant')}) euros dépendant "
                "de notre communauté."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(document, "Fait en quatre exemplaires")
        _add_apporteur_signature_block(document, ctx)
        add_paragraph(document, _conjoint_line(ctx))
        add_italic_instruction(document, _mention_manuscrite(ctx, company, structure))

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _add_company_block(
    document,
    company: Company,
    ctx: DocumentGenerationContext,
) -> None:
    siege = required_address(company.siege, "societe.siege")
    add_company_identity_block(
        document,
        [
            required_text(company.denomination, "societe.denomination"),
            _company_forme_sociale_header(company, ctx),
            f"Au capital de {company_capital_social(company)} €",
            street_line(siege),
            city_line(siege),
        ],
        first_line_bold=True,
        space_after_pt=2,
    )


def _add_conjoint_block(document, ctx: DocumentGenerationContext) -> None:
    address = _conjoint_address(ctx)
    add_right_aligned_lines(
        document,
        [
            _conjoint_line(ctx),
            street_line(address),
            city_line(address),
        ],
        space_after_pt=2,
    )


def _conjoint_address(ctx: DocumentGenerationContext) -> Address:
    if ctx.personne_signataire.adresse_perso is not None:
        return required_address(
            ctx.personne_signataire.adresse_perso,
            "personne_signataire.adresse_perso",
        )
    conjoint = _required_conjoint(ctx)
    return required_address(conjoint.adresse_perso, "conjoint.adresse")


def _required_conjoint(ctx: DocumentGenerationContext) -> Person:
    conjoint = ctx.conjoint
    if conjoint is None:
        raise ValueError("conjoint est obligatoire pour CODE-RC-001.")
    return conjoint


def _conjoint_appel(ctx: DocumentGenerationContext) -> str:
    return f"{_conjoint_line(ctx)},"


def _conjoint_line(ctx: DocumentGenerationContext) -> str:
    conjoint = _required_conjoint(ctx)
    civilite = required_text(conjoint.civilite, "conjoint.civilite_affichage")
    nom = required_text(conjoint.nom, "conjoint.nom")
    return f"{civilite} {nom}"


def _company_forme_sociale_header(
    company: Company,
    ctx: DocumentGenerationContext,
) -> str:
    base = _known_forme_sociale_header(company) or required_text(
        company.forme_sociale_complete
        or company.forme_sociale_libelle_long
        or company.forme_sociale_affichage
        or company_forme_sociale(company),
        "societe.forme_sociale_complete",
    )
    profession = _sel_profession_for_header(company, ctx)
    if profession and not _normalized_contains_profession(base, profession):
        return f"{base} de {profession}"
    return base


def _known_forme_sociale_header(company: Company) -> str | None:
    acronym = (company.forme_sociale_abregee or company.forme_sociale or "").strip().upper()
    if acronym == "SELARL":
        return "Société d’exercice libéral à responsabilité limitée"
    if acronym == "SELAS":
        return "Société d’exercice libéral par actions simplifiée"
    return None


def _sel_profession_for_header(
    company: Company,
    ctx: DocumentGenerationContext,
) -> str | None:
    acronym = (company.forme_sociale_abregee or company.forme_sociale or "").strip().upper()
    if acronym not in {"SELARL", "SELAS"}:
        return None
    for associe in ctx.associes:
        profession = (
            associe.profession_reglementee
            or associe.profession
            or associe.qualification_principale
        )
        if profession and profession.strip():
            return profession.strip()
    if ctx.statuts_sel is not None and ctx.statuts_sel.profession:
        return ctx.statuts_sel.profession.strip()
    profession = ctx.personne_signataire.qualification_principale
    return profession.strip() if profession and profession.strip() else None


def _normalized_contains_profession(base: str, profession: str) -> bool:
    return _normalize_for_match(base).endswith(
        f" de {_normalize_for_match(profession)}"
    )


def _normalize_for_match(value: str) -> str:
    return " ".join(value.casefold().replace("’", "'").split())


def _add_apporteur_signature_block(document, ctx: DocumentGenerationContext) -> None:
    apporteur = ctx.personne_signataire
    civilite = required_text(apporteur.civilite, "apporteur.civilite_affichage")
    prenom = required_text(apporteur.prenom, "apporteur.prenom")
    nom = required_text(apporteur.nom, "apporteur.nom")
    fonction = required_text(apporteur.fonction_dirigeant, "apporteur.fonction_dirigeant")
    add_paragraph(document, f"{civilite} {prenom} {nom}")
    add_italic_instruction(document, f"Agissant en qualité de futur {fonction}")


def _mention_manuscrite(
    ctx: DocumentGenerationContext,
    company: Company,
    structure: str,
) -> str:
    apporteur = ctx.personne_signataire
    apport = required_apport(ctx.apport)
    apporteur_label = (
        f"{required_text(apporteur.civilite, 'apporteur.civilite_affichage')} "
        f"{required_text(apporteur.prenom, 'apporteur.prenom')} "
        f"{required_text(apporteur.nom, 'apporteur.nom')}"
    )
    montant = required_text(apport.montant, "apport.montant")
    denomination = required_text(company.denomination, "societe.denomination")
    if structure == SELARL_STRUCTURE:
        destination = f"à la Société {denomination}"
    else:
        destination = f"à la {company_forme_sociale_abregee(company)} {denomination}"
    return (
        "(Faire précéder de la mention « j’atteste avoir été informé de l’apport de "
        f"{montant} euros par {apporteur_label} {destination} »)"
    )
