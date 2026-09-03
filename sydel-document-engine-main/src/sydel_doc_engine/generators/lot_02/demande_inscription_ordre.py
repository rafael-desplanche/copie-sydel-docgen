from __future__ import annotations

from datetime import date
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    Company,
    DocumentGenerationContext,
    DossierOptions,
    Mandataire,
    OrdreAddress,
    OrdreProfessionnel,
    Person,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_letter_place_date,
    add_paragraph,
    add_right_aligned_lines,
    add_right_indented_block,
    add_spacer,
    add_subject_heading,
    new_document,
)

OUTPUT_FILENAME = "demande_inscription_ordre.docx"
DOCUMENT_CODE = "CODE-ORDRE-001"

SELARL_SELAS_STRUCTURES = {"SELARL", "SELAS"}
SPFPL_STRUCTURES = {"SPFPL cession", "SPFPL apport", "SPFPL_CESSION", "SPFPL_APPORT"}
SCM_STRUCTURES = {"SCM"}
SUPPORTED_STRUCTURES = SELARL_SELAS_STRUCTURES | SPFPL_STRUCTURES | SCM_STRUCTURES

OVERLAY_SEL = "SELARL_SELAS"
OVERLAY_SPFPL = "SPFPL"
OVERLAY_SCM = "SCM"


class DemandeInscriptionOrdreGenerator:
    """Generateur from-scratch de la demande d'inscription a l'ordre."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        overlay = _overlay_for_structure(ctx.structure)
        signataire = ctx.personne_signataire
        company = _required_company(ctx.societe)
        ordre = _required_ordre(ctx.ordre)

        signataire_name = _signataire_name(signataire)
        profession_signataire = _required_text(
            ordre.profession_signataire_affichee,
            "ordre.profession_signataire_affichee",
        )
        adresse_personnelle = _required_text(
            signataire.adresse_personnelle_affichee,
            "personne_signataire.adresse_personnelle_affichee",
        )
        profession_ligne_destinataire = _profession_ligne_destinataire(ordre)
        conseil_lines = _conseil_departemental_lines(
            ordre,
            overlay=overlay,
            profession_ligne_destinataire=profession_ligne_destinataire,
        )
        profession_reglementee = _required_text(
            ordre.profession_reglementee_pluriel,
            "ordre.profession_reglementee_pluriel",
        )
        adresse_ordre_lines = _ordre_address_lines(ordre, overlay)
        destinataire_appel = _required_text(
            ordre.destinataire_appel,
            "ordre.destinataire_appel",
        )
        denomination_societe = _required_text(company.denomination, "societe.denomination")
        mandataire_libelle = _mandataire_libelle(ctx.mandataire)
        derogation_suffixe = _derogation_suffixe(ctx.dossier_options, ordre)

        document = new_document()
        _add_header(
            document,
            signataire_name=signataire_name,
            profession_signataire=profession_signataire,
            adresse_personnelle=adresse_personnelle,
            conseil_lines=conseil_lines,
            adresse_ordre_lines=adresse_ordre_lines,
        )
        _add_signature_place_and_subject(document, ctx)
        _add_body(
            document,
            destinataire_appel=destinataire_appel,
            denomination_societe=denomination_societe,
            profession_reglementee=profession_reglementee,
            mandataire_libelle=mandataire_libelle,
            derogation_suffixe=derogation_suffixe,
        )
        _add_final_signature(document, signataire_name)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _overlay_for_structure(structure: str | None) -> str:
    structure_value = _required_text(structure, "dossier.structure")
    if structure_value in SELARL_SELAS_STRUCTURES:
        return OVERLAY_SEL
    if structure_value in SPFPL_STRUCTURES:
        return OVERLAY_SPFPL
    if structure_value in SCM_STRUCTURES:
        return OVERLAY_SCM
    supported = ", ".join(sorted(SUPPORTED_STRUCTURES))
    raise ValueError(f"dossier.structure doit etre dans [{supported}] pour {DOCUMENT_CODE}.")


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return company


def _required_ordre(ordre: OrdreProfessionnel | None) -> OrdreProfessionnel:
    if ordre is None:
        raise ValueError(f"ordre est obligatoire pour {DOCUMENT_CODE}.")
    return ordre


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()


def _format_date(value: date) -> str:
    return value.strftime("%d/%m/%Y")


def _split_display_lines(value: str | None, field_name: str) -> list[str]:
    text = _required_text(value, field_name)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return lines


def _signataire_name(signataire: Person) -> str:
    titre = _required_text(
        signataire.titre_affichage,
        "personne_signataire.titre_affichage",
    )
    prenom = _required_text(signataire.prenom, "personne_signataire.prenom")
    nom = _required_text(signataire.nom, "personne_signataire.nom")
    return f"{titre} {prenom} {nom}"


def _profession_ligne_destinataire(ordre: OrdreProfessionnel) -> str:
    return _required_text(
        ordre.profession_ligne_destinataire or ordre.profession_reglementee_pluriel,
        "ordre.profession_ligne_destinataire",
    )


def _conseil_departemental_lines(
    ordre: OrdreProfessionnel,
    *,
    overlay: str,
    profession_ligne_destinataire: str,
) -> list[str]:
    if overlay == OVERLAY_SEL and ordre.departement_inscription:
        departement = _required_text(
            ordre.departement_inscription,
            "ordre.departement_inscription",
        )
        return [
            (
                "Conseil départemental de l'Ordre des "
                f"{profession_ligne_destinataire} de {departement}"
            )
        ]
    conseil_libelle = _required_text(
        ordre.conseil_departemental_libelle,
        "ordre.conseil_departemental_libelle",
    )
    return [conseil_libelle, f"Des {profession_ligne_destinataire}"]


def _ordre_address_lines(ordre: OrdreProfessionnel, overlay: str) -> list[str]:
    if overlay == OVERLAY_SEL:
        return _selarl_selas_address_lines(ordre)
    if overlay == OVERLAY_SPFPL:
        return _split_display_lines(
            ordre.adresse_affichee or ordre.adresse_bloc_affiche,
            "ordre.adresse_affichee",
        )
    return _scm_address_lines(ordre)


def _selarl_selas_address_lines(ordre: OrdreProfessionnel) -> list[str]:
    if ordre.adresse is not None:
        ligne_1 = _required_text(ordre.adresse.ligne_1, "ordre.adresse.ligne_1")
        cp = _required_text(ordre.adresse.cp, "ordre.adresse.cp")
        ville = _required_text(ordre.adresse.ville, "ordre.adresse.ville")
        return [ligne_1, f"{cp} {ville}"]
    return _split_display_lines(
        ordre.adresse_bloc_affiche or ordre.adresse_affichee,
        "ordre.adresse_bloc_affiche",
    )


def _scm_address_lines(ordre: OrdreProfessionnel) -> list[str]:
    if ordre.adresse_bloc_affiche or ordre.adresse_affichee:
        return _split_display_lines(
            ordre.adresse_bloc_affiche or ordre.adresse_affichee,
            "ordre.adresse_bloc_affiche",
        )
    return _address_lines_from_structured(ordre.adresse)


def _address_lines_from_structured(address: OrdreAddress | None) -> list[str]:
    if address is None:
        raise ValueError(f"ordre.adresse est obligatoire pour {DOCUMENT_CODE}.")
    ligne_1 = _required_text(address.ligne_1, "ordre.adresse.ligne_1")
    cp = _required_text(address.cp, "ordre.adresse.cp")
    ville = _required_text(address.ville, "ordre.adresse.ville")
    return [ligne_1, f"{cp} {ville}"]


def _mandataire_libelle(mandataire: Mandataire | None) -> str:
    if mandataire is None:
        raise ValueError(f"mandataire est obligatoire pour {DOCUMENT_CODE}.")
    if mandataire.libelle_affiche is not None and mandataire.libelle_affiche.strip():
        return mandataire.libelle_affiche.strip()
    civilite = _required_text(mandataire.civilite_affichage, "mandataire.civilite_affichage")
    prenom = _required_text(mandataire.prenom, "mandataire.prenom")
    nom = _required_text(mandataire.nom, "mandataire.nom")
    fonction = _required_text(mandataire.fonction, "mandataire.fonction")
    cabinet = _required_text(mandataire.cabinet, "mandataire.cabinet")
    return f"{civilite} {prenom} {nom}, {fonction} du cabinet {cabinet}"


def _derogation_suffixe(
    dossier_options: DossierOptions | None,
    ordre: OrdreProfessionnel,
) -> str:
    options = dossier_options or DossierOptions()
    if not options.derogation:
        return ""
    mention = _required_text(
        ordre.derogation_mention_manuelle,
        "ordre.derogation_mention_manuelle",
    )
    return f" {mention}"


def _add_header(
    document,
    *,
    signataire_name: str,
    profession_signataire: str,
    adresse_personnelle: str,
    conseil_lines: list[str],
    adresse_ordre_lines: list[str],
) -> None:
    _add_lines(document, [signataire_name, profession_signataire])
    _add_lines(document, _split_display_lines(adresse_personnelle, "adresse_personnelle"))
    add_spacer(document, space_after_pt=10)
    add_right_indented_block(
        document,
        conseil_lines,
        left_indent_cm=8.7,
        first_line_indent_cm=1.2,
        space_after_pt=2,
    )
    add_right_indented_block(
        document,
        adresse_ordre_lines,
        left_indent_cm=9.7,
        space_after_pt=2,
    )
    add_spacer(document, space_after_pt=12)


def _add_signature_place_and_subject(document, ctx: DocumentGenerationContext) -> None:
    lieu_signature = _required_text(ctx.signature.lieu, "signature.lieu")
    add_letter_place_date(
        document,
        f"{lieu_signature}, le {_format_date(ctx.signature.date)}",
        space_after_pt=12,
    )
    add_subject_heading(
        document,
        "Objet : Demande d’inscription au tableau de l’Ordre",
        space_after_pt=12,
    )


def _add_body(
    document,
    *,
    destinataire_appel: str,
    denomination_societe: str,
    profession_reglementee: str,
    mandataire_libelle: str,
    derogation_suffixe: str,
) -> None:
    add_paragraph(document, f"{destinataire_appel},")
    document.add_paragraph()
    _add_body_paragraph(
        document,
        (
            "Vous trouverez ci-joint le dossier de constitution de ma société dénommée "
            f"{denomination_societe}."
        ),
    )
    _add_body_paragraph(
        document,
        (
            "Je sollicite l’inscription de ma société au tableau de l’Ordre des "
            f"{profession_reglementee}. Je précise que je ne serai associé et praticien et "
            f"exerçant que dans une seule structure.{derogation_suffixe}"
        ),
    )
    _add_body_paragraph(
        document,
        f"Je donne pouvoir à {mandataire_libelle} pour effectuer les formalités.",
    )
    _add_body_paragraph(
        document,
        (
            f"Je vous prie d’agréer, {destinataire_appel}, l’expression de mes sentiments "
            "dévoués."
        ),
    )


def _add_final_signature(document, signataire_name: str) -> None:
    add_spacer(document, space_after_pt=12)
    add_right_aligned_lines(document, [signataire_name], space_after_pt=0)


def _add_lines(document, lines: list[str]) -> None:
    for line in lines:
        add_paragraph(document, line, space_after_pt=2)


def _add_body_paragraph(document, text: str) -> None:
    add_paragraph(document, text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
