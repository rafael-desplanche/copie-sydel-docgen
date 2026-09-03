from __future__ import annotations

from pathlib import Path
from unicodedata import normalize

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import DocumentGenerationContext, Person
from sydel_doc_engine.generators.lot_02.regime_communautaire_common import (
    company_forme_sociale_complete,
    format_display_date,
    required_apport,
    required_company,
    required_regime_communautaire,
    required_text,
    validate_batch_enabled,
)
from sydel_doc_engine.rendering.docx_builder import (
    LETTER_WIDE_STYLE_PROFILE,
    add_paragraph,
    add_right_aligned_lines,
    add_spacer,
    add_subject_heading,
    new_document,
)

OUTPUT_FILENAME = "lettre_renonciation_associe.docx"


class LettreRenonciationAssocieGenerator:
    """Generateur from-scratch de la lettre de renonciation du conjoint."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_batch_enabled(ctx)
        company = required_company(ctx.societe)
        apport = required_apport(ctx.apport)
        regime = required_regime_communautaire(ctx.regime_communautaire)
        if regime.renonciation is None:
            raise ValueError(
                "regime_communautaire.renonciation est obligatoire pour CODE-RC-001."
            )

        date_courrier = _date_courrier_avertissement(ctx)
        lieu_signature = required_text(
            regime.renonciation.lieu_signature,
            "regime_communautaire.renonciation.lieu_signature",
        )
        nombre_exemplaires = required_text(
            regime.renonciation.nombre_exemplaires_lettres,
            "regime_communautaire.renonciation.nombre_exemplaires_lettres",
        )

        document = new_document(style_profile=LETTER_WIDE_STYLE_PROFILE)
        add_right_aligned_lines(
            document,
            [f"À {lieu_signature}"],
            space_after_pt=2,
        )
        add_spacer(document, space_after_pt=12)
        add_subject_heading(
            document,
            "Objet : Lettre de renonciation à revendiquer la qualité d'associé",
            space_after_pt=12,
        )
        add_paragraph(document, _apporteur_appel(ctx))
        denomination = required_text(company.denomination, "societe.denomination")
        regime_matrimonial = _regime_matrimonial_display(
            required_text(
                regime.regime_matrimonial,
                "regime_communautaire.regime_matrimonial",
            )
        )
        add_paragraph(
            document,
            (
                f"Par courrier en date du {date_courrier}, tu m’as fait part du projet de "
                f"constitution de la société {denomination}, "
                f"{company_forme_sociale_complete(company)}, à laquelle tu souhaites t'associer "
                f"en apportant {required_text(apport.montant, 'apport.montant')} "
                f"({required_text(apport.montant_lettres, 'apport.montant_lettres')}) euros "
                f"dépendant de notre {regime_matrimonial}."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        qualite_renoncee = required_text(
            regime.qualite_renoncee,
            "regime_communautaire.qualite_renoncee",
        )
        add_paragraph(
            document,
            (
                "Je te notifie, par la présente, mon intention de renoncer à la faculté de "
                f"devenir personnellement {qualite_renoncee} de cette société."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(
            document,
            (
                "En tout état de cause, et conformément aux dispositions du Code civil, je "
                "déclare donner mon consentement à l'apport effectué par mon conjoint."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(document, "Fait pour servir et valoir ce que de droit.")
        add_paragraph(document, f"En {nombre_exemplaires} exemplaires", space_before_pt=5)
        add_right_aligned_lines(document, [_conjoint_signature(ctx)], space_after_pt=0)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        document.save(output_path)
        return output_path


def _date_courrier_avertissement(ctx: DocumentGenerationContext) -> str:
    regime = required_regime_communautaire(ctx.regime_communautaire)
    if regime.date_courrier_avertissement is not None:
        return format_display_date(
            regime.date_courrier_avertissement,
            "regime_communautaire.date_courrier_avertissement",
        )
    if regime.avertissement is not None and regime.avertissement.date_signature is not None:
        return format_display_date(
            regime.avertissement.date_signature,
            "regime_communautaire.avertissement.date_signature",
        )
    raise ValueError(
        "regime_communautaire.date_courrier_avertissement ou "
        "regime_communautaire.avertissement.date_signature est obligatoire pour CODE-RC-001."
    )


def _regime_matrimonial_display(value: str) -> str:
    normalized = (
        normalize("NFKD", value).encode("ascii", "ignore").decode("ascii").lower()
    )
    if "communaute" in " ".join(normalized.split()):
        return "communauté"
    for prefix in ("sous le régime de ", "sous le regime de ", "régime de ", "regime de "):
        if value.lower().startswith(prefix):
            return value[len(prefix) :].strip()
    return value.strip()


def _apporteur_appel(ctx: DocumentGenerationContext) -> str:
    apporteur = ctx.personne_signataire
    civilite = required_text(apporteur.civilite, "apporteur.civilite_affichage")
    prenom = required_text(apporteur.prenom, "apporteur.prenom")
    nom = required_text(apporteur.nom, "apporteur.nom")
    return f"{civilite} {prenom} {nom},"


def _required_conjoint(ctx: DocumentGenerationContext) -> Person:
    conjoint = ctx.conjoint
    if conjoint is None:
        raise ValueError("conjoint est obligatoire pour CODE-RC-001.")
    return conjoint


def _conjoint_signature(ctx: DocumentGenerationContext) -> str:
    conjoint = _required_conjoint(ctx)
    prenom = required_text(conjoint.prenom, "conjoint.prenom")
    nom = required_text(conjoint.nom, "conjoint.nom")
    return f"{prenom} {nom}"
