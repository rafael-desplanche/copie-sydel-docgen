# ruff: noqa: E501

from __future__ import annotations

import glob
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Literal

from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.domain.models import (
    Address,
    CessionAcquereur,
    CessionBailProfessionnel,
    CessionCabinet,
    CessionConjoint,
    CessionContext,
    CessionCreditVendeur,
    CessionExercice,
    CessionFinancement,
    CessionPret,
    CessionPrix,
    CessionRepresentant,
    CessionSalarie,
    CessionValidations,
    CessionVendeur,
    DocumentContext,
    DocumentGenerationContext,
)
from sydel_doc_engine.utils.grammar import apply_gender_pairs

DOCUMENT_CODE = "CODE-CESSION-CAB-001"

ACTE = "acte"
COMPROMIS = "compromis"
MEDICAL = "medical"
DENTAIRE = "dentaire"
SUPPORTED_STRUCTURES = {"SELARL", "SELAS"}
SUPPORTED_ETAPES = {ACTE, COMPROMIS}
SUPPORTED_CABINET_TYPES = {MEDICAL, DENTAIRE}

# Origine de propriete (regle NotebookLM) : la clause decrit le VENDEUR (cedant),
# comment il est devenu proprietaire. Deux variantes standard ; defaut = "cree"
# si le praticien n'a pas achete son cabinet. Tout autre cas = COMPLEXE -> texte
# libre saisi a la main + validation explicite (souplesse / relecture humaine).
ORIGINE_MODE_CREE = "cree"
ORIGINE_MODE_ACHETE = "achete"
SUPPORTED_ORIGINE_MODES = {ORIGINE_MODE_CREE, ORIGINE_MODE_ACHETE}

# Convention systeme : une liste vide (0 element) se rend "Néant", a l'image des
# apports en nature inexistants. Utilisee pour la reprise des salaries (0/1/N).
NEANT = "Néant."

# Dossier des modeles Word tokenises, resolu independamment du cwd.
# parents[4] depuis src/sydel_doc_engine/generators/lot_03/ = racine du repo.
_SOURCE_MODELS_DIR = (
    Path(__file__).resolve().parents[4] / "project" / "source_documents" / "lot_03"
)

# Motif glob par variante (etape, type_cabinet) -> motif robuste aux accents/apostrophes.
_MODEL_GLOB_BY_VARIANT: dict[tuple[str, str], str] = {
    (ACTE, MEDICAL): "Acte*cession*m*dical*.docx",
    (ACTE, DENTAIRE): "Acte*cession*dentaire*.docx",
    (COMPROMIS, MEDICAL): "Compromis*cession*m*dical*.docx",
    (COMPROMIS, DENTAIRE): "Compromis*cession*dentaire*.docx",
}

# Mois francais accentues pour un rendu fidele "10 mars 1975".
_MONTHS_FR = (
    "",
    "janvier",
    "février",
    "mars",
    "avril",
    "mai",
    "juin",
    "juillet",
    "août",
    "septembre",
    "octobre",
    "novembre",
    "décembre",
)

_ISO_DATE_RE = re.compile(r"^(\d{4})-(\d{2})-(\d{2})$")
_TOKEN_RE = re.compile(r"\[[^\]\[]+\]")


@dataclass(frozen=True)
class CessionCabinetVariant:
    etape: Literal["acte", "compromis"]
    type_cabinet: Literal["medical", "dentaire"]
    output_filename: str


def generate_cession_cabinet_docx(
    ctx: DocumentGenerationContext,
    output_dir: Path,
    variant: CessionCabinetVariant,
) -> Path:
    # Validation metier conservee (regles credit-vendeur / SCM = acte medical uniquement, etc.).
    _validate_context(ctx, variant)
    model_path = _resolve_model_path(variant)
    replacements = _build_cession_replacements(ctx)
    gender_pairs = _build_cession_gender_pairs(ctx)
    output_path = output_dir / variant.output_filename
    return render_cession_from_template(
        model_path,
        replacements,
        output_path,
        gender_pairs=gender_pairs,
    )


# Paires d'accord en genre des modeles de cession, pilotees par la BONNE personne.
# Chaines EXACTES figees relevees dans project/source_documents/lot_03/ :
#  - Acte dentaire : fige au FEMININ ("née le", "Inscrite au tableau",
#    "domiciliée en cette qualité").
#  - Compromis dentaire / medical : fige au MASCULIN ("né le", "inscrit au tableau",
#    "domicilié en cette qualité").
#  - Acte medical : "né(e) le" inclusif (non touche : aucune paire ne le matche).
# JAMAIS de regex de terminaison : uniquement ces chaines litterales ancrees.
# "désigné" (role/invariant) et "Représentée" (la societe, toujours feminin) ne
# sont PAS dans les paires : on n'y touche pas.
_CESSION_VENDEUR_PAIRS: list[tuple[str, str]] = [
    ("né le ", "née le "),
    ("Inscrit au tableau", "Inscrite au tableau"),
    ("inscrit au tableau", "inscrite au tableau"),
]
_CESSION_REPRESENTANT_PAIRS: list[tuple[str, str]] = [
    ("domicilié en cette qualité", "domiciliée en cette qualité"),
]


def _build_cession_gender_pairs(
    ctx: DocumentGenerationContext,
) -> list[tuple[Gender, list[tuple[str, str]]]]:
    """Construit les couples (genre, paires) d'accord pour la cession.

    Le genre vendeur pilote l'identification (« né le », « inscrit au tableau »).
    Le genre du representant de l'acquereur pilote « domicilié en cette qualite ».
    Un genre absent (None, non capture cote front) -> on n'accorde pas cette
    personne et le modele source reste fige tel quel (pas de devinette).
    """
    pairs: list[tuple[Gender, list[tuple[str, str]]]] = []
    cession = ctx.cession
    if cession is None:
        return pairs
    vendeur = cession.vendeur
    if vendeur is not None and vendeur.genre is not None:
        pairs.append((vendeur.genre, _CESSION_VENDEUR_PAIRS))
    acquereur = cession.acquereur
    representant = acquereur.representant if acquereur is not None else None
    if representant is not None and representant.genre is not None:
        pairs.append((representant.genre, _CESSION_REPRESENTANT_PAIRS))
    return pairs


def render_cession_from_template(
    model_path: Path,
    replacements: dict[str, str],
    output_path: Path,
    *,
    gender_pairs: list[tuple[Gender, list[tuple[str, str]]]] | None = None,
) -> Path:
    """Charge le modele tokenise et remplace chaque token [xxx] run par run.

    `gender_pairs` (optionnel) : liste de couples `(genre, paires)` appliques
    APRES le remplacement des tokens et AVANT la securite anti-token-residuel,
    via `grammar.apply_gender_pairs` (corps des paragraphes + cellules de
    tableaux). Chaque entree accorde des chaines EXACTES figees du modele selon
    le `genre` de la BONNE personne (vendeur, representant...). C'est le
    generateur qui pilote les paires : aucune normalisation magique globale.

    Securite anti-trou : si un token [...] subsiste apres remplacement, leve
    ValueError en listant les tokens residuels (un token oublie = un test rouge).
    """
    document = Document(str(model_path))

    for paragraph in _iter_all_paragraphs(document):
        for run in paragraph.runs:
            text = run.text
            if "[" not in text:
                continue
            for token, value in replacements.items():
                if token in text:
                    text = text.replace(token, value)
            if text != run.text:
                run.text = text

    if gender_pairs:
        for paragraph in _iter_all_paragraphs(document):
            _apply_gender_pairs_to_paragraph(paragraph, gender_pairs)

    residual = _collect_residual_tokens(document)
    if residual:
        joined = ", ".join(sorted(residual))
        raise ValueError(
            f"Tokens non remplaces dans {model_path.name} pour {DOCUMENT_CODE} : {joined}."
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _iter_all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _apply_gender_pairs_to_paragraph(
    paragraph,
    gender_pairs: list[tuple[Gender, list[tuple[str, str]]]],
) -> None:
    """Accorde en genre les chaines figees d'un paragraphe (corps + cellules).

    Accord d'abord run par run (preserve la mise en forme). Si une forme a
    accorder est eclatee sur plusieurs runs (le texte attendu du paragraphe n'est
    pas atteint), on reecrit le texte fusionne sur le premier run.
    """
    if not paragraph.runs:
        return

    original_paragraph_text = paragraph.text

    for run in paragraph.runs:
        text = run.text
        for genre, pairs in gender_pairs:
            text = apply_gender_pairs(text, genre, pairs)
        if text != run.text:
            run.text = text

    expected_text = original_paragraph_text
    for genre, pairs in gender_pairs:
        expected_text = apply_gender_pairs(expected_text, genre, pairs)

    if paragraph.text != expected_text:
        runs = paragraph.runs
        runs[0].text = expected_text
        for run in runs[1:]:
            run.text = ""


def _collect_residual_tokens(document) -> set[str]:
    residual: set[str] = set()
    for paragraph in _iter_all_paragraphs(document):
        for match in _TOKEN_RE.findall(paragraph.text):
            residual.add(match)
    return residual


def _resolve_model_path(variant: CessionCabinetVariant) -> Path:
    pattern = _MODEL_GLOB_BY_VARIANT[(variant.etape, variant.type_cabinet)]
    matches = glob.glob(str(_SOURCE_MODELS_DIR / pattern))
    if not matches:
        raise ValueError(
            f"Modele introuvable pour {variant.etape}/{variant.type_cabinet} "
            f"(motif {pattern}) dans {_SOURCE_MODELS_DIR}."
        )
    return Path(matches[0])


# ---------------------------------------------------------------------------
# Construction du dictionnaire token -> valeur
# ---------------------------------------------------------------------------


def _build_cession_replacements(ctx: DocumentGenerationContext) -> dict[str, str]:
    cession = ctx.cession
    if cession is None:
        raise ValueError(f"cession est obligatoire pour {DOCUMENT_CODE}.")
    vendeur = cession.vendeur or CessionVendeur()
    conjoint = vendeur.conjoint or CessionConjoint()
    acquereur = cession.acquereur or CessionAcquereur()
    representant = acquereur.representant or CessionRepresentant()
    cabinet = cession.cabinet or CessionCabinet()
    precedent = cabinet.precedent_proprietaire
    bail = cession.bail_professionnel or CessionBailProfessionnel()
    prix = cession.prix or CessionPrix()
    financement = cession.financement or CessionFinancement()
    credit_vendeur = financement.credit_vendeur or CessionCreditVendeur()
    pret = financement.pret or CessionPret()
    document = ctx.document or DocumentContext()
    signature = ctx.signature

    replacements: dict[str, str] = {}

    def put(token: str, value: object | None) -> None:
        # Les valeurs None ne sont PAS injectees -> token preserve -> anti-trou (etape 1.3).
        if value is None:
            return
        replacements[token] = str(value)

    # --- Vendeur ---
    put("[civilite_vendeur]", vendeur.civilite_affichage)
    put("[prenom_vendeur]", vendeur.prenom)
    put("[nom_vendeur]", vendeur.nom)
    put("[profession_vendeur]", vendeur.profession)
    put("[date_naissance_vendeur]", _french_date(vendeur.date_naissance))
    put("[ville_naissance_vendeur]", vendeur.ville_naissance)
    put("[departement_naissance_vendeur]", vendeur.departement_naissance)
    put("[cp_naissance_vendeur]", vendeur.cp_naissance)
    put("[pays_naissance_vendeur]", vendeur.pays_naissance)
    put("[nationalite_vendeur]", vendeur.nationalite)
    put("[adresse_vendeur]", vendeur.adresse_affichee)
    put("[adresse_exercice_vendeur]", vendeur.adresse_exercice_affichee)
    put("[numero_siren_vendeur]", vendeur.numero_siren)
    put("[numero_ordre_vendeur]", vendeur.numero_ordre)
    put("[numero_rpps_vendeur]", vendeur.numero_rpps)
    put("[ordre_departemental_vendeur]", vendeur.ordre_departemental)
    put("[situation_maritale_vendeur]", vendeur.situation_maritale)
    put("[regime_matrimonial_vendeur]", vendeur.regime_matrimonial)
    put("[civilite_conjoint_vendeur]", conjoint.civilite_affichage)
    put("[prenom_conjoint_vendeur]", conjoint.prenom)
    put("[nom_conjoint_vendeur]", conjoint.nom)

    # --- Acquereur ---
    put("[denomination_societe_acquereur]", acquereur.denomination_societe)
    put("[forme_sociale_acquereur]", acquereur.forme_sociale)
    put("[capital_social_acquereur]", acquereur.capital_social)
    put("[adresse_siege_acquereur]", _address_label(acquereur.siege))
    put("[ville_rcs_acquereur]", acquereur.rcs_ville)
    put("[numero_rcs_acquereur]", acquereur.numero_rcs)
    put("[numero_siret_acquereur]", acquereur.numero_siret)
    put("[date_immatriculation_acquereur]", _french_date(acquereur.date_immatriculation))
    put("[date_inscription_ordre_acquereur]", _french_date(acquereur.date_inscription_ordre))
    put("[civilite_acquereur_representant]", representant.civilite_affichage)
    put("[prenom_acquereur_representant]", representant.prenom)
    put("[nom_acquereur_representant]", representant.nom)
    put("[fonction_acquereur_representant]", representant.fonction)

    # --- Cabinet ---
    put("[adresse_cabinet]", cabinet.adresse_affichee)
    put("[adresse_locaux]", cabinet.adresse_locaux_affichee)
    put("[telephone_cabinet]", cabinet.telephone)
    put("[superficie_local]", cabinet.superficie_local)
    put("[nature_fonds_liberal]", cabinet.nature_fonds_liberal)
    put("[description_origine_propriete]", cabinet.description_origine_propriete)
    put("[date_origine_propriete]", _french_date(cabinet.date_origine_propriete))
    put("[annees_acquisition_patientele]", cabinet.annees_acquisition_patientele)
    put("[prix_origine_propriete]", cabinet.prix_origine_propriete)
    if precedent is not None:
        put("[civilite_precedent_proprietaire]", precedent.civilite_affichage)
        put("[prenom_precedent_proprietaire]", precedent.prenom)
        put("[nom_precedent_proprietaire]", precedent.nom)
    # Origine de propriete (modeles MEDICAUX) : phrase decrivant le VENDEUR,
    # variante creee/achetee (defaut "cree"), ou texte libre pour un cas complexe.
    put("[origine_propriete_phrase]", _build_origine_propriete_phrase(cession))

    # --- Bail professionnel ---
    put("[date_bail]", _french_date(bail.date_bail))
    put("[duree_bail]", bail.duree)
    put("[date_debut_bail]", _french_date(bail.date_debut))
    put("[date_fin_bail]", _french_date(bail.date_fin))
    put("[date_reconduction_bail_1]", _french_date(bail.date_reconduction_1))
    put("[date_reconduction_bail_2]", _french_date(bail.date_reconduction_2))
    put("[loyer_mensuel]", bail.loyer_mensuel)

    # --- Prix ---
    put("[prix_cession]", prix.total)
    put("[prix_cession_lettres]", prix.total_lettres)
    put("[prix_elements_corporels]", prix.elements_corporels)
    put("[prix_elements_corporels_lettres]", prix.elements_corporels_lettres)
    put("[prix_elements_incorporels]", prix.elements_incorporels)
    put("[prix_elements_incorporels_lettres]", prix.elements_incorporels_lettres)

    # --- Financement : credit-vendeur (acte medical) et pret (compromis) ---
    put("[montant_credit_vendeur]", credit_vendeur.montant)
    put("[duree_credit_vendeur]", credit_vendeur.duree)
    put("[taux_credit_vendeur]", credit_vendeur.taux)
    put("[majoration_interet_retard]", credit_vendeur.majoration_interet_retard)
    put("[montant_pret]", pret.montant)
    put("[taux_pret]", pret.taux)
    put("[duree_pret]", pret.duree)

    # --- SCM (acte medical) ---
    if cession.scm is not None:
        put("[nb_parts_scm_a_ceder]", cession.scm.nb_parts_a_ceder)

    # --- Conditions suspensives (compromis) ---
    put("[date_realisation_limite]", _french_date(cession.date_limite_realisation))

    # --- Salaries (acte dentaire) : reprise 0 / 1 / N (regle NotebookLM) ---
    # 0 salarie -> "Néant" (convention systeme) ; 1..N -> liste nom/prenom/poste.
    put("[clause_reprise_salaries]", _build_clause_reprise_salaries(cession.salaries))
    # [date_entree_jouissance] (dentaire) : source choisie = date de debut du bail
    # professionnel (entree en jouissance des locaux). A confirmer cote metier.
    put("[date_entree_jouissance]", _french_date(bail.date_debut))

    # --- Exercices ---
    for index in (0, 1, 2):
        if index < len(cession.exercices):
            exercice = cession.exercices[index]
            put(f"[exercice_{index + 1}]", exercice.periode)
            put(f"[chiffre_affaires_{index + 1}]", exercice.chiffre_affaires)
            put(f"[resultat_{index + 1}]", exercice.resultat)

    # --- Document / signature ---
    put("[lieu_signature]", signature.lieu)
    put("[date_signature]", _french_date(signature.date))
    put("[nombre_exemplaires_lettres]", document.nombre_exemplaires_lettres)
    put("[nombre_pages_lettres]", document.nombre_pages_lettres)
    put("[signature_vendeur]", _person_label(vendeur.civilite_affichage, vendeur.prenom, vendeur.nom))
    put(
        "[signature_acquereur]",
        _person_label(representant.civilite_affichage, representant.prenom, representant.nom),
    )

    return replacements


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _french_date(value: date | str | None) -> str | None:
    """Formate une date en francais long (ex. "10 mars 1975").

    - date -> jour mois annee en francais ;
    - str ISO "YYYY-MM-DD" -> parsee puis formatee FR ;
    - autre str -> renvoyee telle quelle ;
    - None -> None (laisse le token en place pour l'anti-trou).
    """
    if value is None:
        return None
    if isinstance(value, date):
        return f"{value.day} {_MONTHS_FR[value.month]} {value.year}"
    text = value.strip()
    match = _ISO_DATE_RE.match(text)
    if match is not None:
        year, month, day = (int(part) for part in match.groups())
        try:
            parsed = date(year, month, day)
        except ValueError:
            return text
        return f"{parsed.day} {_MONTHS_FR[parsed.month]} {parsed.year}"
    return text


def _person_label(
    civilite: str | None,
    prenom: str | None,
    nom: str | None,
) -> str | None:
    parts = [part for part in (civilite, prenom, nom) if part]
    if not parts:
        return None
    return " ".join(parts)


def _build_origine_propriete_phrase(cession: CessionContext) -> str | None:
    """Construit la clause d'origine de propriete des modeles MEDICAUX.

    Regle NotebookLM : la clause decrit le VENDEUR (cedant) — comment il est
    devenu proprietaire. Deux variantes standard, defaut "cree" :
      - "cree"   -> "... pour l'avoir regulierement cree le <date>."
      - "achete" -> "... pour l'avoir regulierement acquis aupres de <precedent>,
                     le <date> au prix de <prix> euros."
    Un cas COMPLEXE (mode non standard) est porte par le texte libre
    `cabinet.description_origine_propriete`, sous garde-fou de validation manuelle
    (cf. `_validate_origine_propriete`). On reutilise le wording deja valide des
    modeles dentaires (meme sujet vendeur) : aucune reecriture libre.
    """
    cabinet = cession.cabinet or CessionCabinet()
    vendeur = cession.vendeur or CessionVendeur()

    mode = (cabinet.origine_propriete_mode or ORIGINE_MODE_CREE).strip().lower()
    sujet = _person_label(vendeur.civilite_affichage, vendeur.prenom, vendeur.nom)
    description = (cabinet.description_origine_propriete or "").strip()

    # Cas COMPLEXE / non standard -> texte libre saisi a la main (relecture humaine).
    if mode not in SUPPORTED_ORIGINE_MODES:
        return description or None

    if sujet is None:
        # Identite vendeur incomplete -> on ne devine pas, on laisse le token
        # en place (anti-trou) sauf si un texte libre a ete fourni.
        return description or None

    date_origine = _french_date(cabinet.date_origine_propriete)

    if mode == ORIGINE_MODE_CREE:
        if date_origine is None:
            return description or None
        phrase = (
            f"{sujet} est propriétaire des éléments constitutifs du cabinet "
            f"pour l’avoir régulièrement créé le {date_origine}."
        )
    else:  # ORIGINE_MODE_ACHETE
        precedent = cabinet.precedent_proprietaire
        precedent_label = (
            _person_label(precedent.civilite_affichage, precedent.prenom, precedent.nom)
            if precedent is not None
            else None
        )
        prix = (cabinet.prix_origine_propriete or "").strip()
        if date_origine is None or precedent_label is None or not prix:
            return description or None
        phrase = (
            f"{sujet} est propriétaire des éléments constitutifs du cabinet "
            f"pour les avoir régulièrement acquis auprès de {precedent_label}, "
            f"le {date_origine} au prix de {prix} euros."
        )

    # Complement libre eventuel (precisions metier) appose tel quel.
    if description:
        phrase = f"{phrase} {description}"
    return phrase


def _build_clause_reprise_salaries(salaries: list[CessionSalarie]) -> str:
    """Construit la clause de reprise des contrats de travail (acte dentaire).

    Regle NotebookLM : 0 salarie -> "Néant" (convention systeme) ; 1..N salaries
    -> "De reprendre les contrats de travail de <liste>." ou chaque salarie est
    "Civilite Prenom Nom" (+ ", en qualite de <poste>" si le poste est saisi).
    Reutilise le wording de clause existant du modele ; "Néant" applique la
    convention systeme (aucune clause "néant" dediee dans le modele source).
    """
    if not salaries:
        return NEANT

    labels: list[str] = []
    for index, salarie in enumerate(salaries):
        label = _salarie_label(salarie, index)
        poste = (salarie.poste or "").strip()
        if poste:
            label = f"{label}, en qualité de {poste}"
        labels.append(label)

    if len(labels) == 1:
        liste = labels[0]
    else:
        liste = ", ".join(labels[:-1]) + f" et de {labels[-1]}"
    return f"De reprendre les contrats de travail de {liste}."


def _address_label(address: Address | None) -> str | None:
    if address is None:
        return None
    if address.adresse_affichee:
        return address.adresse_affichee
    parts = [address.num_voie, address.voie, address.cp, address.ville]
    joined = " ".join(part for part in parts if part)
    return joined or None


# ---------------------------------------------------------------------------
# Validation metier (inchangee dans son intention : conserve les regles ratifiees)
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class _CessionData:
    ctx: DocumentGenerationContext
    cession: CessionContext
    vendeur: CessionVendeur
    acquereur: CessionAcquereur
    representant: CessionRepresentant
    cabinet: CessionCabinet
    bail: CessionBailProfessionnel
    prix: CessionPrix
    exercices: list[CessionExercice]
    financement: CessionFinancement
    document: DocumentContext
    validations: CessionValidations


def _validate_context(
    ctx: DocumentGenerationContext,
    variant: CessionCabinetVariant,
) -> _CessionData:
    if ctx.structure not in SUPPORTED_STRUCTURES:
        supported = ", ".join(sorted(SUPPORTED_STRUCTURES))
        raise ValueError(f"dossier.structure doit etre dans [{supported}] pour {DOCUMENT_CODE}.")
    if ctx.dossier_options is None or not ctx.dossier_options.cession:
        raise ValueError(f"dossier.options.cession doit etre vrai pour {DOCUMENT_CODE}.")
    cession = _required_cession(ctx)
    _validate_selection(cession, variant)

    vendeur = _required_vendeur(cession.vendeur)
    acquereur = _required_acquereur(cession.acquereur)
    representant = _required_representant(acquereur.representant)
    cabinet = _required_cabinet(cession.cabinet)
    bail = _required_bail(cession.bail_professionnel)
    prix = _required_prix(cession.prix)
    financement = cession.financement or CessionFinancement()
    document = _required_document(ctx.document)
    validations = cession.validations or CessionValidations()

    exercices = _required_exercices(cession.exercices)
    _validate_arbitrage_blocks(cession, variant, validations)
    _validate_financement(cession, variant, financement)
    _validate_salaries(cession, variant, validations)
    _validate_origine_propriete(cession, variant, cabinet, validations)

    return _CessionData(
        ctx=ctx,
        cession=cession,
        vendeur=vendeur,
        acquereur=acquereur,
        representant=representant,
        cabinet=cabinet,
        bail=bail,
        prix=prix,
        exercices=exercices,
        financement=financement,
        document=document,
        validations=validations,
    )


def _validate_selection(cession: CessionContext, variant: CessionCabinetVariant) -> None:
    type_cabinet = _required_text(cession.type_cabinet, "cession.type_cabinet").lower()
    if type_cabinet not in SUPPORTED_CABINET_TYPES:
        supported = ", ".join(sorted(SUPPORTED_CABINET_TYPES))
        raise ValueError(f"cession.type_cabinet doit etre dans [{supported}] pour {DOCUMENT_CODE}.")
    if type_cabinet != variant.type_cabinet:
        raise ValueError(
            f"cession.type_cabinet doit etre {variant.type_cabinet} pour {variant.output_filename}."
        )

    etape = _required_text(cession.etape, "cession.etape").lower()
    if etape not in SUPPORTED_ETAPES:
        supported = ", ".join(sorted(SUPPORTED_ETAPES))
        raise ValueError(f"cession.etape doit etre dans [{supported}] pour {DOCUMENT_CODE}.")
    if etape != variant.etape:
        raise ValueError(f"cession.etape doit etre {variant.etape} pour {variant.output_filename}.")


def _validate_arbitrage_blocks(
    cession: CessionContext,
    variant: CessionCabinetVariant,
    validations: CessionValidations,
) -> None:
    if variant.type_cabinet == MEDICAL and not validations.mentions_bail_medical_validees:
        raise ValueError(
            "cession.validations.mentions_bail_medical_validees doit etre vrai pour "
            f"{DOCUMENT_CODE}."
        )
    if (
        variant.etape == COMPROMIS
        and variant.type_cabinet == MEDICAL
        and not validations.origine_compromis_medical_validee
    ):
        raise ValueError(
            "cession.validations.origine_compromis_medical_validee doit etre vrai pour "
            f"{DOCUMENT_CODE}."
        )
    if variant.etape == COMPROMIS and not validations.date_realisation_compromis_validee:
        raise ValueError(
            "cession.validations.date_realisation_compromis_validee doit etre vrai pour "
            f"{DOCUMENT_CODE}."
        )
    if (
        variant.etape == ACTE
        and variant.type_cabinet == MEDICAL
        and not validations.ligne_contrats_travail_medical_supprimee
    ):
        raise ValueError(
            "cession.validations.ligne_contrats_travail_medical_supprimee doit etre vrai "
            f"pour {DOCUMENT_CODE}."
        )


def _validate_financement(
    cession: CessionContext,
    variant: CessionCabinetVariant,
    financement: CessionFinancement,
) -> None:
    credit_vendeur = financement.credit_vendeur
    if credit_vendeur is not None and credit_vendeur.actif:
        if not (variant.etape == ACTE and variant.type_cabinet == MEDICAL):
            raise ValueError(
                "cession.financement.credit_vendeur.actif est autorise uniquement pour "
                f"l'acte medical {DOCUMENT_CODE}."
            )
        _required_text(credit_vendeur.montant, "cession.financement.credit_vendeur.montant")
        _required_text(credit_vendeur.duree, "cession.financement.credit_vendeur.duree")
        _required_text(credit_vendeur.taux, "cession.financement.credit_vendeur.taux")
        _required_text(
            credit_vendeur.majoration_interet_retard,
            "cession.financement.credit_vendeur.majoration_interet_retard",
        )

    if cession.scm is not None and cession.scm.actif:
        if not (variant.etape == ACTE and variant.type_cabinet == MEDICAL):
            raise ValueError("cession.scm.actif est autorise uniquement pour l'acte medical.")
        _required_text(cession.scm.nb_parts_a_ceder, "cession.scm.nb_parts_a_ceder")


def _validate_salaries(
    cession: CessionContext,
    variant: CessionCabinetVariant,
    validations: CessionValidations,
) -> None:
    # Reprise des salaries rendue uniquement par l'acte dentaire (seul modele
    # portant la clause). Regle NotebookLM : 0 -> "Néant" ; 1..N -> liste.
    if variant.etape == ACTE and variant.type_cabinet == DENTAIRE:
        # 0..N accepte. Chaque salarie liste doit avoir une identite complete
        # (civilite/prenom/nom). Le poste reste optionnel.
        for index, salarie in enumerate(cession.salaries):
            _salarie_label(salarie, index)
        return
    if cession.salaries:
        raise ValueError(
            f"cession.salaries est rendu uniquement pour l'acte dentaire {DOCUMENT_CODE}."
        )


def _validate_origine_propriete(
    cession: CessionContext,
    variant: CessionCabinetVariant,
    cabinet: CessionCabinet,
    validations: CessionValidations,
) -> None:
    """Garde-fou origine de propriete pour les modeles MEDICAUX (token construit).

    Souplesse cas COMPLEXE : un mode d'origine non standard (ni "cree" ni
    "achete", ex. succession / apport / demembrement) DOIT etre saisi en texte
    libre (`cabinet.description_origine_propriete`) ET valide a la main
    (`validations.origine_propriete_complexe_validee`). Le moteur n'emet pas une
    origine devinee. Les modeles dentaires gardent leur clause figee (non
    concernes par le token construit).
    """
    if variant.type_cabinet != MEDICAL:
        return

    mode = (cabinet.origine_propriete_mode or ORIGINE_MODE_CREE).strip().lower()
    if mode in SUPPORTED_ORIGINE_MODES:
        return

    # Cas complexe : exiger texte libre + validation manuelle (relecture humaine).
    description = (cabinet.description_origine_propriete or "").strip()
    if not description or not validations.origine_propriete_complexe_validee:
        raise ValueError(
            "cession.cabinet.origine_propriete_mode non standard "
            f"({cabinet.origine_propriete_mode!r}) : fournir "
            "cession.cabinet.description_origine_propriete ET "
            "cession.validations.origine_propriete_complexe_validee=True pour "
            f"{DOCUMENT_CODE}."
        )


def _required_cession(ctx: DocumentGenerationContext) -> CessionContext:
    if ctx.cession is None:
        raise ValueError(f"cession est obligatoire pour {DOCUMENT_CODE}.")
    return ctx.cession


def _required_vendeur(vendeur: CessionVendeur | None) -> CessionVendeur:
    if vendeur is None:
        raise ValueError(f"cession.vendeur est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.vendeur.civilite_affichage", vendeur.civilite_affichage),
        ("cession.vendeur.prenom", vendeur.prenom),
        ("cession.vendeur.nom", vendeur.nom),
        ("cession.vendeur.profession", vendeur.profession),
        ("cession.vendeur.date_naissance", vendeur.date_naissance),
        ("cession.vendeur.ville_naissance", vendeur.ville_naissance),
        ("cession.vendeur.nationalite", vendeur.nationalite),
        ("cession.vendeur.adresse_affichee", vendeur.adresse_affichee),
        ("cession.vendeur.situation_maritale", vendeur.situation_maritale),
    ]:
        _required_value(value, field_name)
    return vendeur


def _required_acquereur(acquereur: CessionAcquereur | None) -> CessionAcquereur:
    if acquereur is None:
        raise ValueError(f"cession.acquereur est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.acquereur.denomination_societe", acquereur.denomination_societe),
        ("cession.acquereur.forme_sociale", acquereur.forme_sociale),
        ("cession.acquereur.capital_social", acquereur.capital_social),
        ("cession.acquereur.rcs_ville", acquereur.rcs_ville),
    ]:
        _required_text(value, field_name)
    _required_text(_address_label(acquereur.siege), "cession.acquereur.siege.adresse_affichee")
    return acquereur


def _required_representant(representant: CessionRepresentant | None) -> CessionRepresentant:
    if representant is None:
        raise ValueError(f"cession.acquereur.representant est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.acquereur.representant.civilite_affichage", representant.civilite_affichage),
        ("cession.acquereur.representant.prenom", representant.prenom),
        ("cession.acquereur.representant.nom", representant.nom),
        ("cession.acquereur.representant.fonction", representant.fonction),
    ]:
        _required_text(value, field_name)
    return representant


def _required_cabinet(cabinet: CessionCabinet | None) -> CessionCabinet:
    if cabinet is None:
        raise ValueError(f"cession.cabinet est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.cabinet.adresse_affichee", cabinet.adresse_affichee),
        ("cession.cabinet.adresse_locaux_affichee", cabinet.adresse_locaux_affichee),
        ("cession.cabinet.telephone", cabinet.telephone),
    ]:
        _required_value(value, field_name)
    # description_origine_propriete n'est plus un token autonome : la clause
    # d'origine medicale est construite a partir des donnees vendeur (mode
    # cree/achete). Le texte libre n'est exige que pour un cas COMPLEXE
    # (cf. _validate_origine_propriete).
    return cabinet


def _required_bail(bail: CessionBailProfessionnel | None) -> CessionBailProfessionnel:
    if bail is None:
        raise ValueError(f"cession.bail_professionnel est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.bail_professionnel.date_bail", bail.date_bail),
        ("cession.bail_professionnel.duree", bail.duree),
        (
            "cession.bail_professionnel.activite_autorisee_affichee",
            bail.activite_autorisee_affichee,
        ),
    ]:
        _required_value(value, field_name)
    return bail


def _required_prix(prix: CessionPrix | None) -> CessionPrix:
    if prix is None:
        raise ValueError(f"cession.prix est obligatoire pour {DOCUMENT_CODE}.")
    for field_name, value in [
        ("cession.prix.total", prix.total),
        ("cession.prix.total_lettres", prix.total_lettres),
        ("cession.prix.elements_corporels", prix.elements_corporels),
        ("cession.prix.elements_corporels_lettres", prix.elements_corporels_lettres),
        ("cession.prix.elements_incorporels", prix.elements_incorporels),
        ("cession.prix.elements_incorporels_lettres", prix.elements_incorporels_lettres),
    ]:
        _required_text(value, field_name)
    return prix


def _required_document(document: DocumentContext | None) -> DocumentContext:
    if document is None:
        raise ValueError(f"document est obligatoire pour {DOCUMENT_CODE}.")
    _required_text(document.nombre_pages_lettres, "document.nombre_pages_lettres")
    _required_text(document.nombre_exemplaires_lettres, "document.nombre_exemplaires_lettres")
    return document


def _required_exercices(exercices: list[CessionExercice]) -> list[CessionExercice]:
    if len(exercices) != 3:
        raise ValueError("cession.exercices doit contenir exactement trois lignes.")
    for index, exercice in enumerate(exercices):
        prefix = f"cession.exercices[{index}]"
        _required_text(exercice.periode, f"{prefix}.periode")
        _required_text(exercice.chiffre_affaires, f"{prefix}.chiffre_affaires")
        _required_text(exercice.resultat, f"{prefix}.resultat")
    return exercices


def _salarie_label(salarie: CessionSalarie, index: int) -> str:
    field_name = f"cession.salaries[{index}]"
    return (
        f"{_required_text(salarie.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{_required_text(salarie.prenom, f'{field_name}.prenom')} "
        f"{_required_text(salarie.nom, f'{field_name}.nom')}"
    )


def _required_value(value: date | str | None, field_name: str) -> date | str:
    if value is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    if isinstance(value, str) and not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value


def _required_text(value: str | None, field_name: str) -> str:
    if value is None or not value.strip():
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    return value.strip()
