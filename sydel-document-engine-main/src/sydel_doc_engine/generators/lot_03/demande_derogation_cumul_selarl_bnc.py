from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import Company, DocumentGenerationContext, Person
from sydel_doc_engine.generators.lot_03.derogations_common import (
    CUMUL_SEL_BNC,
    MANUAL_BLANK,
    format_display_date,
    require_company,
    require_company_inscription,
    require_derogation_context,
    require_person_contact,
    require_structure,
    required_text,
)
from sydel_doc_engine.rendering.docx_builder import (
    DEROGATION_CUMUL_STYLE_PROFILE,
    add_checkbox_line,
    add_form_section_heading,
    add_italic_instruction,
    add_notice_box,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "demande_derogation_cumul_selarl_bnc_formulaire_a_completer.docx"


class DemandeDerogationCumulSelarlBncGenerator:
    """Generateur partiel de la demande de cumul SELARL / BNC."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        require_structure(ctx, selarl_only=True)
        require_derogation_context(ctx, CUMUL_SEL_BNC)
        company = require_company(ctx)

        docx = new_document(style_profile=DEROGATION_CUMUL_STYLE_PROFILE)
        _add_header(docx)
        _add_principle_notice(docx)
        _add_declarant(docx, ctx.personne_signataire, company)
        _add_company(docx, company)
        _add_lieux_exercice(docx, company)
        _add_motifs(docx)
        _add_certification(docx, ctx)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _add_header(docx) -> None:
    add_paragraph(
        docx,
        "Demande de cumul d'exercices en societe d'exercice liberal (SEL)",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    add_paragraph(docx, "et a titre individuel", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(
        docx,
        "(Articles R.4113-3 et R.4127-85 du Code de la sante publique)",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _add_declarant(docx, signataire: Person, company: Company) -> None:
    contact = require_person_contact(signataire)
    numero_ordre = required_text(
        signataire.numero_inscription_ordre,
        "personne_signataire.numero_inscription_ordre",
    )
    qualification = required_text(
        signataire.qualification_principale,
        "personne_signataire.qualification_principale",
    )
    add_form_section_heading(docx, "Identification du declarant")
    add_paragraph(docx, "Demande formulee par le Docteur :")
    add_paragraph(docx, f"Nom : {required_text(signataire.nom, 'personne_signataire.nom')}")
    add_paragraph(
        docx,
        f"Prenom : {required_text(signataire.prenom, 'personne_signataire.prenom')}",
    )
    inscription = require_company_inscription(company)
    add_paragraph(
        docx,
        (
            "Inscrit au Tableau du Conseil departemental de : "
            f"{required_text(inscription.ville, 'societe.inscription_ordre.ville')}"
        ),
    )
    add_paragraph(docx, f"Sous le numero : {numero_ordre}")
    add_paragraph(docx, f"Qualification principale : {qualification}")
    add_paragraph(
        docx,
        (
            "Autres disciplines exercees (Competences, DESC du groupe 1, VAE ordinale, "
            f"Capacites, Orientations) : {MANUAL_BLANK}"
        ),
    )
    siege = company.siege
    add_paragraph(docx, f"Adresse de correspondance : {_siege_address(company)}")
    add_paragraph(
        docx,
        f"Code postal : {required_text(siege.cp if siege else None, 'societe.siege.cp')}",
    )
    add_paragraph(
        docx,
        f"Commune : {required_text(siege.ville if siege else None, 'societe.siege.ville')}",
    )
    add_paragraph(docx, "Coordonnees :")
    add_paragraph(
        docx,
        (
            "N° de telephone : "
            f"{required_text(contact.telephone, 'personne_signataire.contact.telephone')} ; "
            "|__|__|__|__|__|__|__|__|__|__|"
        ),
    )
    add_paragraph(
        docx,
        (
            "Adresse electronique : "
            f"{required_text(contact.email, 'personne_signataire.contact.email')}"
        ),
    )


def _add_company(docx, company: Company) -> None:
    inscription = require_company_inscription(company)
    add_form_section_heading(docx, "Identification de la societe (SEL)")
    add_paragraph(
        docx,
        f"Denomination sociale : {required_text(company.denomination, 'societe.denomination')}",
    )
    add_paragraph(
        docx,
        (
            "Inscrite au Tableau du Conseil departemental de : "
            f"{required_text(inscription.ville, 'societe.inscription_ordre.ville')}"
        ),
    )
    add_paragraph(
        docx,
        (
            "Sous le numero : "
            f"{required_text(inscription.numero, 'societe.inscription_ordre.numero')}"
        ),
    )
    add_paragraph(docx, f"Adresse du siege social : {_siege_address(company)}")


def _add_lieux_exercice(docx, company: Company) -> None:
    add_form_section_heading(docx, "Lieux d'exercices")
    add_paragraph(docx, "Concernant votre exercice a titre individuel :")
    add_paragraph(docx, "Type d'activite :         Salariee        □ Liberale")
    add_paragraph(docx, f"Adresse : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        f"Temps hebdomadaire consacre (nombre de demi-journees) : {MANUAL_BLANK}",
    )
    add_paragraph(docx, "Concernant votre exercice en SEL :")
    add_paragraph(
        docx,
        (
            "Adresse de la residence professionnelle de votre SEL (activite principale) : "
            f"{_siege_address(company)}"
        ),
    )
    add_paragraph(
        docx,
        f"Temps hebdomadaire consacre (nombre de demi-journees) : {MANUAL_BLANK}",
    )
    add_paragraph(docx, "Autre(s) site(s) d'exercice deja declare(s) (activite(s) secondaire(s)) :")
    add_checkbox_line(docx, "Aucun")
    add_checkbox_line(docx, f"Oui - nombre de sites : {MANUAL_BLANK}")
    add_paragraph(docx, "1er site distinct :")
    add_paragraph(docx, f"Adresse du site : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        f"Temps hebdomadaire consacre (nombre de demi-journees) : {MANUAL_BLANK}",
    )
    add_paragraph(
        docx,
        (
            "Autres sites distincts (indiquer l'adresse et le temps hebdomadaire "
            f"consacre) : {MANUAL_BLANK}"
        ),
    )
    add_paragraph(docx, "Continuite des soins sur l'ensemble de vos lieux d'exercices :")
    add_paragraph(docx, f"A l'adresse de votre activite a titre individuel : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        (
            "A l'adresse de la residence professionnelle de votre SEL "
            f"(activite principale) : {MANUAL_BLANK}"
        ),
    )
    add_paragraph(
        docx,
        f"A l'adresse du 1er site distinct de votre SEL (activite secondaire) : {MANUAL_BLANK}",
    )
    add_paragraph(
        docx,
        f"A l'adresse des autres sites de votre SEL (activite(s) secondaire(s)) : {MANUAL_BLANK}",
    )


def _add_motifs(docx) -> None:
    add_form_section_heading(
        docx,
        "Critere(s) sur le(s)quel(s) est fondee la demande de cumul",
    )
    add_italic_instruction(docx, "Toute case cochee doit etre accompagnee d'une explication :")
    add_checkbox_line(
        docx,
        (
            "L'exercice dans votre SEL est lie a des techniques medicales necessitant "
            "un regroupement ou un travail en equipe (motif non applicable dans le cadre "
            "d'une SEL unipersonnelle, si vous etes le seul associe)"
        ),
    )
    add_checkbox_line(
        docx,
        (
            "L'exercice dans votre SEL est lie a l'acquisition d'equipements ou de "
            "materiels lourds soumis a autorisation"
        ),
    )
    add_checkbox_line(
        docx,
        (
            "L'exercice dans votre SEL necessite l'acquisition d'equipements ou de "
            "materiels qui justifient des utilisations multiples"
        ),
    )


def _add_certification(docx, ctx: DocumentGenerationContext) -> None:
    prenom = required_text(ctx.personne_signataire.prenom, "personne_signataire.prenom")
    nom = required_text(ctx.personne_signataire.nom, "personne_signataire.nom")
    add_paragraph(docx, f"Je soussigne(e) Dr {prenom} {nom}certifie :", space_before_pt=10)
    add_paragraph(
        docx,
        (
            "L'exactitude de l'ensemble des informations fournies ou jointes au present "
            "formulaire et que toute modification de mes conditions d'exercice sera "
            "communiquee au conseil departemental de ma residence professionnelle,"
        ),
    )
    add_italic_instruction(
        docx,
        (
            "(Le Conseil departemental vous informe que toute declaration volontairement "
            "inexacte ou incomplete faite au Conseil de l'Ordre par un medecin peut "
            "donner lieu a des poursuites disciplinaires, conformement a l'article "
            "R. 4127-110 du Code de la sante publique)"
        ),
    )
    add_paragraph(
        docx,
        (
            "Que l'ouverture du site n'est pas contraire aux dispositions legislatives "
            "et reglementaires."
        ),
    )
    add_paragraph(docx, f"Fait le {format_display_date(ctx.signature.date, 'signature.date')}")
    add_paragraph(docx, f"a {required_text(ctx.signature.lieu, 'signature.lieu')}")
    add_paragraph(docx, "Signature :")
    add_notice_box(
        docx,
        [
            "PIECES A JOINDRE AU PRESENT FORMULAIRE DE DECLARATION",
            "Projet d'acte constitutif ou justificatif utile selon la demande.",
        ],
        style_profile=DEROGATION_CUMUL_STYLE_PROFILE,
    )


def _add_principle_notice(docx) -> None:
    add_notice_box(
        docx,
        [
            (
                "En principe, lorsqu'un medecin decide d'exercer en SEL, il ne peut "
                "cumuler cette activite avec un exercice a titre individuel."
            ),
            (
                "Cependant, une derogation peut etre demandee dans les cas prevus par "
                "les textes applicables."
            ),
        ],
        style_profile=DEROGATION_CUMUL_STYLE_PROFILE,
    )


def _siege_address(company: Company) -> str:
    siege = company.siege
    return required_text(
        siege.adresse_affichee if siege else None,
        "societe.siege.adresse_affichee",
    )
