from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    BailContext,
    BailParty,
    Company,
    DocumentContext,
    DocumentGenerationContext,
)
from sydel_doc_engine.generators.lot_03.bail_appel_common import (
    DOCUMENT_CODE,
    format_display_date,
    required_text,
    validate_avenant_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    BAIL_COMPACT_STYLE_PROFILE,
    add_article_heading,
    add_framed_title,
    add_paragraph,
    add_party_marker,
    add_signature_table,
    new_document,
)

OUTPUT_FILENAME = "avenant_contrat_bail.docx"


class AvenantContratBailGenerator:
    """Generateur from-scratch de l'avenant au contrat de bail."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_avenant_context(ctx)
        bail = _required_bail(ctx.bail)
        company = _required_company(ctx.societe)
        document_context = _required_document_context(ctx.document)
        bailleur = _required_party(bail.bailleur, "bail.bailleur")
        locataire = _required_party(bail.locataire, "bail.locataire")

        if not bail.societe_en_cours_immatriculation:
            raise ValueError(
                "bail.societe_en_cours_immatriculation doit etre confirme pour "
                f"{DOCUMENT_CODE}."
            )
        if not bail.bailleur_accepte_changement_locataire:
            raise ValueError(
                "bail.bailleur_accepte_changement_locataire doit etre confirme pour "
                f"{DOCUMENT_CODE}."
            )

        docx = new_document(style_profile=BAIL_COMPACT_STYLE_PROFILE)
        add_framed_title(
            docx,
            [
                (
                    "Avenant n°1 au bail du "
                    f"{format_display_date(bail.date_avenant, 'bail.date_avenant')}"
                )
            ],
            style_profile=BAIL_COMPACT_STYLE_PROFILE,
        )
        _add_parties(docx, bailleur, locataire)
        _add_article_1(docx, bail, locataire, company)
        _add_article_2(docx, locataire)
        _add_article_3(docx)
        nombre_exemplaires = required_text(
            document_context.nombre_exemplaires_lettres,
            "document.nombre_exemplaires_lettres",
        )
        add_paragraph(
            docx,
            (
                f"Fait à {required_text(ctx.signature.lieu, 'signature.lieu')} en "
                f"{nombre_exemplaires} exemplaires, le "
                f"{format_display_date(ctx.signature.date, 'signature.date')}"
            ),
        )
        _add_signature_table(docx)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _required_bail(bail: BailContext | None) -> BailContext:
    if bail is None:
        raise ValueError(f"bail est obligatoire pour {DOCUMENT_CODE}.")
    return bail


def _required_company(company: Company | None) -> Company:
    if company is None:
        raise ValueError(f"societe est obligatoire pour {DOCUMENT_CODE}.")
    return company


def _required_document_context(document_context: DocumentContext | None) -> DocumentContext:
    if document_context is None:
        raise ValueError(f"document est obligatoire pour {DOCUMENT_CODE}.")
    return document_context


def _required_party(party: BailParty | None, field_name: str) -> BailParty:
    if party is None:
        raise ValueError(f"{field_name} est obligatoire pour {DOCUMENT_CODE}.")
    required_text(party.civilite_affichage, f"{field_name}.civilite_affichage")
    required_text(party.prenom, f"{field_name}.prenom")
    required_text(party.nom, f"{field_name}.nom")
    required_text(party.profession, f"{field_name}.profession")
    format_display_date(party.date_naissance, f"{field_name}.date_naissance")
    required_text(party.ville_naissance, f"{field_name}.ville_naissance")
    required_text(party.nationalite, f"{field_name}.nationalite")
    required_text(party.adresse_affichee, f"{field_name}.adresse_affichee")
    return party


def _party_identity(party: BailParty, field_name: str) -> str:
    return (
        f"{required_text(party.civilite_affichage, f'{field_name}.civilite_affichage')} "
        f"{required_text(party.prenom, f'{field_name}.prenom')} "
        f"{required_text(party.nom, f'{field_name}.nom')}"
    )


def _party_full_line(party: BailParty, field_name: str) -> str:
    return (
        f"{_party_identity(party, field_name)}, "
        f"{required_text(party.profession, f'{field_name}.profession')}, né le "
        f"{format_display_date(party.date_naissance, f'{field_name}.date_naissance')}, à "
        f"{required_text(party.ville_naissance, f'{field_name}.ville_naissance')} de nationalité "
        f"{required_text(party.nationalite, f'{field_name}.nationalite')}, demeurant "
        f"{required_text(party.adresse_affichee, f'{field_name}.adresse_affichee')},"
    )


def _add_article_1(
    docx,
    bail: BailContext,
    locataire: BailParty,
    company: Company,
) -> None:
    _add_article_title(docx, "ARTICLE 1 : changement de locataire")
    add_paragraph(
        docx,
        (
            "Le bail signé en date du "
            f"{format_display_date(bail.date_signature_origine, 'bail.date_signature_origine')}, "
            f"a pour locataire {_party_identity(locataire, 'bail.locataire')}, "
            f"({required_text(locataire.profession, 'bail.locataire.profession')})."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    siege = company.siege
    adresse_siege = required_text(
        siege.adresse_affichee if siege else None,
        "societe.siege.adresse_affichee",
    )
    add_paragraph(
        docx,
        (
            "Le présent avenant donne bail à la société "
            f"{required_text(company.denomination, 'societe.denomination')} en cours "
            "d’immatriculation au RCS "
            f"{required_text(company.ville_rcs, 'societe.rcs_ville')}, domiciliée au "
            f"{adresse_siege}."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_article_2(docx, locataire: BailParty) -> None:
    _add_article_title(docx, "ARTICLE 2 : Responsabilité pour une société en cours de formation")
    add_paragraph(
        docx,
        (
            f"Le {required_text(locataire.civilite_courte, 'bail.locataire.civilite_courte')} "
            f"{required_text(locataire.prenom, 'bail.locataire.prenom')} "
            f"{required_text(locataire.nom, 'bail.locataire.nom')}, domicilié "
            f"{required_text(locataire.adresse_affichee, 'bail.locataire.adresse_affichee')}, "
            "engage sa responsabilité pour tous les actes passés au nom de la société jusqu’à "
            "l’immatriculation au RCS."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    add_paragraph(
        docx,
        (
            f"{_party_identity(locataire, 'bail.locataire')} s’engage à fournir au Bailleur un "
            "extrait KBIS une fois que les démarches seront finies."
        ),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_article_3(docx) -> None:
    _add_article_title(docx, "ARTICLE 3 : Clauses du bail")
    add_paragraph(
        docx,
        "Le présent avenant ne modifie pas les clauses du bail en cours.",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )


def _add_parties(docx, bailleur: BailParty, locataire: BailParty) -> None:
    add_paragraph(docx, "Entre les soussign\u00e9s :", style_profile=BAIL_COMPACT_STYLE_PROFILE)
    add_paragraph(
        docx,
        _party_full_line(bailleur, "bail.bailleur"),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )
    add_party_marker(
        docx,
        "Ci-apr\u00e8s d\u00e9sign\u00e9 \u00ab le Bailleur \u00bb",
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )
    add_paragraph(docx, "ET :", bold=True, style_profile=BAIL_COMPACT_STYLE_PROFILE)
    add_paragraph(
        docx,
        _party_full_line(locataire, "bail.locataire"),
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )
    add_party_marker(
        docx,
        "Ci-apr\u00e8s d\u00e9sign\u00e9 \u00ab le Locataire \u00bb",
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )
    add_paragraph(
        docx,
        "Les parties conviennent de ce qui suit :",
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )


def _add_article_title(docx, title: str) -> None:
    add_article_heading(docx, title, style_profile=BAIL_COMPACT_STYLE_PROFILE)


def _add_signature_table(docx) -> None:
    add_signature_table(
        docx,
        [
            ["Le Bailleur", "L\u2019ancien locataire", "Le nouveau locataire"],
        ],
        style_profile=BAIL_COMPACT_STYLE_PROFILE,
    )
