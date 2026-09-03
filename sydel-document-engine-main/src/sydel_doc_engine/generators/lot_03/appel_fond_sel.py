from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    CessionAcquereur,
    CessionCabinet,
    CessionDestinataire,
    CessionFinancement,
    CessionVendeur,
    DocumentContext,
    DocumentGenerationContext,
    DocumentSignataire,
)
from sydel_doc_engine.generators.lot_03.bail_appel_common import (
    CABINET_DENTAIRE,
    CABINET_MEDICAL,
    DOCUMENT_CODE,
    cabinet_type,
    format_display_date,
    required_cession,
    required_text,
    validate_appel_fonds_context,
)
from sydel_doc_engine.rendering.docx_builder import (
    add_centered_amount,
    add_header_logo,
    add_italic_instruction,
    add_paragraph,
    add_right_aligned_lines,
    add_subject_heading,
    new_document,
)

OUTPUT_FILENAME = "appel_fond_sel.docx"

# Libelle d'affichage du type de cabinet dans le corps de la lettre. Le type interne est
# normalise ("dentaire"/"medical") ; ici on rend la forme accentuee attendue dans le texte.
_CABINET_TYPE_LABELS = {
    CABINET_DENTAIRE: "dentaire",
    CABINET_MEDICAL: "médical",
}


class AppelFondSelGenerator:
    """Generateur from-scratch de l'appel de fonds SEL."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        validate_appel_fonds_context(ctx)
        cession = required_cession(ctx)
        cabinet_type_label = _CABINET_TYPE_LABELS[cabinet_type(ctx)]
        financement = _required_financement(cession.financement)
        destinataire = _required_destinataire(financement.destinataire)
        cabinet = _required_cabinet(cession.cabinet)
        vendeur = _required_vendeur(cession.vendeur)
        acquereur = _required_acquereur(cession.acquereur)
        signataire = _required_signataire(_required_document_context(ctx.document).signataire)

        docx = new_document()
        # Logo SYDEL en header, aligne a DROITE (retour UAT Rafael DOC-008).
        add_header_logo(docx, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        nom_banque = required_text(
            financement.banque.nom if financement.banque else None,
            "cession.financement.banque.nom",
        )
        destinataire_label = _destinataire_label(destinataire)
        cabinet_label = required_text(
            cabinet.denomination_ou_adresse_affichee,
            "cession.cabinet.denomination_ou_adresse_affichee",
        )
        vendeur_label = _vendeur_label(vendeur)
        acquereur_label = required_text(
            acquereur.denomination_societe,
            "cession.acquereur.denomination_societe",
        )

        # Bloc banque + lieu/date aligne a DROITE (retour UAT Rafael DOC-008).
        add_paragraph(docx, nom_banque, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        add_paragraph(
            docx,
            (
                f"{required_text(ctx.signature.lieu, 'signature.lieu')}, le "
                f"{format_display_date(ctx.signature.date, 'signature.date')}"
            ),
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        add_italic_instruction(
            docx,
            f"A l’attention de {destinataire_label}",
        )
        add_subject_heading(docx, "Objet : demande de déblocage des fonds")
        add_paragraph(docx, "Cher Monsieur,")
        add_paragraph(
            docx,
            (
                "Nous vous remercions de bien vouloir procéder, ce jour, au déblocage des fonds "
                "d’un montant de :"
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_centered_amount(
            docx,
            [
                required_text(
                    financement.montant_deblocage,
                    "cession.financement.montant_deblocage",
                ),
                "€",
            ],
        )
        add_paragraph(
            docx,
            (
                f"pour la cession du cabinet {cabinet_type_label} exploité au "
                f"{cabinet_label} de {vendeur_label} à la Société {acquereur_label}."
            ),
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        )
        add_paragraph(
            docx,
            "Nous vous prions d’agréer, Cher Monsieur, nos salutations distinguées.",
        )
        add_right_aligned_lines(
            docx,
            [
                (
                    f"{required_text(signataire.prenom, 'document.signataire.prenom')} "
                    f"{required_text(signataire.nom, 'document.signataire.nom')}"
                )
            ],
        )

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _required_financement(financement: CessionFinancement | None) -> CessionFinancement:
    if financement is None:
        raise ValueError(f"cession.financement est obligatoire pour {DOCUMENT_CODE}.")
    if financement.banque is None:
        raise ValueError(f"cession.financement.banque est obligatoire pour {DOCUMENT_CODE}.")
    return financement


def _required_destinataire(
    destinataire: CessionDestinataire | None,
) -> CessionDestinataire:
    if destinataire is None:
        raise ValueError(
            f"cession.financement.destinataire est obligatoire pour {DOCUMENT_CODE}."
        )
    return destinataire


def _required_cabinet(cabinet: CessionCabinet | None) -> CessionCabinet:
    if cabinet is None:
        raise ValueError(f"cession.cabinet est obligatoire pour {DOCUMENT_CODE}.")
    return cabinet


def _required_vendeur(vendeur: CessionVendeur | None) -> CessionVendeur:
    if vendeur is None:
        raise ValueError(f"cession.vendeur est obligatoire pour {DOCUMENT_CODE}.")
    return vendeur


def _required_acquereur(acquereur: CessionAcquereur | None) -> CessionAcquereur:
    if acquereur is None:
        raise ValueError(f"cession.acquereur est obligatoire pour {DOCUMENT_CODE}.")
    return acquereur


def _required_document_context(document_context: DocumentContext | None) -> DocumentContext:
    if document_context is None:
        raise ValueError(f"document est obligatoire pour {DOCUMENT_CODE}.")
    return document_context


def _required_signataire(signataire: DocumentSignataire | None) -> DocumentSignataire:
    if signataire is None:
        raise ValueError(f"document.signataire est obligatoire pour {DOCUMENT_CODE}.")
    return signataire


def _destinataire_label(destinataire: CessionDestinataire) -> str:
    civilite = required_text(
        destinataire.civilite_affichage,
        "cession.financement.destinataire.civilite_affichage",
    )
    prenom = required_text(destinataire.prenom, "cession.financement.destinataire.prenom")
    nom = required_text(destinataire.nom, "cession.financement.destinataire.nom")
    return f"{civilite} {prenom} {nom}"


def _vendeur_label(vendeur: CessionVendeur) -> str:
    return (
        f"{required_text(vendeur.civilite_affichage, 'cession.vendeur.civilite_affichage')} "
        f"{required_text(vendeur.prenom, 'cession.vendeur.prenom')} "
        f"{required_text(vendeur.nom, 'cession.vendeur.nom')}"
    )
