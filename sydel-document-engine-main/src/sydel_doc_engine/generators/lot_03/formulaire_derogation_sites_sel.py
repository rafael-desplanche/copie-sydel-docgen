from __future__ import annotations

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH

from sydel_doc_engine.domain.models import (
    Company,
    DerogationContext,
    DerogationRole,
    DocumentGenerationContext,
)
from sydel_doc_engine.generators.lot_03.derogations_common import (
    MANUAL_BLANK,
    MULTI_SITES_SEL,
    optional_display_date,
    require_company,
    require_company_inscription,
    require_derogation_context,
    require_role,
    require_structure,
    required_text,
)
from sydel_doc_engine.rendering.docx_builder import (
    DEROGATION_FORM_STYLE_PROFILE,
    add_checkbox_line,
    add_form_section_heading,
    add_italic_instruction,
    add_paragraph,
    new_document,
)

OUTPUT_FILENAME = "formulaire_derogation_sites_sel_formulaire_a_completer.docx"


class FormulaireDerogationSitesSelGenerator:
    """Generateur partiel du formulaire multi-sites SEL pre-rempli."""

    def generate(self, ctx: DocumentGenerationContext, output_dir: Path) -> Path:
        require_structure(ctx)
        derogation = require_derogation_context(ctx, MULTI_SITES_SEL)
        company = require_company(ctx)
        representant = require_role(
            derogation.representant_legal,
            "derogation.representant_legal",
        )
        associe = require_role(derogation.associe_exercant, "derogation.associe_exercant")

        docx = new_document(style_profile=DEROGATION_FORM_STYLE_PROFILE)
        _add_header(docx)
        _add_identification(docx, company, representant, associe)
        _add_site_declare(docx, ctx)
        _add_activity_sections(docx)
        _add_sites_existants(docx, ctx, derogation)
        _add_conditions(docx)
        _add_certification(docx, representant)

        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / OUTPUT_FILENAME
        docx.save(output_path)
        return output_path


def _add_header(docx) -> None:
    add_paragraph(
        docx,
        (
            "Declaration prealable d'ouverture d'un site distinct de la residence "
            "professionnelle d'une SEL"
        ),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
    )
    add_italic_instruction(
        docx,
        (
            "A adresser au conseil departemental du lieu ou se situe le site au plus tard "
            "deux mois avant le debut d'activite"
        ),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        docx,
        "Article R4113- 23 du code de la sante publique",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _add_identification(
    docx,
    company: Company,
    representant: DerogationRole,
    associe: DerogationRole,
) -> None:
    inscription = require_company_inscription(company)
    add_form_section_heading(docx, "I - Identification du declarant")
    add_paragraph(docx, "Societe", bold=True)
    add_paragraph(
        docx,
        f"Denomination de la SEL : {required_text(company.denomination, 'societe.denomination')}",
    )
    add_paragraph(
        docx,
        (
            "Departement d'inscription de la SEL : "
            f"{required_text(inscription.departement, 'societe.inscription_ordre.departement')}"
        ),
    )
    add_paragraph(
        docx,
        (
            "N° departemental d'inscription de la SEL : "
            f"{required_text(inscription.numero, 'societe.inscription_ordre.numero')}"
        ),
    )
    add_paragraph(docx, f"Adresse du siege social : {_siege_address(company)}")
    add_paragraph(
        docx,
        (
            "SEL mono disciplinaire de (preciser la qualification principale exercee et/ou "
            f"les autres disciplines exercees) : {MANUAL_BLANK}"
        ),
    )
    add_paragraph(
        docx,
        (
            "SEL pluri disciplinaire de (preciser les qualifications principales exercees "
            f"et/ou les autres disciplines exercees) : {MANUAL_BLANK}"
        ),
    )
    add_paragraph(docx, "Representant legal de la societe", bold=True)
    add_paragraph(
        docx,
        (
            f"Nom : {required_text(representant.nom, 'derogation.representant_legal.nom')}      "
            f"Prenom : {required_text(representant.prenom, 'derogation.representant_legal.prenom')}"
        ),
    )
    add_paragraph(
        docx,
        (
            "Mandat (gerant/president/...) : "
            f"{required_text(representant.fonction, 'derogation.representant_legal.fonction')}"
        ),
    )
    add_paragraph(docx, "N° departemental d'inscription au Tableau de l'Ordre :")
    add_paragraph(docx, "N° de telephone")
    add_paragraph(docx, "Fixe                                 Mobile")
    email = representant.contact.email if representant.contact else None
    add_paragraph(
        docx,
        (
            "Adresse electronique : "
            f"{required_text(email, 'derogation.representant_legal.contact.email')}"
        ),
    )
    add_paragraph(
        docx,
        "Identification de l'associe/des associes qui exercera/ont sur le nouveau site",
        bold=True,
    )
    add_paragraph(docx, f"Nom : {required_text(associe.nom, 'derogation.associe_exercant.nom')}")
    add_paragraph(
        docx,
        f"Prenom : {required_text(associe.prenom, 'derogation.associe_exercant.prenom')}",
    )
    add_paragraph(docx, "N° departemental d'inscription au Tableau de l'Ordre :")
    add_paragraph(docx, "Conseil departemental d'inscription :")
    qualification = required_text(
        associe.qualification_principale,
        "derogation.associe_exercant.qualification_principale",
    )
    add_paragraph(docx, f"Qualification : {qualification}")


def _add_site_declare(docx, ctx: DocumentGenerationContext) -> None:
    add_form_section_heading(
        docx,
        "II - Adresse complete du site pour lequel la declaration est faite :",
    )
    if ctx.site_declare and ctx.site_declare.adresse_affichee:
        add_paragraph(docx, ctx.site_declare.adresse_affichee)
    else:
        add_paragraph(docx, MANUAL_BLANK)
    add_paragraph(docx, f"Date previsionnelle de debut d'activite : {_site_declare_date(ctx)}")
    add_italic_instruction(
        docx,
        (
            "(Attention dans le choix de la date, car le Conseil departemental dispose "
            "d'un delai deux mois a compter de la reception de la declaration pour vous "
            "faire connaitre une eventuelle opposition par une decision motivee)."
        ),
    )


def _site_declare_date(ctx: DocumentGenerationContext) -> str:
    if ctx.site_declare is None:
        return MANUAL_BLANK
    return optional_display_date(ctx.site_declare.date_debut_activite)


def _add_activity_sections(docx) -> None:
    add_form_section_heading(
        docx,
        "III- Nature de l'activite envisagee sur le nouveau site :",
    )
    add_paragraph(docx, f"- consultations (decrire): {MANUAL_BLANK}")
    add_paragraph(docx, f"- actes medico techniques (decrire) : {MANUAL_BLANK}")
    add_paragraph(docx, f"- actes chirurgicaux (decrire) : {MANUAL_BLANK}")
    add_paragraph(docx, f"autres : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        f"Temps hebdomadaire consacre (nombre de jours/demi-journees) : {MANUAL_BLANK}",
    )


def _add_sites_existants(
    docx,
    ctx: DocumentGenerationContext,
    derogation: DerogationContext,
) -> None:
    add_form_section_heading(
        docx,
        (
            "IV - Renseignements sur l'activite au lieu de la residence professionnelle "
            "et le cas echeant, sur les autres sites deja autorises"
        ),
    )
    add_paragraph(docx, "Adresse de la residence professionnelle :")
    add_paragraph(docx, "Autres sites d'exercice :")
    present = _sites_existants_present(derogation)
    add_checkbox_line(docx, "NON", checked=not present)
    add_checkbox_line(docx, "OUI", checked=present)
    nombre_sites = str(len(ctx.sites_existants)) if present else MANUAL_BLANK
    add_paragraph(docx, f"Nombre de sites : {nombre_sites}")
    _add_first_site(docx, ctx, present)
    for index in range(2, 5):
        add_paragraph(docx, f"{index}e site")
        add_paragraph(docx, "Date du debut d'activite : ___|___|/|___|___|/|___|___|___|___|")
        add_paragraph(docx, f"Adresse du site : {MANUAL_BLANK}")
        add_paragraph(
            docx,
            f"Temps hebdomadaire consacre (nombre de jours/demi-journees) : {MANUAL_BLANK}",
        )


def _sites_existants_present(derogation: DerogationContext) -> bool:
    if derogation.sites_existants_present is None:
        raise ValueError(
            "derogation.sites_existants_present est obligatoire pour CODE-DEROG-CORE-001."
        )
    return derogation.sites_existants_present


def _add_first_site(docx, ctx: DocumentGenerationContext, present: bool) -> None:
    first_site = ctx.sites_existants[0] if present and ctx.sites_existants else None
    if present and first_site is None:
        raise ValueError(
            "sites_existants[0] est obligatoire lorsque "
            "derogation.sites_existants_present est vrai."
        )
    add_paragraph(docx, "1er site")
    add_paragraph(
        docx,
        (
            "Date du debut d'activite : "
            f"{optional_display_date(first_site.date_debut_activite if first_site else None)}"
        ),
    )
    add_paragraph(docx, f"Adresse du site : {_site_address(first_site)}")
    add_paragraph(
        docx,
        "Temps hebdomadaire consacre (nombre de jours/demi-journees) : "
        f"{_site_temps(first_site)}",
    )


def _add_conditions(docx) -> None:
    add_form_section_heading(docx, "V- Conditions de l'exercice")
    add_paragraph(docx, "Qualite et securite des soins")
    add_paragraph(docx, "Pour les consultations :")
    add_paragraph(docx, f"- moyens en personnel : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        f"- materiels (decrire le type de materiel existant et/ou prevu) : {MANUAL_BLANK}",
    )
    add_paragraph(docx, "Pour les autres actes :")
    add_paragraph(docx, f"- moyens en personnel : {MANUAL_BLANK}")
    add_paragraph(
        docx,
        f"- materiels (decrire le type de materiel existant et/ou prevu) : {MANUAL_BLANK}",
    )
    add_paragraph(docx, "Continuite des soins")
    add_italic_instruction(
        docx,
        (
            "- dispositions prises pour assurer la continuite des soins sur les differents "
            f"sites : {MANUAL_BLANK}"
        ),
    )
    add_paragraph(docx, "Respect des dispositions du code de deontologie medicale :")
    add_paragraph(docx, f"- informations sur l'environnement de travail : {MANUAL_BLANK}")


def _add_certification(docx, representant: DerogationRole) -> None:
    prenom = required_text(representant.prenom, "derogation.representant_legal.prenom")
    nom = required_text(representant.nom, "derogation.representant_legal.nom")
    add_paragraph(docx, f"Je soussigne Monsieur {prenom} {nom} certifie :", space_before_pt=10)
    add_paragraph(
        docx,
        (
            "l'exactitude de l'ensemble des informations fournies ou jointes au present "
            "formulaire et que toute modification de mes conditions d'exercice sera "
            "communiquee au conseil departemental de la residence professionnelle de la SEL,"
        ),
    )
    add_paragraph(
        docx,
        (
            "que l'ouverture du site n'est pas contraire aux dispositions legislatives "
            "et reglementaires."
        ),
    )
    add_paragraph(docx, "Fait le ___|___|/|___|___|/|___|___|___|___| a")
    add_paragraph(docx, "Pieces a joindre :")
    add_paragraph(docx, "- toute piece utile a l'examen de la declaration")
    add_paragraph(docx, "- le(s) projet(s) de contrat(s) relatifs aux locaux ou aux materiels")


def _siege_address(company: Company) -> str:
    siege = company.siege
    return required_text(
        siege.adresse_affichee if siege else None,
        "societe.siege.adresse_affichee",
    )


def _site_address(first_site) -> str:
    if first_site is None or not first_site.adresse_affichee:
        return MANUAL_BLANK
    return first_site.adresse_affichee


def _site_temps(first_site) -> str:
    if first_site is None or not first_site.temps_hebdomadaire:
        return MANUAL_BLANK
    return first_site.temps_hebdomadaire
