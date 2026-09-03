from __future__ import annotations

from collections.abc import Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Literal

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


@dataclass(frozen=True)
class SydelDocxStyleProfile:
    font_name: str = "Roboto"
    font_size_pt: int = 10
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5
    margin_left_cm: float = 2.5
    margin_right_cm: float = 2.5
    standard_space_after_pt: int = 6
    compact_space_after_pt: int = 2
    legal_reminder_space_after_pt: int = 3
    notable_space_before_pt: int = 10
    signature_width_cm: float = 7.0
    signature_image_width_cm: float = 4.0


DEFAULT_STYLE_PROFILE = SydelDocxStyleProfile()
LETTER_WIDE_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_left_cm=3.17,
    margin_right_cm=3.17,
)
STATUTS_STANDARD_STYLE_PROFILE = DEFAULT_STYLE_PROFILE
STATUTS_SPFPL_COMPACT_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_top_cm=2.8,
    margin_bottom_cm=1.6,
    margin_left_cm=2.0,
    margin_right_cm=2.0,
)
STATUTS_CIVIL_COMPACT_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_top_cm=2.8,
    margin_bottom_cm=1.9,
    margin_left_cm=2.35,
    margin_right_cm=2.2,
)
BAIL_COMPACT_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_top_cm=1.75,
    margin_bottom_cm=0.5,
)
DEROGATION_FORM_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_top_cm=2.0,
)
DEROGATION_CUMUL_STYLE_PROFILE = SydelDocxStyleProfile(
    margin_top_cm=3.25,
    margin_bottom_cm=2.0,
)


def new_document(
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    """Create a clean DOCX document with the shared SYDEL style profile applied."""
    document = Document()
    apply_style_profile(document, style_profile)
    return document


_SYDEL_LOGO_PATH = Path(__file__).resolve().parent.parent / "assets" / "logo_sydel.png"


def add_header_logo(
    document: Any,
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    width_cm: float = 4.5,
) -> None:
    """Insere le logo SYDEL (en-tete de marque) dans le header de la page.

    Les generateurs from-scratch reconstruisent le corps du document en code et ont perdu
    le logo qui vivait dans le header des modeles .docx d'origine (retour UAT Rafael) :
    ce helper le retablit. L'image source est `assets/logo_sydel.png` (extrait du modele).
    """
    header = document.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.alignment = alignment
    paragraph.add_run().add_picture(str(_SYDEL_LOGO_PATH), width=Cm(width_cm))


def apply_style_profile(
    document: Any,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> None:
    section = document.sections[0]
    section.top_margin = Cm(style_profile.margin_top_cm)
    section.bottom_margin = Cm(style_profile.margin_bottom_cm)
    section.left_margin = Cm(style_profile.margin_left_cm)
    section.right_margin = Cm(style_profile.margin_right_cm)

    style = document.styles["Normal"]
    style.font.name = style_profile.font_name
    style.font.size = Pt(style_profile.font_size_pt)
    r_fonts = style.element.rPr.rFonts
    for font_attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(font_attribute), style_profile.font_name)


def add_paragraph(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    space_before_pt: int = 0,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(
        style_profile.standard_space_after_pt if space_after_pt is None else space_after_pt
    )
    if alignment is not None:
        paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    return paragraph


def add_subject_heading(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    space_before_pt: int = 0,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=alignment,
        bold=True,
        underline=True,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
        style_profile=style_profile,
    )


def add_party_marker(
    document: Any,
    text: str,
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        bold=True,
        underline=True,
        style_profile=style_profile,
    )


def add_article_heading(
    document: Any,
    text: str,
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold=True,
        underline=True,
        space_before_pt=style_profile.notable_space_before_pt,
        style_profile=style_profile,
    )


def add_form_section_heading(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=alignment,
        bold=True,
        underline=True,
        space_before_pt=style_profile.notable_space_before_pt,
        style_profile=style_profile,
    )


def add_letter_place_date(
    document: Any,
    text: str,
    *,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after_pt=space_after_pt,
        style_profile=style_profile,
    )


def add_right_aligned_lines(
    document: Any,
    lines: Sequence[str],
    *,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    return [
        add_paragraph(
            document,
            line,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            space_after_pt=space_after_pt,
            style_profile=style_profile,
        )
        for line in lines
    ]


def add_form_field(
    document: Any,
    label: str,
    value: str,
    *,
    underline_label: bool = False,
    left_indent_cm: float = 0.0,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    label_run = paragraph.add_run(f"{label} : ")
    label_run.underline = underline_label
    paragraph.add_run(value)
    return paragraph


def add_form_field_pair(
    document: Any,
    left_label: str,
    left_value: str,
    right_label: str,
    right_value: str,
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    paragraph.add_run(f"{left_label} : ")
    paragraph.add_run(left_value)
    paragraph.add_run("      ")
    paragraph.add_run(f"{right_label} : ")
    paragraph.add_run(right_value)
    return paragraph


def add_checkbox_line(
    document: Any,
    label: str,
    *,
    checked: bool = False,
    left_indent_cm: float = 0.7,
    hanging_indent_cm: float = 0.35,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    marker = "\u2612" if checked else "\u2610"
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    paragraph.paragraph_format.first_line_indent = Cm(-hanging_indent_cm)
    paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    paragraph.add_run(f"{marker} {label}")
    return paragraph


def add_right_indented_block(
    document: Any,
    lines: Sequence[str],
    *,
    left_indent_cm: float = 8.5,
    first_line_indent_cm: float | None = None,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    paragraphs = []
    for line in lines:
        paragraph = add_paragraph(
            document,
            line,
            space_after_pt=space_after_pt,
            style_profile=style_profile,
        )
        paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
        if first_line_indent_cm is not None:
            paragraph.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
        paragraphs.append(paragraph)
    return paragraphs


def add_centered_amount(
    document: Any,
    lines: Sequence[str],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    return [
        add_paragraph(
            document,
            line,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_after_pt=style_profile.compact_space_after_pt,
            style_profile=style_profile,
        )
        for line in lines
    ]


def add_company_identity_block(
    document: Any,
    lines: Sequence[str],
    *,
    first_line_bold: bool = True,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    paragraphs = []
    for index, line in enumerate(lines):
        paragraphs.append(
            add_paragraph(
                document,
                line,
                alignment=alignment,
                bold=first_line_bold and index == 0,
                space_after_pt=(
                    style_profile.compact_space_after_pt
                    if space_after_pt is None
                    else space_after_pt
                ),
                style_profile=style_profile,
            )
        )
    return paragraphs


def add_italic_instruction(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    return add_paragraph(
        document,
        text,
        alignment=alignment,
        italic=True,
        space_after_pt=space_after_pt,
        style_profile=style_profile,
    )


def add_hyphen_list_item(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    italic: bool = False,
    space_after_pt: int | None = None,
    left_indent_cm: float = 0.7,
    hanging_indent_cm: float = 0.35,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    paragraph.paragraph_format.first_line_indent = Cm(-hanging_indent_cm)
    paragraph.paragraph_format.space_after = Pt(
        style_profile.standard_space_after_pt if space_after_pt is None else space_after_pt
    )
    if alignment is not None:
        paragraph.alignment = alignment
    paragraph.add_run("- ")
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return paragraph


def add_statuts_title_box(
    document: Any,
    text: str,
    *,
    bordered: bool = True,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if bordered:
        table.style = "Table Grid"
        _set_table_borders(table)
    else:
        _clear_table_borders(table)

    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = style_profile.font_name
    run.font.size = Pt(style_profile.font_size_pt)
    add_spacer(document, space_after_pt=style_profile.standard_space_after_pt)
    return table


def add_statuts_article_heading(
    document: Any,
    text: str,
    *,
    underline: bool = True,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    left_indent_cm: float | None = None,
    space_before_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = add_paragraph(
        document,
        text,
        alignment=alignment,
        bold=True,
        underline=underline,
        space_before_pt=(
            style_profile.notable_space_before_pt
            if space_before_pt is None
            else space_before_pt
        ),
        style_profile=style_profile,
    )
    if left_indent_cm is not None:
        paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    return paragraph


def add_statuts_part_heading(
    document: Any,
    text: str,
    *,
    mode: Literal["paragraph", "boxed"] = "paragraph",
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    if mode == "boxed":
        return add_statuts_title_box(document, text, bordered=True, style_profile=style_profile)
    return add_paragraph(
        document,
        text,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        bold=True,
        space_before_pt=style_profile.notable_space_before_pt,
        style_profile=style_profile,
    )


def add_statuts_body_paragraph(
    document: Any,
    text: str,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = WD_ALIGN_PARAGRAPH.JUSTIFY,
    indent_profile: Literal["none", "left", "first_line", "hanging"] = "none",
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = add_paragraph(
        document,
        text,
        alignment=alignment,
        style_profile=style_profile,
    )
    if indent_profile == "left":
        paragraph.paragraph_format.left_indent = Cm(0.5)
    elif indent_profile == "first_line":
        paragraph.paragraph_format.first_line_indent = Cm(0.5)
    elif indent_profile == "hanging":
        paragraph.paragraph_format.left_indent = Cm(0.7)
        paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    return paragraph


def add_statuts_hanging_list_item(
    document: Any,
    text: str,
    *,
    marker: str | None = "-",
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.7)
    paragraph.paragraph_format.first_line_indent = Cm(-0.35)
    paragraph.paragraph_format.space_after = Pt(style_profile.standard_space_after_pt)
    if marker:
        paragraph.add_run(f"{marker} ")
    paragraph.add_run(text)
    return paragraph


def add_statuts_signature_block(
    document: Any,
    lines: Sequence[str],
    *,
    mention_lines: Sequence[str] = (),
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
    bold: bool = False,
    underline: bool = False,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    paragraphs = []
    for line in lines:
        paragraphs.append(
            add_paragraph(
                document,
                line,
                alignment=alignment,
                bold=bold,
                underline=underline,
                space_after_pt=style_profile.compact_space_after_pt,
                style_profile=style_profile,
            )
        )
    for line in mention_lines:
        paragraphs.append(
            add_paragraph(
                document,
                line,
                alignment=alignment,
                italic=True,
                space_after_pt=style_profile.compact_space_after_pt,
                style_profile=style_profile,
            )
        )
    return paragraphs


def add_statuts_signature_grid(
    document: Any,
    signers: Sequence[str],
    *,
    mention: str | None = None,
    columns: int = 2,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    if columns <= 0:
        raise ValueError("columns doit etre strictement positif.")
    rows = max(1, (len(signers) + columns - 1) // columns)
    table = document.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    for index, signer in enumerate(signers):
        cell = table.cell(index // columns, index % columns)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if mention:
            mention_run = paragraph.add_run(mention)
            mention_run.italic = True
            paragraph.add_run("\n")
        name_run = paragraph.add_run(signer)
        name_run.bold = True
    add_spacer(document, space_after_pt=style_profile.standard_space_after_pt)
    return table


def add_statuts_annex_heading(
    document: Any,
    title: str,
    subtitle: str | None = None,
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    paragraphs = [
        add_paragraph(
            document,
            title,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            space_before_pt=style_profile.notable_space_before_pt,
            style_profile=style_profile,
        )
    ]
    if subtitle:
        paragraphs.append(
            add_paragraph(
                document,
                subtitle,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                style_profile=style_profile,
            )
        )
    return paragraphs


def add_statuts_matrix_table(
    document: Any,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)
    for cell, header in zip(table.rows[0].cells, headers, strict=True):
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
            paragraph.add_run(value)
    add_spacer(document, space_after_pt=style_profile.standard_space_after_pt)
    return table


def add_spacer(document: Any, *, space_after_pt: int = 0) -> Any:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    return paragraph


def add_framed_title(
    document: Any,
    lines: Sequence[str],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)

    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.bold = True
        run.font.name = style_profile.font_name
        run.font.size = Pt(style_profile.font_size_pt)

    add_spacer(document)
    return table


def add_framed_section_title(
    document: Any,
    text: str,
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)

    paragraph = table.cell(0, 0).paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    run = paragraph.add_run(text)
    run.bold = True
    run.underline = True
    run.font.name = style_profile.font_name
    run.font.size = Pt(style_profile.font_size_pt)
    add_spacer(document, space_after_pt=style_profile.compact_space_after_pt)
    return table


def add_notice_box(
    document: Any,
    lines: Sequence[str],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)
    cell = table.cell(0, 0)
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
        paragraph.add_run(line)
    add_spacer(document, space_after_pt=style_profile.compact_space_after_pt)
    return table


def add_bordered_data_table(
    document: Any,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_borders(table)
    for index, header in enumerate(headers):
        paragraph = table.rows[0].cells[index].paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
        paragraph.add_run(header).bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
            paragraph.add_run(value)
    return table


def add_centered_block(
    document: Any,
    lines: Sequence[tuple[str, bool, bool] | str],
    *,
    space_after_pt: int | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    paragraphs = []
    for line in lines:
        if isinstance(line, str):
            text = line
            bold = False
            italic = False
        else:
            text, bold, italic = line
        paragraphs.append(
            add_paragraph(
                document,
                text,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                bold=bold,
                italic=italic,
                space_after_pt=(
                    style_profile.compact_space_after_pt
                    if space_after_pt is None
                    else space_after_pt
                ),
                style_profile=style_profile,
            )
        )
    return paragraphs


def add_signature_block(
    document: Any,
    lines: Sequence[str],
    *,
    image_path: Path | None = None,
    framed: bool = False,
    width_cm: float | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    if framed:
        return add_framed_signature_block(
            document,
            lines,
            image_path=image_path,
            width_cm=width_cm,
            style_profile=style_profile,
        )
    return add_simple_signature_block(
        document,
        lines,
        image_path=image_path,
        width_cm=width_cm,
        style_profile=style_profile,
    )


def add_simple_signature_block(
    document: Any,
    lines: Sequence[str],
    *,
    image_path: Path | None = None,
    width_cm: float | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    add_spacer(document)
    paragraphs = [
        add_paragraph(
            document,
            line,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            space_after_pt=style_profile.compact_space_after_pt,
            style_profile=style_profile,
        )
        for line in lines
    ]

    signature_paragraph = document.add_paragraph()
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signature_paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
    if image_path is not None:
        if not image_path.exists():
            raise ValueError(f"signature.image_optionnelle est introuvable : {image_path}")
        signature_paragraph.add_run().add_picture(
            str(image_path),
            width=Cm(width_cm or style_profile.signature_image_width_cm),
        )
    else:
        signature_paragraph.add_run("\n\n\n")
    paragraphs.append(signature_paragraph)
    return paragraphs


def add_framed_signature_block(
    document: Any,
    lines: Sequence[str],
    *,
    image_path: Path | None = None,
    width_cm: float | None = None,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    add_spacer(document)
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    _set_table_borders(table)
    cell = table.cell(0, 0)
    cell.width = Cm(width_cm or style_profile.signature_width_cm)
    _add_signature_cell_content(cell, lines, image_path, style_profile)
    return table


def add_signature_lines(
    document: Any,
    names: Sequence[str],
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    bold: bool = False,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> list[Any]:
    return [
        add_paragraph(
            document,
            name,
            alignment=alignment,
            bold=bold,
            style_profile=style_profile,
        )
        for name in names
    ]


def add_signature_table(
    document: Any,
    labels: Sequence[Sequence[str]],
    *,
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> Any:
    if not labels or not labels[0]:
        raise ValueError("labels doit contenir au moins une cellule de signature.")
    column_count = len(labels[0])
    table = document.add_table(rows=len(labels), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)
    for row_index, row in enumerate(labels):
        if len(row) != column_count:
            raise ValueError("Toutes les lignes de signature doivent avoir la meme largeur.")
        for cell_index, label in enumerate(row):
            cell = table.rows[row_index].cells[cell_index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
            paragraph.add_run(label)
            cell.add_paragraph("\n\n\n")
    return table


def add_legal_reminder(
    document: Any,
    *,
    title: str,
    title_suffix: str,
    paragraphs: Sequence[str],
    style_profile: SydelDocxStyleProfile = DEFAULT_STYLE_PROFILE,
) -> None:
    add_spacer(document)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(style_profile.legal_reminder_space_after_pt)
    title_paragraph.style = document.styles["Normal"]
    reminder = title_paragraph.add_run(title)
    reminder.italic = True
    reminder.underline = True
    suffix = title_paragraph.add_run(title_suffix)
    suffix.italic = True

    for text in paragraphs:
        add_paragraph(
            document,
            text,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            italic=True,
            space_after_pt=style_profile.legal_reminder_space_after_pt,
            style_profile=style_profile,
        )


def _add_signature_cell_content(
    cell: Any,
    lines: Sequence[str],
    image_path: Path | None,
    style_profile: SydelDocxStyleProfile,
) -> None:
    first_paragraph = cell.paragraphs[0]
    for index, text in enumerate(lines):
        paragraph = first_paragraph if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(style_profile.compact_space_after_pt)
        paragraph.add_run(text)

    signature_paragraph = cell.add_paragraph()
    signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if image_path is not None:
        if not image_path.exists():
            raise ValueError(f"signature.image_optionnelle est introuvable : {image_path}")
        signature_paragraph.add_run().add_picture(
            str(image_path),
            width=Cm(style_profile.signature_image_width_cm),
        )
    else:
        signature_paragraph.add_run("\n\n\n")


def _set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    existing_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    tbl_pr.append(borders)


def _clear_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    existing_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "FFFFFF")
        borders.append(element)
    tbl_pr.append(borders)
