# ruff: noqa: E501

"""Remplissage générique de modèles .docx tokenisés.

Stratégie de fidélité : on charge le modèle Word d'origine (texte juridique figé)
et on remplace uniquement les tokens `[variable]` par leur valeur. Aucun texte
juridique n'est paraphrasé ni inventé : tout ce qui n'est pas un token reste
exactement tel que dans le modèle source.

Un token peut être réparti sur plusieurs runs Word (Word fragmente parfois un
même mot en plusieurs runs pour des raisons de mise en forme). Le remplacement
se fait donc au niveau du paragraphe : si tous les tokens d'un paragraphe tiennent
chacun dans un seul run, on remplace run par run (préserve au mieux la mise en
forme) ; sinon on fusionne le texte des runs du paragraphe pour pouvoir remplacer
un token éclaté, puis on réécrit le paragraphe sur son premier run.

Sécurité anti-trou : après remplissage, on ré-inspecte le document ; s'il reste un
token `[...]`, on lève ValueError en listant les tokens résiduels (un token oublié
= une erreur visible, jamais un trou silencieux dans un acte juridique).
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from sydel_doc_engine.domain.enums import Gender
from sydel_doc_engine.utils.grammar import apply_gender_pairs

_TOKEN_RE = re.compile(r"\[[^\]\[]+\]")


def fill_docx_template(
    template_path: Path,
    replacements: dict[str, str],
    output_path: Path,
    *,
    gender_pairs: list[tuple[Gender, list[tuple[str, str]]]] | None = None,
) -> Path:
    """Charge un modèle .docx tokenisé, remplace les tokens et sauvegarde.

    - `template_path` : modèle Word d'origine (texte fixe + tokens `[variable]`).
    - `replacements` : dictionnaire `token -> valeur` (le token inclut les crochets,
      ex. `"[denomination_societe]"`). Une clé absente laisse le token en place,
      ce qui déclenche la sécurité anti-trou ci-dessous.
    - `output_path` : chemin du .docx généré (le dossier parent est créé au besoin).
    - `gender_pairs` : optionnel. Liste de couples `(genre, paires)` appliqués
      APRÈS le remplissage des tokens et AVANT la sécurité anti-token-résiduel,
      via `grammar.apply_gender_pairs`. Chaque entrée accorde des chaînes EXACTES
      figées du modèle selon le `genre` de la BONNE personne (signataire, vendeur,
      représentant...). Aucune normalisation magique globale : c'est le générateur
      qui décide quelles paires poser et pour qui.

    Lève `ValueError` si un token `[...]` subsiste après remplissage.
    """
    document = Document(str(template_path))

    for paragraph in _iter_all_paragraphs(document):
        _fill_paragraph(paragraph, replacements)

    if gender_pairs:
        for paragraph in _iter_all_paragraphs(document):
            _apply_gender_pairs_to_paragraph(paragraph, gender_pairs)

    residual = _collect_residual_tokens(document)
    if residual:
        joined = ", ".join(sorted(residual))
        raise ValueError(
            f"Tokens non remplacés dans {template_path.name} : {joined}."
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def _fill_paragraph(paragraph, replacements: dict[str, str]) -> None:
    if "[" not in paragraph.text:
        return

    runs = paragraph.runs
    if not runs:
        return

    # Cas simple : chaque token présent tient dans un unique run -> remplacement
    # run par run (préserve la mise en forme d'origine au plus près).
    if _every_token_within_single_run(paragraph, runs):
        for run in runs:
            text = run.text
            if "[" not in text:
                continue
            for token, value in replacements.items():
                if token in text:
                    text = text.replace(token, value)
            if text != run.text:
                run.text = text
        return

    # Cas d'un token éclaté sur plusieurs runs : on fusionne le texte du paragraphe,
    # on remplace, puis on réécrit tout sur le premier run et on vide les suivants.
    merged = "".join(run.text for run in runs)
    for token, value in replacements.items():
        if token in merged:
            merged = merged.replace(token, value)
    runs[0].text = merged
    for run in runs[1:]:
        run.text = ""


def _apply_gender_pairs_to_paragraph(
    paragraph,
    gender_pairs: list[tuple[Gender, list[tuple[str, str]]]],
) -> None:
    """Accorde en genre les chaînes figées d'un paragraphe (corps + cellules).

    On tente d'abord un accord run par run (préserve la mise en forme). Si une
    forme à accorder est éclatée sur plusieurs runs (le texte du paragraphe
    change alors que les runs pris isolément ne changent pas), on bascule sur un
    accord du texte fusionné réécrit sur le premier run, comme `_fill_paragraph`.
    """
    if not paragraph.runs:
        return

    original_paragraph_text = paragraph.text

    for run in paragraph.runs:
        text = run.text
        for genre, pairs in gender_pairs:
            text = apply_gender_pairs(text, genre, pairs)
        if text != run.text:
            run.text = text

    expected_text = original_paragraph_text
    for genre, pairs in gender_pairs:
        expected_text = apply_gender_pairs(expected_text, genre, pairs)

    # Si l'accord attendu au niveau du paragraphe n'est pas atteint, c'est qu'une
    # forme était éclatée sur plusieurs runs : on réécrit le texte fusionné.
    if paragraph.text != expected_text:
        runs = paragraph.runs
        runs[0].text = expected_text
        for run in runs[1:]:
            run.text = ""


def _every_token_within_single_run(paragraph, runs) -> bool:
    """Vrai si chaque INSTANCE de token tient entièrement dans un run.

    On compare le nombre d'occurrences de chaque token dans le texte complet du
    paragraphe au nombre d'occurrences réparties dans les runs individuels. Si un
    token apparaît plusieurs fois et qu'une seule de ses instances est éclatée sur
    plusieurs runs (cas réel du modèle : `[ville_siege] ... [ville_siege]`), les
    deux comptes diffèrent -> on bascule sur le remplacement fusionné.
    """
    tokens = set(_TOKEN_RE.findall(paragraph.text))
    full_text = paragraph.text
    for token in tokens:
        occurrences_in_text = full_text.count(token)
        occurrences_in_runs = sum(run.text.count(token) for run in runs)
        if occurrences_in_runs != occurrences_in_text:
            return False
    return True


def _iter_all_paragraphs(document):
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _collect_residual_tokens(document) -> set[str]:
    residual: set[str] = set()
    for paragraph in _iter_all_paragraphs(document):
        for match in _TOKEN_RE.findall(paragraph.text):
            residual.add(match)
    return residual
